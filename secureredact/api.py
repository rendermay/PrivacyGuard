"""
SecureRedact 业务 API 层 (PR-C4 v1.1.14)

对外门面模块,7 个核心函数覆盖「扫描 / 执行脱敏 / 命中过滤 / 文档哈希 / 批量替换」。

设计原则:
- 不破坏向后兼容:GUI 老路径 `from secureredact import OCRWorker` 仍可用
- API 层接受/返回 `dict`,不暴露 PyQt6 类型(QRectF 等)
- 实现走 `secureredact/` 子包 wrapper 模式,`from main import` = 0
- 100% 单测覆盖,便于 CLI / 自动化场景复用

来源:`docs/planning/gui-api-refactor-plan.md` §2.3 任务 1.1 / 1.2
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, List, Optional

# 主链路:直接走 secureredact/ 子包,无临时 main.py 依赖
# 注:QThread 子类 worker + QEventLoop 都有 PyQt6 DLL 依赖;为让 api.py 在无 Qt 环境
# 也可 import,所有 PyQt6 符号 + WordBatchReplaceWorker 全部 lazy(函数体内 import)
from secureredact.redaction.doc_hash import compute_doc_hash as _compute_doc_hash_impl
from secureredact.redaction.override_store import HitOverrideStore
from secureredact.utils.exceptions import WorkerCancelledError
from secureredact.ocr import text_pdf as _text_pdf_module


__all__ = [
    "compute_doc_hash",
    "scan_pdf",
    "scan_word",
    "redact_pdf",
    "redact_word",
    "filter_hits_by_overrides",
    "batch_redact_word",
]


# === 1. 文档哈希 ===
def compute_doc_hash(file_path: str | Path) -> str:
    """8 位文档标识,基于路径+size+mtime,与 `secureredact.redaction.doc_hash` 对齐。

    Args:
        file_path: 文档路径,接受 `str` 或 `pathlib.Path`。内部用 `os.fspath()` 归一化。

    Returns:
        8 位 hex 字符串。

    Raises:
        OSError: 文件不存在或无 stat 权限。
    """
    return _compute_doc_hash_impl(os.fspath(file_path))


# === 2. PDF 命中扫描(不执行脱敏,只产 hit 列表)===
def scan_pdf(
    pdf_path: str | Path,
    *,
    rules: Dict[str, Any],
    custom_keywords: str = "",
    options: Optional[Dict[str, Any]] = None,
) -> Dict[int, List[Dict[str, Any]]]:
    """逐页扫描 PDF,产出文本层命中(简化版,不含 image-block OCR 合并)。

    Args:
        pdf_path: PDF 路径,接受 `str` 或 `pathlib.Path`。
        rules: 规则字典(键为规则名,值为正则 pattern)。也接受 `{"patterns": [...]}` 形式。
        custom_keywords: 用户自定义关键词(空格分隔,扫描时附加到 patterns)。
        options: 可选行为控制(详见 plan §2.7 Options 契约)。当前未使用。

    Returns:
        {page_num(0-based int): [{"rect": (x, y, w, h) tuple,
          "source": str, "text": str, "rule_name": str}, ...]}

    Raises:
        FileNotFoundError: PDF 文件不存在。
    """
    import fitz  # 延迟导入避免非 PDF 场景的依赖

    pdf_path = os.fspath(pdf_path)
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    # 提取规则 patterns:支持 dict[str,pattern] 或 dict 含 patterns key
    patterns = []
    if isinstance(rules, dict):
        if "patterns" in rules:
            patterns.extend(rules["patterns"])
        else:
            patterns.extend(p for p in rules.values() if isinstance(p, str))
    if custom_keywords:
        import re as _re
        for kw in custom_keywords.split():
            if kw.strip():
                patterns.append(_re.escape(kw.strip()))

    # 逐页扫描
    result: Dict[int, List[Dict[str, Any]]] = {}
    with fitz.open(pdf_path) as doc:
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            hits = []
            for x0, y0, w, h, text, rule_name in _text_pdf_module.collect_text_pdf_hit_boxes(
                page, patterns
            ):
                hits.append({
                    "rect": (x0, y0, w, h),
                    "source": "rule",
                    "text": text,
                    "rule_name": rule_name,
                })
            if hits:
                result[page_idx] = hits
    return result


# === 3. Word 命中扫描(方案 B:一次性同步)===
def scan_word(
    word_path: str | Path,
    *,
    rules: Dict[str, Any],
    custom_keywords: str = "",
    replacement_text: str = "[已脱敏]",
    options: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """扫描 Word 段落,产出 matches 列表。

    Args:
        word_path: Word 文档路径(.docx)。
        rules: 规则列表(每项为 dict,含 `find`/`mode`/`enabled` 字段)。
        custom_keywords: 用户自定义关键词(空格分隔)。
        replacement_text: 统一替换文本(传给 build_word_rule_matches)。
        options: 当前未使用。

    Returns:
        [{start, end, text, source, rule_name, replacement}, ...]
        (每个 match 是 build_word_rule_matches 的输出格式)

    Raises:
        FileNotFoundError: Word 文件不存在。
    """
    from docx import Document
    from secureredact.redaction.word_rules import (
        build_word_rule_matches,
        normalize_word_replace_rules,
    )

    word_path = os.fspath(word_path)
    if not os.path.isfile(word_path):
        raise FileNotFoundError(f"Word 文件不存在: {word_path}")

    # 归一化 rules 为 build_word_rule_matches 接受的格式
    normalized = normalize_word_replace_rules(
        list(rules) if isinstance(rules, (list, tuple)) else rules,
        replacement_text,
    )

    # 合并自定义关键词到 normalized(转为简单 exact 规则)
    if custom_keywords:
        for kw in custom_keywords.split():
            if kw.strip():
                normalized.append({"find": kw.strip(), "mode": "exact", "enabled": True})

    doc = Document(word_path)
    matches: List[Dict[str, Any]] = []
    cursor = 0
    for para in doc.paragraphs:
        text = "".join(run.text for run in para.runs)
        if not text:
            continue
        para_matches = build_word_rule_matches(text, normalized, replacement_text)
        for m in para_matches:
            # 平移 start/end 到全文档坐标
            if "start" in m and "end" in m:
                m = dict(m)
                m["start"] = m["start"] + cursor
                m["end"] = m["end"] + cursor
            matches.append(m)
        cursor += len(text)
    return matches


# === 4. PDF 一站式脱敏 ===
def redact_pdf(
    pdf_path: str | Path,
    output_path: str | Path,
    *,
    rules: Dict[str, Any],
    custom_keywords: str = "",
    doc_hash: Optional[str] = None,
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """完整 PDF 脱敏链路:扫描 → 黑遮罩写入 → 输出。

    Args:
        pdf_path: 输入 PDF 路径。
        output_path: 输出 PDF 路径(若存在则覆盖)。
        rules: 规则字典。
        custom_keywords: 用户自定义关键词。
        doc_hash: 文档哈希(用于 override 关联,当前未使用)。
        options: 当前未使用。

    Returns:
        {"output": str, "pages": int, "hits": int, "elapsed_sec": float}

    Raises:
        FileNotFoundError: 输入或输出路径不存在。
    """
    import time as _time

    import fitz

    pdf_path = os.fspath(pdf_path)
    output_path = os.fspath(output_path)
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"输入 PDF 不存在: {pdf_path}")

    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.isdir(output_dir):
        raise FileNotFoundError(f"输出目录不存在: {output_dir}")

    t0 = _time.time()
    hits_by_page = scan_pdf(
        pdf_path,
        rules=rules,
        custom_keywords=custom_keywords,
        options=options,
    )
    total_hits = sum(len(h) for h in hits_by_page.values())

    # 用 PyMuPDF 的 redact annotation 做遮罩
    with fitz.open(pdf_path) as doc:
        for page_idx, hits in hits_by_page.items():
            page = doc[page_idx]
            for hit in hits:
                x, y, w, h = hit["rect"]
                rect = fitz.Rect(x, y, x + w, y + h)
                page.add_redact_annot(rect, text="", fill=(0, 0, 0))
            page.apply_redactions()
        doc.save(output_path)

    return {
        "output": output_path,
        "pages": len(hits_by_page) if hits_by_page else 0,
        "hits": total_hits,
        "elapsed_sec": _time.time() - t0,
    }


# === 5. Word 一站式脱敏 ===
def redact_word(
    word_path: str | Path,
    output_path: str | Path,
    *,
    rules: Dict[str, Any],
    custom_keywords: str = "",
    replacement_text: str = "[已脱敏]",
    doc_hash: Optional[str] = None,
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """完整 Word 替换链路。

    Args:
        word_path: 输入 Word 路径(.docx)。
        output_path: 输出 Word 路径(若存在则覆盖)。
        rules: 规则列表。
        custom_keywords: 用户自定义关键词。
        replacement_text: 统一替换文本。
        doc_hash: 文档哈希(用于 override 关联,当前未使用)。
        options: 当前未使用。

    Returns:
        {"output": str, "hits": int, "elapsed_sec": float}

    Raises:
        FileNotFoundError: 输入或输出路径不存在。
    """
    import time as _time

    from docx import Document
    from secureredact.redaction.word_rules import (
        build_word_rule_matches,
        normalize_word_replace_rules,
        replace_matches_in_paragraph,
    )

    word_path = os.fspath(word_path)
    output_path = os.fspath(output_path)
    if not os.path.isfile(word_path):
        raise FileNotFoundError(f"输入 Word 不存在: {word_path}")

    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.isdir(output_dir):
        raise FileNotFoundError(f"输出目录不存在: {output_dir}")

    t0 = _time.time()
    normalized = normalize_word_replace_rules(
        list(rules) if isinstance(rules, (list, tuple)) else rules,
        replacement_text,
    )
    if custom_keywords:
        for kw in custom_keywords.split():
            if kw.strip():
                normalized.append({"find": kw.strip(), "mode": "exact", "enabled": True})

    doc = Document(word_path)
    total_hits = 0
    for para in doc.paragraphs:
        text = "".join(run.text for run in para.runs)
        if not text:
            continue
        matches = build_word_rule_matches(text, normalized, replacement_text)
        if matches:
            total_hits += len(matches)
            replace_matches_in_paragraph(para, matches, text_offset=0,
                                         fallback_replacement_text=replacement_text)
    doc.save(output_path)

    return {
        "output": output_path,
        "hits": total_hits,
        "elapsed_sec": _time.time() - t0,
    }


# === 6. 命中 override 过滤 ===
def filter_hits_by_overrides(
    hits: List[Dict[str, Any]],
    *,
    location: str,
    doc_hash: Optional[str] = None,
    doc_path: Optional[str | Path] = None,
) -> List[Dict[str, Any]]:
    """包装 `HitOverrideStore.instance().filtered_hits()`。

    Args:
        hits: 待过滤的命中列表。
        location: 命中位置(`f"page_{i}"` / `f"paragraph_{idx}"` 等)。
        doc_hash: 文档 8 位哈希。若提供,优先使用;否则从 `doc_path` 计算。
        doc_path: 文档路径。仅当 `doc_hash` 未提供时生效。

    Raises:
        ValueError: `doc_hash` 和 `doc_path` 都未提供时。
    """
    if doc_hash is None:
        if doc_path is None:
            raise ValueError("doc_hash or doc_path required")
        doc_hash = compute_doc_hash(doc_path)
    store = HitOverrideStore.instance()
    return store.filtered_hits(list(hits), location=location, doc_hash=doc_hash)


# === 7. 批量 Word 替换(同步化 wrapper)===
def batch_redact_word(
    word_paths: List[str | Path],
    output_dir: str | Path,
    *,
    rules: Dict[str, Any],
    custom_keywords: str = "",
    replacement_text: str = "[已脱敏]",
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """批量替换多份 Word,同步返回结果。

    Args:
        word_paths: 输入 Word 文档路径列表。
        output_dir: 输出目录(必须存在,否则抛 `FileNotFoundError`)。
        rules: 规则字典。
        custom_keywords: 用户自定义关键词(空格分隔,当前未使用)。
        replacement_text: 统一替换文本。
        options: 可选行为控制(详见 plan §2.7 Options 契约)。

    Returns:
        {
            "success": [{"input": str, "output": str, "total_replacements": int,
                          "rule_counts": [...]}],
            "failed": [{"input": str, "error": str}],
            "stopped": bool,
            "total": int
        }

    Raises:
        FileNotFoundError: `output_dir` 不存在。
        WorkerCancelledError: 批处理超时(`options.timeout_sec` 默认 300s)。
    """
    # Lazy import:QThread 子类 + PyQt6 依赖,延迟到函数调用时
    from PyQt6.QtCore import QEventLoop, QTimer  # noqa: F401
    from secureredact.workers.word_batch_replace_worker import WordBatchReplaceWorker

    output_dir = os.fspath(output_dir)
    if not os.path.isdir(output_dir):
        raise FileNotFoundError(f"output_dir 不存在: {output_dir}")

    # 解析 options(plan §2.7 Options 契约)
    timeout_sec = 300
    if options and "timeout_sec" in options:
        try:
            timeout_sec = int(options["timeout_sec"])
        except (TypeError, ValueError) as e:
            raise ValueError(f"options.timeout_sec 必须为 int,得到 {options['timeout_sec']!r}") from e

    # 路径统一化
    normalized_paths = [os.fspath(p) for p in word_paths]

    # 实例化 worker
    worker = WordBatchReplaceWorker(
        file_paths=normalized_paths,
        rules=list(rules) if isinstance(rules, (list, tuple)) else rules,
        default_replacement_text=replacement_text,
    )

    # 启动 worker + QEventLoop 等待 finished_signal
    # finished_signal.emit(dict) 是 worker.run() 在最后 emit 的 summary
    collected: Dict[str, Any] = {}
    loop = QEventLoop()

    def _on_finished(summary_dict):
        collected.update(summary_dict)
        loop.quit()

    worker.finished_signal.connect(_on_finished)
    worker.start()

    # REVIEWS-v2 C-2 修复:批处理某文件出错时,worker 线程阻塞在
    # _wait_for_error_decision,主线程阻塞在 QEventLoop → 死锁。
    # 解决:用 QTimer 每 5 秒主动提供 "skip" 决策(兜底防卡死)。
    # worker 内部 _wait_for_error_decision 看到 _decision_event.set() 后
    # 取决策继续;无出错时该调用被 worker 内部忽略(决策读不到)。
    auto_skip_timer = QTimer()
    auto_skip_timer.setInterval(5000)

    def _auto_provide_skip():
        if worker.isRunning():
            worker.provide_error_decision("skip")

    auto_skip_timer.timeout.connect(_auto_provide_skip)
    auto_skip_timer.start()

    # REVIEWS-v2 C-3 修复:文件粒度超时兜底(timeout_sec ≥ N × 30s)
    QTimer.singleShot(timeout_sec * 1000, loop.quit)
    loop.exec()

    auto_skip_timer.stop()

    # 等 worker 线程真正结束
    if worker.isRunning():
        worker.wait(5000)

    # 组装返回 summary(plan §2.3 任务 1.1 schema)
    if not collected:
        # QTimer 兜底超时退出,worker 仍在跑(被中断);collected 为空
        raise WorkerCancelledError(
            f"batch_redact_word timeout {timeout_sec}s (file-grained, see plan §2.3 任务 1.2 C-3)"
        )

    return {
        "total": int(collected.get("total", len(normalized_paths))),
        "success": list(collected.get("success", [])),
        "failed": list(collected.get("failed", [])),
        "stopped": bool(collected.get("stopped", False)),
    }