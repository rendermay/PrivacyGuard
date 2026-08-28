"""对话框/Worker 模块 — 从 main.py 迁出 (PR-B3/B4)

公共 API 不变。MainWindow 通过 `from main import X` 或对应模块 re-export 导入。
"""
from __future__ import annotations

import sys
import os
import re
import time
import shutil
import threading
import traceback
from pathlib import Path
from io import BytesIO

from PIL import Image
from bs4 import BeautifulSoup

# PyQt6 — 完整集合,覆盖所有 5 个新模块需要
from PyQt6.QtCore import (
    Qt, QSize, QPoint, QPointF, QRect, QRectF, QTimer, QThread, QObject,
    QUrl, QStandardPaths, pyqtSignal, pyqtSlot,
)
from PyQt6.QtGui import (
    QAction, QBrush, QColor, QFont, QFontMetrics, QIcon, QImage, QKeySequence,
    QPainter, QPalette, QPen, QPixmap, QStandardItem, QStandardItemModel,
    QTextCursor, QTextDocument, QWheelEvent,
)
from PyQt6.QtWidgets import (
    QAbstractItemView, QApplication, QCheckBox, QComboBox, QCompleter,
    QDialog, QDialogButtonBox, QDoubleSpinBox, QFileDialog, QFormLayout,
    QFrame, QGridLayout, QGroupBox, QHBoxLayout, QHeaderView, QInputDialog,
    QLabel, QLineEdit, QListView, QListWidget, QListWidgetItem, QMainWindow,
    QMenu, QMenuBar, QMessageBox, QPlainTextEdit, QProgressBar, QPushButton,
    QRadioButton, QScrollArea, QScrollBar, QSizePolicy, QSlider, QSpinBox,
    QSplitter, QStackedWidget, QStatusBar, QStyle, QTabWidget, QTableWidget,
    QTableWidgetItem, QTextBrowser, QTextEdit, QToolBar, QToolButton,
    QTreeView, QTreeWidget, QTreeWidgetItem, QVBoxLayout, QWidget,
)

# 项目内 import(对话框类可能用到的)
from secureredact.utils.security import validate_safe_path, resource_path
from secureredact.utils.exceptions import PrivacyAppError
from secureredact.redaction.hit_ref import HitRef
class WordBatchReplaceWorker(QThread):
    """Word 批量替换线程（同一套规则应用全部文件）。"""

    progress_signal = pyqtSignal(int, int, str)      # processed, total, current_file
    file_done_signal = pyqtSignal(str, str)          # input_path, output_path
    file_error_signal = pyqtSignal(int, str, str)    # index, input_path, error_msg
    finished_signal = pyqtSignal(dict)               # summary

    def __init__(self, file_paths, rules, default_replacement_text):
        super().__init__()
        self.file_paths = list(file_paths or [])
        self.rules = normalize_word_replace_rules(rules or [], default_replacement_text)
        self.default_replacement_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
        self._error_decision = None
        self._decision_event = threading.Event()
        self._decision_lock = threading.Lock()
        self._temp_dirs = []

    def provide_error_decision(self, decision):
        """主线程调用：为当前错误设置决策（skip/stop）。"""
        with self._decision_lock:
            self._error_decision = decision
        self._decision_event.set()

    def _wait_for_error_decision(self):
        """等待主线程选择错误处理策略。"""
        self._decision_event.clear()
        while not self._decision_event.wait(0.1):
            if self.isInterruptionRequested():
                return "stop"
        with self._decision_lock:
            decision = self._error_decision or "skip"
            self._error_decision = None
        return decision

    def run(self):
        summary = {
            "total": len(self.file_paths),
            "processed": 0,
            "success": [],
            "failed": [],
            "stopped": False,
            "rules": list(self.rules),
        }

        try:
            total = len(self.file_paths)
            for idx, file_path in enumerate(self.file_paths):
                if self.isInterruptionRequested():
                    summary["stopped"] = True
                    break

                current_name = os.path.basename(file_path)
                self.progress_signal.emit(idx, total, current_name)

                try:
                    output_path, replace_stats = self._process_single_file(file_path)
                    summary["success"].append({
                        "input": file_path,
                        "output": output_path,
                        "total_replacements": max(0, int(replace_stats.get("total_replacements", 0) or 0))
                        if isinstance(replace_stats, dict) else 0,
                        "rule_counts": replace_stats.get("rule_counts", []) if isinstance(replace_stats, dict) else [],
                    })
                    self.file_done_signal.emit(file_path, output_path)
                except (OSError, IOError, RuntimeError, ValueError, ConversionError, PermissionError) as e:
                    error_msg = str(e)
                    summary["failed"].append({
                        "input": file_path,
                        "error": error_msg
                    })
                    self.file_error_signal.emit(idx, file_path, error_msg)
                    decision = self._wait_for_error_decision()
                    if decision == "stop":
                        summary["stopped"] = True
                        break
                finally:
                    self.progress_signal.emit(idx + 1, total, current_name)

            summary["processed"] = len(summary["success"]) + len(summary["failed"])
            self.finished_signal.emit(summary)
        finally:
            self._cleanup_temp_dirs()

    def _process_single_file(self, file_path):
        from docx import Document

        is_safe, error_msg = validate_safe_path(file_path, allowed_extensions=[".doc", ".docx"])
        if not is_safe:
            raise ConversionError("输入文件路径不安全", error_msg)

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            source_docx = file_path
        elif ext == ".doc":
            source_docx = self._convert_doc_to_docx(file_path)
        else:
            raise ConversionError("不支持的文件格式", f"{file_path}")

        doc = Document(source_docx)
        replace_stats = self._apply_rules_to_document(doc)
        output_path = self._build_output_path(file_path)
        doc.save(output_path)
        return output_path, replace_stats

    def _apply_rules_to_document(self, doc):
        rule_counts = {}
        total_replacements = 0

        def _collect_match_counts(matches):
            nonlocal total_replacements
            if not matches:
                return
            total_replacements += len(matches)
            for match in matches:
                try:
                    rule_index = int(match.get("rule_index", -1))
                except (TypeError, ValueError, AttributeError):
                    continue
                if rule_index < 0:
                    continue
                rule_counts[rule_index] = rule_counts.get(rule_index, 0) + 1

        for para in doc.paragraphs:
            text = ''.join(run.text for run in para.runs)
            if not text:
                continue
            matches = build_word_rule_matches(text, self.rules, self.default_replacement_text)
            if matches:
                _collect_match_counts(matches)
                replace_matches_in_paragraph(para, matches, text_offset=0,
                                             fallback_replacement_text=self.default_replacement_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text or ""
                    if not cell_text:
                        continue
                    cell_matches = build_word_rule_matches(cell_text, self.rules, self.default_replacement_text)
                    if not cell_matches:
                        continue
                    _collect_match_counts(cell_matches)

                    para_offset = 0
                    paragraphs = list(cell.paragraphs)
                    for idx, para in enumerate(paragraphs):
                        original_para_len = len(''.join(run.text for run in para.runs))
                        replace_matches_in_paragraph(
                            para,
                            cell_matches,
                            text_offset=para_offset,
                            fallback_replacement_text=self.default_replacement_text
                        )
                        para_offset += original_para_len
                        if idx < len(paragraphs) - 1:
                            para_offset += 1

        return {
            "total_replacements": total_replacements,
            "rule_counts": [
                {"rule_index": rule_index, "count": count}
                for rule_index, count in sorted(rule_counts.items())
            ],
        }

    def _build_output_path(self, file_path):
        base_path = os.path.splitext(file_path)[0]
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        candidate = f"{base_path}__replaced_{timestamp}.docx"
        if not os.path.exists(candidate):
            return candidate

        suffix = 1
        while True:
            candidate_with_suffix = f"{base_path}__replaced_{timestamp}_{suffix}.docx"
            if not os.path.exists(candidate_with_suffix):
                return candidate_with_suffix
            suffix += 1

    def _convert_doc_to_docx(self, doc_path):
        """v1.1.11: 委托给共享转换模块。"""
        docx_path, temp_dir = _shared_convert_doc_to_docx(doc_path)
        self._temp_dirs.append(temp_dir)
        return docx_path

    def _cleanup_temp_dirs(self):
        for temp_dir in list(self._temp_dirs):
            try:
                if os.path.isdir(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except (OSError, IOError):
                pass
        self._temp_dirs.clear()
