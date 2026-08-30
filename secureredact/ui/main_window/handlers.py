"""
UI 处理函数 mixin — MainWindow 大型处理方法 (PR-B2.8 迁出)

PR-B2.8 目标:推进 main.py < 5000 行。
10 个 MainWindow 内大型 UI 处理方法(共 822 行)整体搬迁到独立 mixin 模块:
- 高亮类:_highlight_sensitive_info / _highlight_exact_match
- 文件保存:_save_word / save_pdf
- 生命周期:_cleanup_before_open / _set_ui_mode / _show_doc_install_guide
- OCR + 键盘:start_ocr / keyPressEvent
- Word 打开:_open_word_docx

物理迁移策略(同 PR-B2.6 / B2.7 模式):逐字搬迁,逻辑零改动。

依赖 MainWindow 上的属性(由 __init__ / setup_ui 初始化):
    - self.btn_* / self.lbl_* / self.word_doc / self.doc
    - self.theme / self.Theme / self.app_state / self.file_path
    - self._current_doc_hash / self._override_store / self.ocr_engine
"""
from __future__ import annotations

import fitz  # PyMuPDF

from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont, QKeyEvent
from PyQt6.QtWidgets import (
    QApplication, QFileDialog, QFrame, QHBoxLayout, QLabel, QMenu, QMessageBox,
    QPushButton, QWidget,
)


class MainWindowHandlersMixin:
    """MainWindow 大型 UI 处理方法集 — 高亮 / 保存 / 生命周期 / OCR / 键盘 / Word 打开。

    方法签名与实现与原 MainWindow 内一致,直接被 MainWindow 继承使用。
    """

    def _cleanup_before_open(self):
        """v1.1.11: 打开新文档前的完整资源清理

        解决问题：
        - 打开新文档时卡顿/未响应
        - 文件选择窗口内容不显示
        """
        self._invalidate_word_scroll_sync()
        # 1. 停止并等待活跃的 worker 线程
        if self.active_worker and self.active_worker.isRunning():
            self.active_worker.requestInterruption()
            # 等待线程结束（最多 2 秒）
            self.active_worker.wait(2000)
            self.active_worker = None

        # 2. 清理 QWebEngineView（Word 预览）
        # QWebEngineView 占用大量资源，需要正确清理
        if hasattr(self, 'word_preview') and self.word_preview:
            try:
                # 停止加载
                self.word_preview.stop()
                # 清空内容
                self.word_preview.setHtml('')
                # 隐藏
                self.word_preview.hide()
            except Exception as e:
                print(f"[清理] 清理 word_preview 时出错: {e}")
        if hasattr(self, 'word_preview_replaced') and self.word_preview_replaced:
            try:
                self.word_preview_replaced.stop()
                self.word_preview_replaced.setHtml('')
                self.word_preview_replaced.hide()
            except Exception as e:
                print(f"[清理] 清理 word_preview_replaced 时出错: {e}")

        # 3. 关闭 PDF 文档
        if self.doc:
            try:
                self.doc.close()
            except Exception as e:
                print(f"[清理] 关闭 PDF 文档时出错: {e}")
            self.doc = None

        # 4. 重置状态变量
        self.word_doc = None
        self.word_data = {}
        self._reset_word_preview_cache()
        self.page_data = {}
        self._ocr_processed_pages = set()
        self.current_page = None
        self.doc_type = None
        self.file_path = None
        self.word_compare_mode = False
        self.word_compare_user_hidden = False
        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self._reset_batch_session_state()
        self._clear_info_bar_message()

        # 5. v1.1.11: 重置 canvas 显示状态（不删除固定实例）
        # canvas_left 和 canvas_right 是固定实例，在 setup_ui() 中创建
        # 只需清除显示内容，不需要删除
        # v1.1.11: 修复属性名错误 - 使用正确的 rects_manual 和 rects_ocr
        if hasattr(self, 'canvas_left') and self.canvas_left:
            try:
                # 检查 C++ 对象是否仍然有效
                _ = self.canvas_left.size()  # 如果对象已删除，这里会抛出异常
                self.canvas_left.clear()  # 清除显示
                self.canvas_left.page_index = 0  # 重置页面索引
                self.canvas_left.rects_manual = []  # 清除手动脱敏区域（正确的属性名）
                self.canvas_left.rects_ocr = []  # 清除 OCR 区域（正确的属性名）
            except RuntimeError:
                print("[清理] canvas_left 的 C++ 对象已被删除，跳过清理")

        if hasattr(self, 'canvas_right') and self.canvas_right:
            try:
                _ = self.canvas_right.size()
                self.canvas_right.clear()
                self.canvas_right.page_index = 1
                self.canvas_right.rects_manual = []
                self.canvas_right.rects_ocr = []
            except RuntimeError:
                print("[清理] canvas_right 的 C++ 对象已被删除，跳过清理")

        if hasattr(self, 'word_compare_container') and self.word_compare_container:
            self.word_compare_container.hide()
        if hasattr(self, 'canvas_container') and self.canvas_container:
            self.canvas_container.show()

        # 6. 处理待处理的 Qt 事件，确保 UI 响应
        QApplication.processEvents()
        self._sync_ui_mode()

        print("[清理] 打开新文档前的资源清理完成")

    def _set_ui_mode(self, mode):
        """集中管理不同业务模式下的工具显隐，避免 PDF / Word 控件混杂。"""
        self.current_ui_mode = mode
        is_idle = mode == "idle"
        is_pdf = mode == "pdf"
        is_word = mode == "word"
        is_batch = mode == "batch"
        is_image_merge = mode == "image_merge"

        pdf_widgets = [
            self.rb_black, self.rb_white, self.cb_dual, self.btn_fit_utility,
            self.lbl_zoom, self.btn_zoom_out, self.btn_zoom_in,
            self.btn_go_first, self.btn_prev_page, self.lbl_page,
            self.btn_next_page, self.btn_go_last,
        ]
        for widget in pdf_widgets:
            widget.setVisible(is_pdf)
        self.btn_fit.setVisible(False)

        word_widgets = [self.btn_compare_toggle]
        for widget in word_widgets:
            widget.setVisible(is_word)

        show_scan = is_pdf or is_word
        show_save = is_pdf or is_word
        self.btn_open.setVisible(not is_idle)
        self.btn_scan.setVisible(show_scan)
        self.btn_save.setVisible(show_save)

        if is_batch:
            self.btn_scan.setVisible(False)
            self.btn_save.setVisible(False)

        if is_idle:
            self.btn_scan.setVisible(False)
            self.btn_save.setVisible(False)

        if hasattr(self, "idle_workspace_container"):
            self.idle_workspace_container.setVisible(is_idle)
        if hasattr(self, "batch_workspace_container"):
            self.batch_workspace_container.setVisible(is_batch)
        if hasattr(self, "merge_workspace_container"):
            self.merge_workspace_container.setVisible(is_image_merge)
        if hasattr(self, "canvas_container"):
            self.canvas_container.setVisible(is_pdf)
        if hasattr(self, "word_compare_container"):
            self.word_compare_container.setVisible(is_word)
        if hasattr(self, "workbench_panel"):
            self.workbench_panel.setVisible(not is_idle)
        if hasattr(self, "toolbar"):
            self.toolbar.setVisible(not is_idle)
        if hasattr(self, "progress_shell"):
            self.progress_shell.setVisible(not is_idle)

        self._refresh_mode_badge()
        self._refresh_word_compare_toggle()
        self._refresh_toolbar_responsiveness()
        self._refresh_workbench_context()
        self._refresh_info_bar_visibility()

    def _open_word_docx(self, fname):
        """打开 DOCX 文件"""
        try:
            from docx import Document

            # 打开新文件时失效缓存，避免复用旧 HTML
            self._reset_word_preview_cache()
            self.image_merge_in_progress = False
            self.image_merge_total_images = 0
            self._reset_batch_session_state()
            self.word_compare_mode = False
            self.word_compare_user_hidden = False
            self.word_doc = Document(fname)
            self.file_path = fname
            self.doc_type = 'docx'
            self.doc = None  # 清空 PDF 文档对象
            self.page_data = {}

            # 初始化 word_data 结构
            self.word_data = {}
            for idx, para in enumerate(self.word_doc.paragraphs):
                self.word_data[f'paragraph_{idx}'] = {
                    'type': 'paragraph',
                    'index': idx,
                    'text': para.text,
                    'ocr': [],
                    'manual': []
                }

            # 扫描表格
            for table_idx, table in enumerate(self.word_doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        key = f'table_{table_idx}_cell_{row_idx}_{cell_idx}'
                        self.word_data[key] = {
                            'type': 'table_cell',
                            'table': table_idx,
                            'row': row_idx,
                            'cell': cell_idx,
                            'text': cell.text,
                            'ocr': [],
                            'manual': []
                        }

            # 启用按钮
            self.btn_scan.setEnabled(True)
            self.btn_save.setEnabled(True)

            # 根据规则与现有脱敏结果决定是否默认显示双栏预览
            self.word_compare_mode = self._has_word_replacement_candidates()
            self._sync_ui_mode()

            # 显示 HTML 预览
            self.render_word_preview()

            self._clear_info_bar_message()

        except (IOError, OSError, ValueError, KeyError) as e:
            QMessageBox.critical(self, "错误", f"打开 Word 文档失败: {str(e)}")

    def _show_doc_install_guide(self):
        """显示 .doc 转换工具安装指南（v1.1.11 新增）"""
        import platform
        system = platform.system()

        if system == 'Darwin':  # macOS
            guide = """
<h3>安装 LibreOffice（推荐）</h3>
<p>在终端执行：</p>
<code style="background:#f5f5f5;padding:8px;display:block;">brew install --cask libreoffice</code>

<h4 style="margin-top:16px;">或使用轻量级方案 antiword</h4>
<p>在终端执行：</p>
<code style="background:#f5f5f5;padding:8px;display:block;">brew install antiword</code>
<p style="color:#666;margin-top:8px;">注：antiword 只能提取纯文本，会丢失格式</p>
"""
        elif system == 'Windows':
            guide = """
<h3>安装 LibreOffice（推荐）</h3>
<p>请从官网下载安装：</p>
<a href="https://www.libreoffice.org/download/download/">https://www.libreoffice.org/download/</a>
<p style="margin-top:8px;">选择 Windows 版本下载并安装</p>

<h4 style="margin-top:16px;">安装后重启本软件即可</h4>
"""
        else:  # Linux
            guide = """
<h3>安装 LibreOffice（推荐）</h3>
<p>根据您的发行版执行：</p>
<code style="background:#f5f5f5;padding:8px;display:block;">
# Debian/Ubuntu
sudo apt install libreoffice

# Fedora
sudo dnf install libreoffice

# Arch Linux
sudo pacman -S libreoffice
</code>

<h4 style="margin-top:16px;">或使用轻量级方案 antiword</h4>
<code style="background:#f5f5f5;padding:8px;display:block;">
# Debian/Ubuntu
sudo apt install antiword

# Fedora
sudo dnf install antiword
</code>
<p style="color:#666;margin-top:8px;">注：antiword 只能提取纯文本，会丢失格式</p>
"""

        msg = QMessageBox(self)
        msg.setWindowTitle("缺少 .doc 转换工具")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(guide)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.setStyleSheet("""
            QMessageBox {
                min-width: 450px;
            }
            QLabel {
                min-width: 400px;
            }
        """)
        msg.exec()

    def keyPressEvent(self, event):
        """键盘快捷键处理（v1.1.11 新增）

        快捷键列表：
        - PageUp: 上一页
        - PageDown: 下一页
        - Home: 首页
        - End: 尾页
        - Space/Shift+Space: 翻页
        - Ctrl/Cmd + +/-: 缩放
        - Ctrl/Cmd + 0: 适应页面
        """
        key = event.key()
        modifiers = event.modifiers()

        # PageUp: 上一页
        if key == Qt.Key.Key_PageUp:
            if modifiers == Qt.KeyboardModifier.ShiftModifier:
                self.change_page(-5)  # Shift+PageUp 快速翻页
            else:
                self.change_page(-1)
        # PageDown: 下一页
        elif key == Qt.Key.Key_PageDown:
            if modifiers == Qt.KeyboardModifier.ShiftModifier:
                self.change_page(5)  # Shift+PageDown 快速翻页
            else:
                self.change_page(1)
        # Home: 首页
        elif key == Qt.Key.Key_Home:
            self.go_first()
        # End: 尾页
        elif key == Qt.Key.Key_End:
            self.go_last()
        # Space: 下一页（Shift+Space: 上一页）
        elif key == Qt.Key.Key_Space:
            if modifiers == Qt.KeyboardModifier.ShiftModifier:
                self.change_page(-1)
            else:
                self.change_page(1)
        # Ctrl/Cmd + Plus: 放大
        elif key == Qt.Key.Key_Plus or key == Qt.Key.Key_Equal:
            if modifiers in [Qt.KeyboardModifier.ControlModifier, Qt.KeyboardModifier.MetaModifier]:
                self.zoom_in()
            else:
                super().keyPressEvent(event)
        # Ctrl/Cmd + Minus: 缩小
        elif key == Qt.Key.Key_Minus:
            if modifiers in [Qt.KeyboardModifier.ControlModifier, Qt.KeyboardModifier.MetaModifier]:
                self.zoom_out()
            else:
                super().keyPressEvent(event)
        # Ctrl/Cmd + 0: 重置缩放
        elif key == Qt.Key.Key_0:
            if modifiers in [Qt.KeyboardModifier.ControlModifier, Qt.KeyboardModifier.MetaModifier]:
                self.zoom_level = 1.0
                self.render_view()
            else:
                super().keyPressEvent(event)
        else:
            super().keyPressEvent(event)

    def start_ocr(self):
        """智能扫描 - 支持 PDF 和 Word（v1.1.11: 增强错误处理）"""
        from main import DEFAULT_RULES, DEFAULT_RULES_META, config, OCRWorker, WordWorker  # PR-B5.2: 延迟导入
        # 线程安全检查：防止重复启动
        if self.active_worker is not None:
            if self.active_worker.isRunning():
                QMessageBox.warning(self, "提示", "正在处理中，请稍候...")
                return

        self._set_info_bar_message("🔍 正在扫描敏感信息...")
        self.btn_scan.setEnabled(False)
        self.btn_cancel_scan.setVisible(True)  # 显示取消按钮（v1.1.11）
        self.btn_cancel_scan.setEnabled(True)
        self.active_task_type = "scan"

        # v1.1.12: 按 doc_type 过滤规则(USCC 等仅作用于 Word 路径,完全隔离 PDF)
        # region USCC_ISO_FILTER
        try:
            pdf_excluded_names = config.get("redaction.pdf_excluded_rules", []) if config else []
        except Exception:
            pdf_excluded_names = []
        if isinstance(pdf_excluded_names, list) and pdf_excluded_names:
            excluded_patterns = set()
            for rule_name in pdf_excluded_names:
                pat = DEFAULT_RULES.get(rule_name, "")
                if pat:
                    excluded_patterns.add(pat)
            pdf_rules = [r for r in self.active_rules if r not in excluded_patterns]
        else:
            pdf_rules = self.active_rules
        # endregion USCC_ISO_FILTER

        # PDF 处理
        if self.doc:
            # v1.1.11: 检测是否启用印章检测(用原始 self.active_rules,不被 pdf_rules 过滤影响)
            seal_detection_enabled = "__SEAL_DETECTION__" in self.active_rules
            print(f"[OCR] active_rules: {self.active_rules}")
            print(f"[OCR] pdf_rules: {pdf_rules}")
            print(f"[OCR] 印章检测启用: {seal_detection_enabled}")
            self._ocr_processed_pages = set()
            # v1.1.11: 只使用 RapidOCR，移除 use_char_level_ocr 参数
            # v1.1.12: PDF 路径使用过滤后的 pdf_rules,排除仅 Word 规则
            # v1.1.14: 注入 name_context_extra_tokens (config.json →
            # redaction.name_context.extra_tokens),让 PDF 图片通道 OCR 出来的
            # '甲方/乙方/原告/... 与 A、B、C 之间' 类并列名单中的姓名也能被识别.
            _name_ctx_extra = config.get(
                "redaction.name_context.extra_tokens", [],
            ) or []  # JSON 值为 null 时 SimpleConfig.get 返回 None → 归一化为 []
            self.worker = OCRWorker(self.file_path, pdf_rules, self.use_enhance, self.custom_keywords,
                                    self.scan_level, self.offset_x, self.offset_w,
                                    seal_detection_enabled=seal_detection_enabled,
                                    enable_name_recognition=self.enable_name_recognition,
                                    name_context_extra_tokens=_name_ctx_extra)
            self.active_worker = self.worker  # 追踪线程
            self.worker.progress_signal.connect(self.progress.setValue)
            # v1.1.11: connect 到 _receive_page_hits(新签名,接 list[dict])
            # 旧 _on_ocr_page_result 仍存在但已废弃,仅作 QRectF 路径的 fallback
            self.worker.page_result_signal.connect(self._receive_page_hits)
            # v1.1.11: 连接错误信号
            self.worker.error_signal.connect(self._on_ocr_error)
            # 先连接原有的完成处理，再连接清理
            self.worker.finished_signal.connect(self._on_ocr_finished_safe)
            self.worker.finished_signal.connect(self._on_worker_finished)
            self.worker.start()
        # Word 处理
        elif self.word_doc:
            # v1.1.12: Word 路径使用完整 self.active_rules(包含 USCC 等)
            # v1.1.14: 注入 name_context_extra_tokens (config.json →
            # redaction.name_context.extra_tokens),让 '甲方/乙方/原告/... 与 A、B、C 之间'
            # 类并列名单中的姓名也能被识别.
            _name_ctx_extra = config.get(
                "redaction.name_context.extra_tokens", [],
            ) or []  # JSON 值为 null 时 SimpleConfig.get 返回 None → 归一化为 []
            self.worker = WordWorker(self.word_doc, self.word_data, self.active_rules,
                                     self.custom_keywords, self.replacement_text,
                                     enable_name_recognition=self.enable_name_recognition,
                                     default_rules=DEFAULT_RULES,
                                     default_rules_meta=DEFAULT_RULES_META,
                                     name_context_extra_tokens=_name_ctx_extra)
            self.active_worker = self.worker  # 追踪线程
            self.worker.progress_signal.connect(self.progress.setValue)
            # 先连接原有的完成处理，再连接清理
            self.worker.finished_signal.connect(self.word_scan_finished)
            self.worker.finished_signal.connect(self._on_worker_finished)
            self.worker.start()

    def _highlight_exact_match(self, html, match):
        """使用 BeautifulSoup 精确高亮指定位置的文本

        Args:
            html: HTML 字符串
            match: 包含 key, start, end, text 的匹配信息

        Returns:
            修改后的 HTML 字符串
        """
        from bs4 import BeautifulSoup, NavigableString

        key = match['key']
        start = match['start']
        end = match['end']
        text = match['text']

        try:
            soup = BeautifulSoup(html, 'html.parser')

            # 找到对应的容器元素
            container = soup.find(attrs={'data-key': key})
            if not container:
                print(f"[警告] 未找到 data-key={key} 的元素")
                return html

            # 遍历所有文本节点，找到对应位置
            current_pos = 0
            for child in list(container.descendants):
                if isinstance(child, NavigableString) and not isinstance(child, str):
                    continue
                if isinstance(child, str):
                    node_text = str(child)
                    node_start = current_pos
                    node_end = current_pos + len(node_text)

                    # 检查目标范围是否在这个节点内
                    if start >= node_start and end <= node_end:
                        # 计算在当前节点内的相对位置
                        rel_start = start - node_start
                        rel_end = end - node_start

                        # 分割文本并插入 mark 标签
                        before = node_text[:rel_start]
                        highlighted = node_text[rel_start:rel_end]
                        after = node_text[rel_end:]

                        # 创建新的标记
                        mark_tag = soup.new_tag('mark')
                        mark_tag['class'] = 'manual-highlight'
                        mark_tag['data-key'] = key
                        mark_tag['data-start'] = start
                        mark_tag['data-end'] = end
                        mark_tag['style'] = 'background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone;'
                        mark_tag.string = highlighted

                        # 替换原文本节点
                        parent = child.parent
                        # 创建新的节点序列
                        new_nodes = []
                        if before:
                            new_nodes.append(NavigableString(before))
                        new_nodes.append(mark_tag)
                        if after:
                            new_nodes.append(NavigableString(after))

                        # 替换原节点
                        child.replace_with(*new_nodes)

                        return str(soup)

                    current_pos = node_end

            print(f"[警告] 精确模式未找到位置: key={key}, start={start}, end={end}")
            return html

        except (ImportError, AttributeError, TypeError, ValueError, IndexError) as e:
            print(f"[错误] 精确高亮失败: {e}")
            return html

    def _highlight_sensitive_info(self, html):
        """在 HTML 中高亮显示敏感信息（使用 JavaScript 进行高亮）"""
        import json
        from html import escape
        import re

        # 构建所有匹配数据的 JSON
        matches_data = []
        # 构建文本块数据（用于标记段落/单元格）
        text_blocks = {}

        for key, data in self.word_data.items():
            text = data['text']
            if not text:
                continue

            # 记录文本块信息
            text_blocks[key] = {
                'text': text,
                'escaped': escape(text)
            }

            # OCR 匹配
            for match in data['ocr']:
                matches_data.append({
                    'key': key,
                    'start': match['start'],
                    'end': match['end'],
                    'text': match['text'],
                    'type': 'ocr',
                    'rule_name': match.get('rule_name', 'OCR')
                })

            # 手动脱敏
            for match in data['manual']:
                matches_data.append({
                    'key': key,
                    'start': match['start'],
                    'end': match['end'],
                    'text': match['text'],
                    'type': 'manual',
                    'mode': match.get('mode', 'exact')  # 添加模式标识，默认为 exact
                })

        # === 新方法：直接在 HTML 字符串中进行高亮替换 ===
        from html import escape as html_escape

        # 按文本长度降序排序，先处理长的匹配
        matches_data.sort(key=lambda x: len(x['text']), reverse=True)

        # 分离精确模式和全局模式
        exact_matches = [m for m in matches_data if m.get('mode') == 'exact']
        global_matches = [m for m in matches_data if m.get('mode') != 'exact']

        # === 第一步：为所有文本块添加 data-key 属性（使用 BeautifulSoup 替代正则表达式）===
        html = self._add_data_key_attributes(html, text_blocks)

        # 处理全局模式匹配（OCR 和全局手动脱敏）
        # 去重：相同文本只处理一次
        processed_global_texts = set()
        for match in global_matches:
            text = match['text']
            if text in processed_global_texts:
                continue
            processed_global_texts.add(text)

            # 需要转义 HTML 特殊字符
            escaped_text = html_escape(text)
            is_ocr = match['type'] == 'ocr'

            # 构建替换标记
            if is_ocr:
                replacement = f'<mark class="ocr-highlight" data-key="{match["key"]}" data-start="{match["start"]}" data-end="{match["end"]}" title="{match.get("rule_name", "")}" style="background-color: #ffeb3b; color: #000; display: inline; box-decoration-break: clone; -webkit-box-decoration-break: clone;">{escaped_text}</mark>'
            else:
                replacement = f'<mark class="manual-highlight" data-key="{match["key"]}" data-start="{match["start"]}" data-end="{match["end"]}" style="background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone;">{escaped_text}</mark>'

            # 使用正则表达式替换所有匹配
            # 使用正向预查避免匹配已经在标记中的文本
            pattern = re.compile(re.escape(escaped_text) + r'(?![^<]*>)')
            html = pattern.sub(replacement, html)

        # 处理精确模式匹配（只高亮特定位置的文本）
        # 使用 BeautifulSoup 进行精确位置定位
        for match in exact_matches:
            html = self._highlight_exact_match(html, match)

        # 为了兼容性，仍然生成 matches_json（但不再用于高亮）
        matches_json = json.dumps(matches_data, ensure_ascii=False)
        text_blocks_json = json.dumps(text_blocks, ensure_ascii=False)

        # 添加简化版的 JavaScript 高亮脚本（仅用于调试）
        highlight_script = '''
        <script>
        (function() {
            console.log('[Highlight] HTML 预高亮已完成');
            // 高亮已在 Python 端完成，这里不需要再做
        })();
        </script>
        '''

        # 添加样式
        style = f"""
        <style>
            body {{ font-family: {PREVIEW_FONT_STACK}; padding: 20px; line-height: 1.6; }}
            p:empty {{ display: none; margin: 0; }}
            mark.ocr-highlight {{ background-color: #ffeb3b; color: #000; display: inline; box-decoration-break: clone; -webkit-box-decoration-break: clone; }}
            mark.manual-highlight {{ background-color: #ff6b6b; color: #fff; display: inline; cursor: pointer; box-shadow: 0 0 0 1px #e03131; box-decoration-break: clone; -webkit-box-decoration-break: clone; }}
            mark.manual-highlight:hover {{ box-shadow: 0 0 0 1px #e03131, 0 0 4px rgba(225, 49, 49, 0.5); }}
            table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; }}
        </style>
        """

        # 插入样式和脚本
        if '<head>' in html:
            html = html.replace('<head>', '<head>' + style, 1)
        else:
            html = style + html

        # 在 </body> 前插入脚本，或者直接追加到末尾
        if '</body>' in html:
            html = html.replace('</body>', highlight_script + '</body>', 1)
        else:
            html = html + highlight_script

        return html

    def save_pdf(self):
        """保存脱敏后的文档 - 支持 PDF 和 Word - v1.1.11: 安全加固，脱敏区域永久化"""
        # v36: 应用文件对话框样式
        # v1.1.11: PDF 脱敏安全加固 - 脱敏区域永久嵌入，不可编辑
        app = QApplication.instance()
        original_style = app.styleSheet()

        # PDF 保存
        if self.doc:
            app.setStyleSheet(self._get_file_dialog_style())
            try:
                fname, _ = QFileDialog.getSaveFileName(self, "保存 PDF", "", "PDF Files (*.pdf)")
            finally:
                app.setStyleSheet(original_style)

            if fname:
                doc_save = None
                try:
                    QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
                    doc_save = fitz.open(self.file_path)
                    fill_col = (0, 0, 0) if self.current_color.name() == "#000000" else (1, 1, 1)

                    for i in range(len(doc_save)):
                        page = doc_save[i]

                        # v1.1.11: 修复内部编辑功能 - 使用副本避免修改原始数据
                        # 从 page_data 中获取脱敏区域列表
                        # v1.1.11: 走 _rects_for_page 应用 store 过滤(已 ignore 剔除)
                        ocr_list = self._rects_for_page(i)
                        manual_list = self.page_data[i].get('manual', [])

                        # 1. 添加脱敏注释
                        # v1.1.11: 重建 QRectF 确保不修改原始对象
                        for r in ocr_list + manual_list:
                            # 从 QRectF 提取坐标并重建，避免引用问题
                            x, y, w, h = r.x(), r.y(), r.width(), r.height()
                            rect = fitz.Rect(x, y, x + w, y + h)
                            annot = page.add_redact_annot(rect)
                            annot.set_colors(stroke=fill_col, fill=fill_col)
                            annot.update()

                        # v1.1.11: 安全加固 - 修改图像像素，彻底销毁原始内容
                        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)

                        # v1.1.11: 安全加固 - 删除所有注释对象，防止被 PDF 编辑器修改
                        for annot in page.annots():
                            page.delete_annot(annot)

                    # v1.1.11: 安全加固 - 使用垃圾回收和压缩彻底删除未引用对象
                    doc_save.save(
                        fname,
                        garbage=4,        # 最大垃圾回收级别
                        deflate=True,     # 压缩内容流
                        clean=True        # 清理未引用对象
                    )

                    QApplication.restoreOverrideCursor()
                    QMessageBox.information(self, "成功", f"文件已安全保存至：\n{fname}")
                except (IOError, OSError, ValueError, RuntimeError) as e:
                    QApplication.restoreOverrideCursor()
                    QMessageBox.critical(self, "失败", str(e))
                except Exception as e:
                    # v1.1.11: 兜底捕获所有异常 (含 pymupdf.mupdf.FzErrorSystem 等非标准异常)
                    QApplication.restoreOverrideCursor()
                    err_msg = str(e)
                    if "Permission denied" in err_msg or "cannot remove" in err_msg:
                        err_msg += "\n\n可能原因: 目标 PDF 正在被其他程序 (PDF 阅读器/浏览器) 占用.\n请关闭后再试, 或另存为新文件名."
                    QMessageBox.critical(self, "失败", err_msg)
                finally:
                    if doc_save:
                        doc_save.close()
        # Word 保存
        elif self.word_doc:
            app.setStyleSheet(self._get_file_dialog_style())
            try:
                fname, _ = QFileDialog.getSaveFileName(self, "保存 Word", "", "Word 文档 (*.docx)")
            finally:
                app.setStyleSheet(original_style)

            if fname:
                self._save_word(fname)

    def _save_word(self, fname):
        """保存 Word 文档 - v24 改进版：详细错误处理 + 使用 TempFileManager + 合并 OCR 和 Manual 脱敏

        v1.1.11: 导出前先走 HitOverrideStore.filtered_hits 把 ignored OCR hit
        从 ocr_matches 中剔除(manual 永远保留;confirm 保留)。
        """
        try:
            import shutil
            from docx import Document

            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)

            # 使用 TempFileManager 管理临时文件
            temp_file = self.temp_manager.create_temp_file()
            shutil.copy2(self.file_path, temp_file)

            # 打开副本进行修改
            new_doc = Document(temp_file)

            store = self._override_store
            doc_hash = self._current_doc_hash

            # 遍历段落进行 run 级别的文本替换
            for para_idx, para in enumerate(new_doc.paragraphs):
                key = f'paragraph_{para_idx}'
                if key in self.word_data:
                    data = self.word_data[key]
                    source_text = data.get("text", "")
                    # v1.1.11: 导出前 store 过滤 (manual 永远保留)
                    filtered_ocr = store.filtered_hits(
                        list(data.get("ocr", [])),
                        location=key,
                        doc_hash=doc_hash,
                    )
                    merged_matches = merge_word_matches_with_priority(
                        source_text,
                        self.word_replace_rules,
                        self.replacement_text,
                        manual_matches=data.get("manual", []),
                        ocr_matches=filtered_ocr
                    )
                    if merged_matches:
                        replace_matches_in_paragraph(
                            para,
                            merged_matches,
                            text_offset=0,
                            fallback_replacement_text=self.replacement_text
                        )

            # 遍历表格进行 run 级别的文本替换
            for table_idx, table in enumerate(new_doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        key = f'table_{table_idx}_cell_{row_idx}_{cell_idx}'
                        if key in self.word_data:
                            data = self.word_data[key]
                            source_text = data.get("text", "")
                            # v1.1.11: 导出前 store 过滤 (manual 永远保留)
                            filtered_ocr = store.filtered_hits(
                                list(data.get("ocr", [])),
                                location=key,
                                doc_hash=doc_hash,
                            )
                            merged_matches = merge_word_matches_with_priority(
                                source_text,
                                self.word_replace_rules,
                                self.replacement_text,
                                manual_matches=data.get("manual", []),
                                ocr_matches=filtered_ocr
                            )

                            if merged_matches:
                                # 处理单元格内的所有段落（按 cell.text 的偏移映射）
                                para_offset = 0
                                paragraphs = list(cell.paragraphs)
                                for idx, para in enumerate(paragraphs):
                                    original_para_len = len(''.join(run.text for run in para.runs))
                                    replace_matches_in_paragraph(
                                        para,
                                        merged_matches,
                                        text_offset=para_offset,
                                        fallback_replacement_text=self.replacement_text
                                    )
                                    para_offset += original_para_len
                                    if idx < len(paragraphs) - 1:
                                        # python-docx 的 cell.text 使用换行拼接段落
                                        para_offset += 1

            # 保存文档
            new_doc.save(fname)

            QApplication.restoreOverrideCursor()
            QMessageBox.information(self, "成功", f"文件已保存至：\n{fname}")

        except PermissionError:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self, "保存失败",
                f"没有写入权限：\n{fname}\n\n"
                "建议：\n"
                "1. 检查文件是否被其他程序打开\n"
                "2. 检查文件夹权限\n"
                "3. 尝试保存到其他位置"
            )

        except (OSError, IOError, RuntimeError, ValueError, KeyError, AttributeError) as e:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(
                self, "保存失败",
                f"保存 Word 文档时出错：\n{str(e)}\n\n"
                "请尝试：\n"
                "1. 重启应用\n"
                "2. 在 Word 中手动打开文件\n"
                "3. 导出错误日志以供分析"
            )
