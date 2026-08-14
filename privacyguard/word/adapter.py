"""Phase 3 Word 文档 → PII 引擎 TextUnit 流适配器（D-04 / D-17）。"""
from typing import Dict, List, Tuple


class WordAdapter:
    """Word 文档 → TextUnit 流适配器（D-04 / D-17）。

    collect_units 返回 (units, key_index)：
    - units: List[TextUnit]，每个非空段落 / 非空表格单元格 = 一个 unit（page_index 单调递增）
    - key_index: Dict[int, str]，page_index → main.py::_open_word_docx 中的 key
      （paragraph_{idx} 或 table_{t}_cell_{r}_{c}，与既有 word_data 命名严格对齐）
    """

    @staticmethod
    def collect_units(docx_path: str) -> Tuple[List, Dict[int, str]]:
        """收集 docx 文件的段落 + 表格单元格文本，返回 (TextUnit 列表, key_index 映射)。

        函数体内 lazy import `docx.Document` 与 `privacyguard.pii.hits.TextUnit`，
        保持 OPS-03 懒加载纪律（import privacyguard.word 不拉起 python-docx / PII 引擎）。
        """
        from docx import Document
        from privacyguard.pii.hits import TextUnit

        doc = Document(docx_path)
        units: List = []
        key_index: Dict[int, str] = {}
        idx = 0

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text or ''
            if not text.strip():
                continue
            units.append(TextUnit(page_index=idx, text=text, source='text'))
            key_index[idx] = f'paragraph_{para_idx}'
            idx += 1

        for table_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text or ''
                    if not text.strip():
                        continue
                    units.append(TextUnit(page_index=idx, text=text, source='text'))
                    key_index[idx] = f'table_{table_idx}_cell_{r_idx}_{c_idx}'
                    idx += 1

        return units, key_index


__all__ = ['WordAdapter']
