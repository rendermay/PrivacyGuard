# ⚠️ v1.1.13+ (PR-B0) deprecation notice
# ====================================================================
# 本文件已转为 thin shim,真正的运行时入口在 `secureredact/main.py`。
#
# 禁止继续往本文件添加新业务代码。所有新功能请写到 `secureredact/` 对应子包。
# 阶段 B 重构路线图:详见 `frontend-refactor-plan.md`(OpenDesign 工作目录)
#                   与 `docs/refactor/b0-report.md`(项目内)。
# PR-B5 收口时本文件将被彻底移除,所有打包入口同步切换到 `secureredact.main:main`。
#
# 过渡期启动方式:
#     python -m secureredact.main    # 新入口(推荐)
#     python main.py                  # 兼容 shim(过渡期,PR-B5 后停用)
# ====================================================================

import sys
import os
import fitz  # PyMuPDF
import re
import cv2
import numpy as np
import time
import shutil
import threading  # v1.1.11: 线程锁支持
import atexit  # v1.1.11: 用于确保临时文件清理
import tempfile  # v1.1.11: 临时文件管理
import traceback  # v1.1.11: 异常追踪
from pathlib import Path
from io import BytesIO
from PIL import Image
from bs4 import BeautifulSoup
from secureredact.ocr.mixed_pdf import (
    collect_embedded_image_clip_rects,
    collect_image_block_ocr_hits,
)
from secureredact.ocr.text_pdf import collect_text_pdf_hit_boxes
from secureredact.utils.security import validate_safe_path, resource_path
from secureredact.utils.exceptions import (
    PrivacyAppError,
    ConversionError,
    FileFormatError,
    SecurityError,
    MemoryLimitError,
    WorkerCancelledError,
)
from secureredact.utils.temp_manager import TempFileManager
from secureredact.workers.image_merge import ImageMergeWorker
from secureredact.workers.word_worker import WordWorker as _ModularWordWorker
from secureredact.workers.ocr_worker import OCRWorker as _ModularOCRWorker
from secureredact.utils.doc_converter import convert_doc_to_docx as _shared_convert_doc_to_docx
from secureredact.redaction.hit_ref import HitRef  # v1.1.11: 人工干预
from secureredact.ui.main_window import SinglePageCanvas, WebViewBridge  # PR-B2.0 re-export
from secureredact.ui.main_window.toolbar import MainWindowToolbarMixin  # PR-B2.1
from secureredact.ui.main_window.workbench import MainWindowWorkbenchMixin  # PR-B2.2
from secureredact.ui.main_window.word_preview import MainWindowWordPreviewMixin  # PR-B2.3
from secureredact.ui.main_window.pdf_render import MainWindowPdfRenderMixin  # PR-B2.4
from secureredact.ui.main_window.batch_replace import MainWindowBatchReplaceMixin  # PR-B2.5
from secureredact.ui.main_window.density import MainWindowDensityMixin  # PR-B2.6
from secureredact.ui.main_window.setup_ui import MainWindowSetupMixin  # PR-B2.7
from secureredact.ui.main_window.handlers import MainWindowHandlersMixin  # PR-B2.8
from secureredact.redaction.override_store import HitOverrideStore  # v1.1.11: override store 单例
from secureredact.redaction.doc_hash import compute_doc_hash  # v1.1.11: 文档 hash
from secureredact.redaction.black_white_list_store import BlackWhiteListStore  # v1.1.11: 黑/白名单 store

# v1.1.11: 延迟导入 OCR 模块，便于错误处理
RapidOCR = None
OCR_INIT_ERROR = None

def init_ocr_engine():
    """v1.1.11: 安全初始化 OCR 引擎，捕获所有可能的错误"""
    global RapidOCR, OCR_INIT_ERROR
    if RapidOCR is not None:
        return True

    try:
        from rapidocr_onnxruntime import RapidOCR as _RapidOCR
        RapidOCR = _RapidOCR
        # 预热：创建一个测试实例验证 DLL 加载
        _ = _RapidOCR()
        print("[OCR] 引擎初始化成功")
        return True
    except ImportError as e:
        OCR_INIT_ERROR = f"OCR 模块未安装: {e}\n请运行: pip install rapidocr-onnxruntime"
        print(f"[OCR ERROR] {OCR_INIT_ERROR}")
        return False
    except OSError as e:
        OCR_INIT_ERROR = f"OCR DLL 加载失败: {e}\n可能缺少 Visual C++ 运行库"
        print(f"[OCR ERROR] {OCR_INIT_ERROR}")
        return False
    except Exception as e:
        OCR_INIT_ERROR = f"OCR 初始化失败: {type(e).__name__}: {e}"
        print(f"[OCR ERROR] {OCR_INIT_ERROR}")
        traceback.print_exc()
        return False

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QBoxLayout, QGridLayout,
                             QHBoxLayout, QPushButton, QLabel, QFileDialog,
                             QScrollArea, QMessageBox, QProgressBar, QFrame,
                             QDialog, QCheckBox, QGroupBox, QTextEdit, QSpinBox,
                             QRadioButton, QButtonGroup, QComboBox, QSizePolicy,
                             QTextBrowser, QLineEdit, QListWidget, QListWidgetItem,
                             QAbstractItemView, QSlider, QTableWidget, QMenu,
                             QTableWidgetItem, QHeaderView, QStyle,
                             QDockWidget)
from PyQt6.QtGui import QPixmap, QImage, QPainter, QPen, QColor, QWheelEvent, QCursor, QIcon, QDesktopServices
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QRectF, QPointF, QSettings, QMutex, QMutexLocker, QObject, pyqtSlot, QSize, QTimer, QUrl
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebChannel import QWebChannel
from PyQt6 import sip

# 导入主题系统
from theme import Theme

# v1.1.11: 简化配置系统 - 直接从 JSON 文件加载
import json


def read_app_version():
    """从统一版本文件读取基础版本号。"""
    version_file = Path(__file__).resolve().parent / "version.txt"
    try:
        return version_file.read_text(encoding="utf-8").strip()
    except OSError:
        return "1.1.12"

class SimpleConfig:
    """简化配置管理器 - 直接从 config.json 读取"""

    def __init__(self, config_path=None):
        self._config = {}
        if config_path is None:
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.json')
        self._config_path = config_path
        self.load()

    @staticmethod
    def _deep_merge(base: dict, override: dict) -> dict:
        """v1.1.12: 深度合并 — override 中有的字段覆盖 base, base 中独有字段保留。

        用于 SimpleConfig.save() 字段保护, 避免磁盘上其他版本/扩展字段被擦除。
        对 dict 类型递归合并, 非 dict 类型(标量/列表)直接覆盖。
        """
        result = dict(base)  # 浅拷贝作为基线
        for k, v in override.items():
            if k in result and isinstance(result[k], dict) and isinstance(v, dict):
                result[k] = SimpleConfig._deep_merge(result[k], v)
            else:
                result[k] = v
        return result

    def load(self):
        """加载配置文件并补齐默认键。

        v1.1.12: 额外从 DEFAULT_CONFIG(代码内嵌)合并缺失的 default_rules 字段,
        避免 disk 配置被外部操作擦除 mask_* 等扩展字段时丢失功能。
        """
        try:
            if os.path.exists(self._config_path):
                with open(self._config_path, 'r', encoding='utf-8') as f:
                    self._config = json.load(f)
        except (OSError, IOError, json.JSONDecodeError) as e:
            print(f"[配置系统] 加载配置失败: {e}")

        # v1.1.11: 补齐 override 相关默认键
        red = self._config.setdefault("redaction", {})
        red.setdefault("enable_hit_override", True)
        overrides = red.setdefault("overrides", {})
        overrides.setdefault("permanent", [])

        # v1.1.12: 字段保护 — 强制覆盖内置规则的 mask_* 字段(用代码内嵌的 DEFAULT_CONFIG)
        # 原因:disk 上的 mask 字段可能被外部操作(测试脚本/版本控制)擦除,
        # 但代码内嵌的 DEFAULT_CONFIG 是 v1.1.12 单一权威源, 始终以它为准。
        try:
            from secureredact.utils.config import DEFAULT_CONFIG as _DEFAULT_CONFIG
            disk_rules = self._config.setdefault("redaction", {}).setdefault("default_rules", {})
            default_rules = _DEFAULT_CONFIG.get("redaction", {}).get("default_rules", {})
            for rule_name, default_meta in default_rules.items():
                if rule_name not in disk_rules:
                    # 整条规则缺失,补齐整条
                    disk_rules[rule_name] = default_meta
                elif isinstance(default_meta, dict) and isinstance(disk_rules[rule_name], dict):
                    # 规则存在
                    # 强制覆盖 mask_* 字段(代码内嵌为准) — 已被外部擦除的 mask 字段
                    for k, v in default_meta.items():
                        if k.startswith("mask_") and k in default_meta:
                            disk_rules[rule_name][k] = v
                    # 强制覆盖 pattern 字段(代码内嵌为准) — 已被外部擦除的 pattern
                    # 关键: 若 disk pattern 字符类与 code 不一致(如缺少 - / \w / \s*),
                    #       强制用代码内嵌的最新 pattern
                    if "pattern" in default_meta and "pattern" in disk_rules[rule_name]:
                        code_pat = default_meta["pattern"]
                        disk_pat = disk_rules[rule_name]["pattern"]
                        if code_pat != disk_pat:
                            # 简化判断: code pattern 应比 disk 新(更长/字符类更丰富)
                            # 若 code 含 disk 没有的字符(去掉转义后), 用 code
                            code_normalized = code_pat.replace("\\", "").replace("\\\\", "")
                            disk_normalized = disk_pat.replace("\\", "").replace("\\\\", "")
                            if len(code_normalized) > len(disk_normalized):
                                disk_rules[rule_name]["pattern"] = code_pat

            # v1.1.12: whitelist 字段保护 — 强制 append 代码内嵌默认白名单缺失的项
            # 原因:disk 上的 whitelist 可能被外部操作擦除,导致法律主体标签(甲方/乙方/丙方/丁方/戊方)
            # 被 jieba X3 误识为人名后无法被白名单保护
            default_whitelist = _DEFAULT_CONFIG.get("redaction", {}).get("whitelist", [])
            disk_whitelist = self._config.setdefault("redaction", {}).setdefault("whitelist", [])
            for item in default_whitelist:
                if item not in disk_whitelist:
                    disk_whitelist.append(item)
        except Exception as _exc:
            print(f"[配置系统] 强制覆盖 mask 字段失败: {_exc}")

    def _load_config(self):
        """[兼容保留] 旧版加载入口,委托给 load()."""
        self.load()

    def save(self):
        """原子写回磁盘 — 先写 .tmp,再 os.replace 原子替换.

        v1.1.12: 保护磁盘上 self._config 没有的字段(如 mask_* 等其他版本/扩展字段),
        避免在 SettingsDialog 保存时把它们擦除。流程:
          1. 从磁盘读 current_disk_config(若存在)
          2. 深度合并:self._config 中已有的字段覆盖 disk 字段
          3. disk 中有但 self._config 没有的字段保留
          4. 把合并结果写回磁盘
        这样既保证用户当前改动生效,又不会丢字段。
        """
        tmp_path = self._config_path + ".tmp"
        try:
            # v1.1.12: 字段保护 — 从磁盘读 + 合并
            try:
                if os.path.exists(self._config_path):
                    with open(self._config_path, 'r', encoding='utf-8') as f:
                        disk_config = json.load(f)
                else:
                    disk_config = {}
            except (OSError, IOError, json.JSONDecodeError):
                disk_config = {}

            # 深度合并:self._config 覆盖 disk_config,disk_config 中有但 self._config 没有的字段保留
            merged = self._deep_merge(disk_config, self._config)
            with open(tmp_path, 'w', encoding='utf-8') as f:
                json.dump(merged, f, indent=2, ensure_ascii=False)
            os.replace(tmp_path, self._config_path)
            return True
        except (OSError, IOError, TypeError) as e:
            print(f"[配置系统] 保存配置失败: {e}")
            # 清理可能残留的 .tmp
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except OSError:
                pass
            return False

    def get(self, key, default=None):
        """获取配置值（支持点分隔路径）"""
        keys = key.split('.')
        value = self._config
        for k in keys:
            if isinstance(value, dict) and k in value:
                value = value[k]
            else:
                return default
        return value

    def set(self, key, value, persist=True):
        """设置配置值（支持点分隔路径）

        Args:
            key: 配置键，支持点分隔路径
            value: 配置值
            persist: 是否立即保存到文件
        """
        keys = key.split('.')
        config = self._config
        # 遍历到倒数第二层，创建缺失的字典
        for k in keys[:-1]:
            if k not in config:
                config[k] = {}
            elif not isinstance(config[k], dict):
                config[k] = {}
            config = config[k]
        # 设置最终值
        config[keys[-1]] = value

        # 保存到文件
        if persist:
            self.save()

    def get_redaction_rules(self):
        """获取脱敏规则"""
        return self.get('redaction.default_rules', {})

# 初始化配置
config = SimpleConfig()

# === 核心防崩溃设置 ===
cv2.setNumThreads(0)
os.environ["OMP_NUM_THREADS"] = "1"

# === 软件配置 ===
# v1.1.11: 从配置读取，失败时使用硬编码后备
APP_NAME = config.get("app.name", "SecureRedact 信息脱敏助手") if config else "SecureRedact 信息脱敏助手"
APP_VERSION = read_app_version()
VERSION = APP_VERSION
WORD_PREVIEW_IMAGE_EXTENSION_MAP = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/tiff": ".tiff",
    "image/bmp": ".bmp",
    "image/webp": ".webp",
    "image/svg+xml": ".svg",
}
WORD_PREVIEW_BROKEN_IMAGE_DATA_URI = (
    "data:image/gif;base64,R0lGODlhAQABAIAAAAAAAP///ywAAAAAAQABAAACAUwAOw=="
)

# === 常量定义 ===
# v1.1.11: 从配置读取，失败时使用硬编码后备
MIN_RECT_WIDTH = config.get("ocr.min_rect_width", 5) if config else 5
PROGRESS_UPDATE_INTERVAL = config.get("ocr.progress_update_interval", 0.05) if config else 0.05
ZOOM_MIN = config.get("ocr.zoom_min", 0.5) if config else 0.5
ZOOM_MAX = config.get("ocr.zoom_max", 4.0) if config else 4.0
DEBUG_MODE = os.getenv('PRIVACYGUARD_DEBUG', 'False').lower() == 'true' if not config else config.get("advanced.debug_mode", False)

# === 默认规则库 + Mask 元数据 ===
# v1.1.11: 从配置读取,支持新旧两种格式
# v1.1.12: 同时构建 DEFAULT_RULES_META,提供每条规则的 mask_mode / mask_keep_prefix / mask_keep_suffix / mask_char
DEFAULT_RULES = {}
DEFAULT_RULES_META = {}


def _v113_apply_rule_overrides():
    """v1.1.13 强制代码层默认 — 不依赖磁盘 config.json 漂移.

    修复场景: 即使磁盘 config.json 中 '日期时间.enabled=true' 或
    '法定代表人.pattern' 被改回旧版, 代码层仍强制以下行为:
      1) '日期时间' 规则强制禁用 (用户决策: 选 B 不脱敏, 过度脱敏风险高于收益)
      2) '法定代表人' pattern 强制覆盖为带正向 lookahead 的版本,
         防止贪婪匹配把 '继续主张'/'继承' 等普通动词当人名 mask
    """
    DEFAULT_RULES.pop("日期时间", None)
    DEFAULT_RULES_META.pop("日期时间", None)
    DEFAULT_RULES["法定代表人"] = (
        r"法定代表人\s*[::：]?\s*[一-龥]{2,4}(?:·[一-龥]{2,4})?"
        r"(?=[的之及与和按于在跟同向对为由被让等,，。；;）)\]】\s]|$)"
    )


if config:
    _rules_from_config = config.get_redaction_rules()
    for name, rule in _rules_from_config.items():
        if isinstance(rule, dict):
            DEFAULT_RULES[name] = rule.get("pattern", "")
            DEFAULT_RULES_META[name] = {
                "mask_mode": rule.get("mask_mode", "default"),
                "mask_keep_prefix": int(rule.get("mask_keep_prefix", 0) or 0),
                "mask_keep_suffix": int(rule.get("mask_keep_suffix", 0) or 0),
                "mask_char": str(rule.get("mask_char", "*") or "*"),
            }
        else:
            # 旧格式兼容: 仅 pattern 字符串, 无 mask 信息
            DEFAULT_RULES[name] = str(rule)
            DEFAULT_RULES_META[name] = {}
    # v1.1.12: 启动诊断, 帮助用户确认 mask 配置是否从 config.json 正确加载
    _mask_diag = []
    for _name, _meta in DEFAULT_RULES_META.items():
        if not _meta:
            _mask_diag.append(f"{_name}=<空>")
        else:
            _pf = _meta.get("mask_keep_prefix", 0)
            _ps = _meta.get("mask_keep_suffix", 0)
            _md = _meta.get("mask_mode", "default")
            _mask_diag.append(f"{_name}={_md} {_pf}+{_ps}")
    print(f"[v1.1.12 启动诊断] DEFAULT_RULES_META 加载: {len(DEFAULT_RULES_META)} 条")
    for _line in _mask_diag:
        print(f"  - {_line}")
    # v1.1.13: 强制代码层默认 — 不依赖磁盘 config.json 漂移
    _v113_apply_rule_overrides()
else:
    DEFAULT_RULES = {
        "身份证号": r"(?<!\d)([1-9]\d{5}(19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[\dXx]|\d{15})(?!\d)",
        "手机号码": r"(?<!\d)(1[3-9]\d{9})(?!\d)",
        "日期时间": r"\d{4}[年\-\.]\d{1,2}[月\-\.]\d{1,2}[日]?",
        "电子邮箱": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
        "银行卡号": r"(?<!\d)([1-9]\d{12,18})(?!\d)",
        "印章": "__SEAL_DETECTION__",  # v1.1.11: 印章检测特殊标记
        # v1.1.11: 扩展规则 — 起诉讼书场景下常见漏脱敏字段
        "地址（含门牌号）": r"(?:[一-龥]{0,15}?)(?:省|市|自治区|特别行政区)[一-龥\d\s,()（）\-\w]{4,60}?\d+\s*号",
        "固定电话": r"(?<!\d)0[\w]{2,3}[-\s]?[\w]{7,8}(?!\d)",
        "法定代表人": r"法定代表人\s*[::：]?\s*[一-龥]{2,4}(?:·[一-龥]{2,4})?(?=[的之及与和按于在跟同向对为由被让等,，。；;）)\]】\s]|$)",
        # v1.1.12: 统一社会信用代码 - 仅 Word 路径生效,通过 pdf_excluded_rules 隔离 PDF
        # 首字符数字,后续 16-17 位 GB 32100-2015 字符集(排除 I/O/Z/S/V)
        "统一社会信用代码": r"(?<![A-Z0-9])([0-9][0-9A-HJ-NPQRTUWXY]{16,17})(?![A-Za-z0-9])",
        # v1.1.12: 公司名 - 中文公司/企业名称, jieba 不能识别所有公司字号
        "公司名": r"[一-龥]{2,40}(?:有限公司|股份有限公司|有限责任公司|集团公司|控股公司|合伙企业|公司|中心)",
    }
    DEFAULT_RULES_META = {
        "身份证号": {"mask_mode": "default", "mask_keep_prefix": 6, "mask_keep_suffix": 4, "mask_char": "*"},
        "手机号码": {"mask_mode": "default", "mask_keep_prefix": 3, "mask_keep_suffix": 4, "mask_char": "*"},
        "日期时间": {"mask_mode": "default", "mask_keep_prefix": 0, "mask_keep_suffix": 0, "mask_char": "*"},
        "电子邮箱": {"mask_mode": "email",   "mask_keep_prefix": 0, "mask_keep_suffix": 0, "mask_char": "*"},
        "银行卡号": {"mask_mode": "default", "mask_keep_prefix": 4, "mask_keep_suffix": 4, "mask_char": "*"},
        "印章":     {"mask_mode": "default", "mask_keep_prefix": 0, "mask_keep_suffix": 0, "mask_char": "*"},
        "地址（含门牌号）": {"mask_mode": "default", "mask_keep_prefix": 8, "mask_keep_suffix": 2, "mask_char": "*"},
        "固定电话": {"mask_mode": "default", "mask_keep_prefix": 0, "mask_keep_suffix": 4, "mask_char": "*"},
        "法定代表人": {"mask_mode": "name",    "mask_keep_prefix": 1, "mask_keep_suffix": 0, "mask_char": "*"},
        "统一社会信用代码": {"mask_mode": "default", "mask_keep_prefix": 4, "mask_keep_suffix": 4, "mask_char": "*"},
        "公司名": {"mask_mode": "default", "mask_keep_prefix": 0, "mask_keep_suffix": 0, "mask_char": "*"},
    }
    # v1.1.13: 强制代码层默认 — 不依赖磁盘 config.json 漂移
    _v113_apply_rule_overrides()

WORD_RULE_SCHEMA_VERSION = 1


def normalize_word_replace_rules(rules, default_replacement_text="[已脱敏]"):
    """规范化多字段替换规则（会话级规则，不自动持久化）。"""
    normalized = []
    if not isinstance(rules, list):
        return normalized

    fallback_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
    mode_alias = {
        "exact": "exact",
        "regex": "regex",
        "精确": "exact",
        "正则": "regex"
    }

    for item in rules:
        if not isinstance(item, dict):
            continue
        enabled = bool(item.get("enabled", True))
        raw_mode = str(item.get("mode", "exact")).strip().lower()
        mode = mode_alias.get(raw_mode, "exact")

        find_text = str(item.get("find", "")).strip()
        if not find_text:
            continue

        replace_text = item.get("replace")
        if replace_text is None or str(replace_text) == "":
            replace_text = fallback_text
        else:
            replace_text = str(replace_text)

        normalized.append({
            "enabled": enabled,
            "mode": mode,
            "find": find_text,
            "replace": replace_text
        })

    return normalized


def resolve_word_preview_image_suffix(content_type):
    """根据 Mammoth 图片内容类型推导本地文件后缀。"""
    if not isinstance(content_type, str):
        return ".img"

    normalized = content_type.strip().lower()
    if normalized in WORD_PREVIEW_IMAGE_EXTENSION_MAP:
        return WORD_PREVIEW_IMAGE_EXTENSION_MAP[normalized]

    if "/" not in normalized:
        return ".img"

    subtype = normalized.split("/", 1)[1].split(";", 1)[0].strip()
    subtype = subtype.replace("+xml", "").replace("+zip", "")
    subtype = re.sub(r"[^a-z0-9]+", "", subtype)
    if not subtype:
        return ".img"
    return f".{subtype}"


def _range_overlaps(start, end, ranges):
    """判断区间是否与已有区间重叠。"""
    for s, e in ranges:
        if not (end <= s or start >= e):
            return True
    return False


def build_word_rule_matches(text, rules, default_replacement_text="[已脱敏]"):
    """根据规则查找文本匹配，执行策略：exact 优先于 regex，重叠先到先得。"""
    if not isinstance(text, str) or not text:
        return []

    normalized_rules = normalize_word_replace_rules(rules, default_replacement_text)
    selected = []
    occupied_ranges = []

    for target_mode in ("exact", "regex"):
        for rule_index, rule in enumerate(normalized_rules):
            if not rule.get("enabled", True) or rule.get("mode") != target_mode:
                continue

            pattern = re.escape(rule["find"]) if target_mode == "exact" else rule["find"]
            try:
                compiled = re.compile(pattern)
            except re.error:
                continue

            for matched in compiled.finditer(text):
                start = matched.start()
                end = matched.end()
                if start >= end:
                    continue
                if _range_overlaps(start, end, occupied_ranges):
                    continue

                selected.append({
                    "start": start,
                    "end": end,
                    "text": matched.group(0),
                    "replacement": rule["replace"],
                    "mode": target_mode,
                    "rule_index": rule_index,
                    "source": "rule"
                })
                occupied_ranges.append((start, end))

    selected.sort(key=lambda item: item["start"])
    return selected


def apply_rule_matches_to_text(text, matches):
    """将匹配区间应用到文本（倒序替换避免索引偏移）。"""
    if not isinstance(text, str) or not matches:
        return text

    output = text
    for match in sorted(matches, key=lambda item: item.get("start", 0), reverse=True):
        start = match.get("start")
        end = match.get("end")
        replacement = match.get("replacement", "")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        if start < 0 or end > len(output) or start >= end:
            continue
        if replacement is None:
            replacement = ""
        if not isinstance(replacement, str):
            replacement = str(replacement)
        output = output[:start] + replacement + output[end:]
    return output


def apply_word_rules_to_text(text, rules, default_replacement_text="[已脱敏]"):
    """直接按规则替换文本并返回替换结果。"""
    matches = build_word_rule_matches(text, rules, default_replacement_text)
    return apply_rule_matches_to_text(text, matches)


def build_replaced_preview_segments(text, matches, default_replacement_text="[已脱敏]"):
    """根据匹配区间生成替换后文本分段（用于右侧预览高亮）。"""
    if not isinstance(text, str):
        return [{"type": "text", "value": ""}]
    if not matches:
        return [{"type": "text", "value": text}]

    fallback_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
    segments = []
    cursor = 0

    for match in sorted(matches, key=lambda item: item.get("start", 0)):
        start = match.get("start")
        end = match.get("end")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        if start < cursor or start < 0 or end > len(text) or start >= end:
            continue

        if start > cursor:
            segments.append({
                "type": "text",
                "value": text[cursor:start]
            })

        replacement = match.get("replacement", fallback_text)
        if replacement is None:
            replacement = fallback_text
        if not isinstance(replacement, str):
            replacement = str(replacement)

        segments.append({
            "type": "replacement",
            "value": replacement,
            "source": match.get("source", "rule"),
            "mode": match.get("mode", ""),
            "rule_name": match.get("rule_name", "")
        })
        cursor = end

    if cursor < len(text):
        segments.append({
            "type": "text",
            "value": text[cursor:]
        })

    if not segments:
        return [{"type": "text", "value": text}]
    return segments


def build_highlight_preview_segments(text, matches):
    """根据匹配区间生成原文高亮分段（用于左侧预览）。"""
    if not isinstance(text, str):
        return [{"type": "text", "value": ""}]
    if not matches:
        return [{"type": "text", "value": text}]

    segments = []
    cursor = 0
    for match in sorted(matches, key=lambda item: item.get("start", 0)):
        start = match.get("start")
        end = match.get("end")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        if start < cursor or start < 0 or end > len(text) or start >= end:
            continue

        if start > cursor:
            segments.append({"type": "text", "value": text[cursor:start]})

        segments.append({
            "type": "highlight",
            "value": text[start:end],
            "source": match.get("source", "manual"),
            "mode": match.get("mode", ""),
            "rule_name": match.get("rule_name", ""),
            "start": start,
            "end": end,
        })
        cursor = end

    if cursor < len(text):
        segments.append({"type": "text", "value": text[cursor:]})

    if not segments:
        return [{"type": "text", "value": text}]
    return segments


WORD_PREVIEW_BLOCK_SELECTOR = '[data-word-block="1"][data-key]'


def build_word_panel_update_script(block_updates):
    """构建仅更新正文块的 Word 预览增量刷新脚本。"""
    payload = json.dumps(block_updates or {}, ensure_ascii=False)
    return f"""
        (function() {{
            const updates = {payload};
            const elements = document.querySelectorAll('{WORD_PREVIEW_BLOCK_SELECTOR}');
            elements.forEach(function(el) {{
                const key = el.dataset.key;
                if (Object.prototype.hasOwnProperty.call(updates, key)) {{
                    const nextHtml = updates[key];
                    if (el.innerHTML !== nextHtml) {{
                        el.innerHTML = nextHtml;
                    }}
                }}
            }});
        }})();
    """


def should_reload_word_panel(source_changed, loaded_source_path, current_file_path, panel_ready):
    """判断 Word 预览面板是否需要重新加载完整文档。"""
    if source_changed:
        return True
    if not panel_ready:
        return True
    return loaded_source_path != current_file_path


def format_signed_percent(value):
    """将百分比格式化为适合界面展示的文案。"""
    try:
        number = int(value)
    except (TypeError, ValueError):
        number = 0
    return f"{number:+d}%" if number else "0%"


def build_settings_nav_labels(enabled_rules, keyword_count, precision_is_default, ocr_adjust_value, blacklist_count=0, whitelist_count=0):
    """构建设置中心左侧导航标签。"""
    return [
        f"1 通用规则 · {max(0, int(enabled_rules))}项启用",
        f"2 自定义关键词 · {max(0, int(keyword_count))}条",
        f"3 黑名单 · {max(0, int(blacklist_count))}条",
        f"4 白名单 · {max(0, int(whitelist_count))}条",
        f"5 扫描与微调 · {'默认' if precision_is_default else '已微调'}",
        f"6 OCR 检测框 · {format_signed_percent(ocr_adjust_value)}",
    ]


def build_settings_hero_tags(enabled_rules, keyword_count, enabled_word_rules, precision_is_default, ocr_adjust_value, scan_label):
    """构建设置页顶部的动态摘要标签。"""
    common_tag = (
        f"常用：规则 {max(0, int(enabled_rules))} 项 · "
        f"关键词 {max(0, int(keyword_count))} 条 · "
        f"Word {max(0, int(enabled_word_rules))} 条"
    )
    if precision_is_default:
        advanced_tag = f"高级：扫描推荐值 · OCR {format_signed_percent(ocr_adjust_value)}"
    else:
        normalized_scan_label = str(scan_label or "-").strip() or "-"
        advanced_tag = (
            f"高级：{normalized_scan_label} · "
            f"OCR {format_signed_percent(ocr_adjust_value)} · 已微调"
        )
    return common_tag, advanced_tag


def build_batch_result_rows(summary):
    """将批量替换 summary 转成结果表格行。"""
    if not isinstance(summary, dict):
        return []

    rows = []
    failed_items = summary.get("failed", []) if isinstance(summary.get("failed", []), list) else []
    success_items = summary.get("success", []) if isinstance(summary.get("success", []), list) else []

    for item in failed_items:
        if not isinstance(item, dict):
            continue
        input_path = str(item.get("input", "") or "")
        rows.append({
            "status": "失败",
            "status_key": "failed",
            "document": os.path.basename(input_path) if input_path else "未知文档",
            "detail": str(item.get("error", "") or "处理失败"),
            "action": "双击定位原文件",
            "open_path": input_path,
            "fallback_dir": os.path.dirname(input_path) if input_path else "",
        })

    for item in success_items:
        if not isinstance(item, dict):
            continue
        input_path = str(item.get("input", "") or "")
        output_path = str(item.get("output", "") or "")
        rows.append({
            "status": "成功",
            "status_key": "success",
            "document": os.path.basename(input_path) if input_path else "未知文档",
            "detail": os.path.basename(output_path) if output_path else "已生成输出文件",
            "action": "双击打开输出",
            "open_path": output_path,
            "fallback_dir": os.path.dirname(output_path) if output_path else "",
        })

    return rows


def summarize_batch_result_rows(rows):
    """汇总批量结果行数量。"""
    summary = {"total": 0, "success": 0, "failed": 0}
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        summary["total"] += 1
        status_key = row.get("status_key")
        if status_key == "success":
            summary["success"] += 1
        elif status_key == "failed":
            summary["failed"] += 1
    return summary


def build_batch_filter_labels(summary_counts, show_counts=False):
    """构建批量结果筛选按钮文案。"""
    counts = summary_counts if isinstance(summary_counts, dict) else {}
    total = max(0, int(counts.get("total", 0) or 0))
    success = max(0, int(counts.get("success", 0) or 0))
    failed = max(0, int(counts.get("failed", 0) or 0))
    if not show_counts:
        return {"all": "全部", "success": "成功", "failed": "失败"}
    return {
        "all": f"全部 {total}",
        "success": f"成功 {success}",
        "failed": f"失败 {failed}",
    }


def filter_batch_result_rows(rows, filter_mode):
    """按筛选模式过滤批量结果行。"""
    if filter_mode not in {"all", "success", "failed"}:
        filter_mode = "all"

    filtered = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        if filter_mode == "all":
            filtered.append(row)
        elif row.get("status_key") == filter_mode:
            filtered.append(row)
    return filtered


def build_batch_rule_summary_lines(rules, success_items, default_replacement_text="[已脱敏]"):
    """按规则生成批量替换摘要明细。"""
    normalized_rules = normalize_word_replace_rules(rules, default_replacement_text)
    if not normalized_rules:
        return []

    def _extract_rule_count(item, target_rule_index):
        if not isinstance(item, dict):
            return 0

        counts = item.get("rule_counts", [])
        if isinstance(counts, dict):
            try:
                return max(0, int(counts.get(str(target_rule_index), counts.get(target_rule_index, 0)) or 0))
            except (TypeError, ValueError):
                return 0

        if not isinstance(counts, list):
            return 0

        for entry in counts:
            if not isinstance(entry, dict):
                continue
            try:
                rule_index = int(entry.get("rule_index", -1))
                count = int(entry.get("count", 0) or 0)
            except (TypeError, ValueError):
                continue
            if rule_index == target_rule_index:
                return max(0, count)
        return 0

    lines = []
    for rule_index, rule in enumerate(normalized_rules, start=1):
        doc_parts = []
        replacement_text = rule.get("replace") or default_replacement_text

        for item in success_items or []:
            count = _extract_rule_count(item, rule_index - 1)
            if count <= 0:
                continue
            input_name = os.path.basename(str(item.get("input", "") or "")) or "未知文档"
            doc_parts.append(f"{input_name} 成功替换 {count} 条")

        if doc_parts:
            lines.append(
                f"{rule_index}、“{rule.get('find', '')}”替换为“{replacement_text}”，"
                + "，".join(doc_parts)
                + "；"
            )
        else:
            lines.append(
                f"{rule_index}、“{rule.get('find', '')}”替换为“{replacement_text}”，本轮未命中；"
            )

    return lines


def build_workbench_guidance(mode, batch_stage="rule_setup", has_results=False, compare_mode=False):
    """按当前模式生成顶部工作台的下一步引导标签。"""
    if mode == "pdf":
        first_step = "下一步：人工复核并导出" if has_results else "下一步：先点智能脱敏"
        return [
            first_step,
            "黑 / 白遮罩可立即切换",
            "支持手动画框补充脱敏",
        ]
    if mode == "word":
        compare_tip = "当前可隐藏对比预览" if compare_mode else "需要时可打开对比预览"
        first_step = "下一步：先检查替换规则" if not has_results else "下一步：复核替换结果"
        return [
            first_step,
            "原文预览与替换预览分开显示",
            compare_tip,
        ]
    if mode == "batch":
        if batch_stage == "running":
            return [
                "当前：正在批量替换文档",
                "可随时停止并保留已完成结果",
                "完成后可筛选成功 / 失败清单",
            ]
        if batch_stage in ("finished", "stopped"):
            return [
                "下一步：先看失败文档和原因",
                "可仅重试失败文档",
                "双击结果可打开输出或定位原文件",
            ]
        return [
            "下一步：确认规则后再开始执行",
            "这一步不会改动任何原文件",
            "建议至少启用一条 Word 替换规则",
        ]
    if mode == "image_merge":
        return [
            "下一步：确认图片顺序后开始合并",
            "支持多张图片自动合成为 PDF",
            "合并完成后会直接进入 PDF 脱敏",
        ]
    return [
        "支持拖拽导入，系统会自动分流",
        "PDF 走脱敏，Word 走替换",
        "多个 Word 会先进入批量规则确认",
        "多张图片可直接合并为 PDF",
    ]


def build_toolbar_mode_labels(mode, density_mode, has_results=False, enabled_word_rules=0):
    """构建工具栏在不同模式下的主动作文案。"""
    compact = density_mode != "wide"

    if mode == "pdf":
        if has_results:
            scan_text = "重脱" if compact else "重新脱敏"
            scan_tooltip = "重新执行 PDF 智能脱敏扫描"
        else:
            scan_text = "脱敏" if compact else "智能脱敏"
            scan_tooltip = "执行 PDF 智能脱敏扫描"
        save_text = "导出" if compact else "导出 PDF"
        save_tooltip = "导出当前 PDF 脱敏结果"
    elif mode == "word":
        if has_results:
            scan_text = "重替" if compact else "重新替换"
            scan_tooltip = "重新执行 Word 智能替换扫描"
        else:
            scan_text = "替换" if compact else "智能替换"
            scan_tooltip = "执行 Word 智能替换扫描"
        save_text = "导出" if compact else "导出 Word"
        save_tooltip = "导出当前 Word 替换结果"
    else:
        scan_text = "脱敏" if compact else "智能脱敏"
        save_text = "导出"
        scan_tooltip = "执行智能脱敏扫描"
        save_tooltip = "导出处理结果"

    if enabled_word_rules > 0:
        word_rules_text = f"规则 {enabled_word_rules}" if compact else f"替换规则 {enabled_word_rules}"
        word_rules_tooltip = f"打开 Word 替换规则（当前启用 {enabled_word_rules} 条）"
    else:
        word_rules_text = "规则" if compact else "替换规则"
        word_rules_tooltip = "打开 Word 替换规则"

    return {
        "open_text": "打开",
        "open_tooltip": "打开 PDF、Word 或图片文件",
        "scan_text": scan_text,
        "scan_tooltip": scan_tooltip,
        "save_text": save_text,
        "save_tooltip": save_tooltip,
        "word_rules_text": word_rules_text,
        "word_rules_tooltip": word_rules_tooltip,
    }


def _shift_density_mode(mode, order, step):
    """在既定密度序列里前后移动一档。"""
    if mode not in order:
        return mode
    index = order.index(mode)
    target = max(0, min(len(order) - 1, index + step))
    return order[target]


def resolve_workspace_density_mode(mode, width, height=0, scale=1.0):
    """解析主工作区工具栏密度档位，兼顾 Windows DPI 与窗口高度。"""
    width = max(int(width or 0), 1)
    height = max(int(height or 0), 0)
    scale = max(1.0, float(scale or 1.0))

    if mode == "pdf":
        wide_threshold = 1500
        compact_threshold = 1260
    elif mode == "word":
        wide_threshold = 1220
        compact_threshold = 980
    else:
        wide_threshold = 1080
        compact_threshold = 860

    if scale >= 1.5:
        wide_threshold += 90
        compact_threshold += 60
    elif scale >= 1.25:
        wide_threshold += 50
        compact_threshold += 30

    if height:
        if height >= 980:
            wide_threshold -= 50
            compact_threshold -= 30
        elif height <= 760:
            wide_threshold += 70
            compact_threshold += 40

    wide_threshold = max(compact_threshold + 80, wide_threshold)
    compact_threshold = max(720, min(compact_threshold, wide_threshold - 80))

    if width >= wide_threshold:
        return "wide"
    if width >= compact_threshold:
        return "compact"
    return "narrow"


def resolve_settings_density_mode(width, height=0, scale=1.0):
    """解析高级设置页密度档位，优先为 Windows 高 DPI 和不同窗口高度收口。"""
    width = max(int(width or 0), 1)
    height = max(int(height or 0), 0)
    scale = max(1.0, float(scale or 1.0))
    order = ["narrow", "compact", "roomy", "wide"]

    if width >= 1700:
        density_mode = "wide"
    elif width >= 1450:
        density_mode = "roomy"
    elif width >= 1260:
        density_mode = "compact"
    else:
        density_mode = "narrow"

    if scale >= 1.5:
        density_mode = _shift_density_mode(density_mode, order, -1)
        if width < 1360:
            density_mode = "narrow"
    elif scale >= 1.25 and density_mode == "wide" and width < 1760:
        density_mode = "roomy"

    if height:
        if height <= 820:
            density_mode = _shift_density_mode(density_mode, order, -1)
        elif height >= 980:
            if density_mode == "compact" and width >= 1380:
                density_mode = "roomy"
            elif density_mode == "roomy" and width >= 1600:
                density_mode = "wide"

    return density_mode


def merge_word_matches_with_priority(text, rules, default_replacement_text,
                                     manual_matches=None, ocr_matches=None):
    """合并规则替换、手动脱敏、OCR 脱敏区间，优先级：规则 > 手动 > OCR。"""
    manual_matches = manual_matches or []
    ocr_matches = ocr_matches or []
    text_len = len(text) if isinstance(text, str) else 0
    fallback_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"

    merged = []
    occupied_ranges = []

    def _append_candidates(candidates, source_name):
        for item in candidates:
            start = item.get("start")
            end = item.get("end")
            if not isinstance(start, int) or not isinstance(end, int):
                continue
            if start < 0 or end > text_len or start >= end:
                continue
            if _range_overlaps(start, end, occupied_ranges):
                continue

            replacement = item.get("replacement", fallback_text)
            if replacement is None:
                replacement = fallback_text
            if not isinstance(replacement, str):
                replacement = str(replacement)

            merged.append({
                "start": start,
                "end": end,
                "text": item.get("text", text[start:end] if isinstance(text, str) else ""),
                "replacement": replacement,
                "source": source_name,
                "mode": item.get("mode", "global"),
                "rule_name": item.get("rule_name", "")
            })
            occupied_ranges.append((start, end))

    _append_candidates(build_word_rule_matches(text, rules, fallback_text), "rule")
    _append_candidates(manual_matches, "manual")
    _append_candidates(ocr_matches, "ocr")
    merged.sort(key=lambda item: item["start"])
    return merged


def apply_range_to_runs(para, start, end, replacement):
    """在段落 run 列表上应用一次区间替换。"""
    if start >= end:
        return
    if not para.runs:
        return

    run_ranges = []
    cursor = 0
    for idx, run in enumerate(para.runs):
        text = run.text or ''
        run_start = cursor
        run_end = cursor + len(text)
        run_ranges.append((idx, run_start, run_end))
        cursor = run_end

    total_len = cursor
    if start < 0 or end > total_len:
        return

    start_run_idx = None
    start_offset = 0
    for idx, run_start, run_end in run_ranges:
        if start < run_end:
            start_run_idx = idx
            start_offset = start - run_start
            break
    if start_run_idx is None:
        start_run_idx = run_ranges[-1][0]
        start_offset = len(para.runs[start_run_idx].text or '')

    end_run_idx = None
    end_offset = 0
    for idx, run_start, run_end in run_ranges:
        if end <= run_end:
            end_run_idx = idx
            end_offset = end - run_start
            break
    if end_run_idx is None:
        end_run_idx = run_ranges[-1][0]
        end_offset = len(para.runs[end_run_idx].text or '')

    start_run = para.runs[start_run_idx]
    end_run = para.runs[end_run_idx]
    start_text = start_run.text or ''
    end_text = end_run.text or ''

    prefix = start_text[:start_offset]
    suffix = end_text[end_offset:]
    start_run.text = prefix + replacement + suffix

    if end_run_idx > start_run_idx:
        for idx in range(start_run_idx + 1, end_run_idx + 1):
            para.runs[idx].text = ''


def replace_matches_in_paragraph(para, matches, text_offset=0, fallback_replacement_text="[已脱敏]"):
    """按匹配区间替换段落文本，避免同词误替换和跨 run 漏替换。"""
    if not matches or not para.runs:
        return

    paragraph_text = ''.join(run.text for run in para.runs)
    if not paragraph_text:
        return

    text_len = len(paragraph_text)
    ranges = []
    seen = set()

    for match in matches:
        start = match.get('start')
        end = match.get('end')
        if not isinstance(start, int) or not isinstance(end, int):
            continue

        local_start = start - text_offset
        local_end = end - text_offset
        if local_start < 0 or local_end > text_len or local_start >= local_end:
            continue

        replacement = match.get('replacement', fallback_replacement_text)
        if replacement is None:
            replacement = fallback_replacement_text
        if not isinstance(replacement, str):
            replacement = str(replacement)

        key = (local_start, local_end, replacement)
        if key in seen:
            continue
        seen.add(key)
        ranges.append({
            'start': local_start,
            'end': local_end,
            'replacement': replacement
        })

    if not ranges:
        return

    ranges.sort(key=lambda item: (item['start'], -(item['end'] - item['start'])))
    filtered = []
    last_end = -1
    for item in ranges:
        if item['start'] < last_end:
            continue
        filtered.append(item)
        last_end = item['end']

    for item in reversed(filtered):
        apply_range_to_runs(para, item['start'], item['end'], item['replacement'])


# === 设置对话框 ===
class SettingsDialog(QDialog):
    """设置对话框 - v1.1.11: 支持配置持久化"""

    def __init__(self, parent=None, current_rules=None, use_enhance=False, custom_keywords="",
                 scan_level=2.0, offset_x=0, offset_w=0, replacement_text="[已脱敏]",
                 word_replace_rules=None,
                 config_manager=None,
                 enable_name_recognition=False,
                 current_blacklist=None,
                 current_whitelist=None):
        super().__init__(parent)
        self.config = config_manager

        # v1.1.11: 修复 Windows 深色模式下对话框显示问题
        # 设置窗口标志，确保对话框在深色系统主题下使用浅色样式
        self.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        # v1.1.11: 从配置读取窗口尺寸，并结合当前屏幕做自适应约束
        if self.config:
            dialog_width = self.config.get("app.window.dialog_settings_width", 550)
            dialog_height = self.config.get("app.window.dialog_settings_height", 700)
        else:
            dialog_width, dialog_height = 550, 700

        self.setWindowTitle("高级设置")
        screen = QApplication.primaryScreen()
        if screen:
            available = screen.availableGeometry()
            max_w = min(max(760, int(available.width() * 0.92)), available.width())
            max_h = min(max(620, int(available.height() * 0.90)), available.height())
            preferred_w = min(max_w, 1180)
            preferred_h = min(max_h, 900)
            min_w = min(900, max_w)
            min_h = min(680, max_h)
            width = max(min(min(max(dialog_width, 960), preferred_w), max_w), min_w)
            height = max(min(min(max(dialog_height, 760), preferred_h), max_h), min_h)
            self.resize(width, height)
            self.setMinimumSize(min_w, min_h)
            self.setMaximumSize(max_w, max_h)
        else:
            self.resize(980, 760)
            self.setMinimumSize(820, 620)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.setSizeGripEnabled(True)
        self._settings_bound_window_handle = None

        # v1.1.11: 应用对话框主题样式
        self._apply_dialog_theme()

        self.selected_rules = []
        self.use_enhance = use_enhance
        self.custom_keywords = custom_keywords
        self.scan_level = scan_level
        self.offset_x = offset_x
        self.offset_w = offset_w
        self.replacement_text = replacement_text if isinstance(replacement_text, str) and replacement_text else "[已脱敏]"
        self.word_replace_rules = normalize_word_replace_rules(word_replace_rules or [], self.replacement_text)
        self.default_replacement_text = "[已脱敏]"
        self.recommended_rule_names = ["身份证号", "手机号码"]
        # v1.1.11: 中文姓名启发式识别开关(默认 False,向后兼容)
        self.enable_name_recognition = bool(enable_name_recognition)
        # v1.1.11: 黑/白名单(来自 config.json 的 redaction.blacklist/whitelist)
        self._initial_blacklist = list(current_blacklist or [])
        self._initial_whitelist = list(current_whitelist or [])

        # v1.1.11: 从配置读取范围和标签
        if self.config:
            offset_config = self.config.get("redaction.offset", {})
            x_range = offset_config.get("x_range", [-20, 20])
            w_range = offset_config.get("w_range", [-20, 20])
            scan_config = self.config.get("redaction.scan", {})
            available_levels = scan_config.get("available_levels", [1.5, 2.0, 3.0])
            level_labels = scan_config.get("level_labels", {
                "1.5": "标准 (1.5x)",
                "2.0": "高精 (2.0x 推荐)",
                "3.0": "超精 (3.0x 最慢)"
            })
        else:
            x_range = [-20, 20]
            w_range = [-20, 20]
            available_levels = [1.5, 2.0, 3.0]
            level_labels = {
                "1.5": "标准 (1.5x)",
                "2.0": "高精 (2.0x 推荐)",
                "3.0": "超精 (3.0x 最慢)"
            }

        self.x_range = x_range
        self.w_range = w_range
        self.available_levels = available_levels
        self.level_labels = level_labels

        outer_layout = QVBoxLayout(self)
        outer_layout.setContentsMargins(18, 16, 18, 18)
        outer_layout.setSpacing(14)

        hero = QFrame(self)
        hero.setObjectName("settingsHero")
        hero_layout = QVBoxLayout(hero)
        hero_layout.setContentsMargins(18, 16, 18, 16)
        hero_layout.setSpacing(6)
        self.settings_hero_layout = hero_layout
        hero_title = QLabel("高级设置中心")
        hero_title.setObjectName("settingsTitle")
        self.settings_hero_title = hero_title
        hero_subtitle = QLabel("这里保存长期配置。PDF 的黑 / 白、单双页、缩放等即时操作，仍然放在主工作台里。")
        hero_subtitle.setObjectName("settingsSubtitle")
        hero_subtitle.setWordWrap(True)
        self.settings_hero_subtitle = hero_subtitle
        hero_layout.addWidget(hero_title)
        hero_layout.addWidget(hero_subtitle)
        hero_tag_layout = QHBoxLayout()
        hero_tag_layout.setSpacing(8)
        self.settings_hero_tag_layout = hero_tag_layout
        self.lbl_settings_common_tag = QLabel("常用设置：通用规则 / 关键词 / Word 替换")
        self.lbl_settings_common_tag.setObjectName("settingsHeroTag")
        self.lbl_settings_advanced_tag = QLabel("高级微调：扫描 / 覆盖 / OCR 检测框")
        self.lbl_settings_advanced_tag.setObjectName("settingsHeroTag")
        hero_tag_layout.addWidget(self.lbl_settings_common_tag)
        hero_tag_layout.addWidget(self.lbl_settings_advanced_tag)
        hero_tag_layout.addStretch()
        hero_layout.addLayout(hero_tag_layout)
        outer_layout.addWidget(hero)

        body_layout = QHBoxLayout()
        body_layout.setSpacing(18)
        self.settings_body_layout = body_layout
        outer_layout.addLayout(body_layout, stretch=1)

        sidebar = QFrame(self)
        sidebar.setObjectName("settingsSidebar")
        sidebar.setFixedWidth(236)
        self.settings_sidebar = sidebar
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(16, 16, 16, 16)
        sidebar_layout.setSpacing(12)
        self.settings_sidebar_layout = sidebar_layout

        nav_hint = QLabel("设置导航")
        nav_hint.setObjectName("settingsHint")
        self.settings_nav_hint = nav_hint
        sidebar_layout.addWidget(nav_hint)
        nav_subtitle = QLabel("点击左侧即可快速跳转到对应模块。")
        nav_subtitle.setObjectName("settingsSidebarSubtle")
        nav_subtitle.setWordWrap(True)
        self.settings_nav_subtitle = nav_subtitle
        sidebar_layout.addWidget(nav_subtitle)
        sidebar_layout.addWidget(self._create_settings_divider())

        self.settings_nav = QListWidget()
        self.settings_nav.setObjectName("settingsNav")
        self.settings_nav.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        self._settings_nav_base_titles = ["1 通用规则", "2 自定义关键词", "3 黑名单", "4 白名单", "5 扫描与微调", "6 OCR 检测框"]
        for title in self._settings_nav_base_titles:
            self.settings_nav.addItem(title)
        sidebar_layout.addWidget(self.settings_nav, stretch=1)

        sidebar_meta_card = QFrame()
        sidebar_meta_card.setObjectName("settingsSidebarMetaCard")
        self.settings_sidebar_meta_card = sidebar_meta_card
        sidebar_meta_layout = QVBoxLayout(sidebar_meta_card)
        sidebar_meta_layout.setContentsMargins(12, 12, 12, 12)
        sidebar_meta_layout.setSpacing(10)
        sidebar_note = QLabel("建议：第一次使用只改规则和关键词；OCR 微调只在识别偏移时再调整。")
        sidebar_note.setWordWrap(True)
        sidebar_note.setObjectName("settingsSidebarNote")
        self.settings_sidebar_note = sidebar_note
        sidebar_meta_layout.addWidget(sidebar_note)
        sidebar_meta_layout.addWidget(self._create_settings_divider())
        self.settings_sidebar_status = QLabel("")
        self.settings_sidebar_status.setWordWrap(True)
        self.settings_sidebar_status.setObjectName("settingsSidebarStatus")
        sidebar_meta_layout.addWidget(self.settings_sidebar_status)
        sidebar_layout.addWidget(sidebar_meta_card)
        body_layout.addWidget(sidebar)

        content_scroll = QScrollArea(self)
        content_scroll.setWidgetResizable(True)
        content_scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.content_scroll = content_scroll

        content_widget = QWidget()
        self.content_widget = content_widget
        layout = QVBoxLayout(content_widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(14)

        overview_card = QFrame()
        overview_card.setObjectName("settingsOverview")
        self.settings_overview_card = overview_card
        overview_layout = QVBoxLayout(overview_card)
        overview_layout.setContentsMargins(18, 16, 18, 16)
        overview_layout.setSpacing(12)
        self.settings_overview_layout = overview_layout
        overview_title = QLabel("当前配置概览")
        overview_title.setObjectName("settingsOverviewTitle")
        self.settings_overview_title = overview_title
        overview_text = QLabel("先看这里，就能快速知道当前默认规则、关键词、Word 替换和 OCR 微调处于什么状态。")
        overview_text.setObjectName("settingsOverviewText")
        overview_text.setWordWrap(True)
        self.settings_overview_text = overview_text
        overview_layout.addWidget(overview_title)
        overview_layout.addWidget(overview_text)

        overview_metrics_layout = QGridLayout()
        overview_metrics_layout.setHorizontalSpacing(10)
        overview_metrics_layout.setVerticalSpacing(10)
        self.settings_overview_metrics_layout = overview_metrics_layout
        self.lbl_metric_rules = None
        self.lbl_metric_keywords = None
        self.lbl_metric_word_rules = None
        self.lbl_metric_ocr = None
        self.settings_metric_cards = []
        for key, title in [
            ("rules", "通用规则"),
            ("keywords", "自定义关键词"),
            ("word_rules", "Word 规则"),
            ("ocr", "OCR 调节"),
        ]:
            metric_card, metric_value = self._create_settings_metric_card(title)
            setattr(self, f"lbl_metric_{key}", metric_value)
            self.settings_metric_cards.append(metric_card)
            overview_metrics_layout.addWidget(metric_card, 0, len(self.settings_metric_cards) - 1)
        overview_layout.addLayout(overview_metrics_layout)

        overview_actions_layout = QGridLayout()
        overview_actions_layout.setHorizontalSpacing(8)
        overview_actions_layout.setVerticalSpacing(8)
        self.settings_overview_actions_layout = overview_actions_layout
        self._settings_quick_jump_titles = [
            ("去看通用规则", "通用规则"),
            ("去看关键词", "关键词"),
            ("去看扫描微调", "扫描微调"),
            ("去看 OCR", "OCR"),
        ]
        self.settings_quick_jump_buttons = []
        for row, (title, _short_title) in enumerate(self._settings_quick_jump_titles):
            jump_btn = QPushButton(title)
            jump_btn.setObjectName("settingsQuickJumpButton")
            jump_btn.clicked.connect(lambda _checked=False, target_row=row: self.settings_nav.setCurrentRow(target_row))
            jump_btn.setMinimumHeight(32)
            self.settings_quick_jump_buttons.append(jump_btn)
            overview_actions_layout.addWidget(jump_btn, 0, row)
        overview_layout.addLayout(overview_actions_layout)
        layout.addWidget(overview_card)

        # 1. 规则
        box_rules = QFrame()
        box_rules.setObjectName("settingsSectionCard")
        v_box = QVBoxLayout(box_rules)
        v_box.setContentsMargins(16, 16, 16, 16)
        v_box.setSpacing(12)
        rules_lead = QLabel("勾选后的规则会作为默认智能识别规则。第一次使用建议至少保留身份证号和手机号。")
        rules_lead.setObjectName("settingsSectionLead")
        rules_lead.setWordWrap(True)
        self.lbl_rules_summary = QLabel("")
        self.lbl_rules_summary.setObjectName("settingsSectionSummary")
        self.lbl_rules_summary.setWordWrap(True)
        v_box.addWidget(self._create_settings_section_header("1. 通用规则", rules_lead, self.lbl_rules_summary))
        rules_actions = QHBoxLayout()
        rules_actions.setSpacing(8)
        rules_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_rules_recommended = QPushButton("恢复推荐勾选")
        btn_rules_recommended.setObjectName("settingsInlineButton")
        btn_rules_recommended.clicked.connect(self._apply_recommended_rules)
        rules_actions.addWidget(btn_rules_recommended)
        btn_rules_all = QPushButton("全部勾选")
        btn_rules_all.setObjectName("settingsInlineButton")
        btn_rules_all.clicked.connect(self._select_all_rules)
        rules_actions.addWidget(btn_rules_all)
        btn_rules_clear = QPushButton("全部清空")
        btn_rules_clear.setObjectName("settingsInlineButton")
        btn_rules_clear.clicked.connect(self._clear_all_rules)
        rules_actions.addWidget(btn_rules_clear)
        rules_actions.addStretch()
        v_box.addLayout(rules_actions)
        rules_columns = QHBoxLayout()
        rules_columns.setSpacing(24)
        rules_left_col = QVBoxLayout()
        rules_left_col.setSpacing(6)
        rules_right_col = QVBoxLayout()
        rules_right_col.setSpacing(6)
        self.checks = {}
        # v1.1.12: 读取 PDF 排除规则列表,用于在规则面板给"仅 Word"规则加标识
        try:
            pdf_excluded_names_ui = config.get("redaction.pdf_excluded_rules", []) if config else []
        except Exception:
            pdf_excluded_names_ui = []
        if not isinstance(pdf_excluded_names_ui, list):
            pdf_excluded_names_ui = []
        pdf_excluded_set = set(pdf_excluded_names_ui)
        rule_items = list(DEFAULT_RULES.items())
        split_index = (len(rule_items) + 1) // 2
        for index, (name, pattern) in enumerate(rule_items):
            # v1.1.12: 若规则在 PDF 排除列表中,UI 名称追加 "📝 仅 Word" 标识
            display_name = f"{name}  📝 仅 Word" if name in pdf_excluded_set else name
            cb = QCheckBox(display_name)
            if current_rules and pattern in current_rules: cb.setChecked(True)
            elif not current_rules and name in ["身份证号", "手机号码", "地址（含门牌号）", "固定电话", "法定代表人"]: cb.setChecked(True)
            self.checks[name] = cb
            cb.toggled.connect(self._refresh_rule_summary)
            if index < split_index:
                rules_left_col.addWidget(cb)
            else:
                rules_right_col.addWidget(cb)
        rules_left_col.addStretch()
        rules_right_col.addStretch()
        rules_columns.addLayout(rules_left_col, stretch=1)
        rules_columns.addLayout(rules_right_col, stretch=1)
        v_box.addLayout(rules_columns)
        layout.addWidget(box_rules)

        # 2. 关键词 + 统一替换文本 + Word 替换规则入口
        box_custom = QFrame()
        box_custom.setObjectName("settingsSectionCard")
        v_custom = QVBoxLayout(box_custom)
        v_custom.setContentsMargins(16, 16, 16, 16)
        v_custom.setSpacing(12)
        custom_lead = QLabel("这里适合录入姓名、单位、案号等固定内容；Word 替换规则适合更精细的定向替换。")
        custom_lead.setObjectName("settingsSectionLead")
        custom_lead.setWordWrap(True)
        self.lbl_custom_summary = QLabel("")
        self.lbl_custom_summary.setObjectName("settingsSectionSummary")
        self.lbl_custom_summary.setWordWrap(True)
        v_custom.addWidget(self._create_settings_section_header("2. 自定义关键词", custom_lead, self.lbl_custom_summary))
        custom_row = QGridLayout()
        custom_row.setHorizontalSpacing(28)
        custom_row.setVerticalSpacing(14)
        self.settings_custom_row = custom_row

        left_panel_widget = QFrame()
        left_panel_widget.setObjectName("settingsFieldCard")
        self.settings_left_field_card = left_panel_widget
        left_panel_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        left_panel = QVBoxLayout(left_panel_widget)
        left_panel.setContentsMargins(16, 16, 16, 16)
        left_panel.setSpacing(10)
        left_title = QLabel("关键词列表")
        left_title.setObjectName("settingsFieldTitle")
        left_note = QLabel("按行输入固定敏感词。适合姓名、机构、案号等直接命中的内容。")
        left_note.setObjectName("settingsFieldNote")
        left_note.setWordWrap(True)
        left_panel.addWidget(left_title)
        left_panel.addWidget(left_note)
        left_panel.addWidget(self._create_settings_divider())
        custom_actions = QHBoxLayout()
        custom_actions.setSpacing(8)
        custom_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_clear_keywords = QPushButton("清空关键词")
        btn_clear_keywords.setObjectName("settingsInlineButton")
        btn_clear_keywords.clicked.connect(self._clear_custom_keywords)
        custom_actions.addWidget(btn_clear_keywords)
        custom_actions.addStretch()
        left_panel.addLayout(custom_actions)
        # v1.1.11: 中文姓名启发式识别开关
        self.cb_name_recognition = QCheckBox("启用中文姓名启发式识别 (jieba)")
        self.cb_name_recognition.setObjectName("settingsInlineCheckbox")
        self.cb_name_recognition.setChecked(self.enable_name_recognition)
        self.cb_name_recognition.setToolTip(
            "开启后,工具会调用 jieba 词性标注提取文本中的中文姓名,"
            "追加到 OCR/Word 命中规则。\n"
            "适用于法律文书等含姓名角色的场景;非中文或纯文本场景建议关闭。"
        )
        left_panel.addWidget(self.cb_name_recognition)
        self.txt_custom = QTextEdit()
        self.txt_custom.setPlaceholderText("例如：法院 张三 (支持多行)")
        self.txt_custom.setPlainText(custom_keywords)
        self.txt_custom.setMinimumHeight(120)
        self.txt_custom.setMaximumHeight(190)
        self.txt_custom.textChanged.connect(self._refresh_custom_summary)
        left_panel.addWidget(self.txt_custom)
        custom_row.addWidget(left_panel_widget, 0, 0)

        right_panel_widget = QFrame()
        right_panel_widget.setObjectName("settingsFieldCard")
        right_panel_widget.setMinimumWidth(320)
        self.settings_right_field_card = right_panel_widget
        right_panel_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        right_panel = QVBoxLayout(right_panel_widget)
        right_panel.setContentsMargins(16, 16, 16, 16)
        right_panel.setSpacing(10)
        right_title = QLabel("Word 替换规则")
        right_title.setObjectName("settingsFieldTitle")
        right_note = QLabel("这里适合维护更精细的 exact / regex 替换规则，批量 Word 会直接复用。")
        right_note.setObjectName("settingsFieldNote")
        right_note.setWordWrap(True)
        right_panel.addWidget(right_title)
        right_panel.addWidget(right_note)
        right_panel.addWidget(self._create_settings_divider())
        replacement_row = QHBoxLayout()
        replacement_row.setSpacing(8)
        replacement_row.addWidget(self._create_settings_form_label("统一替换文本:", 102))
        self.input_replacement_text = QLineEdit(self.replacement_text)
        self.input_replacement_text.setPlaceholderText("[已脱敏]")
        self.input_replacement_text.setMinimumWidth(240)
        self.input_replacement_text.textChanged.connect(self._refresh_custom_summary)
        replacement_row.addWidget(self.input_replacement_text)
        right_panel.addLayout(replacement_row)
        replacement_actions = QHBoxLayout()
        replacement_actions.setSpacing(8)
        replacement_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_reset_replacement = QPushButton("恢复默认替换词")
        btn_reset_replacement.setObjectName("settingsInlineButton")
        btn_reset_replacement.clicked.connect(self._reset_replacement_text)
        replacement_actions.addWidget(btn_reset_replacement)
        replacement_actions.addStretch()
        right_panel.addLayout(replacement_actions)

        right_panel.addSpacing(8)
        right_panel.addWidget(self._create_settings_form_label("Word 替换规则:", 120))
        self.btn_edit_word_rules = QPushButton("打开替换规则设置")
        self.btn_edit_word_rules.clicked.connect(self._open_word_rules_editor)
        self.btn_edit_word_rules.setMinimumHeight(34)
        right_panel.addWidget(self.btn_edit_word_rules)

        self.lbl_word_rule_count = QLabel("")
        self.lbl_word_rule_count.setObjectName("settingsFieldNote")
        self.lbl_word_rule_count.setWordWrap(True)
        right_panel.addWidget(self.lbl_word_rule_count)
        word_rules_hint = QLabel("批量 Word 替换会直接复用这里的规则结构，不需要另外维护一套。")
        word_rules_hint.setWordWrap(True)
        word_rules_hint.setObjectName("settingsFieldNote")
        right_panel.addWidget(word_rules_hint)
        right_panel.addStretch()

        custom_row.addWidget(right_panel_widget, 0, 1)
        v_custom.addLayout(custom_row)
        layout.addWidget(box_custom)

        # v1.1.11: 2.5 黑名单 (强制脱敏, 优先级最低 / 实际生效低于规则)
        box_black = QFrame()
        box_black.setObjectName("settingsSectionCard")
        v_black = QVBoxLayout(box_black)
        v_black.setContentsMargins(16, 16, 16, 16)
        v_black.setSpacing(12)
        black_lead = QLabel("每行一条子串,强制进入脱敏命中。即使通用规则和关键词都没命中,只要文本里包含这里的内容就会被遮盖。")
        black_lead.setObjectName("settingsSectionLead")
        black_lead.setWordWrap(True)
        self.lbl_black_summary = QLabel("")
        self.lbl_black_summary.setObjectName("settingsSectionSummary")
        self.lbl_black_summary.setWordWrap(True)
        v_black.addWidget(self._create_settings_section_header(
            "3. 黑名单 (强制脱敏)", black_lead, self.lbl_black_summary))
        black_actions = QHBoxLayout()
        black_actions.setSpacing(8)
        black_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_clear_black = QPushButton("清空黑名单")
        btn_clear_black.setObjectName("settingsInlineButton")
        btn_clear_black.clicked.connect(self._clear_blacklist)
        black_actions.addWidget(btn_clear_black)
        black_actions.addStretch()
        v_black.addLayout(black_actions)
        black_card = QFrame()
        black_card.setObjectName("settingsFieldCard")
        black_card_layout = QVBoxLayout(black_card)
        black_card_layout.setContentsMargins(16, 16, 16, 16)
        black_card_layout.setSpacing(10)
        black_card_title = QLabel("黑名单条目")
        black_card_title.setObjectName("settingsFieldTitle")
        black_card_note = QLabel("子串匹配 (不需完整词)。每行一条,实时生效,关闭对话框也会保留。")
        black_card_note.setObjectName("settingsFieldNote")
        black_card_note.setWordWrap(True)
        black_card_layout.addWidget(black_card_title)
        black_card_layout.addWidget(black_card_note)
        black_card_layout.addWidget(self._create_settings_divider())
        self.txt_blacklist = QTextEdit()
        self.txt_blacklist.setPlaceholderText("例如: 盖章 / 内部资料 / 签字")
        self.txt_blacklist.setPlainText("\n".join(self._initial_blacklist))
        self.txt_blacklist.setMinimumHeight(120)
        self.txt_blacklist.setMaximumHeight(190)
        self.txt_blacklist.textChanged.connect(self._on_black_white_changed)
        black_card_layout.addWidget(self.txt_blacklist)
        v_black.addWidget(black_card)
        layout.addWidget(box_black)

        # v1.1.11: 2.6 白名单 (永不脱敏, 优先级最高)
        box_white = QFrame()
        box_white.setObjectName("settingsSectionCard")
        v_white = QVBoxLayout(box_white)
        v_white.setContentsMargins(16, 16, 16, 16)
        v_white.setSpacing(12)
        white_lead = QLabel("每行一条子串,即使被规则或黑名单命中也不会被脱敏。注意: 这里的\"永不脱敏\"会覆盖其它所有机制。")
        white_lead.setObjectName("settingsSectionLead")
        white_lead.setWordWrap(True)
        self.lbl_white_summary = QLabel("")
        self.lbl_white_summary.setObjectName("settingsSectionSummary")
        self.lbl_white_summary.setWordWrap(True)
        v_white.addWidget(self._create_settings_section_header(
            "4. 白名单 (永不脱敏)", white_lead, self.lbl_white_summary))
        white_actions = QHBoxLayout()
        white_actions.setSpacing(8)
        white_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_clear_white = QPushButton("清空白名单")
        btn_clear_white.setObjectName("settingsInlineButton")
        btn_clear_white.clicked.connect(self._clear_whitelist)
        white_actions.addWidget(btn_clear_white)
        white_actions.addStretch()
        v_white.addLayout(white_actions)
        white_card = QFrame()
        white_card.setObjectName("settingsFieldCard")
        white_card_layout = QVBoxLayout(white_card)
        white_card_layout.setContentsMargins(16, 16, 16, 16)
        white_card_layout.setSpacing(10)
        white_card_title = QLabel("白名单条目")
        white_card_title.setObjectName("settingsFieldTitle")
        white_card_note = QLabel("子串匹配 (不需完整词)。每行一条,实时生效。")
        white_card_note.setObjectName("settingsFieldNote")
        white_card_note.setWordWrap(True)
        white_card_layout.addWidget(white_card_title)
        white_card_layout.addWidget(white_card_note)
        white_card_layout.addWidget(self._create_settings_divider())
        self.txt_whitelist = QTextEdit()
        self.txt_whitelist.setPlaceholderText("例如: 公司全称 / 产品代号")
        self.txt_whitelist.setPlainText("\n".join(self._initial_whitelist))
        self.txt_whitelist.setMinimumHeight(120)
        self.txt_whitelist.setMaximumHeight(190)
        self.txt_whitelist.textChanged.connect(self._on_black_white_changed)
        white_card_layout.addWidget(self.txt_whitelist)
        v_white.addWidget(white_card)
        layout.addWidget(box_white)

        # 3. 精度与微调
        box_enhance = QFrame()
        box_enhance.setObjectName("settingsSectionCard")
        v_enhance = QVBoxLayout(box_enhance)
        v_enhance.setContentsMargins(16, 16, 16, 16)
        v_enhance.setSpacing(12)
        precision_lead = QLabel("只有当扫描偏移、覆盖范围不理想时再调整这里；默认设置更适合大多数文档。")
        precision_lead.setObjectName("settingsSectionLead")
        precision_lead.setWordWrap(True)
        self.lbl_precision_summary = QLabel("")
        self.lbl_precision_summary.setObjectName("settingsSectionSummary")
        self.lbl_precision_summary.setWordWrap(True)
        v_enhance.addWidget(self._create_settings_section_header("3. 精度与微调", precision_lead, self.lbl_precision_summary))
        precision_actions = QHBoxLayout()
        precision_actions.setSpacing(8)
        precision_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_reset_precision = QPushButton("恢复推荐值")
        btn_reset_precision.setObjectName("settingsInlineButton")
        btn_reset_precision.clicked.connect(self._reset_precision_defaults)
        precision_actions.addWidget(btn_reset_precision)
        precision_actions.addStretch()
        v_enhance.addLayout(precision_actions)

        precision_cards = QGridLayout()
        precision_cards.setHorizontalSpacing(16)
        precision_cards.setVerticalSpacing(14)
        self.settings_precision_cards_layout = precision_cards

        scan_card = QFrame()
        scan_card.setObjectName("settingsFieldCard")
        self.settings_scan_card = scan_card
        scan_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        scan_card_layout = QVBoxLayout(scan_card)
        scan_card_layout.setContentsMargins(16, 16, 16, 16)
        scan_card_layout.setSpacing(10)
        scan_title = QLabel("扫描模式")
        scan_title.setObjectName("settingsFieldTitle")
        scan_note = QLabel("正常情况下保持默认即可。模式越高越细，但速度也会更慢。")
        scan_note.setObjectName("settingsFieldNote")
        scan_note.setWordWrap(True)
        scan_card_layout.addWidget(scan_title)
        scan_card_layout.addWidget(scan_note)
        scan_card_layout.addWidget(self._create_settings_divider())
        h_precision = QHBoxLayout()
        h_precision.addWidget(self._create_settings_form_label("当前模式:", 86))
        self.combo_precision = QComboBox()
        # v1.1.11: 从配置动态添加扫描级别选项
        for level in self.available_levels:
            label = self.level_labels.get(str(level), f"{level}x")
            self.combo_precision.addItem(label, level)
        idx = self.combo_precision.findData(scan_level)
        self.combo_precision.setCurrentIndex(idx if idx >=0 else 1)
        self.combo_precision.currentIndexChanged.connect(self._refresh_precision_summary)
        h_precision.addWidget(self.combo_precision)
        scan_card_layout.addLayout(h_precision)
        scan_card_layout.addStretch()
        precision_cards.addWidget(scan_card, 0, 0)

        calibrate_card = QFrame()
        calibrate_card.setObjectName("settingsFieldCard")
        self.settings_calibrate_card = calibrate_card
        calibrate_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        calibrate_card_layout = QVBoxLayout(calibrate_card)
        calibrate_card_layout.setContentsMargins(16, 16, 16, 16)
        calibrate_card_layout.setSpacing(10)
        calibrate_title = QLabel("覆盖微调")
        calibrate_title.setObjectName("settingsFieldTitle")
        calibrate_note = QLabel("只有在遮罩范围明显偏左、偏宽时再调，默认值更适合多数 PDF。")
        calibrate_note.setObjectName("settingsFieldNote")
        calibrate_note.setWordWrap(True)
        calibrate_card_layout.addWidget(calibrate_title)
        calibrate_card_layout.addWidget(calibrate_note)
        calibrate_card_layout.addWidget(self._create_settings_divider())
        h_calibrate = QHBoxLayout()
        v_cal_1 = QVBoxLayout()
        label_offset_x = QLabel("向左修正(px):")
        label_offset_x.setObjectName("settingsFieldLabel")
        v_cal_1.addWidget(label_offset_x)
        self.spin_offset_x = QSpinBox()
        self.spin_offset_x.setRange(x_range[0], x_range[1])
        self.spin_offset_x.setValue(offset_x)
        self.spin_offset_x.valueChanged.connect(self._refresh_precision_summary)
        v_cal_1.addWidget(self.spin_offset_x)

        v_cal_2 = QVBoxLayout()
        label_offset_w = QLabel("宽度收缩(px):")
        label_offset_w.setObjectName("settingsFieldLabel")
        v_cal_2.addWidget(label_offset_w)
        self.spin_offset_w = QSpinBox()
        self.spin_offset_w.setRange(w_range[0], w_range[1])
        self.spin_offset_w.setValue(offset_w)
        self.spin_offset_w.valueChanged.connect(self._refresh_precision_summary)
        v_cal_2.addWidget(self.spin_offset_w)

        h_calibrate.addLayout(v_cal_1)
        h_calibrate.addLayout(v_cal_2)
        calibrate_card_layout.addLayout(h_calibrate)
        calibrate_tip = QLabel("提示：对扫描区域 / 嵌入图片区域生效")
        calibrate_tip.setObjectName("settingsFieldNote")
        calibrate_tip.setWordWrap(True)
        calibrate_card_layout.addWidget(calibrate_tip)

        self.cb_enhance = QCheckBox("开启图像增强 (仅针对手写体)")
        self.cb_enhance.setChecked(use_enhance)
        self.cb_enhance.toggled.connect(self._refresh_precision_summary)
        calibrate_card_layout.addWidget(self.cb_enhance)
        calibrate_card_layout.addStretch()
        precision_cards.addWidget(calibrate_card, 0, 1)

        v_enhance.addLayout(precision_cards)
        layout.addWidget(box_enhance)

        # v1.1.11: 4. OCR 检测框调节（移除引擎选择，只保留 RapidOCR）
        box_ocr = QFrame()
        box_ocr.setObjectName("settingsSectionCard")
        v_ocr = QVBoxLayout(box_ocr)
        v_ocr.setContentsMargins(16, 16, 16, 16)
        v_ocr.setSpacing(12)
        ocr_lead = QLabel("检测框只在 OCR 结果偏大或偏小时再调。负值扩大框，正值收缩框，0 表示保持原样。")
        ocr_lead.setObjectName("settingsSectionLead")
        ocr_lead.setWordWrap(True)
        self.lbl_ocr_summary = QLabel("")
        self.lbl_ocr_summary.setObjectName("settingsSectionSummary")
        self.lbl_ocr_summary.setWordWrap(True)
        v_ocr.addWidget(self._create_settings_section_header("4. OCR 检测框调节", ocr_lead, self.lbl_ocr_summary))
        ocr_actions = QHBoxLayout()
        ocr_actions.setSpacing(8)
        ocr_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        btn_reset_ocr = QPushButton("恢复 0%")
        btn_reset_ocr.setObjectName("settingsInlineButton")
        btn_reset_ocr.clicked.connect(self._reset_ocr_adjustment)
        ocr_actions.addWidget(btn_reset_ocr)
        ocr_actions.addStretch()
        v_ocr.addLayout(ocr_actions)

        adjust_card = QFrame()
        adjust_card.setObjectName("settingsFieldCard")
        self.settings_adjust_card = adjust_card
        adjust_card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        adjust_card_layout = QVBoxLayout(adjust_card)
        adjust_card_layout.setContentsMargins(16, 16, 16, 16)
        adjust_card_layout.setSpacing(10)
        adjust_title = QLabel("检测框调节")
        adjust_title.setObjectName("settingsFieldTitle")
        adjust_note = QLabel("当 OCR 框偏大或偏小时再调。负值扩大，正值收缩。")
        adjust_note.setObjectName("settingsFieldNote")
        adjust_note.setWordWrap(True)
        adjust_card_layout.addWidget(adjust_title)
        adjust_card_layout.addWidget(adjust_note)
        adjust_card_layout.addWidget(self._create_settings_divider())

        # v1.1.11: 检测框大小调节（支持负值扩大、正值收缩）
        h_adjust = QHBoxLayout()
        h_adjust.addWidget(self._create_settings_form_label("检测框调节:", 96))

        # 读取当前配置值（新配置名）
        adjust_ratio = config.get("ocr.box_adjust_ratio", 0.0) if config else 0.0

        self.slider_adjust = QSlider(Qt.Orientation.Horizontal)
        self.slider_adjust.setRange(-30, 50)  # -30% 到 +50%
        self.slider_adjust.setValue(int(adjust_ratio * 100))
        self.slider_adjust.valueChanged.connect(self._on_adjust_changed)
        h_adjust.addWidget(self.slider_adjust)

        self.label_adjust_value = QLabel(f"{int(adjust_ratio * 100)}%")
        self.label_adjust_value.setMinimumWidth(40)
        h_adjust.addWidget(self.label_adjust_value)

        adjust_card_layout.addLayout(h_adjust)

        adjust_info = QLabel("提示：负值扩大，正值收缩，0保持原样（默认0%）")
        adjust_info.setObjectName("settingsFieldNote")
        adjust_info.setWordWrap(True)
        adjust_card_layout.addWidget(adjust_info)

        # 说明标签
        info_text = (
            "\n引擎说明：\n"
            "• RapidOCR：默认 OCR 引擎，速度快，适合大文档批量处理\n"
        )

        info = QLabel(info_text)
        info.setObjectName("settingsFieldNote")
        info.setWordWrap(True)
        adjust_card_layout.addWidget(info)
        v_ocr.addWidget(adjust_card)
        layout.addWidget(box_ocr)

        # v1.1.11: 5. 永久 override 维护
        box_overrides = QFrame()
        box_overrides.setObjectName("settingsSectionCard")
        v_overrides = QVBoxLayout(box_overrides)
        v_overrides.setContentsMargins(16, 16, 16, 16)
        v_overrides.setSpacing(12)
        ov_lead = QLabel("维护永久 ignore / confirm 条目。永久 override 会跨会话生效,建议定期清理过期记录。")
        ov_lead.setObjectName("settingsSectionLead")
        ov_lead.setWordWrap(True)
        self.lbl_overrides_summary = QLabel("")
        self.lbl_overrides_summary.setObjectName("settingsSectionSummary")
        self.lbl_overrides_summary.setWordWrap(True)
        v_overrides.addWidget(
            self._create_settings_section_header(
                "5. 永久 override 名单", ov_lead, self.lbl_overrides_summary
            )
        )
        ov_actions = QHBoxLayout()
        ov_actions.setSpacing(8)
        ov_actions.addWidget(self._create_settings_action_hint("快捷操作"))
        clean_btn = QPushButton("清理 30 天前失效的 permanent overrides")
        clean_btn.setObjectName("settingsInlineButton")
        clean_btn.clicked.connect(self._on_clean_stale_overrides)
        ov_actions.addWidget(clean_btn)
        ov_actions.addStretch()
        v_overrides.addLayout(ov_actions)
        ov_card = QFrame()
        ov_card.setObjectName("settingsFieldCard")
        ov_card_layout = QVBoxLayout(ov_card)
        ov_card_layout.setContentsMargins(16, 16, 16, 16)
        ov_card_layout.setSpacing(8)
        ov_card_title = QLabel("Permanent 列表")
        ov_card_title.setObjectName("settingsFieldTitle")
        ov_card_note = QLabel("永久条目存于 config.json 的 redaction.overrides.permanent,共 N 条。")
        ov_card_note.setObjectName("settingsFieldNote")
        ov_card_note.setWordWrap(True)
        ov_card_layout.addWidget(ov_card_title)
        ov_card_layout.addWidget(ov_card_note)
        v_overrides.addWidget(ov_card)
        layout.addWidget(box_overrides)

        layout.addStretch(1)
        content_scroll.setWidget(content_widget)
        body_layout.addWidget(content_scroll, stretch=1)

        footer = QFrame(self)
        footer.setObjectName("settingsFooter")
        footer_layout = QHBoxLayout(footer)
        footer_layout.setContentsMargins(18, 14, 18, 14)
        footer_layout.setSpacing(12)
        self.settings_footer_layout = footer_layout

        footer_note = QLabel("修改完成后点击“保存设置”，取消不会影响当前已生效配置。")
        footer_note.setObjectName("settingsFooterNote")
        footer_note.setWordWrap(True)
        self.settings_footer_note = footer_note
        footer_layout.addWidget(footer_note, stretch=1)

        footer_actions = QHBoxLayout()
        footer_actions.setSpacing(10)
        btn_cancel = QPushButton("取消")
        btn_cancel.setObjectName("settingsSecondaryButton")
        btn_cancel.clicked.connect(self.reject)
        btn_cancel.setMinimumHeight(40)
        btn_cancel.setMinimumWidth(116)
        self.btn_settings_cancel = btn_cancel
        footer_actions.addWidget(btn_cancel)

        btn_ok = QPushButton("保存设置")
        btn_ok.setObjectName("settingsPrimaryButton")
        btn_ok.clicked.connect(self.save_settings)
        btn_ok.setMinimumHeight(40)
        btn_ok.setMinimumWidth(164)
        self.btn_settings_save = btn_ok
        footer_actions.addWidget(btn_ok)
        footer_layout.addLayout(footer_actions)
        outer_layout.addWidget(footer)
        self._settings_sections = [box_rules, box_custom, box_black, box_white, box_enhance, box_ocr, box_overrides]
        self._settings_nav_syncing = False
        self.settings_nav.currentRowChanged.connect(self._scroll_to_settings_section)
        self.content_scroll.verticalScrollBar().valueChanged.connect(self._sync_settings_nav_from_scroll)
        self.settings_nav.setCurrentRow(0)
        self._refresh_word_rule_summary()
        self._refresh_rule_summary()
        self._refresh_precision_summary()
        self._refresh_ocr_summary()
        self._refresh_overrides_summary()
        self._refresh_settings_layout_density()

    def _on_adjust_changed(self, value):
        """v1.1.11: 检测框调节滑块值变化回调"""
        self.label_adjust_value.setText(f"{value}%")
        self._refresh_ocr_summary()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._refresh_settings_layout_density()

    def showEvent(self, event):
        super().showEvent(event)
        self._bind_settings_window_handle_signals()
        QTimer.singleShot(0, self._refresh_settings_layout_density)

    def _get_settings_display_scale_factor(self):
        """返回设置窗口当前屏幕缩放，仅在 Windows 下生效。"""
        import platform

        if platform.system() != "Windows":
            return 1.0

        screen = None
        try:
            handle = self.windowHandle()
            if handle:
                screen = handle.screen()
        except Exception:
            screen = None

        if screen is None:
            screen = QApplication.primaryScreen()
        if screen is None:
            return 1.0

        try:
            scale = screen.logicalDotsPerInch() / 96.0
        except Exception:
            scale = 1.0

        return max(1.0, min(scale, 2.0))

    def _bind_settings_window_handle_signals(self):
        """绑定设置窗口跨屏信号，保证 DPI 切换后密度立即刷新。"""
        try:
            handle = self.windowHandle()
        except Exception:
            handle = None

        if handle is None or handle is self._settings_bound_window_handle:
            return

        try:
            handle.screenChanged.connect(self._on_settings_screen_changed)
        except Exception:
            pass
        self._settings_bound_window_handle = handle

    def _on_settings_screen_changed(self, _screen):
        """设置窗口切换屏幕后刷新密度。"""
        QTimer.singleShot(0, self._refresh_settings_layout_density)

    def _create_settings_metric_card(self, title):
        """创建设置页顶部概览指标卡。"""
        card = QFrame()
        card.setObjectName("settingsMetricCard")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(4)

        title_label = QLabel(title)
        title_label.setObjectName("settingsMetricLabel")
        value_label = QLabel("--")
        value_label.setObjectName("settingsMetricValue")
        value_label.setWordWrap(True)

        layout.addWidget(title_label)
        layout.addWidget(value_label)
        layout.addStretch()
        return card, value_label

    def _rebuild_settings_grid(self, layout, widgets, columns, stretches=None):
        """按列数重排设置页网格布局。"""
        if not layout:
            return
        while layout.count():
            layout.takeAt(0)

        max_slots = max(len(widgets), 4)
        for column in range(max_slots):
            layout.setColumnStretch(column, 0)
        for row in range(max_slots):
            layout.setRowStretch(row, 0)

        columns = max(1, columns)
        for index, widget in enumerate(widgets):
            row = index // columns
            column = index % columns
            layout.addWidget(widget, row, column)

        if stretches:
            for column, stretch in enumerate(stretches):
                layout.setColumnStretch(column, stretch)
        else:
            for column in range(columns):
                layout.setColumnStretch(column, 1)

    def _refresh_settings_layout_density(self):
        """按当前窗口宽度微调设置页整体比例。"""
        width = max(self.width(), 1)
        height = max(self.height(), 1)
        scale = self._get_settings_display_scale_factor()
        density_mode = resolve_settings_density_mode(width, height, scale)
        is_tall_window = height >= 980
        is_short_window = height <= 820
        is_very_wide_window = width >= 1620
        is_ultra_wide_window = width >= 1920
        is_cinema_wide_window = width >= 2140

        if density_mode == "wide":
            sidebar_width = 264
            body_spacing = 22
            hero_margins = (22, 18, 22, 18)
            overview_margins = (22, 18, 22, 18)
            footer_margins = (22, 14, 22, 14)
            hero_tag_spacing = 10
            metric_spacing = 12
            action_spacing = 10
            custom_spacing = 24
            precision_spacing = 18
            jump_min_width = 134
            quick_jump_compact = False
            overview_metric_columns = 4
            overview_action_columns = 4
            custom_columns = 2
            precision_columns = 2
            custom_stretches = (3, 2)
            precision_stretches = (1, 1)
            right_panel_max_width = 680
            right_panel_min_width = 430
            left_card_min_height = 348
            right_card_min_height = 348
            precision_card_min_height = 238
            adjust_card_min_height = 234
            cancel_min_width = 122
            save_min_width = 172
        elif density_mode == "roomy":
            sidebar_width = 248
            body_spacing = 18
            hero_margins = (18, 16, 18, 16)
            overview_margins = (18, 16, 18, 16)
            footer_margins = (18, 14, 18, 14)
            hero_tag_spacing = 8
            metric_spacing = 10
            action_spacing = 8
            custom_spacing = 22
            precision_spacing = 16
            jump_min_width = 124
            quick_jump_compact = False
            overview_metric_columns = 4
            overview_action_columns = 4
            custom_columns = 2
            precision_columns = 2
            custom_stretches = (3, 2)
            precision_stretches = (1, 1)
            right_panel_max_width = 620
            right_panel_min_width = 392
            left_card_min_height = 330
            right_card_min_height = 330
            precision_card_min_height = 224
            adjust_card_min_height = 220
            cancel_min_width = 116
            save_min_width = 164
        elif density_mode == "compact":
            sidebar_width = 232
            body_spacing = 16
            hero_margins = (16, 15, 16, 15)
            overview_margins = (16, 15, 16, 15)
            footer_margins = (16, 12, 16, 12)
            hero_tag_spacing = 8
            metric_spacing = 8
            action_spacing = 8
            custom_spacing = 18
            precision_spacing = 14
            jump_min_width = 114
            quick_jump_compact = False
            overview_metric_columns = 2
            overview_action_columns = 2
            custom_columns = 2
            precision_columns = 1
            custom_stretches = (7, 5)
            precision_stretches = (1,)
            right_panel_max_width = 540
            right_panel_min_width = 344
            left_card_min_height = 302
            right_card_min_height = 302
            precision_card_min_height = 208
            adjust_card_min_height = 204
            cancel_min_width = 112
            save_min_width = 156
        else:
            sidebar_width = 220
            body_spacing = 14
            hero_margins = (14, 14, 14, 14)
            overview_margins = (14, 14, 14, 14)
            footer_margins = (14, 12, 14, 12)
            hero_tag_spacing = 6
            metric_spacing = 8
            action_spacing = 6
            custom_spacing = 14
            precision_spacing = 12
            jump_min_width = 102
            quick_jump_compact = True
            overview_metric_columns = 2
            overview_action_columns = 2
            custom_columns = 1
            precision_columns = 1
            custom_stretches = (1,)
            precision_stretches = (1,)
            right_panel_max_width = 9999
            right_panel_min_width = 0
            left_card_min_height = 270
            right_card_min_height = 270
            precision_card_min_height = 194
            adjust_card_min_height = 194
            cancel_min_width = 108
            save_min_width = 148

        if scale >= 1.5:
            body_spacing += 2
            hero_tag_spacing += 1
            metric_spacing += 2
            action_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            jump_min_width += 6
            cancel_min_width += 8
            save_min_width += 10
            left_card_min_height += 14
            right_card_min_height += 14
            precision_card_min_height += 10
            adjust_card_min_height += 10
        elif scale >= 1.25:
            jump_min_width += 4
            cancel_min_width += 6
            save_min_width += 8
            left_card_min_height += 8
            right_card_min_height += 8
            precision_card_min_height += 6
            adjust_card_min_height += 6

        if is_tall_window:
            left_card_min_height += 18
            right_card_min_height += 18
            precision_card_min_height += 10
            adjust_card_min_height += 10
            hero_margins = (
                hero_margins[0] + 2,
                hero_margins[1] + 2,
                hero_margins[2] + 2,
                hero_margins[3] + 2,
            )
            overview_margins = (
                overview_margins[0] + 2,
                overview_margins[1] + 2,
                overview_margins[2] + 2,
                overview_margins[3] + 2,
            )
        elif is_short_window:
            body_spacing = max(12, body_spacing - 2)
            hero_tag_spacing = max(6, hero_tag_spacing - 1)
            metric_spacing = max(8, metric_spacing - 2)
            action_spacing = max(6, action_spacing - 2)
            custom_spacing = max(14, custom_spacing - 2)
            precision_spacing = max(12, precision_spacing - 2)
            left_card_min_height = max(270, left_card_min_height - 18)
            right_card_min_height = max(270, right_card_min_height - 18)
            precision_card_min_height = max(194, precision_card_min_height - 10)
            adjust_card_min_height = max(194, adjust_card_min_height - 10)
            footer_margins = (
                footer_margins[0],
                max(10, footer_margins[1] - 2),
                footer_margins[2],
                max(10, footer_margins[3] - 2),
            )

        if is_very_wide_window:
            sidebar_width += 10
            body_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            right_panel_max_width += 60 if right_panel_max_width < 9000 else 0
            right_panel_min_width += 24 if right_panel_min_width else 0
            left_card_min_height += 10
            right_card_min_height += 10
            precision_card_min_height += 8
            adjust_card_min_height += 8
            hero_margins = (
                hero_margins[0] + 2,
                hero_margins[1] + 2,
                hero_margins[2] + 2,
                hero_margins[3] + 2,
            )
            overview_margins = (
                overview_margins[0] + 2,
                overview_margins[1] + 2,
                overview_margins[2] + 2,
                overview_margins[3] + 2,
            )

        if is_ultra_wide_window:
            sidebar_width += 8
            body_spacing += 2
            hero_tag_spacing += 1
            metric_spacing += 2
            action_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            jump_min_width += 6
            right_panel_max_width += 80 if right_panel_max_width < 9000 else 0
            right_panel_min_width += 28 if right_panel_min_width else 0
            left_card_min_height += 10
            right_card_min_height += 10

        if is_cinema_wide_window:
            sidebar_width += 10
            body_spacing += 2
            hero_tag_spacing += 1
            metric_spacing += 2
            action_spacing += 2
            custom_spacing += 2
            precision_spacing += 2
            jump_min_width += 8
            right_panel_max_width += 120 if right_panel_max_width < 9000 else 0
            right_panel_min_width += 36 if right_panel_min_width else 0
            left_card_min_height += 10
            right_card_min_height += 10
            precision_card_min_height += 8
            adjust_card_min_height += 8
            hero_margins = (
                hero_margins[0] + 4,
                hero_margins[1] + 2,
                hero_margins[2] + 4,
                hero_margins[3] + 2,
            )
            overview_margins = (
                overview_margins[0] + 4,
                overview_margins[1] + 2,
                overview_margins[2] + 4,
                overview_margins[3] + 2,
            )

        quick_jump_height = 34 if scale >= 1.25 else 32
        footer_button_height = 42 if scale >= 1.25 else 40

        if hasattr(self, "settings_sidebar"):
            self.settings_sidebar.setFixedWidth(sidebar_width)
        if hasattr(self, "settings_body_layout"):
            self.settings_body_layout.setSpacing(body_spacing)
        if hasattr(self, "settings_hero_layout"):
            self.settings_hero_layout.setContentsMargins(*hero_margins)
        if hasattr(self, "settings_hero_tag_layout"):
            self.settings_hero_tag_layout.setSpacing(hero_tag_spacing)
        if hasattr(self, "settings_overview_layout"):
            self.settings_overview_layout.setContentsMargins(*overview_margins)
        if hasattr(self, "settings_overview_metrics_layout"):
            self.settings_overview_metrics_layout.setHorizontalSpacing(metric_spacing)
            self.settings_overview_metrics_layout.setVerticalSpacing(metric_spacing)
            self._rebuild_settings_grid(
                self.settings_overview_metrics_layout,
                self.settings_metric_cards,
                overview_metric_columns,
            )
        if hasattr(self, "settings_overview_actions_layout"):
            self.settings_overview_actions_layout.setHorizontalSpacing(action_spacing)
            self.settings_overview_actions_layout.setVerticalSpacing(action_spacing)
            self._rebuild_settings_grid(
                self.settings_overview_actions_layout,
                self.settings_quick_jump_buttons,
                overview_action_columns,
            )
        if hasattr(self, "settings_footer_layout"):
            self.settings_footer_layout.setContentsMargins(*footer_margins)
        for index, button in enumerate(getattr(self, "settings_quick_jump_buttons", [])):
            button.setMinimumWidth(jump_min_width)
            button.setMinimumHeight(quick_jump_height)
            button.setMaximumHeight(quick_jump_height)
            if hasattr(self, "_settings_quick_jump_titles") and index < len(self._settings_quick_jump_titles):
                full_title, short_title = self._settings_quick_jump_titles[index]
                button.setText(short_title if quick_jump_compact else full_title)
        if hasattr(self, "btn_settings_cancel"):
            self.btn_settings_cancel.setMinimumWidth(cancel_min_width)
            self.btn_settings_cancel.setMinimumHeight(footer_button_height)
            self.btn_settings_cancel.setMaximumHeight(footer_button_height)
        if hasattr(self, "btn_settings_save"):
            self.btn_settings_save.setMinimumWidth(save_min_width)
            self.btn_settings_save.setMinimumHeight(footer_button_height)
            self.btn_settings_save.setMaximumHeight(footer_button_height)
        if hasattr(self, "settings_overview_text"):
            overview_text_width = 920 if density_mode == "wide" else 860 if density_mode == "roomy" else 760 if density_mode == "compact" else 680
            if is_very_wide_window:
                overview_text_width += 80
            if is_ultra_wide_window:
                overview_text_width += 80
            if is_cinema_wide_window:
                overview_text_width += 120
            self.settings_overview_text.setMaximumWidth(overview_text_width)
        if hasattr(self, "settings_hero_subtitle"):
            hero_subtitle_width = 1000 if density_mode == "wide" else 930 if density_mode == "roomy" else 840 if density_mode == "compact" else 720
            if is_very_wide_window:
                hero_subtitle_width += 100
            if is_ultra_wide_window:
                hero_subtitle_width += 100
            if is_cinema_wide_window:
                hero_subtitle_width += 140
            self.settings_hero_subtitle.setMaximumWidth(hero_subtitle_width)
        if hasattr(self, "settings_custom_row"):
            self.settings_custom_row.setHorizontalSpacing(custom_spacing)
            self.settings_custom_row.setVerticalSpacing(max(12, custom_spacing - 8))
            self._rebuild_settings_grid(
                self.settings_custom_row,
                [self.settings_left_field_card, self.settings_right_field_card],
                custom_columns,
                custom_stretches,
            )
        if hasattr(self, "settings_precision_cards_layout"):
            self.settings_precision_cards_layout.setHorizontalSpacing(precision_spacing)
            self.settings_precision_cards_layout.setVerticalSpacing(max(12, precision_spacing - 2))
            self._rebuild_settings_grid(
                self.settings_precision_cards_layout,
                [self.settings_scan_card, self.settings_calibrate_card],
                precision_columns,
                precision_stretches,
            )
        if hasattr(self, "settings_right_field_card"):
            self.settings_right_field_card.setMaximumWidth(right_panel_max_width)
            self.settings_right_field_card.setMinimumWidth(right_panel_min_width)
            self.settings_right_field_card.setMinimumHeight(right_card_min_height)
        if hasattr(self, "settings_left_field_card"):
            self.settings_left_field_card.setMinimumHeight(left_card_min_height)
        if hasattr(self, "settings_scan_card"):
            self.settings_scan_card.setMinimumHeight(precision_card_min_height)
        if hasattr(self, "settings_calibrate_card"):
            self.settings_calibrate_card.setMinimumHeight(precision_card_min_height)
        if hasattr(self, "settings_adjust_card"):
            self.settings_adjust_card.setMinimumHeight(adjust_card_min_height)
        self._apply_settings_density_styles(density_mode, scale, is_short_window)

    def _create_settings_section_title(self, title):
        """创建设置模块卡片内标题。"""
        title_label = QLabel(title)
        title_label.setObjectName("settingsSectionTitle")
        return title_label

    def _create_settings_section_header(self, title, lead_label, summary_label):
        """创建设置模块卡片内部统一头部。"""
        header = QFrame()
        header.setObjectName("settingsSectionHeader")
        layout = QVBoxLayout(header)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(10)
        layout.addWidget(self._create_settings_section_title(title))
        layout.addWidget(lead_label)
        layout.addWidget(summary_label)
        return header

    def _create_settings_action_hint(self, text):
        """创建设置分区内的轻量操作提示。"""
        hint = QLabel(text)
        hint.setObjectName("settingsActionHint")
        return hint

    def _create_settings_form_label(self, text, width=96):
        """创建设置子卡中的统一表单标签。"""
        label = QLabel(text)
        label.setObjectName("settingsFieldLabel")
        label.setFixedWidth(width)
        return label

    def _create_settings_divider(self):
        """创建设置子卡内的轻量分隔线。"""
        divider = QFrame()
        divider.setObjectName("settingsFieldDivider")
        divider.setMinimumHeight(1)
        divider.setMaximumHeight(1)
        return divider

    def _apply_settings_density_styles(self, density_mode, scale, is_short_window):
        """按当前设置页密度统一标题、说明、导航与按钮字级。"""
        theme = Theme.LIGHT
        if density_mode == "wide":
            hero_title_size = 21
            hero_subtitle_size = 13
            hero_tag_size = 12
            overview_title_size = 17
            overview_text_size = 12
            section_title_size = 16
            lead_size = 12
            summary_size = 12
            field_title_size = 13
            field_label_size = 12
            field_note_size = 12
            hint_size = 11
            nav_font_size = 12
            nav_padding_v = 12
            nav_padding_h = 14
            nav_margin_v = 3
            metric_label_size = 11
            metric_value_size = 17
            quick_jump_font_size = 12
            footer_note_size = 12
        elif density_mode == "roomy":
            hero_title_size = 20
            hero_subtitle_size = 12
            hero_tag_size = 11
            overview_title_size = 16
            overview_text_size = 11
            section_title_size = 15
            lead_size = 11
            summary_size = 11
            field_title_size = 12
            field_label_size = 11
            field_note_size = 11
            hint_size = 11
            nav_font_size = 12
            nav_padding_v = 11
            nav_padding_h = 13
            nav_margin_v = 2
            metric_label_size = 11
            metric_value_size = 16
            quick_jump_font_size = 12
            footer_note_size = 11
        elif density_mode == "compact":
            hero_title_size = 19
            hero_subtitle_size = 12
            hero_tag_size = 11
            overview_title_size = 15
            overview_text_size = 11
            section_title_size = 15
            lead_size = 11
            summary_size = 11
            field_title_size = 12
            field_label_size = 11
            field_note_size = 11
            hint_size = 10
            nav_font_size = 11
            nav_padding_v = 10
            nav_padding_h = 12
            nav_margin_v = 2
            metric_label_size = 11
            metric_value_size = 15
            quick_jump_font_size = 11
            footer_note_size = 11
        else:
            hero_title_size = 18
            hero_subtitle_size = 11
            hero_tag_size = 10
            overview_title_size = 15
            overview_text_size = 11
            section_title_size = 14
            lead_size = 11
            summary_size = 11
            field_title_size = 12
            field_label_size = 11
            field_note_size = 11
            hint_size = 10
            nav_font_size = 11
            nav_padding_v = 10
            nav_padding_h = 11
            nav_margin_v = 2
            metric_label_size = 10
            metric_value_size = 15
            quick_jump_font_size = 11
            footer_note_size = 11

        if scale >= 1.5:
            hero_title_size += 1
            hero_subtitle_size += 1
            overview_title_size += 1
            section_title_size += 1
            lead_size += 1
            summary_size += 1
            field_title_size += 1
            field_label_size += 1
            field_note_size += 1
            nav_font_size += 1
            metric_value_size += 1
            quick_jump_font_size += 1
            footer_note_size += 1
            nav_padding_v += 1
            nav_padding_h += 1
        elif scale >= 1.25:
            hero_subtitle_size += 1
            overview_text_size += 1
            lead_size += 1
            summary_size += 1
            field_note_size += 1
            nav_padding_h += 1

        if is_short_window:
            hero_title_size = max(18, hero_title_size - 1)
            hero_subtitle_size = max(11, hero_subtitle_size - 1)
            section_title_size = max(14, section_title_size - 1)
            lead_size = max(11, lead_size - 1)
            summary_size = max(11, summary_size - 1)
            nav_padding_v = max(9, nav_padding_v - 1)
            nav_margin_v = 1

        if hasattr(self, "settings_hero_title"):
            self.settings_hero_title.setStyleSheet(
                f"color: {theme['text']}; font-size: {hero_title_size}px; font-weight: 700;"
            )
        if hasattr(self, "settings_hero_subtitle"):
            self.settings_hero_subtitle.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {hero_subtitle_size}px; line-height: 1.6;"
            )
        for label in [getattr(self, "lbl_settings_common_tag", None), getattr(self, "lbl_settings_advanced_tag", None)]:
            if label:
                label.setStyleSheet(
                    f"color: {theme['primary']}; background-color: #E9F1FB; border: 1px solid {theme['border']}; "
                    f"border-radius: 10px; padding: 6px 10px; font-size: {hero_tag_size}px; font-weight: 700;"
                )
        if hasattr(self, "settings_overview_title"):
            self.settings_overview_title.setStyleSheet(
                f"color: {theme['text']}; font-size: {overview_title_size}px; font-weight: 700;"
            )
        if hasattr(self, "settings_overview_text"):
            self.settings_overview_text.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {overview_text_size}px; line-height: 1.6;"
            )
        if hasattr(self, "settings_nav_hint"):
            self.settings_nav_hint.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {hint_size}px; font-weight: 700; letter-spacing: 0.05em;"
            )
        if hasattr(self, "settings_nav_subtitle"):
            self.settings_nav_subtitle.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {field_note_size}px; line-height: 1.6;"
            )
        if hasattr(self, "settings_sidebar_note"):
            self.settings_sidebar_note.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {field_note_size}px; line-height: 1.6;"
            )
        if hasattr(self, "settings_sidebar_status"):
            self.settings_sidebar_status.setStyleSheet(
                f"color: {theme['text_secondary']}; background-color: #F8FBFE; border: 1px solid {theme['border']}; "
                f"border-radius: 10px; padding: 8px 10px; font-size: {field_note_size}px; line-height: 1.6; font-weight: 600;"
            )
        if hasattr(self, "settings_footer_note"):
            self.settings_footer_note.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {footer_note_size}px; line-height: 1.6; padding-right: 8px;"
            )
        if hasattr(self, "settings_nav"):
            self.settings_nav.setStyleSheet(
                f"""
                QListWidget#settingsNav {{
                    background-color: transparent;
                    border: none;
                    outline: none;
                    padding: 0;
                }}
                QListWidget#settingsNav::item {{
                    background-color: transparent;
                    border: 1px solid transparent;
                    border-radius: 10px;
                    padding: {nav_padding_v}px {nav_padding_h}px;
                    margin: {nav_margin_v}px 0;
                    color: {theme["text"]};
                    font-size: {nav_font_size}px;
                    font-weight: 600;
                }}
                QListWidget#settingsNav::item:selected {{
                    background-color: {theme["hover"]};
                    border-color: {theme["border"]};
                    color: {theme["primary"]};
                }}
                QListWidget#settingsNav::item:hover {{
                    background-color: {theme["hover"]};
                }}
                """
            )
        for label in self.findChildren(QLabel, "settingsMetricLabel"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {metric_label_size}px; font-weight: 700;"
            )
        for label in self.findChildren(QLabel, "settingsMetricValue"):
            label.setStyleSheet(
                f"color: {theme['text']}; font-size: {metric_value_size}px; font-weight: 700;"
            )
        for label in self.findChildren(QLabel, "settingsSectionTitle"):
            label.setStyleSheet(
                f"color: {theme['primary']}; font-size: {section_title_size}px; font-weight: 700; background-color: transparent;"
            )
        for label in self.findChildren(QLabel, "settingsSectionLead"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {lead_size}px; line-height: 1.7;"
            )
        for label in self.findChildren(QLabel, "settingsSectionSummary"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; background-color: #F8FBFE; border: 1px solid {theme['border']}; "
                f"border-radius: 10px; padding: 8px 12px; font-size: {summary_size}px; font-weight: 600; line-height: 1.6;"
            )
        for label in self.findChildren(QLabel, "settingsFieldTitle"):
            label.setStyleSheet(
                f"color: {theme['text']}; font-size: {field_title_size}px; font-weight: 700;"
            )
        for label in self.findChildren(QLabel, "settingsFieldLabel"):
            label.setStyleSheet(
                f"color: {theme['text']}; font-size: {field_label_size}px; font-weight: 600;"
            )
        for label in self.findChildren(QLabel, "settingsFieldNote"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {field_note_size}px; line-height: 1.6;"
            )
        for label in self.findChildren(QLabel, "settingsActionHint"):
            label.setStyleSheet(
                f"color: {theme['text_secondary']}; font-size: {hint_size}px; font-weight: 700; letter-spacing: 0.04em;"
            )
        for button in getattr(self, "settings_quick_jump_buttons", []):
            button.setStyleSheet(
                f"""
                QPushButton {{
                    background-color: #FBFCFE;
                    color: {theme["text"]};
                    border: 1px solid {theme["border"]};
                    border-radius: 10px;
                    padding: 8px 14px;
                    font-size: {quick_jump_font_size}px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background-color: {theme["hover"]};
                    border-color: {theme["primary"]};
                    color: {theme["primary"]};
                }}
                """
            )

    def _set_rule_checkbox_states(self, enabled_names):
        """批量更新规则勾选状态，避免多次触发联动刷新。"""
        enabled_names = set(enabled_names)
        for checkbox in self.checks.values():
            checkbox.blockSignals(True)
        try:
            for name, checkbox in self.checks.items():
                checkbox.setChecked(name in enabled_names)
        finally:
            for checkbox in self.checks.values():
                checkbox.blockSignals(False)
        self._refresh_rule_summary()

    def _apply_recommended_rules(self):
        """恢复首次使用推荐的通用规则勾选。"""
        self._set_rule_checkbox_states(self.recommended_rule_names)

    def _select_all_rules(self):
        """勾选全部通用规则。"""
        self._set_rule_checkbox_states(self.checks.keys())

    def _clear_all_rules(self):
        """清空所有通用规则勾选。"""
        self._set_rule_checkbox_states([])

    def _clear_custom_keywords(self):
        """清空自定义关键词文本框。"""
        self.txt_custom.clear()

    def _clear_blacklist(self):
        """清空黑名单文本框。"""
        self.txt_blacklist.clear()

    def _clear_whitelist(self):
        """清空白名单文本框。"""
        self.txt_whitelist.clear()

    @staticmethod
    def _parse_lines(text):
        """按行解析文本, 去空 / 去空白 / 去重保序.

        用于黑/白名单的文本编辑框 → store 列表转换.
        """
        if not isinstance(text, str):
            return []
        seen = set()
        out = []
        for line in text.splitlines():
            stripped = line.strip()
            if stripped and stripped not in seen:
                seen.add(stripped)
                out.append(stripped)
        return out

    def _on_black_white_changed(self):
        """v1.1.11: 黑/白名单文本变化时回写 store + config.json.

        通过公开 API BlackWhiteListStore.update_permanent 写入,
        会话层独立保留, 与 _initial_blacklist / _initial_whitelist 无关.
        """
        try:
            black = self._parse_lines(self.txt_blacklist.toPlainText())
            white = self._parse_lines(self.txt_whitelist.toPlainText())
            store = BlackWhiteListStore.instance()
            store.update_permanent(black, white)
            store.save_permanent()
            if hasattr(self, "lbl_black_summary"):
                self.lbl_black_summary.setText(f"当前条目: {len(black)} 条 (已写入 config.json)")
            if hasattr(self, "lbl_white_summary"):
                self.lbl_white_summary.setText(f"当前条目: {len(white)} 条 (已写入 config.json)")
            self._refresh_settings_sidebar()
        except Exception as exc:
            print(f"[Settings] 黑/白名单回写失败: {exc}")

    def _reset_replacement_text(self):
        """恢复默认统一替换文本。"""
        self.input_replacement_text.setText(self.default_replacement_text)

    def _reset_precision_defaults(self):
        """恢复推荐的扫描与微调设置。"""
        default_index = self.combo_precision.findData(2.0)
        if default_index >= 0:
            self.combo_precision.setCurrentIndex(default_index)
        self.spin_offset_x.setValue(0)
        self.spin_offset_w.setValue(0)
        self.cb_enhance.setChecked(False)
        self._refresh_precision_summary()

    def _reset_ocr_adjustment(self):
        """恢复 OCR 检测框默认调节值。"""
        self.slider_adjust.setValue(0)
        self._refresh_ocr_summary()

    def _scroll_to_settings_section(self, row):
        """左侧导航切换到对应设置区块。"""
        if getattr(self, "_settings_nav_syncing", False):
            return
        if row < 0 or row >= len(getattr(self, "_settings_sections", [])):
            return
        target = self._settings_sections[row]
        if not target or not hasattr(self, "content_scroll"):
            return
        self.content_scroll.verticalScrollBar().setValue(max(0, target.pos().y() - 8))

    def _sync_settings_nav_from_scroll(self, value):
        """滚动右侧内容时，同步高亮左侧设置导航。"""
        sections = getattr(self, "_settings_sections", [])
        if not sections or not hasattr(self, "settings_nav"):
            return

        active_row = 0
        threshold = value + 24
        for index, section in enumerate(sections):
            if section and section.pos().y() <= threshold:
                active_row = index
            else:
                break

        if self.settings_nav.currentRow() == active_row:
            return

        self._settings_nav_syncing = True
        self.settings_nav.blockSignals(True)
        self.settings_nav.setCurrentRow(active_row)
        self.settings_nav.blockSignals(False)
        self._settings_nav_syncing = False

    def _refresh_settings_overview(self):
        """刷新设置页顶部配置概览。"""
        if not all(hasattr(self, attr) for attr in [
            "lbl_metric_rules", "lbl_metric_keywords", "lbl_metric_word_rules", "lbl_metric_ocr"
        ]):
            return

        enabled_rules = len([name for name, cb in self.checks.items() if cb.isChecked()]) if hasattr(self, "checks") else 0
        total_rules = len(self.checks) if hasattr(self, "checks") else 0
        keywords = [line.strip() for line in self.txt_custom.toPlainText().splitlines() if line.strip()] if hasattr(self, "txt_custom") else []
        enabled_word_rules = len([r for r in self.word_replace_rules if r.get("enabled", True) and r.get("find")]) if hasattr(self, "word_replace_rules") else 0
        total_word_rules = len(self.word_replace_rules) if hasattr(self, "word_replace_rules") else 0
        adjust_value = self.slider_adjust.value() if hasattr(self, "slider_adjust") else 0
        scan_label = self.combo_precision.currentText().strip() if hasattr(self, "combo_precision") else "-"
        precision_is_default = True
        if hasattr(self, "combo_precision") and hasattr(self, "spin_offset_x") and hasattr(self, "spin_offset_w") and hasattr(self, "cb_enhance"):
            precision_is_default = (
                self.combo_precision.currentData() == 2.0
                and self.spin_offset_x.value() == 0
                and self.spin_offset_w.value() == 0
                and not self.cb_enhance.isChecked()
            )

        self.lbl_metric_rules.setText(f"{enabled_rules} / {total_rules} 已启用")
        self.lbl_metric_keywords.setText(f"{len(keywords)} 条关键词")
        self.lbl_metric_word_rules.setText(f"{enabled_word_rules} / {total_word_rules} 条规则")
        self.lbl_metric_ocr.setText(f"{adjust_value}% · {scan_label}")
        if hasattr(self, "lbl_settings_common_tag") and hasattr(self, "lbl_settings_advanced_tag"):
            common_tag, advanced_tag = build_settings_hero_tags(
                enabled_rules,
                len(keywords),
                enabled_word_rules,
                precision_is_default,
                adjust_value,
                scan_label,
            )
            self.lbl_settings_common_tag.setText(common_tag)
            self.lbl_settings_common_tag.setToolTip(common_tag)
            self.lbl_settings_advanced_tag.setText(advanced_tag)
            self.lbl_settings_advanced_tag.setToolTip(advanced_tag)
        self._refresh_settings_sidebar()

    def _refresh_settings_sidebar(self):
        """刷新左侧导航和侧栏摘要。"""
        if not hasattr(self, "settings_nav"):
            return

        enabled_rules = len([name for name, cb in self.checks.items() if cb.isChecked()]) if hasattr(self, "checks") else 0
        keyword_count = len([line.strip() for line in self.txt_custom.toPlainText().splitlines() if line.strip()]) if hasattr(self, "txt_custom") else 0
        blacklist_count = len([line.strip() for line in self.txt_blacklist.toPlainText().splitlines() if line.strip()]) if hasattr(self, "txt_blacklist") else 0
        whitelist_count = len([line.strip() for line in self.txt_whitelist.toPlainText().splitlines() if line.strip()]) if hasattr(self, "txt_whitelist") else 0
        precision_is_default = True
        if hasattr(self, "combo_precision") and hasattr(self, "spin_offset_x") and hasattr(self, "spin_offset_w") and hasattr(self, "cb_enhance"):
            precision_is_default = (
                self.combo_precision.currentData() == 2.0
                and self.spin_offset_x.value() == 0
                and self.spin_offset_w.value() == 0
                and not self.cb_enhance.isChecked()
            )
        adjust_value = self.slider_adjust.value() if hasattr(self, "slider_adjust") else 0

        nav_labels = build_settings_nav_labels(enabled_rules, keyword_count, precision_is_default, adjust_value, blacklist_count, whitelist_count)
        for index, text in enumerate(nav_labels):
            item = self.settings_nav.item(index)
            if item and item.text() != text:
                item.setText(text)

        if hasattr(self, "settings_sidebar_status"):
            advanced_text = "扫描保持默认" if precision_is_default else "扫描参数已微调"
            self.settings_sidebar_status.setText(
                f"常用区：规则 {enabled_rules} 项、关键词 {keyword_count} 条。\n"
                f"高级区：{advanced_text} · OCR {format_signed_percent(adjust_value)}。"
            )

    def _refresh_word_rule_summary(self):
        enabled_count = len([r for r in self.word_replace_rules if r.get("enabled", True) and r.get("find")])
        total_count = len(self.word_replace_rules)
        self.lbl_word_rule_count.setText(
            f"当前规则：{enabled_count} 条启用 / {total_count} 条总计\n"
            "点击“打开替换规则设置”可进入原有替换规则弹窗。"
        )
        self._refresh_custom_summary()

    def _refresh_rule_summary(self):
        enabled_names = [name for name, cb in self.checks.items() if cb.isChecked()]
        enabled_count = len(enabled_names)
        preview = "、".join(enabled_names[:3]) if enabled_names else "当前未启用任何通用规则"
        if enabled_count > 3:
            preview = f"{preview} 等 {enabled_count} 项"
        self.lbl_rules_summary.setText(f"当前启用：{preview}")
        self._refresh_settings_overview()

    def _refresh_custom_summary(self):
        keyword_lines = [line.strip() for line in self.txt_custom.toPlainText().splitlines() if line.strip()]
        keyword_count = len(keyword_lines)
        replacement_preview = self.input_replacement_text.text().strip() or "[已脱敏]"
        enabled_rule_count = len([r for r in self.word_replace_rules if r.get("enabled", True) and r.get("find")])
        self.lbl_custom_summary.setText(
            f"自定义关键词 {keyword_count} 条 · 统一替换文本：{replacement_preview} · Word 规则：{enabled_rule_count} 条启用"
        )
        self._refresh_settings_overview()

    def _refresh_precision_summary(self):
        current_label = self.combo_precision.currentText().strip()
        enhance_text = "已开启图像增强" if self.cb_enhance.isChecked() else "图像增强关闭"
        self.lbl_precision_summary.setText(
            f"当前扫描模式：{current_label} · 向左修正 {self.spin_offset_x.value()} px · 宽度收缩 {self.spin_offset_w.value()} px · {enhance_text}"
        )
        self._refresh_settings_overview()

    def _refresh_ocr_summary(self):
        adjust_value = self.slider_adjust.value()
        if adjust_value < 0:
            trend = "扩大检测框"
        elif adjust_value > 0:
            trend = "收缩检测框"
        else:
            trend = "保持原始检测框"
        self.lbl_ocr_summary.setText(f"当前检测框调节：{adjust_value}% · {trend}")
        self._refresh_settings_overview()

    def _refresh_overrides_summary(self):
        """v1.1.11: 读取 config 中的 permanent overrides 数量,刷新概述."""
        if not self.config:
            self.lbl_overrides_summary.setText("当前未挂载配置管理器,无法统计。")
            return
        items = self.config.get("redaction.overrides.permanent", []) or []
        total = len(items) if isinstance(items, list) else 0
        ignore_n = sum(1 for it in items if isinstance(it, dict) and it.get("action") == "ignore")
        confirm_n = sum(1 for it in items if isinstance(it, dict) and it.get("action") == "confirm")
        self.lbl_overrides_summary.setText(
            f"当前 permanent 列表: 共 {total} 条(ignore {ignore_n} / confirm {confirm_n})"
        )
        self._refresh_settings_overview()

    def _on_clean_stale_overrides(self):
        """v1.1.11: 清理 30 天前 promoted 的 permanent override."""
        if not self.config:
            QMessageBox.warning(self, "提示", "未挂载配置管理器,无法清理。")
            return
        from secureredact.redaction.override_store import clean_stale_permanent
        items = self.config.get("redaction.overrides.permanent", []) or []
        if not isinstance(items, list):
            QMessageBox.warning(self, "提示", "permanent 字段格式异常,无法清理。")
            return
        before = len(items)
        cleaned = clean_stale_permanent(items, max_age_days=30)
        removed = before - len(cleaned)
        self.config.set("redaction.overrides.permanent", cleaned)
        try:
            self.config.save()
        except Exception as exc:
            QMessageBox.warning(self, "失败", f"保存配置失败: {exc}")
            return
        # 同步内存中的 store
        mw = self.parent()
        if mw is not None and hasattr(mw, "_override_store"):
            try:
                mw._override_store.replace_permanent(cleaned)
            except Exception:
                pass
        self._refresh_overrides_summary()
        QMessageBox.information(
            self,
            "完成",
            f"已清理 {removed} 条失效 permanent override(剩余 {len(cleaned)} 条)。",
        )

    def _open_word_rules_editor(self):
        default_text = self.input_replacement_text.text().strip() or "[已脱敏]"
        dlg = WordReplaceRulesDialog(
            self,
            rules=self.word_replace_rules,
            default_replacement_text=default_text,
            title="Word 替换规则设置",
            apply_text="应用规则"
        )
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        self.word_replace_rules = dlg.rules
        self.input_replacement_text.setText(dlg.default_replacement_text)
        self._refresh_word_rule_summary()

    def save_settings(self):
        self.selected_rules = [DEFAULT_RULES[name] for name, cb in self.checks.items() if cb.isChecked()]
        # v1.1.11: 添加调试输出
        print(f"[Settings] 保存的规则: {self.selected_rules}")
        print(f"[Settings] 勾选的规则名称: {[name for name, cb in self.checks.items() if cb.isChecked()]}")
        self.use_enhance = self.cb_enhance.isChecked()
        self.custom_keywords = self.txt_custom.toPlainText().strip()
        self.scan_level = self.combo_precision.currentData()
        self.offset_x = self.spin_offset_x.value()
        self.offset_w = self.spin_offset_w.value()
        self.replacement_text = self.input_replacement_text.text().strip() or "[已脱敏]"
        self.word_replace_rules = normalize_word_replace_rules(self.word_replace_rules, self.replacement_text)

        # v1.1.11: 保存检测框调节比例
        self.box_adjust_ratio = self.slider_adjust.value() / 100.0

        # v1.1.11: 中文姓名启发式识别开关
        self.enable_name_recognition = self.cb_name_recognition.isChecked()

        # v1.1.11: 保存到配置文件
        if self.config:
            try:
                self.config.set("redaction.scan.default_level", self.scan_level, persist=False)
                self.config.set("redaction.offset.default_x", self.offset_x, persist=False)
                self.config.set("redaction.offset.default_w", self.offset_w, persist=False)
                # v1.1.11: 移除 OCR 引擎选择，只保留 RapidOCR
                # v1.1.11: 保存检测框调节比例（新配置名）
                self.config.set("redaction.custom_keywords", self.custom_keywords, persist=False)
                self.config.set("redaction.replacement_text", self.replacement_text, persist=False)
                self.config.set("ocr.box_adjust_ratio", self.box_adjust_ratio, persist=False)
                # v1.1.11: 姓名识别开关持久化
                self.config.set("redaction.enable_name_recognition",
                                self.enable_name_recognition, persist=False)
                self.config.save()
            except Exception as e:
                print(f"[设置] 保存配置失败: {e}")

        self.accept()

    def _apply_dialog_theme(self):
        """应用对话框浅色主题样式（v1.1.11: 修复 Windows 深色模式显示问题）"""
        from theme import Theme
        theme = Theme.LIGHT

        self.setStyleSheet(f"""
            QDialog {{
                background-color: {theme["background"]};
                font-family: {Theme.FONT_FAMILY};
            }}
            QWidget {{
                background-color: {theme["background"]};
                color: {theme["text"]};
                font-family: {Theme.FONT_FAMILY};
            }}
            QFrame#settingsHero {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 18px;
            }}
            QFrame#settingsOverview {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 18px;
            }}
            QFrame#settingsMetricCard {{
                background-color: #F8FBFE;
                border: 1px solid {theme["border"]};
                border-radius: 12px;
            }}
            QLabel#settingsTitle {{
                color: {theme["text"]};
                font-size: 20px;
                font-weight: 700;
            }}
            QLabel#settingsOverviewTitle {{
                color: {theme["text"]};
                font-size: 16px;
                font-weight: 700;
            }}
            QLabel#settingsOverviewText {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QLabel#settingsMetricLabel {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                font-weight: 700;
            }}
            QLabel#settingsMetricValue {{
                color: {theme["text"]};
                font-size: 16px;
                font-weight: 700;
            }}
            QPushButton#settingsQuickJumpButton {{
                background-color: #FBFCFE;
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 8px 14px;
                font-size: 12px;
                font-weight: 600;
            }}
            QPushButton#settingsQuickJumpButton:hover {{
                background-color: {theme["hover"]};
                border-color: {theme["primary"]};
                color: {theme["primary"]};
            }}
            QLabel#settingsSubtitle {{
                color: {theme["text_secondary"]};
                font-size: 12px;
                line-height: 1.6;
            }}
            QLabel#settingsHeroTag {{
                color: {theme["primary"]};
                background-color: #E9F1FB;
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 6px 10px;
                font-size: 11px;
                font-weight: 700;
            }}
            QFrame#settingsSidebar {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 18px;
            }}
            QFrame#settingsSidebarMetaCard {{
                background-color: #FBFCFE;
                border: 1px solid {theme["border"]};
                border-radius: 14px;
            }}
            QFrame#settingsFooter {{
                background-color: {theme["surface"]};
                border: 1px solid {theme["border"]};
                border-radius: 16px;
            }}
            QFrame#settingsFieldCard, QWidget#settingsInnerPanel {{
                background-color: #FBFCFE;
                border: 1px solid {theme["border"]};
                border-radius: 14px;
            }}
            QLabel#settingsHint {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                font-weight: 700;
                letter-spacing: 0.05em;
            }}
            QLabel#settingsSidebarNote {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QLabel#settingsSidebarSubtle {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QLabel#settingsSidebarStatus {{
                color: {theme["text_secondary"]};
                background-color: #F8FBFE;
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 8px 10px;
                font-size: 11px;
                line-height: 1.6;
                font-weight: 600;
            }}
            QLabel#settingsFooterNote {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
                padding-right: 8px;
            }}
            QLabel#settingsFieldTitle {{
                color: {theme["text"]};
                font-size: 12px;
                font-weight: 700;
            }}
            QLabel#settingsFieldLabel {{
                color: {theme["text"]};
                font-size: 11px;
                font-weight: 600;
            }}
            QLabel#settingsFieldNote {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.6;
            }}
            QFrame#settingsFieldDivider {{
                background-color: {theme["border"]};
                border: none;
            }}
            QLabel#settingsSectionLead {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                line-height: 1.7;
            }}
            QLabel#settingsSectionSummary {{
                color: {theme["text_secondary"]};
                background-color: #F8FBFE;
                border: 1px solid {theme["border"]};
                border-radius: 10px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: 600;
                line-height: 1.6;
            }}
            QListWidget#settingsNav {{
                background-color: transparent;
                border: none;
                outline: none;
                padding: 0;
            }}
            QListWidget#settingsNav::item {{
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 10px;
                padding: 10px 12px;
                margin: 2px 0;
                color: {theme["text"]};
                font-weight: 600;
            }}
            QListWidget#settingsNav::item:selected {{
                background-color: {theme["hover"]};
                border-color: {theme["border"]};
                color: {theme["primary"]};
            }}
            QListWidget#settingsNav::item:hover {{
                background-color: {theme["hover"]};
            }}
            QFrame#settingsSectionCard {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 16px;
            }}
            QFrame#settingsSectionHeader {{
                background-color: transparent;
                border: none;
            }}
            QLabel#settingsSectionTitle {{
                color: {theme["primary"]};
                font-size: 15px;
                font-weight: 700;
                background-color: transparent;
            }}
            QLabel {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QCheckBox {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QCheckBox::indicator {{
                width: 16px;
                height: 16px;
            }}
            QTextEdit {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 8px;
            }}
            QLineEdit {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 6px 8px;
                min-height: 22px;
            }}
            QComboBox {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 6px;
                min-width: 100px;
            }}
            QComboBox::drop-down {{
                border: none;
                width: 24px;
            }}
            QComboBox QAbstractItemView {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                selection-background-color: {theme["primary"]};
            }}
            QSpinBox {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 6px;
            }}
            QSlider {{
                background-color: transparent;
            }}
            QSlider::groove:horizontal {{
                border: none;
                height: 4px;
                background-color: {theme["border"]};
                border-radius: 2px;
            }}
            QSlider::handle:horizontal {{
                background-color: {theme["primary"]};
                border: none;
                width: 16px;
                height: 16px;
                margin: -6px 0;
                border-radius: 8px;
            }}
            QScrollArea {{
                background-color: transparent;
                border: none;
            }}
            QPushButton {{
                background-color: {theme["primary"]};
                color: white;
                border: none;
                border-radius: 10px;
                padding: 10px 20px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background-color: {theme["primary"]};
                opacity: 0.9;
            }}
            QPushButton:pressed {{
                background-color: {theme["pressed"]};
            }}
            QPushButton#settingsSecondaryButton {{
                background-color: #FBFCFE;
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                padding: 10px 18px;
            }}
            QPushButton#settingsSecondaryButton:hover {{
                background-color: {theme["hover"]};
                border-color: {theme["primary"]};
            }}
            QLabel#settingsActionHint {{
                color: {theme["text_secondary"]};
                font-size: 11px;
                font-weight: 700;
            }}
            QPushButton#settingsInlineButton {{
                background-color: #FBFCFE;
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 9px;
                padding: 7px 12px;
                font-size: 11px;
                font-weight: 600;
            }}
            QPushButton#settingsInlineButton:hover {{
                background-color: {theme["hover"]};
                border-color: {theme["primary"]};
                color: {theme["primary"]};
            }}
        """)


class WordReplaceRulesDialog(QDialog):
    """Word 多字段替换规则对话框（会话级规则，可导入/导出）。"""

    def __init__(self, parent=None, rules=None, default_replacement_text="[已脱敏]",
                 title="替换规则设置", apply_text="应用规则"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(780, 520)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)

        self.rules = normalize_word_replace_rules(rules or [], default_replacement_text)
        self.default_replacement_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"

        main_layout = QVBoxLayout(self)

        header = QLabel("支持精确(exact)和正则(regex)两种模式。执行顺序：精确优先，其次正则；同模式按规则顺序。")
        header.setWordWrap(True)
        main_layout.addWidget(header)

        default_row = QHBoxLayout()
        default_row.addWidget(QLabel("默认替换文本:"))
        self.input_default_text = QLineEdit(self.default_replacement_text)
        self.input_default_text.setPlaceholderText("[已脱敏]")
        default_row.addWidget(self.input_default_text)
        main_layout.addLayout(default_row)

        self.table = QTableWidget(0, 4)
        self.table.setHorizontalHeaderLabels(["启用", "模式(exact/regex)", "查找文本", "替换文本"])
        self.table.verticalHeader().setVisible(False)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        main_layout.addWidget(self.table)

        row_btn_layout = QHBoxLayout()
        btn_add = QPushButton("新增规则")
        btn_del = QPushButton("删除规则")
        btn_up = QPushButton("上移")
        btn_down = QPushButton("下移")
        btn_import = QPushButton("导入JSON")
        btn_export = QPushButton("导出JSON")
        row_btn_layout.addWidget(btn_add)
        row_btn_layout.addWidget(btn_del)
        row_btn_layout.addWidget(btn_up)
        row_btn_layout.addWidget(btn_down)
        row_btn_layout.addStretch()
        row_btn_layout.addWidget(btn_import)
        row_btn_layout.addWidget(btn_export)
        main_layout.addLayout(row_btn_layout)

        btn_add.clicked.connect(self.add_rule_row)
        btn_del.clicked.connect(self.remove_selected_rule)
        btn_up.clicked.connect(lambda: self.move_selected_rule(-1))
        btn_down.clicked.connect(lambda: self.move_selected_rule(1))
        btn_import.clicked.connect(self.import_rules_json)
        btn_export.clicked.connect(self.export_rules_json)

        footer = QHBoxLayout()
        footer.addStretch()
        btn_cancel = QPushButton("取消")
        btn_apply = QPushButton(apply_text)
        footer.addWidget(btn_cancel)
        footer.addWidget(btn_apply)
        main_layout.addLayout(footer)

        btn_cancel.clicked.connect(self.reject)
        btn_apply.clicked.connect(self.apply_rules)

        for rule in self.rules:
            self.add_rule_row(rule)

        if self.table.rowCount() == 0:
            self.add_rule_row()

    def add_rule_row(self, rule=None):
        rule = rule or {"enabled": True, "mode": "exact", "find": "", "replace": ""}
        row = self.table.rowCount()
        self.table.insertRow(row)

        enabled_item = QTableWidgetItem()
        enabled_item.setFlags(
            Qt.ItemFlag.ItemIsEnabled |
            Qt.ItemFlag.ItemIsUserCheckable |
            Qt.ItemFlag.ItemIsSelectable
        )
        enabled_item.setCheckState(Qt.CheckState.Checked if rule.get("enabled", True) else Qt.CheckState.Unchecked)
        self.table.setItem(row, 0, enabled_item)

        mode = str(rule.get("mode", "exact")).strip().lower()
        if mode not in ("exact", "regex"):
            mode = "exact"
        self.table.setItem(row, 1, QTableWidgetItem(mode))
        self.table.setItem(row, 2, QTableWidgetItem(str(rule.get("find", ""))))
        self.table.setItem(row, 3, QTableWidgetItem(str(rule.get("replace", ""))))
        self.table.selectRow(row)

    def remove_selected_rule(self):
        row = self.table.currentRow()
        if row < 0:
            return
        self.table.removeRow(row)
        if self.table.rowCount() == 0:
            self.add_rule_row()

    def move_selected_rule(self, direction):
        current_row = self.table.currentRow()
        if current_row < 0:
            return
        target_row = current_row + direction
        if target_row < 0 or target_row >= self.table.rowCount():
            return

        current_data = []
        target_data = []
        for col in range(self.table.columnCount()):
            current_item = self.table.item(current_row, col)
            target_item = self.table.item(target_row, col)
            current_data.append(current_item.clone() if current_item else QTableWidgetItem())
            target_data.append(target_item.clone() if target_item else QTableWidgetItem())

        for col in range(self.table.columnCount()):
            self.table.setItem(current_row, col, target_data[col])
            self.table.setItem(target_row, col, current_data[col])

        self.table.selectRow(target_row)

    def _collect_rules_from_table(self, validate_regex=True):
        errors = []
        rules = []
        default_text = self.input_default_text.text().strip() or "[已脱敏]"

        for row in range(self.table.rowCount()):
            enabled_item = self.table.item(row, 0)
            mode_item = self.table.item(row, 1)
            find_item = self.table.item(row, 2)
            replace_item = self.table.item(row, 3)

            enabled = enabled_item is not None and enabled_item.checkState() == Qt.CheckState.Checked
            mode = mode_item.text().strip().lower() if mode_item else "exact"
            find_text = find_item.text().strip() if find_item else ""
            replace_text = replace_item.text() if replace_item else ""

            if mode not in ("exact", "regex"):
                errors.append(f"第 {row + 1} 行模式无效：{mode}（仅支持 exact/regex）")
                continue
            if enabled and not find_text:
                errors.append(f"第 {row + 1} 行查找文本不能为空")
                continue
            if enabled and mode == "regex" and validate_regex:
                try:
                    re.compile(find_text)
                except re.error as e:
                    errors.append(f"第 {row + 1} 行正则无效：{e}")
                    continue
            if not replace_text:
                replace_text = default_text

            rules.append({
                "enabled": enabled,
                "mode": mode,
                "find": find_text,
                "replace": replace_text
            })

        normalized = normalize_word_replace_rules(rules, default_text)
        return normalized, default_text, errors

    def apply_rules(self):
        rules, default_text, errors = self._collect_rules_from_table(validate_regex=True)
        if errors:
            QMessageBox.warning(self, "规则校验失败", "\n".join(errors))
            return

        self.rules = rules
        self.default_replacement_text = default_text
        self.accept()

    def import_rules_json(self):
        fname, _ = QFileDialog.getOpenFileName(self, "导入替换规则", "", "JSON 文件 (*.json)")
        if not fname:
            return

        try:
            with open(fname, "r", encoding="utf-8") as f:
                data = json.load(f)

            if not isinstance(data, dict):
                raise ValueError("JSON 根节点必须为对象")

            default_text = str(data.get("default_replacement_text", self.input_default_text.text().strip() or "[已脱敏]"))
            rules = data.get("rules", [])
            if not isinstance(rules, list):
                raise ValueError("rules 字段必须为数组")

            normalized = normalize_word_replace_rules(rules, default_text)

            self.input_default_text.setText(default_text)
            self.table.setRowCount(0)
            for rule in normalized:
                self.add_rule_row(rule)
            if self.table.rowCount() == 0:
                self.add_rule_row()

            QMessageBox.information(self, "导入成功", f"已导入 {len(normalized)} 条规则")
        except (OSError, IOError, ValueError, json.JSONDecodeError) as e:
            QMessageBox.critical(self, "导入失败", f"无法导入规则文件：\n{e}")

    def export_rules_json(self):
        rules, default_text, errors = self._collect_rules_from_table(validate_regex=True)
        if errors:
            QMessageBox.warning(self, "无法导出", "\n".join(errors))
            return

        fname, _ = QFileDialog.getSaveFileName(self, "导出替换规则", "word_replace_rules.json", "JSON 文件 (*.json)")
        if not fname:
            return

        if not fname.lower().endswith(".json"):
            fname = fname + ".json"

        payload = {
            "version": WORD_RULE_SCHEMA_VERSION,
            "default_replacement_text": default_text,
            "rules": rules
        }
        try:
            with open(fname, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "导出成功", f"规则已导出到：\n{fname}")
        except (OSError, IOError, ValueError) as e:
            QMessageBox.critical(self, "导出失败", f"无法导出规则文件：\n{e}")


# === 图片排序对话框 ===
class ImageListDialog(QDialog):
    """图片排序对话框 - 支持拖拽调整图片顺序"""
    def __init__(self, image_paths, parent=None):
        super().__init__(parent)

        # v1.1.11: 修复 Windows 深色模式下对话框显示问题
        self.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        self.setWindowTitle("调整图片顺序")
        # v1.1.11: 从配置读取窗口尺寸
        if config:
            dialog_width = config.get("app.window.dialog_image_list_width", 600)
            dialog_height = config.get("app.window.dialog_image_list_height", 500)
        else:
            dialog_width, dialog_height = 600, 500
        self.resize(dialog_width, dialog_height)
        self.image_paths = image_paths

        layout = QVBoxLayout(self)

        # 说明标签
        info_label = QLabel("拖拽缩略图调整图片顺序，完成后点击「确认合并」")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)

        # 缩略图列表（支持拖拽）
        self.list_widget = QListWidget()
        self.list_widget.setIconSize(QSize(120, 120))
        self.list_widget.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)
        self.list_widget.setSelectionMode(QListWidget.SelectionMode.SingleSelection)

        # 添加缩略图
        for path in image_paths:
            item = QListWidgetItem(os.path.basename(path))
            item.setData(Qt.ItemDataRole.UserRole, path)
            # 生成缩略图
            try:
                pixmap = QPixmap(path).scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio)
                item.setIcon(QIcon(pixmap))
            except (OSError, IOError, ValueError) as e:
                # 如果缩略图生成失败，使用默认图标
                print(f"[ImageMergeDialog] 缩略图生成失败: {path}: {e}")
                pass
            self.list_widget.addItem(item)

        layout.addWidget(self.list_widget)

        # 按钮
        btn_layout = QHBoxLayout()
        btn_ok = QPushButton("确认合并")
        btn_ok.clicked.connect(self.accept)
        btn_cancel = QPushButton("取消")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

        # v1.1.11: 应用对话框主题样式
        self._apply_dialog_theme()

    def _apply_dialog_theme(self):
        """应用对话框浅色主题样式（v1.1.11: 修复 Windows 深色模式显示问题）"""
        from theme import Theme
        theme = Theme.LIGHT

        self.setStyleSheet(f"""
            QDialog {{
                background-color: {theme["background"]};
            }}
            QWidget {{
                background-color: {theme["background"]};
                color: {theme["text"]};
            }}
            QLabel {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QListWidget {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                border: 1px solid {theme["border"]};
                border-radius: 6px;
                padding: 8px;
            }}
            QListWidget::item {{
                background-color: {theme["surface"]};
                color: {theme["text"]};
                padding: 8px;
                border-radius: 4px;
            }}
            QListWidget::item:selected {{
                background-color: {theme["primary"]};
                color: white;
            }}
            QPushButton {{
                background-color: {theme["primary"]};
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 20px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background-color: {theme["primary"]};
                opacity: 0.9;
            }}
        """)

    def get_ordered_paths(self):
        """获取排序后的图片路径"""
        paths = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            paths.append(item.data(Qt.ItemDataRole.UserRole))
        return paths

# === 配置常量 ===

# === 反馈对话框 ===
class FeedbackDialog(QDialog):
    """开发者信息与反馈对话框"""
    def __init__(self, parent=None):
        super().__init__(parent)

        # v1.1.11: 修复 Windows 深色模式下对话框显示问题
        self.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        self.setWindowTitle("关于与反馈")
        # v1.1.11: 从配置读取窗口尺寸
        if config:
            dialog_width = config.get("app.window.dialog_feedback_width", 480)
            dialog_height = config.get("app.window.dialog_feedback_height", 600)
        else:
            dialog_width, dialog_height = 480, 600
        self.resize(dialog_width, dialog_height)
        self.setWindowModality(Qt.WindowModality.ApplicationModal)

        # 获取当前主题
        self.theme = Theme.LIGHT if not parent or not hasattr(parent, 'is_dark') or not parent.is_dark else Theme.DARK

        layout = QVBoxLayout(self)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # === 标题区域 ===
        title_layout = QHBoxLayout()
        logo_label = QLabel()
        logo_label.setFixedSize(64, 64)
        logo_label.setStyleSheet(f"""
            background: {self.theme['primary']};
            border-radius: 12px;
            color: white;
            font-size: 24px;
            font-weight: bold;
        """)
        logo_label.setText("PG")
        logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_layout.addWidget(logo_label)

        title_info = QVBoxLayout()
        title_info.setSpacing(4)
        app_name = QLabel(APP_NAME)
        app_name.setStyleSheet(f"font-size: 20px; font-weight: bold; color: {self.theme['text']};")
        version_label = QLabel(f"版本 {VERSION}")
        version_label.setStyleSheet(f"font-size: 12px; color: {self.theme['text_secondary']};")
        title_info.addWidget(app_name)
        title_info.addWidget(version_label)
        title_layout.addLayout(title_info)
        title_layout.addStretch()
        layout.addLayout(title_layout)

        # === 分隔线 ===
        line1 = QFrame()
        line1.setFrameShape(QFrame.Shape.HLine)
        line1.setStyleSheet(f"background: {self.theme['border']}; max-height: 1px;")
        layout.addWidget(line1)




        # === 分隔线 ===
        line2 = QFrame()
        line2.setFrameShape(QFrame.Shape.HLine)
        line2.setStyleSheet(f"background: {self.theme['border']}; max-height: 1px;")
        layout.addWidget(line2)

        # === 免责声明 ===
        disclaimer = QTextBrowser()
        disclaimer.setOpenExternalLinks(True)
        disclaimer.setMaximumHeight(160)
        disclaimer.setStyleSheet(f"""
            QTextBrowser {{
                background: {self.theme['surface']};
                border: 1px solid {self.theme['border']};
                border-radius: 8px;
                padding: 12px;
                color: {self.theme['text_secondary']};
                font-size: 11px;
            }}
        """)
        disclaimer.setHtml("""
            <p style="font-weight: bold; margin-bottom: 8px; color: #FF9500;">⚠️ 免责声明</p>
            <ol style="margin-left: 16px; line-height: 1.6;">
                <li>本软件免费仅供学习和个人使用，不构成任何法律建议。</li>
                <li>使用本软件进行文档脱敏处理后，用户需自行核实脱敏结果。</li>
                <li>开发者不对因使用本软件而产生的任何直接或间接损失承担责任。</li>
                <li>本软件不收集任何用户数据，所有处理均在本地完成。</li>
                <li>请勿将本软件用于任何违法用途。</li>
            </ol>
        """)
        layout.addWidget(disclaimer)

        # === 关闭按钮 ===
        close_btn = QPushButton("关闭")
        close_btn.setStyleSheet(f"""
            QPushButton {{
                background: {self.theme['hover']};
                color: {self.theme['text']};
                border: none;
                border-radius: 8px;
                padding: 10px;
                font-size: 13px;
            }}
            QPushButton:hover {{
                background: {self.theme['pressed']};
            }}
        """)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

    def _copy_to_clipboard(self, text):
        """复制文本到剪贴板"""
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        # 显示复制成功提示
        if self.parent():
            QMessageBox.information(self.parent(), "复制成功", f"已复制: {text}")


class WordBatchReplaceWorker(QThread):
    """Word 批量替换线程（同一套规则应用全部文件）。"""

    progress_signal = pyqtSignal(int, int, str)      # processed, total, current_file
    file_done_signal = pyqtSignal(str, str)          # input_path, output_path
    file_error_signal = pyqtSignal(int, str, str)    # index, input_path, error_msg
    finished_signal = pyqtSignal(dict)               # summary

    def __init__(self, file_paths, rules, default_replacement_text):
        super().__init__()
        self.file_paths = list(file_paths or [])
        self.rules = normalize_word_replace_rules(rules or [], default_replacement_text)
        self.default_replacement_text = default_replacement_text if isinstance(default_replacement_text, str) and default_replacement_text else "[已脱敏]"
        self._error_decision = None
        self._decision_event = threading.Event()
        self._decision_lock = threading.Lock()
        self._temp_dirs = []

    def provide_error_decision(self, decision):
        """主线程调用：为当前错误设置决策（skip/stop）。"""
        with self._decision_lock:
            self._error_decision = decision
        self._decision_event.set()

    def _wait_for_error_decision(self):
        """等待主线程选择错误处理策略。"""
        self._decision_event.clear()
        while not self._decision_event.wait(0.1):
            if self.isInterruptionRequested():
                return "stop"
        with self._decision_lock:
            decision = self._error_decision or "skip"
            self._error_decision = None
        return decision

    def run(self):
        summary = {
            "total": len(self.file_paths),
            "processed": 0,
            "success": [],
            "failed": [],
            "stopped": False,
            "rules": list(self.rules),
        }

        try:
            total = len(self.file_paths)
            for idx, file_path in enumerate(self.file_paths):
                if self.isInterruptionRequested():
                    summary["stopped"] = True
                    break

                current_name = os.path.basename(file_path)
                self.progress_signal.emit(idx, total, current_name)

                try:
                    output_path, replace_stats = self._process_single_file(file_path)
                    summary["success"].append({
                        "input": file_path,
                        "output": output_path,
                        "total_replacements": max(0, int(replace_stats.get("total_replacements", 0) or 0))
                        if isinstance(replace_stats, dict) else 0,
                        "rule_counts": replace_stats.get("rule_counts", []) if isinstance(replace_stats, dict) else [],
                    })
                    self.file_done_signal.emit(file_path, output_path)
                except (OSError, IOError, RuntimeError, ValueError, ConversionError, PermissionError) as e:
                    error_msg = str(e)
                    summary["failed"].append({
                        "input": file_path,
                        "error": error_msg
                    })
                    self.file_error_signal.emit(idx, file_path, error_msg)
                    decision = self._wait_for_error_decision()
                    if decision == "stop":
                        summary["stopped"] = True
                        break
                finally:
                    self.progress_signal.emit(idx + 1, total, current_name)

            summary["processed"] = len(summary["success"]) + len(summary["failed"])
            self.finished_signal.emit(summary)
        finally:
            self._cleanup_temp_dirs()

    def _process_single_file(self, file_path):
        from docx import Document

        is_safe, error_msg = validate_safe_path(file_path, allowed_extensions=[".doc", ".docx"])
        if not is_safe:
            raise ConversionError("输入文件路径不安全", error_msg)

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            source_docx = file_path
        elif ext == ".doc":
            source_docx = self._convert_doc_to_docx(file_path)
        else:
            raise ConversionError("不支持的文件格式", f"{file_path}")

        doc = Document(source_docx)
        replace_stats = self._apply_rules_to_document(doc)
        output_path = self._build_output_path(file_path)
        doc.save(output_path)
        return output_path, replace_stats

    def _apply_rules_to_document(self, doc):
        rule_counts = {}
        total_replacements = 0

        def _collect_match_counts(matches):
            nonlocal total_replacements
            if not matches:
                return
            total_replacements += len(matches)
            for match in matches:
                try:
                    rule_index = int(match.get("rule_index", -1))
                except (TypeError, ValueError, AttributeError):
                    continue
                if rule_index < 0:
                    continue
                rule_counts[rule_index] = rule_counts.get(rule_index, 0) + 1

        for para in doc.paragraphs:
            text = ''.join(run.text for run in para.runs)
            if not text:
                continue
            matches = build_word_rule_matches(text, self.rules, self.default_replacement_text)
            if matches:
                _collect_match_counts(matches)
                replace_matches_in_paragraph(para, matches, text_offset=0,
                                             fallback_replacement_text=self.default_replacement_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text or ""
                    if not cell_text:
                        continue
                    cell_matches = build_word_rule_matches(cell_text, self.rules, self.default_replacement_text)
                    if not cell_matches:
                        continue
                    _collect_match_counts(cell_matches)

                    para_offset = 0
                    paragraphs = list(cell.paragraphs)
                    for idx, para in enumerate(paragraphs):
                        original_para_len = len(''.join(run.text for run in para.runs))
                        replace_matches_in_paragraph(
                            para,
                            cell_matches,
                            text_offset=para_offset,
                            fallback_replacement_text=self.default_replacement_text
                        )
                        para_offset += original_para_len
                        if idx < len(paragraphs) - 1:
                            para_offset += 1

        return {
            "total_replacements": total_replacements,
            "rule_counts": [
                {"rule_index": rule_index, "count": count}
                for rule_index, count in sorted(rule_counts.items())
            ],
        }

    def _build_output_path(self, file_path):
        base_path = os.path.splitext(file_path)[0]
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        candidate = f"{base_path}__replaced_{timestamp}.docx"
        if not os.path.exists(candidate):
            return candidate

        suffix = 1
        while True:
            candidate_with_suffix = f"{base_path}__replaced_{timestamp}_{suffix}.docx"
            if not os.path.exists(candidate_with_suffix):
                return candidate_with_suffix
            suffix += 1

    def _convert_doc_to_docx(self, doc_path):
        """v1.1.11: 委托给共享转换模块。"""
        docx_path, temp_dir = _shared_convert_doc_to_docx(doc_path)
        self._temp_dirs.append(temp_dir)
        return docx_path

    def _cleanup_temp_dirs(self):
        for temp_dir in list(self._temp_dirs):
            try:
                if os.path.isdir(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except (OSError, IOError):
                pass
        self._temp_dirs.clear()

# === 单页画布 (完全复制 v1.1.11 的 PDFCanvas 实现) ===


# === OCR 线程 ===
# v1.1.11: 改为使用模块化 OCRWorker，自动注入 box_adjust_ratio
class OCRWorker(_ModularOCRWorker):
    """OCR 处理线程（兼容层：自动注入 config 中的 box_adjust_ratio + enable_name_recognition）"""

    def __init__(self, pdf_path, rules, use_enhance, custom_keywords, scan_scale, off_x, off_w,
                 use_char_level_ocr: bool = False, seal_detection_enabled: bool = False,
                 enable_name_recognition: bool = False):
        box_adjust_ratio = config.get("ocr.box_adjust_ratio", 0.0) if config else 0.0
        super().__init__(pdf_path, rules, use_enhance, custom_keywords, scan_scale, off_x, off_w,
                         use_char_level_ocr=use_char_level_ocr,
                         seal_detection_enabled=seal_detection_enabled,
                         box_adjust_ratio=box_adjust_ratio,
                         enable_name_recognition=enable_name_recognition)

# === WebView Bridge：Python 与 JavaScript 通信 ===

# === Word 文档处理线程 ===
# v1.1.11: 改为使用模块化 WordWorker,补充 default_rules 参数
# v1.1.12: 兼容层透传 default_rules + default_rules_meta,支持 partial masking
class WordWorker(_ModularWordWorker):
    """Word 文档智能脱敏线程(兼容层: 自动注入 DEFAULT_RULES + enable_name_recognition)

    v1.1.12: 透传 default_rules_meta 到模块化 WordWorker,让 mask 配置生效。
    兼容层对外仍然只暴露原始 6 个参数,新增 kwargs 默认值与模块化层一致。
    """

    def __init__(self, word_doc, word_data, rules, custom_keywords, replacement_text,
                 enable_name_recognition: bool = False,
                 default_rules=None,
                 default_rules_meta=None):
        super().__init__(word_doc, word_data, rules, custom_keywords,
                         replacement_text,
                         default_rules=default_rules if default_rules is not None else DEFAULT_RULES,
                         enable_name_recognition=enable_name_recognition,
                         default_rules_meta=default_rules_meta if default_rules_meta is not None else DEFAULT_RULES_META)

# === Word 预览交互式 JavaScript 代码常量 ===
# v1.1.11: 提取为常量，避免 _inject_interactive_html 函数过长
_INTERACTIVE_JS_CODE = r"""
<script>
    let pyBridge = null;
    let webChannelReady = false;

    document.addEventListener('DOMContentLoaded', function() {
        // 初始化 QWebChannel
        if (typeof qt !== 'undefined' && qt.webChannelTransport) {
            try {
                new QWebChannel(qt.webChannelTransport, function(channel) {
                    pyBridge = channel.objects.pyBridge;
                    if (pyBridge) {
                        webChannelReady = true;
                    }
                });
            } catch (e) {
                console.error('QWebChannel error:', e);
            }
        }
        setupContextMenu();
    });

    // v1.1.11: 备用右键位置（用于复杂 DOM 结构中 getSelection() 返回空值的情况）
    let lastContextMenuEvent = null;

    function setupContextMenu() {
        // v1.1.11: 使用捕获阶段监听（更可靠，在事件冒泡前捕获）
        document.addEventListener('contextmenu', function(e) {
            e.preventDefault();
            lastContextMenuEvent = {
                clientX: e.clientX,
                clientY: e.clientY,
                target: e.target
            };
            handleContextMenu(e);
        }, true);  // true = 捕获阶段

        // v1.1.11: mousedown 事件预先保存右键位置（备用方案）
        document.addEventListener('mousedown', function(e) {
            if (e.button === 2) {  // 右键
                lastContextMenuEvent = {
                    clientX: e.clientX,
                    clientY: e.clientY,
                    target: e.target
                };
            }
        }, true);

        document.addEventListener('click', function(e) {
            if (!e.target.closest('#redaction-menu')) {
                hideContextMenu();
            }
        });
    }

    // v1.1.11: 增强的右键菜单处理函数
    function handleContextMenu(e) {
        const target = e.target;
        let selection = window.getSelection();
        let selectedText = selection.toString().trim();

        // v1.1.11: OCR / jieba hit 命中 — 弹出 HitOverrideStore 操作菜单
        // 仅当 mark 上有 data-hit-id 时触发;manual-highlight 仍走旧路径。
        let ocrHitElement = target.closest('mark[data-hit-id]');
        if (ocrHitElement) {
            const key = ocrHitElement.getAttribute('data-key') || '';
            const source = ocrHitElement.getAttribute('data-source') || '';
            const text = ocrHitElement.textContent || '';
            const hitId = ocrHitElement.getAttribute('data-hit-id') || '';
            if (pyBridge && webChannelReady &&
                typeof pyBridge.handle_ocr_hit_contextmenu === 'function') {
                try {
                    pyBridge.handle_ocr_hit_contextmenu(
                        key, source, text, hitId,
                        parseInt(e.clientX), parseInt(e.clientY)
                    );
                } catch (err) {
                    console.error('[ContextMenu] handle_ocr_hit_contextmenu failed:', err);
                }
            }
            return;
        }

        // 查找点击的目标是否是手动脱敏标记或其内部元素
        let highlightElement = target.closest('.manual-highlight');

        // 点击手动脱敏标记
        if (highlightElement) {
            console.log('[ContextMenu] 点击了手动脱敏标记');
            showRemoveMenu(e.clientX, e.clientY, highlightElement);
            return;
        }

        // 选择了文本（主要路径）
        if (selectedText.length > 0) {
            // v1.1.11: 移除敏感信息日志，仅记录操作类型
            console.log('[ContextMenu] 选择了文本（已隐藏内容）');
            try {
                const range = selection.getRangeAt(0);
                showAddMenu(e.clientX, e.clientY, selection, selectedText);
            } catch (err) {
                console.warn('[ContextMenu] getRangeAt 失败，尝试备用方案:', err);
                // 备用方案：使用预先保存的位置
                if (lastContextMenuEvent) {
                    tryFallbackTextDetection(e, target);
                }
            }
            return;
        }

        // v1.1.11: 备用方案 - 尝试从点击位置获取文本
        console.log('[ContextMenu] getSelection() 为空，尝试备用检测');
        tryFallbackTextDetection(e, target);
    }

    // v1.1.11: 备用文本检测（当 window.getSelection() 失败时）
    function tryFallbackTextDetection(e, target) {
        // 方案1: 检查目标元素是否包含文本
        let textElement = target;

        // 向上查找包含 data-key 的文本块
        for (let i = 0; i < 10 && textElement; i++) {
            if (textElement.dataset && textElement.dataset.key) {
                console.log('[tryFallbackTextDetection] 找到 data-key 元素:', textElement.dataset.key);
                // 显示一个提示菜单（v1.1.11 安全修复：使用配置对象）
                const menu = createMenu([
                    { text: '请在文本上拖动选择后再右键', disabled: true }
                ]);
                positionMenu(menu, e.clientX, e.clientY);
                setTimeout(hideContextMenu, 2000);
                return;
            }
            textElement = textElement.parentNode;
        }

        // 未找到 data-key，显示提示
        console.warn('[tryFallbackTextDetection] 未找到有效的文本块');
    }

    function showRemoveMenu(x, y, element) {
        const key = element.dataset.key;
        const start = parseInt(element.dataset.start);
        const end = parseInt(element.dataset.end);

        console.log('[showRemoveMenu] key:', key, 'start:', start, 'end:', end);

        // v1.1.11 安全修复：使用配置对象替代 HTML 字符串
        const menu = createMenu([
            { text: '❌ 撤销脱敏', action: 'remove', key: key, start: start, end: end }
        ]);
        positionMenu(menu, x, y);
        attachMenuHandlers();
    }

    function showAddMenu(x, y, selection, selectedText) {
        const range = selection.getRangeAt(0);
        const textInfo = findTextPosition(selectedText, range);

        // 即使 textInfo 为 null 也不直接返回，使用全局模式
        let buttonConfigs = [];

        if (!textInfo || textInfo.mode === 'global' || textInfo.key === '__GLOBAL__') {
            // 降级到全局模式：只显示全文脱敏选项
            console.log('[showAddMenu] 使用全局降级模式');
            // v1.1.11 安全修复：使用配置对象替代 HTML 字符串
            buttonConfigs = [
                { text: '📄 全文脱敏此内容', action: 'add-global-only', textData: selectedText }
            ];
        } else {
            // 正常情况：提供精确和全局两种选项
            buttonConfigs = [
                { text: '🎯 选中区域添加脱敏', action: 'add-exact', key: textInfo.key, start: textInfo.start, end: textInfo.end, textData: selectedText },
                { text: '📄 整篇相同字节添加脱敏', action: 'add-global', key: textInfo.key, textData: selectedText }
            ];
        }

        const menu = createMenu(buttonConfigs);
        positionMenu(menu, x, y);
        attachMenuHandlers();
    }

    function attachMenuHandlers() {
        const menu = document.getElementById('redaction-menu');
        if (!menu) return;

        menu.querySelectorAll('button').forEach(btn => {
            btn.addEventListener('click', function(e) {
                e.preventDefault();
                e.stopPropagation();
                const action = this.dataset.action;
                const key = this.dataset.key;
                const start = parseInt(this.dataset.start);
                const end = parseInt(this.dataset.end);
                const text = this.dataset.text;

                if (action === 'remove') {
                    callRemove(key, start, end);
                } else if (action === 'add-exact') {
                    callAddExact(key, start, end, text);
                } else if (action === 'add-global') {
                    callAddGlobal(key, text);
                } else if (action === 'add-global-only') {
                    // 全局降级模式，不需要 key 参数
                    callAddGlobal(null, text);
                }
            });
        });
    }

    function callRemove(key, start, end) {
        console.log('[callRemove] 调用撤销: key=' + key + ' start=' + start + ' end=' + end);
        console.log('[callRemove] pyBridge=', pyBridge, 'webChannelReady=', webChannelReady);
        if (pyBridge && webChannelReady) {
            try {
                pyBridge.remove_manual_redaction(key, start, end);
                console.log('[callRemove] ✓ 撤销调用成功');
            } catch(e) {
                console.error('[callRemove] ✗ 撤销调用失败:', e);
            }
        } else {
            console.error('[callRemove] ✗ pyBridge 或 webChannel 未就绪');
        }
        hideContextMenu();
    }

    function callAddExact(key, start, end, text) {
        if (pyBridge && webChannelReady) {
            pyBridge.add_manual_redaction(key, start, end, text);
        }
        hideContextMenu();
        const selection = window.getSelection();
        selection.removeAllRanges();
    }

    function callAddGlobal(key, text) {
        // v1.1.11: 移除敏感信息日志
        console.log('[callAddGlobal] 调用全局脱敏（文本已隐藏）');
        if (pyBridge && webChannelReady) {
            // key 为 null 时表示纯全局模式
            pyBridge.add_manual_redaction_global(key || '', text);
        }
        hideContextMenu();
        const selection = window.getSelection();
        selection.removeAllRanges();
    }

    function findTextPosition(selectedText, range) {
        try {
            // v1.1.11: 移除敏感信息日志
            console.log('[findTextPosition] ========== 开始查找 ==========');
            console.log('[findTextPosition] 选中文本（已隐藏内容）');
            console.log('[findTextPosition] Range:', {
                startContainer: range.startContainer?.nodeName,
                startOffset: range.startOffset,
                endContainer: range.endContainer?.nodeName,
                endOffset: range.endOffset
                // 移除 text 字段，避免泄露敏感信息
            });

            let container = range.commonAncestorContainer;
            if (!container) {
                console.warn('[findTextPosition] ✗ 无 commonAncestorContainer，使用全局降级');
                return createGlobalFallbackResult(selectedText);
            }

            console.log('[findTextPosition] commonAncestorContainer:', container.nodeName, container.nodeType);

            // 如果是文本节点，获取其父元素
            if (container.nodeType === 3) {
                container = container.parentNode;
            }

            // 查找包含 data-key 的元素（向上遍历）
            let element = container;
            let maxIterations = 50;
            let iterations = 0;
            let foundKey = null;

            while (element && iterations < maxIterations) {
                iterations++;
                if (element.dataset && element.dataset.key) {
                    foundKey = element.dataset.key;
                    console.log('[findTextPosition] ✓ 找到 data-key:', foundKey);
                    break;
                }
                element = element.parentNode;
                if (element === document.body || element === document.documentElement) {
                    console.warn('[findTextPosition] 到达文档顶部，未找到 data-key，尝试文本内容定位');
                    break;  // 不返回 null，继续尝试其他方法
                }
            }

            // 方法 1: 精确计算（如果有 data-key）
            if (foundKey && element) {
                const key = foundKey;

                // === 方法 1a: 直接使用 Range 计算位置（最可靠）===
                try {
                    const textNodes = [];
                    const walker = document.createTreeWalker(
                        element,
                        NodeFilter.SHOW_TEXT,
                        {
                            acceptNode: function(node) {
                                // 接受所有非空文本节点
                                return NodeFilter.FILTER_ACCEPT;
                            }
                        }
                    );

                    let node;
                    while (node = walker.nextNode()) {
                        textNodes.push(node);
                    }

                    console.log('[findTextPosition] 找到', textNodes.length, '个文本节点');

                    // 打印所有文本节点信息
                    for (let i = 0; i < Math.min(textNodes.length, 10); i++) {
                        const tn = textNodes[i];
                        console.log(`[findTextPosition] 节点[${i}]:`, {
                            text: tn.textContent.substring(0, 30),
                            isStart: tn === range.startContainer,
                            isEnd: tn === range.endContainer
                        });
                    }

                    // 计算起始位置
                    let startOffset = 0;
                    let startFound = false;

                    // 特殊处理：如果 startContainer 是元素节点，找到它的第一个文本节点
                    let startContainer = range.startContainer;
                    if (startContainer.nodeType === 1) {  // 元素节点
                        console.log('[findTextPosition] startContainer 是元素节点，查找第一个文本子节点');
                        for (let i = 0; i < textNodes.length; i++) {
                            if (element.contains(textNodes[i]) || element === textNodes[i].parentNode) {
                                startContainer = textNodes[i];
                                break;
                            }
                        }
                    }

                    for (let i = 0; i < textNodes.length; i++) {
                        const tn = textNodes[i];
                        if (!startFound) {
                            if (tn === startContainer || (startContainer && tn.contains && tn.contains(startContainer))) {
                                startOffset += range.startOffset;
                                startFound = true;
                                console.log('[findTextPosition] ✓ 起始节点匹配, offset:', range.startOffset);
                            } else {
                                startOffset += tn.textContent.length;
                            }
                        }
                    }

                    // 计算结束位置
                    let endOffset = 0;
                    let endFound = false;

                    let endContainer = range.endContainer;
                    if (endContainer.nodeType === 1) {
                        console.log('[findTextPosition] endContainer 是元素节点，查找第一个文本子节点');
                        for (let i = 0; i < textNodes.length; i++) {
                            if (element.contains(textNodes[i]) || element === textNodes[i].parentNode) {
                                endContainer = textNodes[i];
                                break;
                            }
                        }
                    }

                    for (let i = 0; i < textNodes.length; i++) {
                        const tn = textNodes[i];
                        if (!endFound) {
                            if (tn === endContainer || (endContainer && tn.contains && tn.contains(endContainer))) {
                                endOffset += range.endOffset;
                                endFound = true;
                                console.log('[findTextPosition] ✓ 结束节点匹配, offset:', range.endOffset);
                            } else {
                                endOffset += tn.textContent.length;
                            }
                        }
                    }

                    console.log('[findTextPosition] Range 计算结果:', { startFound, endFound, startOffset, endOffset });

                    if (startFound && endFound && startOffset >= 0 && endOffset > startOffset) {
                        console.log('[findTextPosition] ✓✓✓ Range 计算成功 ✓✓✓');
                        return { key, start: startOffset, end: endOffset };
                    }
                } catch (e) {
                    console.error('[findTextPosition] Range 计算出错:', e);
                }

                // === 方法 1b: 文本匹配（后备）===
                console.log('[findTextPosition] 尝试文本匹配方法...');
                let originalText = '';
                if (element.dataset.originalText) {
                    // 安全解码：使用 DOMParser 替代 innerHTML，防止 XSS
                    const parser = new DOMParser();
                    const doc = parser.parseFromString(element.dataset.originalText, 'text/html');
                    originalText = doc.body.textContent || '';
                    console.log('[findTextPosition] 原始文本长度:', originalText.length);
                }

                if (!originalText) {
                    originalText = element.textContent || '';
                }

                // 尝试直接匹配
                let foundStart = originalText.indexOf(selectedText);
                if (foundStart !== -1) {
                    console.log('[findTextPosition] ✓✓✓ 直接匹配成功 ✓✓✓');
                    return { key, start: foundStart, end: foundStart + selectedText.length };
                }

                // 尝试归一化匹配
                const normalizedSelected = selectedText.replace(/\s+/g, ' ').trim();
                const normalizedOriginal = originalText.replace(/\s+/g, ' ');
                foundStart = normalizedOriginal.indexOf(normalizedSelected);
                if (foundStart !== -1) {
                    console.log('[findTextPosition] ✓✓✓ 归一化匹配成功 ✓✓✓');
                    return { key, start: foundStart, end: foundStart + selectedText.length };
                }
            }

            // 方法 2: 文本内容定位（遍历所有 data-key 元素）
            console.log('[findTextPosition] 尝试文本内容定位方法...');
            const textResult = findPositionByTextContent(selectedText);
            if (textResult) {
                console.log('[findTextPosition] ✓✓✓ 文本内容定位成功 ✓✓✓');
                return textResult;
            }

            // 方法 3: 全局脱敏降级
            console.log('[findTextPosition] 所有精确方法失败，使用全局降级模式');
            return createGlobalFallbackResult(selectedText);

        } catch (e) {
            console.error('[findTextPosition] 异常:', e);
            return createGlobalFallbackResult(selectedText);
        }
    }

    function createGlobalFallbackResult(selectedText) {
        console.log('[findTextPosition] 创建全局降级结果，文本长度:', selectedText.length);
        return {
            key: '__GLOBAL__',
            start: 0,
            end: selectedText.length,
            mode: 'global',
            text: selectedText
        };
    }

    function findPositionByTextContent(selectedText) {
        // 遍历所有带有 data-key 的元素，查找包含选中文本的元素
        const elements = document.querySelectorAll('[data-key]');
        const normalizedSelected = selectedText.replace(/\s+/g, ' ').trim();

        for (const el of elements) {
            const key = el.dataset.key;
            if (!key || key === '__GLOBAL__') continue;

            let originalText = '';
            if (el.dataset.originalText) {
                // 安全解码：使用 DOMParser 替代 innerHTML，防止 XSS
                const parser = new DOMParser();
                const doc = parser.parseFromString(el.dataset.originalText, 'text/html');
                originalText = doc.body.textContent || '';
            } else {
                originalText = el.textContent || '';
            }

            const normalizedOriginal = originalText.replace(/\s+/g, ' ');
            const foundStart = normalizedOriginal.indexOf(normalizedSelected);

            if (foundStart !== -1) {
                console.log('[findTextPosition] 文本内容定位找到匹配，key:', key);
                return { key, start: foundStart, end: foundStart + selectedText.length };
            }
        }

        return null;
    }

    // v1.1.11 安全修复：使用 DOM 方法替代 innerHTML，防止 XSS
    function createMenu(buttonConfigs) {
        hideContextMenu();
        const menu = document.createElement('div');
        menu.id = 'redaction-menu';
        menu.style.cssText = 'position:fixed; background:white; border:1px solid #ddd; border-radius:8px; box-shadow:0 4px 12px rgba(0,0,0,0.15); padding:8px 0; z-index:10000; min-width:150px;';

        const buttonStyle = 'display:block; width:100%; padding:8px 16px; border:none; background:none; text-align:left; cursor:pointer; font-size:14px; color:#333;';

        // 安全创建按钮元素，避免使用 innerHTML
        buttonConfigs.forEach(config => {
            const btn = document.createElement('button');
            btn.style.cssText = buttonStyle;
            btn.onmouseover = () => btn.style.backgroundColor = '#f5f5f5';
            btn.onmouseout = () => btn.style.backgroundColor = 'transparent';

            // 设置按钮文本（自动转义 HTML）
            btn.textContent = config.text || '';

            // 设置 data 属性
            if (config.action) btn.setAttribute('data-action', config.action);
            if (config.key !== undefined) btn.setAttribute('data-key', config.key);
            if (config.start !== undefined) btn.setAttribute('data-start', config.start);
            if (config.end !== undefined) btn.setAttribute('data-end', config.end);
            if (config.textData !== undefined) btn.setAttribute('data-text', config.textData);

            // 设置 disabled 状态
            if (config.disabled) {
                btn.disabled = true;
                btn.style.color = '#999';
                btn.style.cursor = 'default';
            }

            menu.appendChild(btn);
        });

        document.body.appendChild(menu);
        return menu;
    }

    function positionMenu(menu, x, y) {
        menu.style.display = 'block';
        menu.style.left = x + 'px';
        menu.style.top = y + 'px';
    }

    function hideContextMenu() {
        const menu = document.getElementById('redaction-menu');
        if (menu) menu.remove();
    }
</script>
"""

# === 主窗口 ===
class MainWindow(MainWindowToolbarMixin, MainWindowWorkbenchMixin, MainWindowWordPreviewMixin, MainWindowPdfRenderMixin, MainWindowBatchReplaceMixin, MainWindowDensityMixin, MainWindowSetupMixin, MainWindowHandlersMixin, QMainWindow):
    def _apply_light_theme(self):
        """应用浅色主题样式(v1.1.11: Windows 强制浅色主题;v1.1.13 PR-B1: QSS 集中化)。

        622 行 QSS 已迁至 secureredact/ui/styles/*.qss;本函数只负责分发样式表到对应 widget。
        """
        from secureredact.ui.styles import StylesheetLoader, get_substitution_map
        loader = StylesheetLoader()
        loader.apply(self, "light", scope="main")
        if hasattr(self, "scroll") and hasattr(self, "scroll_style"):
            self.scroll.setStyleSheet(
                self.scroll_style.format(get_substitution_map("light")["scroll_area"])
            )
        if hasattr(self, "word_compare_header"):
            self.word_compare_header.setStyleSheet("""
                QFrame#wordCompareHeader { background-color: transparent; border: none; }
            """)
        if hasattr(self, "lbl_word_original_header"):
            self.lbl_word_original_header.setStyleSheet(f"""
                QLabel#wordCompareLabel {{
                    color: {Theme.LIGHT["text_secondary"]};
                    background-color: #F7FAFD;
                    border: 1px solid {Theme.LIGHT["border"]};
                    border-radius: 12px;
                    padding: 4px 12px;
                    font-size: 12px;
                    font-weight: 700;
                }}
            """)
        if hasattr(self, "lbl_word_replaced_header"):
            self.lbl_word_replaced_header.setStyleSheet(f"""
                QLabel#wordCompareLabel {{
                    color: {Theme.LIGHT["text_secondary"]};
                    background-color: #F7FAFD;
                    border: 1px solid {Theme.LIGHT["border"]};
                    border-radius: 12px;
                    padding: 4px 12px;
                    font-size: 12px;
                    font-weight: 700;
                }}
            """)
        if hasattr(self, "word_header_divider"):
            self.word_header_divider.setStyleSheet("background-color: transparent; border: none;")
        self._refresh_mode_badge()
        self._refresh_word_compare_toggle()
        self._refresh_workbench_context()

    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} {APP_VERSION}")

        # v1.1.11: 从配置读取窗口尺寸，失败时使用硬编码后备
        # v1.1.11: 读取 OCR 引擎配置
        if config:
            min_width = config.get("app.window.min_width", 900)
            min_height = config.get("app.window.min_height", 600)
            default_width = config.get("app.window.default_width", 1300)
            default_height = config.get("app.window.default_height", 900)
            self.replacement_text = config.get("redaction.replacement_text", "[已脱敏]")
            self.scan_level = config.get("redaction.scan.default_level", 2.0)
            self.offset_x = config.get("redaction.offset.default_x", 0)
            self.offset_w = config.get("redaction.offset.default_w", 0)
            self.custom_keywords = config.get("redaction.custom_keywords", "")
            # v1.1.11: 中文姓名启发式识别开关(默认 False,向后兼容)
            self.enable_name_recognition = config.get(
                "redaction.enable_name_recognition", False)
            # v1.1.11: 黑/白名单(用于设置中心初始化)
            self.current_blacklist = config.get("redaction.blacklist", []) or []
            self.current_whitelist = config.get("redaction.whitelist", []) or []
            # v1.1.11: 移除 OCR 引擎配置，只使用 RapidOCR
        else:
            min_width, min_height = 900, 600
            default_width, default_height = 1300, 900
            self.replacement_text = "[已脱敏]"
            self.scan_level = 2.0
            self.offset_x = 0
            self.offset_w = 0
            self.custom_keywords = ""
            self.enable_name_recognition = False
            self.current_blacklist = []
            self.current_whitelist = []

        # 窗口尺寸设置：最小尺寸 + 默认尺寸
        self.setMinimumSize(min_width, min_height)
        self.resize(default_width, default_height)

        # 窗口状态保存
        self.settings = QSettings("SecureRedact", "App")
        self._restore_window_state()

        self.doc = None
        self.word_doc = None  # Word 文档对象
        self.file_path = ""
        self.current_page = 0
        self.zoom_level = 1.0
        self.page_data = {}
        self._ocr_processed_pages = set()  # OCR 实际处理过的页（用于准确完成状态提示）
        # v1.1.11: 人工干预 override store + 当前文档 hash
        self._override_store = HitOverrideStore.instance()
        self._current_doc_hash = ""
        self.word_data = {}  # Word 文档数据结构
        self.word_replace_rules = []  # 会话级多字段替换规则
        self.word_compare_mode = False  # Word 预览是否开启原文/替换后对比
        self.word_compare_user_hidden = False  # 用户主动隐藏右侧对比预览
        self._word_data_lock = QMutex()  # v1.1.11: 保护 word_data 线程安全
        self.doc_type = None  # 'pdf', 'docx', 'doc'
        self.current_ui_mode = "idle"  # idle / pdf / word / batch / image_merge
        self.batch_stage = "idle"  # idle / rule_setup / running / finished / stopped
        self.batch_selected_files = []
        self.batch_total_files = 0
        self.batch_processed_files = 0
        self.batch_success_count = 0
        self.batch_failed_count = 0
        self.batch_current_file = ""
        self.batch_last_summary = None
        self.batch_result_filter_mode = "all"
        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self.info_bar_message = ""
        self.toolbar_density_mode = "wide"
        self._bound_window_handle = None
        self._button_density_metrics = {}
        # v1.1.12: 动态构造 active_rules — 遍历 DEFAULT_RULES 包含所有内置规则
        # 这样新增规则(如公司名)无需手动添加到这里
        # 但保留历史硬编码顺序作为 fallback, 避免回归
        self.active_rules = [
            DEFAULT_RULES.get("身份证号", ""),
            DEFAULT_RULES.get("手机号码", ""),
            # v1.1.11: 起诉讼书场景常用字段
            DEFAULT_RULES.get("地址（含门牌号）", ""),
            DEFAULT_RULES.get("固定电话", ""),
            DEFAULT_RULES.get("法定代表人", ""),
            # v1.1.12: 统一社会信用代码 - Word 命中,PDF 通过派发过滤隔离
            DEFAULT_RULES.get("统一社会信用代码", ""),
            # v1.1.12: 公司名 - 命中"有限公司/集团/公司"等尾缀, 智能 mask 保留省/市+公司尾缀
            DEFAULT_RULES.get("公司名", ""),
        ]
        # 兜底:遍历 DEFAULT_RULES,把上面列表中缺失的 pattern 补齐,确保新规则不漏
        for _name, _pat in DEFAULT_RULES.items():
            if _pat and _pat not in self.active_rules and _pat != "__SEAL_DETECTION__":
                self.active_rules.append(_pat)
        self.use_enhance = False
        self.current_color = QColor(0, 0, 0)
        self.dual_view = False

        # 预先创建 word_preview
        self.word_preview = None
        self.word_preview_replaced = None
        self.bridge = None
        self.word_web_channel = None
        self._word_scroll_sync_timer = QTimer(self)
        self._word_scroll_sync_timer.setInterval(160)
        self._word_scroll_sync_timer.timeout.connect(self._poll_word_compare_scroll_sync)
        self._word_scroll_sync_polling = False
        self._word_scroll_sync_last_ratios = {"original": None, "replaced": None}
        self._word_scroll_sync_pending_target = None
        self._word_scroll_sync_pending_ratio = None
        self._word_scroll_sync_generation = 0
        self._word_preview_assets_dir = None
        self._word_preview_assets_base_url = QUrl()
        self._reset_word_preview_cache()

        # 线程管理和临时文件管理（v24 稳定性优化）
        self.active_worker = None
        self.batch_worker = None
        self.active_task_type = None  # 'scan', 'batch_replace'
        self.worker_lock = QMutex()
        self.temp_manager = TempFileManager()

        # 注册退出清理
        import atexit
        atexit.register(self._app_exit_cleanup)

        self.setup_ui()

        # v1.1.11: 干预 override store (右键菜单 + 永久 override)
        self._override_store.bind_config(config)
        # 启动时加载 config.json 中的 permanent overrides
        perms = config.get("redaction.overrides.permanent", []) if config else []
        if perms:
            self._override_store.load_permanent(perms)

        # v1.1.11: 黑/白名单 store 引导 (bind_config + load_permanent)
        BlackWhiteListStore.instance().bind_config(config)
        BlackWhiteListStore.instance().load_permanent(
            config.get("redaction.blacklist", []) if config else [],
            config.get("redaction.whitelist", []) if config else [],
        )

        # v1.1.11: 启用拖拽支持
        self.setAcceptDrops(True)
        self._drag_active = False  # 拖拽状态标记
        self._drag_valid = False   # 拖拽文件是否有效

    def _detect_system_theme(self):
        """检测系统主题（v1.1.11 新增）

        Returns:
            str: 'light' 或 'dark'
        """
        import platform

        system = platform.system()

        try:
            if system == 'Darwin':  # macOS
                import subprocess
                result = subprocess.run(
                    ['defaults', 'read', '-g', 'AppleInterfaceStyle'],
                    capture_output=True, text=True
                )
                if result.returncode == 0 and 'Dark' in result.stdout:
                    return 'dark'
                return 'light'

            elif system == 'Windows':
                import winreg
                key = winreg.OpenKey(
                    winreg.HKEY_CURRENT_USER,
                    r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize"
                )
                value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
                winreg.CloseKey(key)
                return 'light' if value == 1 else 'dark'

            else:  # Linux
                import os
                gtk_theme = os.environ.get('GTK_THEME', '').lower()
                if 'dark' in gtk_theme:
                    return 'dark'
                return 'light'

        except (OSError, IOError, KeyError, ValueError, ImportError) as e:
            # 检测失败（注册表读取错误、环境变量异常等），默认浅色
            print(f"[MainWindow] 主题检测失败: {e}")
            return 'light'

    def _restore_window_state(self):
        """恢复窗口状态"""
        geometry = self.settings.value("window_geometry")
        if geometry:
            self.restoreGeometry(geometry)

    def closeEvent(self, event):
        """保存窗口状态并清理临时文件"""
        # v1.1.11: 关闭前把 permanent overrides 落盘
        if hasattr(self, "_override_store"):
            try:
                self._override_store.save_permanent()
            except Exception as exc:
                print(f"[关闭] 保存 permanent override 失败: {exc}")
        self._app_exit_cleanup()
        self.settings.setValue("window_geometry", self.saveGeometry())
        super().closeEvent(event)

    def showEvent(self, event):
        """首次显示和重新显示窗口时，补绑屏幕缩放监听。"""
        super().showEvent(event)
        self._bind_window_handle_signals()
        QTimer.singleShot(0, self._refresh_toolbar_responsiveness)

    def resizeEvent(self, event):
        """窗口大小改变时自动重新适应页面"""
        super().resizeEvent(event)

        if hasattr(self, "toolbar"):
            self._refresh_toolbar_responsiveness()

        # 只在 PDF 模式且文档已加载时处理
        if not self.doc or self.current_page is None:
            return

        # 自动重新适应（保持页面完整显示）
        self.fit_page()

    def _bind_window_handle_signals(self):
        """绑定窗口句柄相关信号，便于跨屏和 DPI 切换时收口工具栏。"""
        try:
            handle = self.windowHandle()
        except Exception:
            handle = None

        if handle is None or handle is self._bound_window_handle:
            return

        try:
            handle.screenChanged.connect(self._on_window_screen_changed)
        except Exception:
            pass
        self._bound_window_handle = handle

    def _on_window_screen_changed(self, _screen):
        """窗口切换到不同屏幕时，重新计算工具栏密度。"""
        QTimer.singleShot(0, self._refresh_toolbar_responsiveness)





    # ============== v1.1.11: 拖拽打开文件功能 ==============

    def dragEnterEvent(self, event):
        """拖拽进入事件 - 验证文件类型并提供视觉反馈"""
        self._drag_active = True
        self._drag_valid = False

        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()

            # 验证所有文件
            valid_exts = ['.pdf', '.doc', '.docx', '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
            all_valid = True
            file_count = 0

            for url in urls:
                if url.isLocalFile():
                    file_count += 1
                    path = url.toLocalFile().lower()
                    if not any(path.endswith(ext) for ext in valid_exts):
                        all_valid = False
                        break

            # 必须有文件且都有效
            if file_count > 0 and all_valid:
                self._drag_valid = True
                event.acceptProposedAction()
                self._update_drag_visual_feedback(True)
            else:
                # 有文件但格式不支持
                if file_count > 0:
                    event.ignore()
                    self._update_drag_visual_feedback(False)
                else:
                    event.ignore()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        """拖拽移动事件 - 持续反馈"""
        if self._drag_active and event.mimeData().hasUrls():
            # 检查鼠标位置是否在预览区域
            pos = event.position().toPoint() if hasattr(event, 'position') else event.pos()
            if self._is_in_preview_area(pos):
                event.acceptProposedAction()
            else:
                event.ignore()
        else:
            event.ignore()

    def dragLeaveEvent(self, event):
        """拖拽离开事件 - 清除视觉反馈"""
        self._drag_active = False
        self._drag_valid = False
        self._update_drag_visual_feedback(None)

    def dropEvent(self, event):
        """拖放事件 - 处理文件"""
        self._drag_active = False
        self._update_drag_visual_feedback(None)

        if not event.mimeData().hasUrls():
            event.ignore()
            return

        urls = event.mimeData().urls()
        file_paths = [url.toLocalFile() for url in urls if url.isLocalFile()]

        if not file_paths:
            QMessageBox.warning(self, "无效文件", "只支持本地文件拖拽")
            event.ignore()
            return

        # 调用处理逻辑
        self._handle_dropped_files(file_paths)
        event.acceptProposedAction()

    def _is_in_preview_area(self, pos):
        """检查坐标是否在预览区域内"""
        # pos 是 MainWindow 坐标，转换到 scroll 局部坐标后做命中测试
        map_pos = self.scroll.mapFrom(self, pos)
        return self.scroll.rect().contains(map_pos)

    def _update_drag_visual_feedback(self, valid):
        """更新拖拽视觉反馈

        Args:
            valid: True(有效), False(无效), None(清除)
        """
        if valid is True:
            # 有效文件 - 绿色边框提示
            self.scroll.setStyleSheet(f"""
                QScrollArea {{
                    background-color: {Theme.LIGHT["scroll_area"]};
                    border-radius: {Theme.BORDER_RADIUS}px;
                    border: 3px solid #34C759;
                }}
            """)
        elif valid is False:
            # 无效文件 - 红色边框提示
            self.scroll.setStyleSheet(f"""
                QScrollArea {{
                    background-color: {Theme.LIGHT["scroll_area"]};
                    border-radius: {Theme.BORDER_RADIUS}px;
                    border: 3px solid #FF3B30;
                }}
            """)
        else:
            # 清除 - 恢复默认
            self.scroll.setStyleSheet(self.scroll_style.format(Theme.LIGHT["scroll_area"]))

    def _handle_dropped_files(self, file_paths):
        """处理拖拽的文件

        Args:
            file_paths: 文件路径列表
        """
        # 清理之前的状态
        self._cleanup_before_open()
        self._cleanup_temp_file()

        if len(file_paths) == 1:
            # 单个文件
            fname = file_paths[0]
            doc_type = self.detect_file_type(fname)

            try:
                if doc_type == 'pdf':
                    self._open_pdf_file(fname)
                elif doc_type == 'docx':
                    self._open_word_docx(fname)
                elif doc_type == 'doc':
                    self._open_word_doc(fname)
                elif doc_type == 'image':
                    self._open_images_merge([fname])
                else:
                    QMessageBox.warning(
                        self, "不支持的格式",
                        f"文件: {os.path.basename(fname)}\n\n"
                        f"请选择 PDF、Word 文档或图片文件"
                    )
            except (IOError, OSError, ValueError, ConversionError) as e:
                QMessageBox.critical(self, "错误", f"打开文件失败: {str(e)}")
        else:
            # 多个文件：支持图片合并或 Word 批量替换
            are_all_images = all(self.detect_file_type(f) == 'image' for f in file_paths)
            are_all_words = all(self.detect_file_type(f) in ('docx', 'doc') for f in file_paths)
            if are_all_images:
                try:
                    self._open_images_merge(file_paths)
                except (IOError, OSError, ValueError, ConversionError) as e:
                    QMessageBox.critical(self, "错误", f"合并图片失败: {str(e)}")
            elif are_all_words:
                self.start_batch_replace(file_paths=file_paths)
            else:
                QMessageBox.warning(
                    self, "不支持的混合拖拽",
                    "同时拖拽多个文件时，仅支持两种场景：\n"
                    "1. 全部是图片（自动合并为PDF）\n"
                    "2. 全部是Word（自动启动批量替换）"
                )

    def _show_drag_tooltip(self, file_paths, valid):
        """显示拖拽文件提示信息（可选增强）

        Args:
            file_paths: 文件路径列表
            valid: 是否有效
        """
        if not file_paths:
            return

        if len(file_paths) == 1:
            fname = os.path.basename(file_paths[0])
            doc_type = self.detect_file_type(file_paths[0])
            type_names = {
                'pdf': 'PDF 文档',
                'docx': 'Word 文档',
                'doc': 'Word 文档(旧版)',
                'image': '图片文件',
                'unknown': '未知格式'
            }
            tooltip = f"{fname}\n类型: {type_names.get(doc_type, '未知')}"
        else:
            tooltip = f"共 {len(file_paths)} 个文件"
            if valid:
                tooltip += "\n将合并为 PDF"
            else:
                tooltip += "\n格式不支持"

        # 使用 QToolTip 显示
        from PyQt6.QtWidgets import QToolTip
        from PyQt6.QtGui import QCursor
        QToolTip.showText(QCursor.pos(), tooltip)

    # ============== 拖拽功能结束 ==============


    def _has_active_open_context(self):
        """判断当前是否已有活跃文档/任务，需要在打开新文件前先做清理。"""
        return bool(
            self.doc
            or self.word_doc
            or self.file_path
            or self.doc_type
            or self.page_data
            or self.word_data
            or self.image_merge_in_progress
            or (self.batch_stage != "idle" and self.batch_selected_files)
            or self.current_ui_mode != "idle"
        )




    def _is_word_web_view_valid(self, web_view):
        """判断 Word 预览 WebView 是否仍然可用。"""
        if web_view is None:
            return False
        if sip.isdeleted(web_view):
            return False
        try:
            page = web_view.page()
            if page is None or sip.isdeleted(page):
                return False
            _ = web_view.isHidden()
            return True
        except RuntimeError:
            return False

    def _invalidate_word_scroll_sync(self):
        """停用并失效当前 Word 双栏滚动同步链，避免异步回调撞上已销毁对象。"""
        if hasattr(self, "_word_scroll_sync_timer"):
            self._word_scroll_sync_timer.stop()
        self._word_scroll_sync_polling = False
        self._word_scroll_sync_pending_target = None
        self._word_scroll_sync_pending_ratio = None
        self._word_scroll_sync_last_ratios = {"original": None, "replaced": None}
        self._word_scroll_sync_generation += 1

    def _app_exit_cleanup(self):
        """应用退出时的清理（v24 稳定性优化）"""
        self._invalidate_word_scroll_sync()
        # 取消正在运行的线程
        if self.active_worker and self.active_worker.isRunning():
            self.active_worker.requestInterruption()
            # 等待线程结束（最多 2 秒）
            self.active_worker.wait(2000)

        # 清理临时文件
        if hasattr(self, 'temp_manager'):
            self.temp_manager.cleanup()

        # 清理旧版临时文件
        self._cleanup_temp_file()

    # v1.1.11: 移除 eventFilter，直接在 SinglePageCanvas.mousePressEvent 中处理


















    def _get_display_scale_factor(self):
        """返回当前显示环境的缩放因子，主要用于 Windows DPI 收口。"""
        import platform

        if platform.system() != "Windows":
            return 1.0

        screen = None
        try:
            handle = self.windowHandle()
            if handle:
                screen = handle.screen()
        except Exception:
            screen = None

        if screen is None:
            screen = QApplication.primaryScreen()
        if screen is None:
            return 1.0

        try:
            scale = screen.logicalDotsPerInch() / 96.0
        except Exception:
            scale = 1.0

        return max(1.0, min(scale, 2.0))


    def _refresh_button_density_styles(self):
        """在 DPI 或密度切换后，重新应用所有按钮的字体和内边距。"""
        for button in self.findChildren(QPushButton):
            style_type = button.property("btn_style")
            if not style_type:
                continue
            if button is getattr(self, "btn_more", None):
                continue
            button.setStyleSheet(self._get_button_style(style_type))

        self._refresh_toolbar_more_button_style()
        self._apply_native_toolbar_icons()









    def _refresh_workflow_steps(self, active_index):
        """刷新顶部流程步骤，让主路径始终可见。"""
        if not hasattr(self, "workflow_step_labels"):
            return
        step_font_size = getattr(self, "_workflow_step_font_size", 11)

        for index, label in enumerate(self.workflow_step_labels):
            if index < active_index:
                fg = Theme.LIGHT["success"]
                bg = "#EAF8F1"
            elif index == active_index:
                fg = Theme.LIGHT["primary"]
                bg = "#E9F1FB"
            else:
                fg = Theme.LIGHT["text_secondary"]
                bg = Theme.LIGHT["hover"]

            label.setStyleSheet(
                f"""
                QLabel#workflowStep {{
                    color: {fg};
                    background-color: {bg};
                    border: 1px solid {Theme.LIGHT["border"]};
                    border-radius: 10px;
                    padding: 5px 10px;
                    font-size: {step_font_size}px;
                    font-weight: 700;
                }}
                """
            )




    def _count_enabled_general_rules(self):
        """统计启用的通用规则数量。"""
        return len([rule for rule in self.active_rules if isinstance(rule, str) and rule])

    def _count_enabled_word_rules(self):
        """统计启用的 Word 替换规则数量。"""
        normalized = normalize_word_replace_rules(self.word_replace_rules, self.replacement_text)
        return len([item for item in normalized if item.get("enabled", True) and item.get("find")])

    def _has_pdf_redactions(self):
        """当前 PDF 是否已有智能/手动涂抹结果。"""
        return any(data.get("ocr") or data.get("manual") for data in self.page_data.values())

    def _has_word_redactions(self):
        """当前 Word 是否已有智能/手动脱敏结果。"""
        return any(data.get("ocr") or data.get("manual") for data in self.word_data.values())















    def _refresh_merge_workspace(self):
        """刷新图片合并模式提示。"""
        if not hasattr(self, "lbl_merge_meta"):
            return

        total_images = self.image_merge_total_images
        in_progress = self.image_merge_in_progress and total_images > 0

        if in_progress:
            self.lbl_merge_title.setText("图片正在合并为 PDF")
            self.lbl_merge_subtitle.setText("系统会按当前顺序生成 PDF，完成后自动进入 PDF 脱敏工作台。")
            self.lbl_merge_meta.setText(
                f"当前共 {total_images} 张图片。合并完成后会自动打开生成的 PDF，继续进入 PDF 脱敏流程。"
            )
            self._set_status_badge_style(self.lbl_merge_stage_badge, Theme.LIGHT["primary"], "#E9F1FB")
            self.lbl_merge_stage_badge.setText("合并中")
            stage_states = ["done", "active", "pending"]
            if hasattr(self, "lbl_merge_metric_images"):
                self.lbl_merge_metric_images.setText(str(total_images))
                self.lbl_merge_metric_images_note.setText("当前已载入待合并的图片数量")
            if hasattr(self, "lbl_merge_metric_status"):
                self.lbl_merge_metric_status.setText("进行中")
                self.lbl_merge_metric_status_note.setText("底部进度条会同步显示合并进度")
            if hasattr(self, "lbl_merge_metric_next"):
                self.lbl_merge_metric_next.setText("PDF 脱敏")
                self.lbl_merge_metric_next_note.setText("完成后自动进入 PDF 脱敏工作台")
        else:
            self.lbl_merge_title.setText("图片合并为 PDF")
            self.lbl_merge_subtitle.setText("系统会先整理图片顺序，再生成 PDF 并自动打开工作台。")
            self.lbl_merge_meta.setText("当前还没有开始合并。")
            self._set_status_badge_style(self.lbl_merge_stage_badge, Theme.LIGHT["warning"], "#FFF3E6")
            self.lbl_merge_stage_badge.setText("等待开始")
            stage_states = ["active", "pending", "pending"]
            if hasattr(self, "lbl_merge_metric_images"):
                self.lbl_merge_metric_images.setText(str(total_images) if total_images else "--")
                self.lbl_merge_metric_images_note.setText("拖入多张图片后会自动进入该工作台")
            if hasattr(self, "lbl_merge_metric_status"):
                self.lbl_merge_metric_status.setText("未开始")
                self.lbl_merge_metric_status_note.setText("当前尚未开始合并图片")
            if hasattr(self, "lbl_merge_metric_next"):
                self.lbl_merge_metric_next.setText("等待图片")
                self.lbl_merge_metric_next_note.setText("拖入图片后开始合并并进入 PDF 工作台")

        for index, state in enumerate(stage_states):
            if index < len(getattr(self, "merge_stage_cards", [])):
                frame, title_label, note_label = self.merge_stage_cards[index]
                self._set_batch_step_style(frame, title_label, note_label, state)





    def _sync_ui_mode(self):
        """根据当前运行状态推导界面模式。"""
        if self.active_task_type == "batch_replace":
            self._set_ui_mode("batch")
        elif self.image_merge_in_progress:
            self._set_ui_mode("image_merge")
        elif self.doc:
            self._set_ui_mode("pdf")
        elif self.word_doc:
            self._set_ui_mode("word")
        elif self.batch_stage != "idle" and self.batch_selected_files:
            self._set_ui_mode("batch")
        else:
            self._set_ui_mode("idle")



    def create_btn(self, text, func, enabled=True, style="primary", width=None, tooltip=""):
        """创建现代化按钮

        Args:
            text: 按钮文本
            func: 点击回调
            enabled: 是否启用
            style: 样式类型 (primary, secondary, success, danger, icon)
            width: 固定宽度（可选）
            tooltip: 工具提示
        """
        btn = QPushButton(text)
        btn.clicked.connect(func)
        btn.setEnabled(enabled)
        if tooltip:
            btn.setToolTip(tooltip)
        if width:
            btn.setFixedWidth(width)

        # 设置游标
        btn.setCursor(Qt.CursorShape.PointingHandCursor)
        btn.setAutoDefault(False)
        btn.setDefault(False)
        btn.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)

        # 应用样式
        btn.setStyleSheet(self._get_button_style(style))

        # 保存样式类型以便主题切换
        btn.setProperty("btn_style", style)

        return btn





    def toggle_dual_view(self, checked):
        """
        v1.1.11: 简化的单/双页切换 - 保持 canvas_container 为固定 widget
        只隐藏/显示内部的 canvas
        """
        self.dual_view = checked

        if checked:
            # 双页模式：显示两个 canvas
            self.canvas_right.show()
        else:
            # 单页模式：只显示 left canvas
            self.canvas_right.hide()

        self.render_view()


    def open_settings(self):
        # v1.1.11: 传递配置管理器以支持配置持久化
        # v1.1.11: 移除 OCR 引擎选择，只保留 RapidOCR

        dlg = SettingsDialog(self, self.active_rules, self.use_enhance, self.custom_keywords,
                            self.scan_level, self.offset_x, self.offset_w, self.replacement_text,
                            self.word_replace_rules,
                            config_manager=config,
                            enable_name_recognition=self.enable_name_recognition,
                            current_blacklist=self.current_blacklist,
                            current_whitelist=self.current_whitelist)
        if dlg.exec():
            self.active_rules = dlg.selected_rules
            self.use_enhance = dlg.use_enhance
            self.custom_keywords = dlg.custom_keywords
            self.scan_level = dlg.scan_level
            self.offset_x = dlg.offset_x
            self.offset_w = dlg.offset_w
            self.replacement_text = dlg.replacement_text
            self.word_replace_rules = dlg.word_replace_rules
            # v1.1.11: 同步姓名识别开关
            self.enable_name_recognition = dlg.enable_name_recognition
            # v1.1.11: 同步黑/白名单(已在文本变化时实时落盘,此处再拉一次 store 的最终值)
            self.current_blacklist = list(BlackWhiteListStore.instance().effective_blacklist())
            self.current_whitelist = list(BlackWhiteListStore.instance().effective_whitelist())
            if self.word_doc:
                if not self._has_word_replacement_candidates():
                    self.word_compare_user_hidden = False
                self.render_word_preview()
            else:
                self._refresh_workbench_context()
            self._clear_info_bar_message()
            msg = self.create_message_box(self, QMessageBox.Icon.Information, "成功", "设置已保存")
            msg.exec()

    def show_feedback(self):
        """显示反馈与开发者信息对话框"""
        dlg = FeedbackDialog(self)
        dlg.exec()

    @staticmethod
    def create_message_box(parent, icon, title, text, buttons=QMessageBox.StandardButton.Ok, default_button=QMessageBox.StandardButton.Ok):
        """创建带有浅色主题样式的消息框（v1.1.11: 修复 Windows 深色模式显示问题）

        Args:
            parent: 父窗口
            icon: QMessageBox.Icon 类型
            title: 标题
            text: 内容文本
            buttons: 按钮类型
            default_button: 默认按钮

        Returns:
            QMessageBox: 配置好的消息框实例
        """
        msg = QMessageBox(parent)
        msg.setIcon(icon)
        msg.setWindowTitle(title)
        msg.setText(text)
        msg.setStandardButtons(buttons)
        msg.setDefaultButton(default_button)

        # 设置窗口标志，防止 Windows 强制应用深色模式
        msg.setWindowFlags(
            Qt.WindowType.Dialog |
            Qt.WindowType.WindowTitleHint |
            Qt.WindowType.WindowCloseButtonHint
        )

        # 应用浅色主题样式
        from theme import Theme
        theme = Theme.LIGHT
        msg.setStyleSheet(f"""
            QMessageBox {{
                background-color: {theme["background"]};
            }}
            QMessageBox QLabel {{
                color: {theme["text"]};
                background-color: transparent;
            }}
            QPushButton {{
                background-color: {theme["primary"]};
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px 16px;
                font-weight: 600;
                min-width: 80px;
            }}
            QPushButton:hover {{
                background-color: {theme["primary"]};
                opacity: 0.9;
            }}
        """)

        return msg

    def _on_replacement_changed(self, text):
        """替换文本变化时的处理"""
        self.replacement_text = text if text.strip() else "[已脱敏]"



    def _has_word_replacement_candidates(self):
        """是否存在可在右侧预览中展示的替换结果（规则/OCR/手动）。"""
        if self._has_enabled_word_replace_rules():
            return True
        for data in self.word_data.values():
            if not isinstance(data, dict):
                continue
            if data.get("manual") or data.get("ocr"):
                return True
        return False


    def _build_replaced_preview_fragment(self, source_text, merged_matches):
        """根据合并后的匹配区间构建右侧预览片段（统一高亮样式）。"""
        from html import escape as html_escape

        source_label_map = {
            "rule": "规则替换",
            "manual": "手动脱敏",
            "ocr": "智能脱敏"
        }
        segments = build_replaced_preview_segments(source_text, merged_matches, self.replacement_text)
        fragment_parts = []

        for segment in segments:
            value = segment.get("value", "")
            if value is None:
                value = ""
            if not isinstance(value, str):
                value = str(value)
            escaped_value = html_escape(value).replace("\n", "<br/>")

            if segment.get("type") != "replacement":
                fragment_parts.append(escaped_value)
                continue

            source = str(segment.get("source", "rule"))
            title = source_label_map.get(source, "文本替换")
            rule_name = str(segment.get("rule_name", "")).strip()
            if rule_name:
                title = f"{title} ({rule_name})"

            title_attr = html_escape(title)
            source_attr = html_escape(source)
            fragment_parts.append(
                f'<mark class="replace-preview-highlight" data-source="{source_attr}" title="{title_attr}">{escaped_value}</mark>'
            )

        return "".join(fragment_parts)







    def detect_file_type(self, fname):
        """检测文件类型"""
        ext = os.path.splitext(fname)[1].lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.doc':
            return 'doc'
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
            return 'image'
        else:
            return 'unknown'

    def _get_file_dialog_style(self):
        """获取文件对话框样式（v1.1.11: 使用系统默认按钮样式确保跨平台可读性）"""
        return f"""
            QFileDialog {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                font-family: {Theme.FONT_FAMILY};
                font-size: 12px;
            }}
            QFileDialog QLabel {{ color: #1D1D1F; }}
            /* 按钮使用系统默认样式，确保跨平台可读性 */
            /* macOS/Windows 会自动应用适合的按钮颜色 */
            QFileDialog QListView, QFileDialog QTreeView {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D1D1D6;
            }}
            QFileDialog QListView::item:selected, QFileDialog QTreeView::item:selected {{
                background-color: #007AFF;
                color: white;
            }}
            QFileDialog QComboBox {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D1D1D6;
                padding: 4px;
            }}
            QFileDialog QLineEdit {{
                background-color: #FFFFFF;
                color: #1D1D1F;
                border: 1px solid #D1D1D6;
                padding: 4px;
            }}
        """

    def open_pdf(self):
        """打开文件（支持图片多选和多 Word 批量替换）"""
        try:
            has_active_context = self._has_active_open_context()

            # 已有文档/任务时，保持原有清理策略，避免新旧上下文互相污染。
            if has_active_context:
                self._cleanup_before_open()
                self._cleanup_temp_file()

            # v1.1.11: 使用原生文件对话框，更稳定
            # 不使用 DontUseNativeDialog，让系统处理渲染
            fnames, _ = QFileDialog.getOpenFileNames(
                self, "选择文件", "",
                "支持的文件 (*.pdf *.doc *.docx *.jpg *.jpeg *.png *.bmp *.tiff)"
            )

            if not fnames:
                return

            # 首次从空首页打开文件时，推迟到真正选中文件后再做清理，
            # 避免文件对话框弹出前触发首页抖动/重排。
            if not has_active_context:
                self._cleanup_before_open()
                self._cleanup_temp_file()

            # 根据选择数量处理
            if len(fnames) == 1:
                # 单个文件，按原有逻辑处理
                fname = fnames[0]
                doc_type = self.detect_file_type(fname)
                if doc_type == 'pdf':
                    self._open_pdf_file(fname)
                elif doc_type == 'docx':
                    self._open_word_docx(fname)
                elif doc_type == 'doc':
                    self._open_word_doc(fname)
                elif doc_type == 'image':
                    self._open_images_merge([fname])
                else:
                    QMessageBox.warning(self, "不支持的格式", "请选择 PDF、Word 文档或图片文件")
            else:
                # 多个文件：支持图片合并或 Word 批量替换
                are_all_images = all(self.detect_file_type(f) == 'image' for f in fnames)
                are_all_words = all(self.detect_file_type(f) in ('doc', 'docx') for f in fnames)
                if are_all_images:
                    self._open_images_merge(fnames)
                elif are_all_words:
                    self.start_batch_replace(file_paths=fnames)
                else:
                    QMessageBox.warning(self, "不支持的混合选择",
                        "同时选择多个文件时，仅支持两种场景：\n"
                        "1. 全部是图片（自动合并为PDF）\n"
                        "2. 全部是Word（自动启动批量替换）")
        except (IOError, OSError, ValueError, ConversionError) as e:
            QMessageBox.critical(self, "错误", f"打开文件失败: {str(e)}")

    def _open_pdf_file(self, fname):
        """内部方法：打开 PDF 文件"""
        # 关闭已打开的PDF文档（防止资源泄露）
        if self.doc:
            self.doc.close()
            self.doc = None

        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self._reset_batch_session_state()
        self._clear_info_bar_message()
        self.file_path = fname
        self.doc = fitz.open(fname)
        self.doc_type = 'pdf'
        total = len(self.doc)
        self.page_data = {i: {'ocr': [], 'manual': []} for i in range(total)}
        # v1.1.11: 计算当前文档 hash,用于 HitOverrideStore 关联 override
        try:
            self._current_doc_hash = compute_doc_hash(fname)
        except OSError:
            self._current_doc_hash = ""
        self.current_page = 0
        self.word_doc = None
        self.word_data = {}
        self.word_compare_mode = False
        self.word_compare_user_hidden = False
        self.btn_scan.setEnabled(True)
        self.btn_save.setEnabled(True)

        # 切换显示：显示 canvas，隐藏 Word 预览
        self.canvas_container.show()
        self.word_compare_container.hide()
        self.word_preview.hide()
        self.word_preview_replaced.hide()
        self._sync_ui_mode()

        self.fit_page()


    def _open_word_doc(self, fname):
        """打开 .doc 文件（通过转换为 .docx）"""
        import shutil
        import subprocess
        import tempfile

        # 检查系统支持（v1.1.11: 使用增强的跨平台检测）
        support_info = self._check_doc_support()
        method = support_info['recommended']

        if not method:
            # 无可用工具，显示安装指南
            self._show_doc_install_guide()
            return

        # 显示格式限制提示
        tool_name = "LibreOffice" if method == 'libreoffice' else "antiword"
        reply = QMessageBox.question(
            self,
            "格式提示",
            f".doc 是旧版 Word 格式，将使用 {tool_name} 转换。\n\n"
            f"{'（antiword 只保留纯文本，会丢失格式）' if method == 'antiword' else '转换后可能丢失部分格式。'}\n\n"
            "建议先在 Word 中另存为 .docx 格式以获得最佳效果。\n\n"
            "是否继续转换？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.No:
            return

        # 转换并打开
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            docx_path = self._convert_doc_to_docx(fname, method)
            if docx_path:
                # 使用转换后的文件打开
                self.converted_temp_file = docx_path  # 保存临时文件路径以便后续清理
                self._open_word_docx(docx_path)
                self._set_info_bar_message("📝 已完成 .doc 转换，现已进入 Word 替换工作台。")
        except (IOError, OSError, ValueError, RuntimeError) as e:
            QMessageBox.critical(self, "转换失败", f"无法转换 .doc 文件:\n{str(e)}")
        finally:
            QApplication.restoreOverrideCursor()

    def _check_doc_support(self):
        """检查系统是否支持 .doc 格式（v1.1.11: 增强跨平台检测）"""
        import shutil
        import platform

        system = platform.system()
        result = {
            'libreoffice': False,
            'antiword': False,
            'recommended': None
        }

        # 检查 LibreOffice
        if system == 'Darwin':  # macOS
            # 检查应用程序目录
            libreoffice_app = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
            if os.path.exists(libreoffice_app):
                result['libreoffice'] = True
            # 也检查 PATH
            elif shutil.which('soffice'):
                result['libreoffice'] = True
        elif system == 'Windows':
            # Windows: 检查常见安装路径
            program_files = os.environ.get('ProgramFiles', 'C:\\Program Files')
            program_files_x86 = os.environ.get('ProgramFiles(x86)', 'C:\\Program Files (x86)')
            libreoffice_paths = [
                os.path.join(program_files, 'LibreOffice', 'program', 'soffice.exe'),
                os.path.join(program_files_x86, 'LibreOffice', 'program', 'soffice.exe'),
            ]
            for path in libreoffice_paths:
                if os.path.exists(path):
                    result['libreoffice'] = True
                    break
        else:  # Linux
            if shutil.which('soffice') or shutil.which('libreoffice'):
                result['libreoffice'] = True

        # 检查 antiword
        if shutil.which('antiword'):
            result['antiword'] = True

        # 确定推荐方案
        if result['libreoffice']:
            result['recommended'] = 'libreoffice'
        elif result['antiword']:
            result['recommended'] = 'antiword'

        return result


    def _convert_doc_to_docx(self, doc_path, method='libreoffice'):
        """v1.1.11: 委托给共享转换模块。"""
        from secureredact.utils.doc_converter import (
            convert_with_libreoffice, convert_with_antiword,
        )
        try:
            if method == 'libreoffice':
                docx_path, temp_dir = _shared_convert_doc_to_docx(doc_path)
                self.converted_temp_file = docx_path
                return docx_path
            elif method == 'antiword':
                temp_dir = self.temp_manager.create_temp_dir()
                return convert_with_antiword(doc_path, temp_dir=temp_dir)
            else:
                raise ValueError(f"不支持的转换方法: {method}")
        except ConversionError:
            raise
        except (OSError, IOError, RuntimeError, ValueError) as e:
            raise ConversionError(f"转换出错: {e}", "请尝试在 Word 中手动另存为 .docx 格式")

    def _open_images_merge(self, image_paths):
        """处理图片合并为PDF"""
        try:
            self.image_merge_in_progress = True
            self.image_merge_total_images = len(image_paths)
            self._set_ui_mode("image_merge")
            # 1. 让用户排序（如果是多张图片）
            if len(image_paths) > 1:
                dlg = ImageListDialog(image_paths, self)
                if dlg.exec() != QDialog.DialogCode.Accepted:
                    self.image_merge_in_progress = False
                    self.image_merge_total_images = 0
                    self._sync_ui_mode()
                    return
                image_paths = dlg.get_ordered_paths()
                self.image_merge_total_images = len(image_paths)

            # 2. 生成输出路径（在第一张图片所在目录）
            base_dir = os.path.dirname(image_paths[0])
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output_filename = f"merged_{timestamp}.pdf"
            final_path = os.path.join(base_dir, output_filename)

            # 3. 创建临时文件
            temp_pdf = self.temp_manager.create_temp_file(suffix='.pdf')

            # 4. 启动Worker合并
            self.merge_worker = ImageMergeWorker(image_paths, temp_pdf)
            self.merge_worker.progress_signal.connect(self.progress.setValue)
            self.merge_worker.finished_signal.connect(lambda p: self._on_merge_finished(p, final_path))
            self.merge_worker.error_signal.connect(self._on_merge_error)
            self.merge_worker.start()

            self._set_info_bar_message(f"正在合并 {len(image_paths)} 张图片...")
            self._refresh_workbench_context()

        except (IOError, OSError, ValueError, RuntimeError) as e:
            self.image_merge_in_progress = False
            self.image_merge_total_images = 0
            self._sync_ui_mode()
            QMessageBox.critical(self, "错误", f"启动图片合并失败: {str(e)}")

    def _on_merge_finished(self, temp_pdf, final_path):
        """合并完成回调"""
        try:
            # 将临时文件移动到最终位置
            shutil.move(temp_pdf, final_path)
            self._set_info_bar_message(f"✓ 合并完成: {final_path}")
            self.image_merge_in_progress = False

            # 自动打开生成的PDF
            self._open_pdf_file(final_path)

        except (IOError, OSError, ValueError) as e:
            self.image_merge_in_progress = False
            self.image_merge_total_images = 0
            QMessageBox.critical(self, "错误", f"保存合并文件失败: {str(e)}")

    def _on_merge_error(self, error_msg):
        """合并错误回调"""
        self.image_merge_in_progress = False
        self.image_merge_total_images = 0
        self._set_info_bar_message("✗ 合并失败")
        self._sync_ui_mode()
        QMessageBox.critical(self, "合并失败", error_msg)

    def _cleanup_temp_file(self):
        """清理转换产生的临时文件（v1.1.11 安全修复：跨平台安全清理）"""
        import tempfile

        if hasattr(self, 'converted_temp_file') and self.converted_temp_file:
            try:
                if os.path.exists(self.converted_temp_file):
                    os.remove(self.converted_temp_file)
                    print(f"[清理] 已删除临时文件: {self.converted_temp_file}")

                # 清理临时目录（安全版本）
                temp_dir = os.path.dirname(self.converted_temp_file)

                # 获取系统临时目录（跨平台：Windows 返回 C:\Users\...\AppData\Local\Temp，macOS/Linux 返回 /tmp 或 /var/tmp）
                system_temp_dir = tempfile.gettempdir()

                # 规范化路径进行比较（处理大小写敏感/不敏感、路径分隔符等）
                norm_temp_dir = os.path.normcase(os.path.abspath(temp_dir))
                norm_system_temp = os.path.normcase(os.path.abspath(system_temp_dir))

                # 安全检查：只有当目录位于系统临时目录下，且不是系统临时目录本身时才删除
                if (norm_temp_dir.startswith(norm_system_temp + os.sep) and
                    norm_temp_dir != norm_system_temp):
                    try:
                        # 检查目录是否为空（防止误删文件）
                        if os.path.isdir(temp_dir) and not os.listdir(temp_dir):
                            os.rmdir(temp_dir)
                            print(f"[清理] 已删除空临时目录: {temp_dir}")
                        else:
                            print(f"[清理] 临时目录非空或不存在，跳过删除: {temp_dir}")
                    except OSError as e:
                        print(f"[清理] 删除临时目录时出错（可能是非空）: {temp_dir} - {e}")
                else:
                    print(f"[清理] 跳过非系统临时目录: {temp_dir}")

            except Exception as e:
                print(f"[清理] 清理临时文件时出错: {e}")
            finally:
                self.converted_temp_file = None





    def on_rect_added(self, page_idx, rect):
        """由于使用共享列表引用，canvas 已直接修改列表，这里只需刷新视图"""
        # self.page_data[page_idx]['manual'].append(rect)  # 不需要，canvas 已经添加
        self.render_view()

    def on_rect_removed(self, page_idx, rect_idx, is_manual):
        """由于使用共享列表引用，canvas 已直接修改列表，这里只需刷新视图"""
        # if is_manual: del self.page_data[page_idx]['manual'][rect_idx]  # 不需要，canvas 已经删除
        # else: del self.page_data[page_idx]['ocr'][rect_idx]
        self.render_view()


    def go_first(self):
        if not self.doc: return
        self.current_page = 0
        self.render_view()
        self.scroll.verticalScrollBar().setValue(0)

    def go_last(self):
        if not self.doc: return
        self.current_page = len(self.doc) - 1
        if self.dual_view and self.current_page % 2 != 0: self.current_page -= 1
        self.render_view()
        self.scroll.verticalScrollBar().setValue(0)



    def zoom_in(self):
        self.handle_zoom_request(0.25)

    def zoom_out(self):
        self.handle_zoom_request(-0.25)


    def cancel_ocr_scan(self):
        """取消智能脱敏扫描（v1.1.11）"""
        if self.active_worker and self.active_worker.isRunning():
            if self.active_task_type == "batch_replace":
                title = "确认停止"
                text = "确定要停止批量替换吗？\n已完成的文件会被保留。"
            else:
                title = "确认取消"
                text = "确定要停止扫描吗？\n已扫描的进度将被保留。"
            reply = QMessageBox.question(
                self,
                title,
                text,
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                if self.active_task_type == "batch_replace":
                    self._set_info_bar_message("⏹️ 正在停止批量替换...")
                    self._append_batch_log("收到停止指令，正在安全结束批量任务。", "warning")
                else:
                    self._set_info_bar_message("⏹️ 正在停止扫描...")
                self.btn_cancel_scan.setEnabled(False)  # 防止重复点击
                self.active_worker.requestInterruption()  # 请求中断
                if self.active_task_type == "batch_replace" and hasattr(self, "batch_worker") and self.batch_worker:
                    self.batch_worker.provide_error_decision("stop")
                # Worker会在完成后通过finished_signal通知主线程

    def _on_worker_finished(self):
        """v1.1.11: 工作线程完成后的清理 + 延迟错误显示"""
        # v1.1.11: 等待线程完全终止，防止死锁
        if self.active_worker and self.active_worker.isRunning():
            self.active_worker.wait(3000)  # 最多等待 3 秒

        was_cancelled = self.active_worker and self.active_worker.isInterruptionRequested()
        self.active_worker = None
        self.active_task_type = None
        self.btn_scan.setEnabled(True)
        self.btn_cancel_scan.setVisible(False)  # 隐藏取消按钮
        self.btn_cancel_scan.setEnabled(True)
        self._sync_ui_mode()

        if was_cancelled:
            self._set_info_bar_message("⏹️ 扫描已取消，已保留部分结果")
        else:
            self._set_info_bar_message("✅ 扫描完成！")

        # v1.1.11: 线程清理完成后，延迟显示错误对话框（非阻塞）
        if hasattr(self, '_pending_error_msg') and self._pending_error_msg:
            error_msg = self._pending_error_msg
            self._pending_error_msg = None
            QTimer.singleShot(100, lambda: self._show_deferred_error(error_msg))

    def _show_deferred_error(self, error_msg: str):
        """v1.1.11: 安全显示错误对话框（在线程清理完成后调用）"""
        QMessageBox.critical(
            self,
            "OCR 错误",
            f"智能脱敏遇到问题：\n\n{error_msg}\n\n"
            "可能的解决方案：\n"
            "1. 重新安装 OCR 依赖: pip install rapidocr-onnxruntime\n"
            "2. 安装 Visual C++ 运行库（Windows）\n"
            "3. 如果是扫描版 PDF，请尝试使用文本版 PDF"
        )

    def ocr_finished(self, results):
        """v1.1.11: 添加去重逻辑，解决重复矩形导致的点击2次问题；v1.1.11: 支持部分结果"""
        total_pages = len(self.page_data)
        scanned_pages = len(results)

        for page, rects in results.items():
            # 去重：移除重复或高度重叠的矩形
            deduped = self._deduplicate_rects(rects)
            self.page_data[page]['ocr'] = deduped
            if len(rects) != len(deduped):
                if DEBUG_MODE:
                    print(f"[DEBUG] 页面{page}: 去重前{len(rects)}个矩形, 去重后{len(deduped)}个")
        self.render_view()
        self.progress.setValue(0)

        # 判断是部分结果还是完整结果（v1.1.11）
        if scanned_pages < total_pages:
            QMessageBox.information(
                self,
                "扫描已取消",
                f"已扫描 {scanned_pages}/{total_pages} 页，\n部分结果已保留，可以继续编辑。"
            )
        else:
            QMessageBox.information(self, "完成", "智能扫描已完成！")

    # v1.1.11: 非阻塞 OCR 错误处理
    def _on_ocr_error(self, error_msg: str):
        """v1.1.11: 非阻塞处理 OCR 错误（存储错误，延迟到线程清理后显示）"""
        print(f"[OCR ERROR] {error_msg}")
        self._set_info_bar_message(f"❌ OCR 错误: {error_msg[:50]}...")
        # v1.1.11: 存储错误消息，延迟到线程清理完成后显示
        # 避免在模态对话框阻塞主线程时形成死锁
        self._pending_error_msg = error_msg



    @staticmethod
    def _filter_hits_to_rects(hits: list, *, store, location: str, doc_hash: str) -> list:
        """v1.1.11: 模块级纯函数 — store 过滤 + 抽 QRectF.

        抽到这里便于直接单测,无需启动 QMainWindow 实例。
        语义: manual 永远保留;ignore 命中剔除;confirm / 未操作保留。
        """
        kept = store.filtered_hits(hits, location=location, doc_hash=doc_hash)
        return [h["rect"] for h in kept]

    def _on_ocr_finished_safe(self, _):
        """v1.1.11: 线程安全 - OCR 完成处理（在主线程执行）

        参数 _ 是空字典，保留以兼容信号签名
        """
        self.render_view()
        self.progress.setValue(0)

        # 统计已扫描的页面
        scanned_pages = len(self._ocr_processed_pages) if self._ocr_processed_pages else 0
        total_pages = len(self.page_data)
        was_cancelled = self.active_worker is not None and self.active_worker.isInterruptionRequested()

        # 判断是部分结果还是完整结果（以取消状态为准，避免“无命中=已取消”误判）
        if was_cancelled:
            QMessageBox.information(
                self,
                "扫描已取消",
                f"已扫描 {scanned_pages}/{total_pages} 页，\n部分结果已保留，可以继续编辑。"
            )
        else:
            QMessageBox.information(self, "完成", "智能扫描已完成！")

    def _deduplicate_rects(self, rects):
        """移除重复或高度重叠的矩形"""
        if not rects:
            return rects

        # 按中心点坐标和尺寸排序，便于去重
        sorted_rects = sorted(rects, key=lambda r: (r.x(), r.y(), r.width(), r.height()))
        deduped = []

        for rect in sorted_rects:
            is_duplicate = False
            for existing in deduped:
                # 检查是否与已有矩形高度重叠（IoU > 0.7 或中心点距离 < 5 像素）
                if self._is_overlapping(rect, existing):
                    is_duplicate = True
                    break

            if not is_duplicate:
                deduped.append(rect)

        return deduped

    def _is_overlapping(self, rect1, rect2, threshold=0.7):
        """检查两个矩形是否高度重叠"""
        # 计算交集
        x_left = max(rect1.x(), rect2.x())
        y_top = max(rect1.y(), rect2.y())
        x_right = min(rect1.x() + rect1.width(), rect2.x() + rect2.width())
        y_bottom = min(rect1.y() + rect1.height(), rect2.y() + rect2.height())

        if x_right < x_left or y_bottom < y_top:
            return False

        # 计算交集面积
        intersection = (x_right - x_left) * (y_bottom - y_top)
        area1 = rect1.width() * rect1.height()
        area2 = rect2.width() * rect2.height()

        # IoU > threshold 视为重复
        union = area1 + area2 - intersection
        iou = intersection / union if union > 0 else 0

        # 或者中心点距离 < 5 像素也视为重复
        center1_x = rect1.x() + rect1.width() / 2
        center1_y = rect1.y() + rect1.height() / 2
        center2_x = rect2.x() + rect2.width() / 2
        center2_y = rect2.y() + rect2.height() / 2
        distance = ((center1_x - center2_x)**2 + (center1_y - center2_y)**2)**0.5

        return iou > threshold or distance < 5

    def word_scan_finished(self, results):
        """Word 文档扫描完成（v1.1.11: 线程安全）"""
        results_copy = dict(results)
        scan_meta = results_copy.pop('__scan_meta__', None)

        # v1.1.11: 使用锁保护 word_data 访问
        with QMutexLocker(self._word_data_lock):
            total_items = len(self.word_data)
            if scan_meta:
                processed_items = int(scan_meta.get('processed_items', total_items))
                total_from_meta = int(scan_meta.get('total_items', total_items))
                if total_from_meta > 0:
                    total_items = total_from_meta
            else:
                processed_items = total_items

            self.word_data = results_copy

            # 统计扫描结果
            total_matches = sum(len(data['ocr']) for data in self.word_data.values())

        self.render_word_preview()
        self.progress.setValue(0)

        was_cancelled = False
        if scan_meta is not None:
            was_cancelled = bool(scan_meta.get('cancelled', False))
        elif self.active_worker is not None:
            was_cancelled = self.active_worker.isInterruptionRequested()

        # 判断是部分结果还是完整结果（以取消状态为准）
        if was_cancelled:
            QMessageBox.information(
                self,
                "扫描已取消",
                f"已处理 {processed_items}/{total_items} 个段落/单元格，\n部分结果已保留，可以继续编辑。"
            )
        else:
            QMessageBox.information(self, "完成", f"智能扫描已完成！\n共发现 {total_matches} 处敏感信息")








    def _apply_word_panel_updates(self, web_view, block_updates):
        if web_view is None or not block_updates:
            return

        script = build_word_panel_update_script(block_updates)
        web_view.page().runJavaScript(script)

    def _configure_word_scroll_sync_panel(self, web_view, panel_id):
        """配置单个 Word 预览面板的滚动同步能力。"""
        if not self._is_word_web_view_valid(web_view):
            return
        sync_enabled = bool(
            self.word_doc
            and self.word_compare_mode
            and not self.word_compare_user_hidden
            and not web_view.isHidden()
        )
        panel_js = json.dumps(str(panel_id))
        enabled_js = "true" if sync_enabled else "false"
        web_view.page().runJavaScript(
            f"""
            if (window.__setWordPreviewPanelId) window.__setWordPreviewPanelId({panel_js});
            if (window.__setWordPreviewScrollSyncEnabled) window.__setWordPreviewScrollSyncEnabled({enabled_js});
            """
        )

    def _configure_word_scroll_sync_panels(self):
        """同步刷新左右 Word 预览面板的滚动联动状态。"""
        if hasattr(self, "word_preview"):
            self._configure_word_scroll_sync_panel(self.word_preview, "original")
        if hasattr(self, "word_preview_replaced"):
            self._configure_word_scroll_sync_panel(self.word_preview_replaced, "replaced")
        self._refresh_word_scroll_sync_timer()

    def _refresh_word_scroll_sync_timer(self):
        """按当前 Word 双栏状态启停滚动同步轮询。"""
        should_run = bool(
            self.word_doc
            and self.word_compare_mode
            and not self.word_compare_user_hidden
            and self._word_preview_ready
            and self._word_replaced_ready
            and self._is_word_web_view_valid(self.word_preview)
            and self._is_word_web_view_valid(self.word_preview_replaced)
            and not self.word_preview.isHidden()
            and not self.word_preview_replaced.isHidden()
        )
        if should_run:
            if not self._word_scroll_sync_timer.isActive():
                self._word_scroll_sync_timer.start()
        else:
            self._word_scroll_sync_timer.stop()
            self._word_scroll_sync_polling = False
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None
            self._word_scroll_sync_last_ratios = {"original": None, "replaced": None}

    def _apply_word_scroll_ratio_to_panel(self, panel_id, ratio):
        """把滚动比例应用到指定 Word 预览面板。"""
        try:
            ratio_value = max(0.0, min(1.0, float(ratio)))
        except (TypeError, ValueError):
            return

        target_view = self.word_preview_replaced if panel_id == "replaced" else self.word_preview
        if not self._is_word_web_view_valid(target_view) or target_view.isHidden():
            return

        self._word_scroll_sync_pending_target = panel_id
        self._word_scroll_sync_pending_ratio = ratio_value
        target_view.page().runJavaScript(
            f"if (window.__applyExternalWordScrollRatio) window.__applyExternalWordScrollRatio({ratio_value:.6f});"
        )





    def _handle_word_scroll_sync_original_ratio(self, original_ratio, generation):
        """获取左侧比例后继续读取右侧比例。"""
        if generation != self._word_scroll_sync_generation:
            self._word_scroll_sync_polling = False
            return
        try:
            normalized_original = max(0.0, min(1.0, float(original_ratio or 0.0)))
        except (TypeError, ValueError):
            normalized_original = 0.0

        if not self._is_word_web_view_valid(self.word_preview_replaced) or self.word_preview_replaced.isHidden():
            self._word_scroll_sync_polling = False
            return

        self.word_preview_replaced.page().runJavaScript(
            "window.__getWordPreviewScrollRatio ? window.__getWordPreviewScrollRatio() : 0;",
            lambda replaced_ratio, sync_generation=generation: self._handle_word_scroll_sync_ratio_pair(normalized_original, replaced_ratio, sync_generation),
        )

    def _handle_word_scroll_sync_ratio_pair(self, original_ratio, replaced_ratio, generation):
        """比较双栏滚动位置，并同步变化更明显的一侧。"""
        if generation != self._word_scroll_sync_generation:
            self._word_scroll_sync_polling = False
            return
        try:
            normalized_original = max(0.0, min(1.0, float(original_ratio or 0.0)))
        except (TypeError, ValueError):
            normalized_original = 0.0
        try:
            normalized_replaced = max(0.0, min(1.0, float(replaced_ratio or 0.0)))
        except (TypeError, ValueError):
            normalized_replaced = 0.0

        previous_original = self._word_scroll_sync_last_ratios.get("original")
        previous_replaced = self._word_scroll_sync_last_ratios.get("replaced")
        self._word_scroll_sync_last_ratios = {
            "original": normalized_original,
            "replaced": normalized_replaced,
        }

        if previous_original is None or previous_replaced is None:
            self._word_scroll_sync_polling = False
            return

        delta_original = abs(normalized_original - previous_original)
        delta_replaced = abs(normalized_replaced - previous_replaced)
        threshold = 0.012

        if (
            self._word_scroll_sync_pending_target == "replaced"
            and abs(normalized_replaced - float(self._word_scroll_sync_pending_ratio or 0.0)) <= 0.02
        ):
            delta_replaced = 0.0
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None
        elif (
            self._word_scroll_sync_pending_target == "original"
            and abs(normalized_original - float(self._word_scroll_sync_pending_ratio or 0.0)) <= 0.02
        ):
            delta_original = 0.0
            self._word_scroll_sync_pending_target = None
            self._word_scroll_sync_pending_ratio = None

        if delta_original >= threshold and delta_original > delta_replaced:
            self._apply_word_scroll_ratio_to_panel("replaced", normalized_original)
        elif delta_replaced >= threshold and delta_replaced > delta_original:
            self._apply_word_scroll_ratio_to_panel("original", normalized_replaced)

        self._word_scroll_sync_polling = False



    def _add_data_key_attributes(self, html, text_blocks):
        """使用 BeautifulSoup 为文本块添加 data-key 属性

        Args:
            html: HTML 字符串
            text_blocks: 文本块字典 {key: {'text': 原始文本, 'escaped': 转义文本}}

        Returns:
            修改后的 HTML 字符串
        """
        try:
            soup = BeautifulSoup(html, 'html.parser')

            # 支持的标签列表（扩展）
            target_tags = ['p', 'td', 'th', 'li', 'span', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'a']

            for key, info in text_blocks.items():
                original_text = info['text']
                if not original_text or not original_text.strip():
                    continue

                # 归一化文本：将多个空白字符合并为单个空格
                normalized_original = ' '.join(original_text.split())

                for tag_name in target_tags:
                    for element in soup.find_all(tag_name):
                        # 跳过已经有 data-key 的元素
                        if element.get('data-key'):
                            continue

                        # 获取元素的文本内容并归一化
                        element_text = element.get_text()
                        normalized_element = ' '.join(element_text.split())

                        # 比较归一化后的文本
                        if normalized_original == normalized_element:
                            element['data-key'] = key
                            element['data-original-text'] = original_text
                            element['data-word-block'] = '1'
                            break

            return str(soup)

        except (ImportError, AttributeError, TypeError, ValueError) as e:
            print(f"[警告] BeautifulSoup 处理失败，使用正则表达式后备方案: {e}")
            return self._add_data_key_regex_fallback(html, text_blocks)

    def _add_data_key_regex_fallback(self, html, text_blocks):
        """使用正则表达式为文本块添加 data-key 属性（后备方案）

        Args:
            html: HTML 字符串
            text_blocks: 文本块字典

        Returns:
            修改后的 HTML 字符串
        """
        from html import escape as html_escape

        for key, info in text_blocks.items():
            original_text = info['text']
            escaped_text = html_escape(original_text)

            # 尝试各种 HTML 标签
            patterns = [
                (
                    f'<p([^>]*)>({re.escape(escaped_text)})</p>',
                    f'<p\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</p>'
                ),
                (
                    f'<td([^>]*)>({re.escape(escaped_text)})</td>',
                    f'<td\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</td>'
                ),
                (
                    f'<li([^>]*)>({re.escape(escaped_text)})</li>',
                    f'<li\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</li>'
                ),
                (
                    f'<span([^>]*)>({re.escape(escaped_text)})</span>',
                    f'<span\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</span>'
                ),
                (
                    f'<div([^>]*)>({re.escape(escaped_text)})</div>',
                    f'<div\\1 data-key="{key}" data-original-text="{escaped_text}" data-word-block="1">\\2</div>'
                ),
            ]

            for pattern, replacement in patterns:
                if f'data-key="{key}"' not in html:
                    if re.search(pattern, html):
                        html = re.sub(pattern, replacement, html, count=1)
                        break

        return html



    def _inject_interactive_html(self, html, scroll_restore=''):
        """注入 JavaScript 交互逻辑用于右键菜单和脱敏操作

        v1.1.11: 重构为使用模块级常量 _INTERACTIVE_JS_CODE，简化函数逻辑

        Args:
            html: 要注入的 HTML
            scroll_restore: 滚动恢复脚本（可选）
        """
        # 包装 HTML 为完整文档（如果不是的话）
        html = self._wrap_html_document(html)

        # 注入脚本
        qwebchannel_js = '<script src="qrc:///qtwebchannel/qwebchannel.js"></script>'

        if '</head>' in html:
            html = html.replace('</head>', qwebchannel_js + _INTERACTIVE_JS_CODE + scroll_restore + '</head>')
        else:
            html = qwebchannel_js + _INTERACTIVE_JS_CODE + scroll_restore + html

        return html

    def _wrap_html_document(self, html):
        """将 HTML 包装成完整文档（如果不是完整文档的话）

        Args:
            html: 输入的 HTML 字符串

        Returns:
            完整的 HTML 文档字符串
        """
        is_full_document = '<html' in html.lower() or '<!doctype' in html.lower()

        if is_full_document:
            return html

        return f'''<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
    body {{ margin: 0; padding: 20px; font-family: {PREVIEW_FONT_STACK}; line-height: 1.6; }}
    img {{ max-width: 100%; height: auto; }}
    p {{ margin: 0 0 10px 0; }}
</style>
</head>
<body>
{html}
</body>
</html>'''



    def _replace_in_paragraph(self, para, matches, text_offset=0):
        """兼容入口：调用通用段落替换实现。"""
        replace_matches_in_paragraph(
            para,
            matches,
            text_offset=text_offset,
            fallback_replacement_text=self.replacement_text
        )

    def _apply_range_to_runs(self, para, start, end, replacement):
        """兼容入口：调用通用 run 区间替换实现。"""
        apply_range_to_runs(para, start, end, replacement)

    def _copy_run_format(self, target_run, source_run):
        """复制 run 的所有格式属性

        Args:
            target_run: 目标 run
            source_run: 源 run
        """
        # 字体属性
        if source_run.bold is not None:
            target_run.bold = source_run.bold
        if source_run.italic is not None:
            target_run.italic = source_run.italic
        if source_run.underline is not None:
            target_run.underline = source_run.underline
        if source_run.strike is not None:
            target_run.strike = source_run.strike

        # 字体名称和大小
        if source_run.font.name:
            try:
                target_run.font.name = source_run.font.name
            except (AttributeError, TypeError) as e:
                print(f"[字体复制] 复制字体名称失败: {e}")
        if source_run.font.size:
            target_run.font.size = source_run.font.size

        # 颜色
        if source_run.font.color and source_run.font.color.rgb:
            try:
                target_run.font.color.rgb = source_run.font.color.rgb
            except (AttributeError, TypeError) as e:
                print(f"[字体复制] 复制字体颜色失败: {e}")

        # 高亮
        if source_run.font.highlight_color:
            try:
                target_run.font.highlight_color = source_run.font.highlight_color
            except (AttributeError, TypeError) as e:
                print(f"[字体复制] 复制高亮颜色失败: {e}")

        # 下标/上标
        if source_run.font.subscript:
            target_run.font.subscript = True
        if source_run.font.superscript:
            target_run.font.superscript = True

if __name__ == "__main__":
    # v1.1.13 (PR-B0): 入口迁移过渡期 shim。
    # 真实入口已迁至 `secureredact.main:main`;本块保留仅为阶段 B5 收口前的兼容。
    # PR-B5 时移除本块,统一走 `python -m secureredact.main`。
    import sys as _sys
    from secureredact.main import main as _entry_main
    _sys.exit(_entry_main())
