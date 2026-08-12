"""Phase 3 合成含 PII 的 docx fixture（D-26 / OPS-05 — Faker 合成数据）。"""
import os
import tempfile
from typing import List, Optional


def build_fake_docx(
    paragraphs: Optional[List[str]] = None,
    tables: Optional[List[List[List[str]]]] = None,
    add_pii: bool = True,
) -> str:
    """合成一个 docx 文件（含可选 PII 段落）并返回临时文件路径。

    Args:
        paragraphs: 普通段落文本列表。
        tables: 表格内容，3 层 list：tables[table_idx][row_idx][cell_idx] = text。
        add_pii: True 时追加 5 个含 PII 段落（身份证 / 手机 / 邮箱 / 银行卡 / USCC）。

    Returns:
        docx 文件的临时路径（调用方负责 os.remove）。
    """
    from docx import Document

    doc = Document()

    if paragraphs:
        for p_text in paragraphs:
            doc.add_paragraph(p_text)

    if tables:
        for table_data in tables:
            if not table_data:
                continue
            row_count = len(table_data)
            col_count = max((len(r) for r in table_data), default=0)
            if row_count == 0 or col_count == 0:
                continue
            tbl = doc.add_table(rows=row_count, cols=col_count)
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    if r_idx < len(tbl.rows) and c_idx < len(tbl.rows[r_idx].cells):
                        tbl.rows[r_idx].cells[c_idx].text = cell_text

    if add_pii:
        from tests.fixtures.fake_pii import (
            fake_bank_card,
            fake_email,
            fake_id_card,
            fake_phone,
            fake_uscc,
        )
        doc.add_paragraph(f"甲方身份证 {fake_id_card()}")
        doc.add_paragraph(f"联系电话 {fake_phone()}")
        doc.add_paragraph(f"邮箱 {fake_email()}")
        doc.add_paragraph(f"卡号 {fake_bank_card()}")
        doc.add_paragraph(f"统一信用代码 {fake_uscc()}")

    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(path)
    return path


__all__ = ['build_fake_docx']
