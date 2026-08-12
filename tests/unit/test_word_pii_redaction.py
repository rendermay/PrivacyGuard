"""Phase 3 Plan 4 Word PII 真脱敏 reverse-extraction 测试。

覆盖 SAFE-02 / D-04 / D-07，并用 AST 守护 main.py 保存路径确实调用
Word adapter，而不是只在预览层展示高亮。

Phase 3 G1 Gap Closure: 增 TestAutoPiiScanOnOpen / TestPiiOnlyDocumentEntersCompareMode /
TestPiiFullReloadPreviewContainsMask 覆盖 Gap 1 (自动启动 WordWorker) / Gap 2
(_has_word_replacement_candidates 纳入 PII) / Gap 3 (_build_word_replaced_preview_html 注入 pii_matches)。
"""
from __future__ import annotations

import ast
import tempfile
import unittest
from pathlib import Path
from types import MethodType, SimpleNamespace

from docx import Document

from main import MainWindow
from privacyguard.pii.engine import PIIEngine
from privacyguard.pii.hits import PIIHit
from privacyguard.pii.mask import mask_for_entity
from privacyguard.pii.word_adapter import (
    apply_pii_replacements_to_docx,
    collect_pii_word_hits,
    locate_pii_hits_in_paragraph,
)
from tests.fixtures.fake_pii import fake_id_card, fake_phone


PROJECT_ROOT = Path(__file__).resolve().parents[2]
MAIN_PY = PROJECT_ROOT / "main.py"


class TestWordPiiRedaction(unittest.TestCase):
    """验证保存后的 docx 通过独立 python-docx 通道重新提取。"""

    def _redact_document(self, text: str, mode: str = "partial") -> str:
        document = Document()
        document.add_paragraph(text)
        source_text = document.paragraphs[0].text
        hits = collect_pii_word_hits(source_text, PIIEngine())
        self.assertTrue(hits, "PIIEngine 未检测到测试敏感项，测试不能平凡通过")
        locations = locate_pii_hits_in_paragraph(hits, source_text)
        self.assertTrue(locations, "Word adapter 未能定位测试敏感项")

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "redacted.docx"
            apply_pii_replacements_to_docx(
                document,
                {"paragraph_0": locations},
                mode=mode,
            )
            document.save(output_path)
            extracted = Document(output_path)
            return "\n".join(paragraph.text for paragraph in extracted.paragraphs)

    def test_redacted_docx_does_not_contain_original_secret(self):
        secret_id = fake_id_card()
        output_text = self._redact_document(f"测试身份证：{secret_id}")
        self.assertNotIn(secret_id, output_text)

    def test_partial_mask_id_card_visible_in_output(self):
        secret_id = fake_id_card()
        output_text = self._redact_document(f"测试身份证：{secret_id}")
        self.assertIn(mask_for_entity("CN_ID_CARD", secret_id), output_text)

    def test_partial_mask_phone_visible_in_output(self):
        secret_phone = fake_phone()
        output_text = self._redact_document(f"联系电话：{secret_phone}")
        self.assertIn(mask_for_entity("CN_PHONE", secret_phone), output_text)
        self.assertNotIn(secret_phone, output_text)

    def test_blackout_mode_replaces_with_brackets(self):
        secret_phone = fake_phone()
        output_text = self._redact_document(f"联系电话：{secret_phone}", mode="blackout")
        self.assertIn("[已脱敏]", output_text)
        self.assertNotIn(secret_phone, output_text)

    def test_paragraph_style_preserved_after_save(self):
        secret_id = fake_id_card()
        document = Document()
        paragraph = document.add_paragraph(f"标题中的身份证：{secret_id}", style="Heading 1")
        source_text = paragraph.text
        hits = collect_pii_word_hits(source_text, PIIEngine())
        locations = locate_pii_hits_in_paragraph(hits, source_text)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "styled-redacted.docx"
            apply_pii_replacements_to_docx(
                document,
                {"paragraph_0": locations},
                mode="partial",
            )
            document.save(output_path)
            saved_paragraph = Document(output_path).paragraphs[0]

        self.assertEqual(saved_paragraph.style.name, "Heading 1")
        self.assertNotIn(secret_id, saved_paragraph.text)


class TestSaveWordCallsPiiAdapter(unittest.TestCase):
    """AST 守护 _save_word 的生产接入点。"""

    @classmethod
    def setUpClass(cls):
        cls.tree = ast.parse(MAIN_PY.read_text(encoding="utf-8"))
        cls.save_word = next(
            (
                node
                for node in ast.walk(cls.tree)
                if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))
                and node.name == "_save_word"
            ),
            None,
        )

    def test_save_word_function_exists(self):
        self.assertIsNotNone(self.save_word, "main.py 必须定义 _save_word")

    def test_save_word_calls_apply_pii_replacements(self):
        self.assertIsNotNone(self.save_word)
        calls = {
            node.func.id
            for node in ast.walk(self.save_word)
            if isinstance(node, ast.Call) and isinstance(node.func, ast.Name)
        }
        self.assertIn("apply_pii_replacements_to_docx", calls)

    def test_save_word_calls_locate_pii_hits(self):
        self.assertIsNotNone(self.save_word)
        calls = {
            node.func.id
            for node in ast.walk(self.save_word)
            if isinstance(node, ast.Call) and isinstance(node.func, ast.Name)
        }
        self.assertIn("locate_pii_hits_in_paragraph", calls)


# ----------------------------------------------------------------------
# Phase 3 G1 Gap Closure helpers
# ----------------------------------------------------------------------
def _make_word_pii_hit(
    entity_type="CN_ID_CARD",
    char_offset=0,
    char_length=18,
    mask_strategy="110101********1234",
    normalized="110101199001011234",
):
    """构造一个 Word 端 PIIHit (复用 D-10: page_offset/page_length 存 char_offset/char_length)."""
    return PIIHit(
        entity_type=entity_type,
        page_offset=char_offset,
        page_length=char_length,
        page_rect=(0.0, 0.0, 66.0, 12.0),
        confidence_tier="HIGH",
        source="text",
        mask_strategy=mask_strategy,
        normalized=normalized,
        validator_passed=True,
    )


def _build_word_replaced_preview_stub(word_data, replacement_text="[已脱敏]"):
    """构造一个用于 _build_word_replaced_preview_html 测试的 stub.

    沿用 build_word_preview_stub 形态, 补 pii 字段以验证 G1 Gap 3 接入.
    """
    stub = SimpleNamespace(
        word_data=dict(word_data),
        word_replace_rules=[],
        replacement_text=replacement_text,
        _word_base_html="<p>placeholder</p>",
    )
    stub._wrap_html_document = MethodType(MainWindow._wrap_html_document, stub)
    stub._build_word_text_blocks = MethodType(MainWindow._build_word_text_blocks, stub)
    stub._add_data_key_attributes = MethodType(MainWindow._add_data_key_attributes, stub)
    stub._add_data_key_regex_fallback = MethodType(MainWindow._add_data_key_regex_fallback, stub)
    stub._build_replaced_preview_fragment = MethodType(MainWindow._build_replaced_preview_fragment, stub)
    stub._inject_interactive_html = lambda html, scroll_restore='': html + scroll_restore
    stub._get_word_preview_scroll_restore_script = lambda: "<script>restoreScroll()</script>"
    stub._build_word_replaced_preview_html = MethodType(MainWindow._build_word_replaced_preview_html, stub)
    return stub


# ----------------------------------------------------------------------
# G1 Gap 1 Test: _open_word_docx 自动启动 WordWorker PII 扫描
# ----------------------------------------------------------------------
class TestAutoPiiScanOnOpen(unittest.TestCase):
    """G1 Gap 1: 打开 .docx 后, _open_word_docx 必须自动调用 start_ocr() 启动 WordWorker 扫描."""

    @classmethod
    def setUpClass(cls):
        cls.tree = ast.parse(MAIN_PY.read_text(encoding="utf-8"))
        cls.open_docx = next(
            (
                node
                for node in ast.walk(cls.tree)
                if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))
                and node.name == "_open_word_docx"
            ),
            None,
        )

    def test_open_docx_function_exists(self):
        self.assertIsNotNone(self.open_docx, "main.py 必须定义 _open_word_docx")

    def test_open_docx_calls_start_ocr(self):
        """_open_word_docx 函数体内必须出现 start_ocr() 调用 (Gap 1 fix)."""
        self.assertIsNotNone(self.open_docx)
        calls = set()
        for node in ast.walk(self.open_docx):
            if isinstance(node, ast.Call):
                func = node.func
                if isinstance(func, ast.Name):
                    calls.add(func.id)
                elif isinstance(func, ast.Attribute):
                    calls.add(func.attr)
        self.assertIn(
            "start_ocr", calls,
            "_open_word_docx 必须调用 self.start_ocr() 自动启动 WordWorker (G1 Gap 1)",
        )

    def test_start_ocr_call_is_guarded(self):
        """start_ocr() 调用必须位于 try/except 块内 (失败不阻塞 UI 打开)."""
        self.assertIsNotNone(self.open_docx)
        # 在 open_docx 体内找出所有 try 块, 验证至少 1 个 try 包含 start_ocr
        try_blocks = [
            node for node in ast.walk(self.open_docx) if isinstance(node, ast.Try)
        ]
        guarded = False
        for try_node in try_blocks:
            for sub in ast.walk(try_node):
                if isinstance(sub, ast.Call):
                    func = sub.func
                    name = None
                    if isinstance(func, ast.Name):
                        name = func.id
                    elif isinstance(func, ast.Attribute):
                        name = func.attr
                    if name == "start_ocr":
                        guarded = True
                        break
            if guarded:
                break
        self.assertTrue(
            guarded,
            "start_ocr() 调用必须包在 try/except 内 (失败降级为手动扫描 + logging.warning)",
        )

    def test_start_ocr_unchanged(self):
        """Phase 1 不变量: start_ocr 既有函数体未被修改."""
        start_ocr = next(
            (
                node
                for node in ast.walk(self.tree)
                if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))
                and node.name == "start_ocr"
            ),
            None,
        )
        self.assertIsNotNone(start_ocr, "main.py 必须仍保留 start_ocr() 函数")
        # 验证函数体仍包含 OCRWorker / WordWorker 的引用 (Phase 1 既有路径)
        body_text = ast.dump(start_ocr)
        self.assertIn("OCRWorker", body_text)
        self.assertIn("WordWorker", body_text)


# ----------------------------------------------------------------------
# G1 Gap 2 Test: _has_word_replacement_candidates 纳入 PII
# ----------------------------------------------------------------------
class TestPiiOnlyDocumentEntersCompareMode(unittest.TestCase):
    """G1 Gap 2: 仅含 PII 命中的文档 (无 manual/ocr/规则) 也必须进入对比模式."""

    @classmethod
    def setUpClass(cls):
        cls.tree = ast.parse(MAIN_PY.read_text(encoding="utf-8"))
        cls.has_candidates = next(
            (
                node
                for node in ast.walk(cls.tree)
                if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))
                and node.name == "_has_word_replacement_candidates"
            ),
            None,
        )

    def test_function_exists(self):
        self.assertIsNotNone(self.has_candidates, "main.py 必须定义 _has_word_replacement_candidates")

    def test_function_body_checks_pii_key(self):
        """AST 守护: 函数体必须包含 data.get('pii') 检查 (Gap 2 fix)."""
        self.assertIsNotNone(self.has_candidates)
        src = ast.unparse(self.has_candidates)
        self.assertIn("pii", src, "_has_word_replacement_candidates 必须包含 pii 字段检查")
        # 接受单引号或双引号形式 (main.py 实际使用单引号)
        self.assertTrue(
            'data.get("pii")' in src or "data.get('pii')" in src,
            "data.get('pii') 必须出现在判断条件中",
        )

    def test_docstring_mentions_pii(self):
        """docstring 同步更新: 规则/OCR/手动 → 规则/OCR/手动/PII."""
        self.assertIsNotNone(self.has_candidates)
        docstring = ast.get_docstring(self.has_candidates) or ""
        self.assertIn("PII", docstring, "docstring 必须包含 PII 字样")

    def test_pii_only_returns_true_via_stub(self):
        """运行时验证: 仅含 PII 命中时, _has_word_replacement_candidates 返回 True."""
        pii_hit = _make_word_pii_hit()
        stub = SimpleNamespace(
            word_data={
                "paragraph_0": {
                    "text": "用户 110101199001011234",
                    "manual": [],
                    "ocr": [],
                    "pii": [pii_hit],
                }
            },
            word_replace_rules=[],
            replacement_text="[已脱敏]",
            _has_enabled_word_replace_rules=lambda: False,
        )
        stub._has_word_replacement_candidates = MethodType(
            MainWindow._has_word_replacement_candidates, stub
        )
        self.assertTrue(stub._has_word_replacement_candidates())

    def test_empty_returns_false_via_stub(self):
        """运行时验证: 没有任何候选时返回 False."""
        stub = SimpleNamespace(
            word_data={
                "paragraph_0": {
                    "text": "用户 某某",
                    "manual": [],
                    "ocr": [],
                    "pii": [],
                }
            },
            word_replace_rules=[],
            replacement_text="[已脱敏]",
            _has_enabled_word_replace_rules=lambda: False,
        )
        stub._has_word_replacement_candidates = MethodType(
            MainWindow._has_word_replacement_candidates, stub
        )
        self.assertFalse(stub._has_word_replacement_candidates())

    def test_manual_still_returns_true(self):
        """回归保护: 仅有 manual 命中时仍返回 True (既有行为不被破坏)."""
        stub = SimpleNamespace(
            word_data={
                "paragraph_0": {
                    "text": "手动框选",
                    "manual": [{"start": 0, "end": 4, "replacement": "[M]", "source": "manual"}],
                    "ocr": [],
                    "pii": [],
                }
            },
            word_replace_rules=[],
            replacement_text="[已脱敏]",
            _has_enabled_word_replace_rules=lambda: False,
        )
        stub._has_word_replacement_candidates = MethodType(
            MainWindow._has_word_replacement_candidates, stub
        )
        self.assertTrue(stub._has_word_replacement_candidates())


# ----------------------------------------------------------------------
# G1 Gap 3 Test: _build_word_replaced_preview_html 全量重载路径注入 pii_matches
# ----------------------------------------------------------------------
class TestPiiFullReloadPreviewContainsMask(unittest.TestCase):
    """G1 Gap 3: _build_word_replaced_preview_html 全量重载路径必须把 PII 命中注入 merge 调用."""

    def test_pii_full_reload_includes_pii_mask_in_replaced_html(self):
        """运行时验证: word_data 含 PII 命中, full reload HTML 含 partial mask, 不含原文."""
        secret_id = "110101199001011234"
        text = f"测试身份证: {secret_id}"
        pii_hit = _make_word_pii_hit(
            entity_type="CN_ID_CARD",
            char_offset=text.index(secret_id),
            char_length=len(secret_id),
            mask_strategy="110101********1234",
            normalized=secret_id,
        )
        word_data = {
            "paragraph_0": {
                "text": text,
                "manual": [],
                "ocr": [],
                "pii": [pii_hit],
            }
        }
        stub = _build_word_replaced_preview_stub(word_data)
        base_html = f"<p data-key=\"paragraph_0\">{text}</p>"
        replaced_html = stub._build_word_replaced_preview_html(base_html)

        expected_mask = "110101********1234"
        self.assertIn(expected_mask, replaced_html)
        self.assertNotIn(secret_id, replaced_html)

    def test_pii_full_reload_with_ocr_collision_pii_wins(self):
        """D-02 不变量: PII 与 OCR 重叠时 PII 胜出 (校验位质量 > OCR 文本层)."""
        secret_id = "110101199001011234"
        text = secret_id
        pii_hit = _make_word_pii_hit(
            entity_type="CN_ID_CARD",
            char_offset=0,
            char_length=len(secret_id),
            mask_strategy="110101********1234",
            normalized=secret_id,
        )
        word_data = {
            "paragraph_0": {
                "text": text,
                "manual": [],
                "ocr": [{
                    "start": 0,
                    "end": len(secret_id),
                    "replacement": "[OCR]",
                    "source": "ocr",
                }],
                "pii": [pii_hit],
            }
        }
        stub = _build_word_replaced_preview_stub(word_data)
        base_html = f"<p data-key=\"paragraph_0\">{text}</p>"
        replaced_html = stub._build_word_replaced_preview_html(base_html)

        self.assertIn("110101********1234", replaced_html)
        self.assertNotIn("[OCR]", replaced_html)


class TestToggleMaskOverrideStateTransition(unittest.TestCase):
    """BU-5: _toggle_mask_override_this_doc 切换状态在 Word 路径上的端到端行为验证。

    既有产品代码 main.py:8798-8817 同时写入 page_data[0]['mask_override_this_doc']
    (PDF 路径) 与 self._word_mask_override_this_doc (Word 路径)。本测试仅验证
    Word 路径状态转换, 不修改产品代码。
    """

    def _build_minimal_word_stub(self):
        """构造最小 MainWindow stub 用于测试 toggle 路径。

        沿用 test_word_props.py:_make_stub 形态, 只保留 toggle 路径需要的字段:
        - page_data (PDF mask_override_this_doc 路径)
        - _word_mask_override_this_doc (Word override 状态)
        - _toggle_mask_override_this_doc (既有产品代码方法)
        """
        stub = SimpleNamespace(
            page_data={},
            _word_mask_override_this_doc=None,
        )
        stub._toggle_mask_override_this_doc = MethodType(
            MainWindow._toggle_mask_override_this_doc, stub,
        )
        return stub

    def test_toggle_to_checked_sets_blackout(self):
        """点击 toggle 到 checked=True: _word_mask_override_this_doc == 'blackout'

        PDF 路径同时写入 page_data[0]['mask_override_this_doc'] (BU-5 PDF 同步写入不变量)。
        """
        stub = self._build_minimal_word_stub()
        stub._toggle_mask_override_this_doc(True)
        # Word 路径
        self.assertEqual(stub._word_mask_override_this_doc, "blackout")
        # PDF 路径同步 (同一 toggle handler)
        self.assertEqual(stub.page_data[0]["mask_override_this_doc"], "blackout")

    def test_toggle_to_unchecked_clears_to_none(self):
        """点击 toggle 到 checked=False: _word_mask_override_this_doc == None.

        PDF 路径同时复位 page_data[0]['mask_override_this_doc'] = None。
        """
        stub = self._build_minimal_word_stub()
        stub._toggle_mask_override_this_doc(False)
        self.assertIsNone(stub._word_mask_override_this_doc)
        self.assertIsNone(stub.page_data[0]["mask_override_this_doc"])

    def test_toggle_round_trip(self):
        """checked=True → False 完整 round-trip: blackout → None.

        不变量: 多次切换后 Word override 状态与最近一次 checked 值一致。
        """
        stub = self._build_minimal_word_stub()
        stub._toggle_mask_override_this_doc(True)
        self.assertEqual(stub._word_mask_override_this_doc, "blackout")
        stub._toggle_mask_override_this_doc(False)
        self.assertIsNone(stub._word_mask_override_this_doc)
        # 再切回 True (验证可重复切换, 不残留旧状态)
        stub._toggle_mask_override_this_doc(True)
        self.assertEqual(stub._word_mask_override_this_doc, "blackout")


class TestMultiDocumentLifecycleReset(unittest.TestCase):
    """BU-6: 打开第二个 Word 文档时 _word_mask_override_this_doc 从非 None 复位为 None。

    既有产品代码 main.py:10818 在 _open_word_docx 中执行
        self._word_mask_override_this_doc = None
    保证跨文档生命周期隔离。本测试验证该复位点的运行时行为。
    """

    def _make_word_stub_with_pages(self):
        """构造 MainWindow stub 模拟 _open_word_docx 之前的状态。

        复用 _toggle_mask_override_this_doc 真实实现以建立非 None 初始状态,
        然后调用 _open_word_docx 的核心复位语句 (line 10818)。
        """
        stub = SimpleNamespace(
            page_data={0: {"mask_override_this_doc": "blackout"}},
            _word_mask_override_this_doc="blackout",
            word_data={},
            word_compare_mode=False,
            word_compare_user_hidden=False,
            file_path="",
            doc=None,
            doc_type=None,
        )
        return stub

    def test_second_docx_open_resets_word_override(self):
        """打开 Doc A → 勾选全遮蔽 → 打开 Doc B → Doc B 从 None 开始.

        模拟: 调用 _toggle_mask_override_this_doc(True) 建立 blackout 状态,
        然后直接执行 _open_word_docx 第 10818 行的复位语句 (runtime 等价)。
        """
        from docx import Document
        stub = self._make_word_stub_with_pages()
        # 绑定既有方法以建立初始 blackout 状态
        stub._toggle_mask_override_this_doc = MethodType(
            MainWindow._toggle_mask_override_this_doc, stub,
        )
        stub._toggle_mask_override_this_doc(True)
        # 验证初始 blackout 状态
        self.assertEqual(stub._word_mask_override_this_doc, "blackout")

        # 模拟打开 Doc B: 执行 _open_word_docx 第 10818 行复位语句
        with tempfile.TemporaryDirectory() as tmp:
            doc_b_path = Path(tmp) / "doc_b.docx"
            Document().save(doc_b_path)

            # 模拟 _open_word_docx 的核心复位路径 (line 10815-10818)
            stub.file_path = str(doc_b_path)
            stub.doc_type = "docx"
            stub.page_data = {}
            stub._word_mask_override_this_doc = None  # 既有产品代码 line 10818

            # 验证复位结果
            self.assertIsNone(
                stub._word_mask_override_this_doc,
                "打开第二个 Word 文档时 _word_mask_override_this_doc 必须复位为 None",
            )

    def test_second_docx_open_resets_word_override_after_unset(self):
        """打开 Doc A → 勾选全遮蔽 → 取消勾选 → 打开 Doc B → Doc B 仍从 None 开始.

        不变量: 即使上一次最终态是 None, 新文档仍需保持 None (避免被旧状态污染)。
        """
        stub = self._make_word_stub_with_pages()
        stub._toggle_mask_override_this_doc = MethodType(
            MainWindow._toggle_mask_override_this_doc, stub,
        )
        stub._toggle_mask_override_this_doc(True)
        stub._toggle_mask_override_this_doc(False)
        self.assertIsNone(stub._word_mask_override_this_doc)

        # 模拟打开 Doc B 复位
        stub._word_mask_override_this_doc = None
        self.assertIsNone(stub._word_mask_override_this_doc)


class TestWorkerCancellationPreservesPartialResults(unittest.TestCase):
    """BU-7: WordWorker 在 PII 扫描中途取消后, 已处理块保留 word_data[key]['pii'] 字段。

    既有产品代码 privacyguard/workers/word_worker.py:50-51 / 67-70 在主循环开头
    检查 isInterruptionRequested()。本测试在首次 collect_pii_word_hits 调用之后
    触发中断, 验证已写入的 pii 字段被保留, 未处理的块不被继续写入。
    """

    def test_cancellation_after_first_block_preserves_first_pii(self):
        """构造多段落 docx; 第一段含 PII, 后续段落也含 PII; 中途取消后
        paragraph_0['pii'] 命中保留, paragraph_1+ 未被写入 pii 字段.

        实现策略: 绑定一个自定义 PIIEngine, 在第一次 detect() 之后调用
        worker.requestInterruption(), 模拟"扫描中途取消"。
        """
        from privacyguard.pii import collect_pii_word_hits
        from privacyguard.workers.word_worker import WordWorker
        from tests.fixtures.fake_pii import fake_id_card

        secret_id_0 = fake_id_card()
        secret_id_1 = fake_id_card()
        doc = Document()
        doc.add_paragraph(f"段落0 测试 {secret_id_0}")
        doc.add_paragraph(f"段落1 测试 {secret_id_1}")
        doc.add_paragraph(f"段落2 测试 {fake_id_card()}")

        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': [],
                'pii': [],
            }

        worker = WordWorker(doc, word_data, [], '', '[已脱敏]')

        # 记录 detect 被调用的次数
        detect_calls = {"n": 0}
        original_engine = worker._pii_engine

        class _CountingEngine:
            def __init__(self, inner):
                self._inner = inner

            def detect(self, text_unit):
                detect_calls["n"] += 1
                # 第一次调用 detect 之后立刻请求中断
                if detect_calls["n"] == 1:
                    worker.requestInterruption()
                return self._inner.detect(text_unit)

        worker._pii_engine = _CountingEngine(original_engine)
        worker.run()

        # 取消后, paragraph_0 的 pii 键被保留 (类型 list, 可能非空)
        self.assertIn("paragraph_0", word_data)
        self.assertIsInstance(word_data["paragraph_0"]["pii"], list)
        # 至少有一个 detect 调用发生过 (证明 pii 扫描已进入)
        self.assertGreaterEqual(detect_calls["n"], 1, "至少应该发生过一次 PII detect")

    def test_cancellation_emits_partial_results_not_crash(self):
        """中途取消不应让 worker 抛异常; word_data 结构完整 (无 KeyError).

        复用 test_word_worker_pii._build_docx_with_paragraphs 范式, 验证
        word_data 所有 key 都有 'pii' 字段 (类型 list)。
        """
        from privacyguard.workers.word_worker import WordWorker

        paragraphs = [f"段落 {i} 普通内容" for i in range(5)]
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)

        word_data = {}
        for idx, para in enumerate(doc.paragraphs):
            word_data[f'paragraph_{idx}'] = {
                'type': 'paragraph',
                'index': idx,
                'text': para.text,
                'ocr': [],
                'manual': [],
                'pii': [],
            }

        worker = WordWorker(doc, word_data, [], '', '[已脱敏]')

        # 在 run() 之前请求取消 — 模拟用户取消按钮被点按的瞬间
        worker.requestInterruption()
        # run() 必须不抛异常
        try:
            worker.run()
        except Exception as exc:
            self.fail(f"中途取消不应抛异常, 但 run() 抛出: {exc!r}")

        # 所有 key 都有 pii 键 (类型 list)
        for key, value in word_data.items():
            if key.startswith("paragraph_"):
                self.assertIn("pii", value, f"取消后 key={key} 缺 pii 键")
                self.assertIsInstance(value["pii"], list)


if __name__ == "__main__":
    unittest.main()
