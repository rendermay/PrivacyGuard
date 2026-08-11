"""Phase 3 (03-word) — word_adapter 三函数 + 懒加载纪律测试 (D-11 / D-13 / OPS-03).

锁:
- `collect_pii_word_hits(paragraph_text, engine)` 复用 PIIEngine.detect (D-11 引擎层无 IO)
- `locate_pii_hits_in_paragraph(hits, paragraph_text)` 用 paragraph_text.find(needle, start)
  顺序扫描, 同文本重复逐个展开 (D-08 + D-09)
- `apply_pii_replacements_to_docx(doc, hit_locations, mode='partial'|'blackout')` 不
  引用 docx top-level namespace; 调用方持有 Document 句柄, partial 走 mask_for_entity,
  blackout 写 "[已脱敏]". 段级 paragraph.style.name 在 replace 前后保持不变 (D-07)
- 三函数经 privacyguard.pii._LAZY_IMPORTS 注册; 通过 privacyguard.pii 顶层访问
  触发 word_adapter 模块加载 (D-13 + OPS-03)
- word_adapter.py 源码不含 `import docx` / `from docx import` 子串 (D-11 + T-03-02)
"""
import importlib
import os
import sys
import tempfile
import unittest

from privacyguard.pii.engine import PIIEngine
from privacyguard.pii.word_adapter import (
    apply_pii_replacements_to_docx,
    collect_pii_word_hits,
    locate_pii_hits_in_paragraph,
)
from tests.fixtures.fake_pii import fake_email, fake_id_card, fake_phone


# ----------------------------------------------------------------------
# Test 1: collect_pii_word_hits
# ----------------------------------------------------------------------
class TestCollectPiiWordHits(unittest.TestCase):
    """D-11: collect_pii_word_hits(paragraph_text, engine) -> List[PIIHit]
    复用 PIIEngine.detect 识别单段文本。
    """

    def test_empty_text_returns_empty_list(self):
        hits = collect_pii_word_hits("", PIIEngine())
        self.assertEqual(hits, [])

    def test_whitespace_only_text_returns_empty_list(self):
        hits = collect_pii_word_hits("   \n\t  ", PIIEngine())
        self.assertEqual(hits, [])

    def test_detects_id_card_in_paragraph(self):
        secret_id = fake_id_card()
        text = f"测试 身份证 {secret_id} 后面"
        hits = collect_pii_word_hits(text, PIIEngine())
        self.assertGreater(len(hits), 0)
        entity_types = {h.entity_type for h in hits}
        self.assertIn("CN_ID_CARD", entity_types)

    def test_detects_multiple_entity_types(self):
        secret_id = fake_id_card()
        secret_phone = fake_phone()
        secret_email = fake_email()
        text = (
            f"身份证 {secret_id} "
            f"电话 {secret_phone} "
            f"邮箱 {secret_email}"
        )
        hits = collect_pii_word_hits(text, PIIEngine())
        entity_types = {h.entity_type for h in hits}
        self.assertIn("CN_ID_CARD", entity_types)
        self.assertIn("CN_PHONE", entity_types)
        self.assertIn(
            "CN_EMAIL", entity_types,
            f"hits 应包含至少一条 CN_EMAIL; actual entities: {entity_types}",
        )

    def test_returns_pii_hit_with_5_locked_fields(self):
        """D-05: PIIHit 字段锁; 返回的每个 hit 必须含 5 个必填字段。"""
        secret_id = fake_id_card()
        text = f"测试 {secret_id} 文本"
        hits = collect_pii_word_hits(text, PIIEngine())
        self.assertGreater(len(hits), 0)
        hit = hits[0]
        for field in (
            "entity_type", "page_offset", "page_length",
            "page_rect", "confidence_tier",
        ):
            self.assertTrue(
                hasattr(hit, field),
                f"PIIHit 缺失字段 {field!r} (D-05 字段锁违反)",
            )


# ----------------------------------------------------------------------
# Test 2: locate_pii_hits_in_paragraph
# ----------------------------------------------------------------------
class TestLocatePiiHitsInParagraph(unittest.TestCase):
    """D-08 + D-09: locate_pii_hits_in_paragraph(hits, paragraph_text) -> [(PIIHit, offset)]
    paragraph_text.find(needle, start_offset) 顺序扫描; 同文本重复逐个展开。
    """

    def _make_hit(self, text):
        """构造一个 PIIHit (用于 locate 测试, 不依赖收集路径)."""
        from privacyguard.pii.hits import PIIHit
        return PIIHit(
            entity_type="CN_ID_CARD",
            page_offset=0,
            page_length=len(text),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy=text[:6] + "*" * 8 + text[14:] if len(text) == 18 else "*" * len(text),
            normalized=text,
            validator_passed=True,
        )

    def test_locate_single_hit_returns_offset(self):
        """单一命中在 paragraph_text 中返回 (hit, offset) 偏移. offset 与 text.find 一致。"""
        text = "前面文本 110101199001011234 后面文本"
        needle = "110101199001011234"
        hit = self._make_hit(needle)
        locations = locate_pii_hits_in_paragraph([hit], text)
        self.assertEqual(len(locations), 1)
        self.assertEqual(locations[0][1], text.find(needle))

    def test_locate_duplicate_text_expands_all_occurrences(self):
        """D-09: 同文本重复逐个展开为多个 (hit, offset) 元组. 偏移严格递增。"""
        text = "110101199001011234 中间 110101199001011234 末尾"
        needle = "110101199001011234"
        hit = self._make_hit(needle)
        locations = locate_pii_hits_in_paragraph([hit], text)
        self.assertEqual(len(locations), 2)
        offsets = [off for _, off in locations]
        self.assertEqual(offsets[0], 0)
        self.assertEqual(offsets[1], text.find(needle, 1))
        # 严格递增
        self.assertLess(offsets[0], offsets[1])

    def test_locate_three_duplicate_text_returns_three(self):
        """D-09: 同文本重复 3 次展开为 3 个独立 (hit, offset)。"""
        needle = "110101"
        # 构造确定偏移: "aa"(2) + needle(6) + "bb"(2) + needle(6) + "cc"(2) + needle(6) + "dd"
        # offsets: 2, 10, 18
        text = "aa" + needle + "bb" + needle + "cc" + needle + "dd"
        hit = self._make_hit(needle)
        locations = locate_pii_hits_in_paragraph([hit], text)
        self.assertEqual(len(locations), 3)
        offsets = [off for _, off in locations]
        self.assertEqual(offsets, [2, 10, 18])
        # 严格递增
        self.assertEqual(offsets, sorted(offsets))

    def test_locate_empty_hits_returns_empty(self):
        text = "some paragraph text"
        self.assertEqual(locate_pii_hits_in_paragraph([], text), [])

    def test_locate_empty_paragraph_text_returns_empty(self):
        hit = self._make_hit("110101")
        self.assertEqual(locate_pii_hits_in_paragraph([hit], ""), [])

    def test_locate_not_found_returns_empty_for_that_hit(self):
        """needle 不在 paragraph_text 中 → 不入 locations (但不抛错)."""
        text = "some paragraph without needle"
        hit = self._make_hit("110101199001011234")
        self.assertEqual(locate_pii_hits_in_paragraph([hit], text), [])


# ----------------------------------------------------------------------
# Test 3: apply_pii_replacements_to_docx
# ----------------------------------------------------------------------
class TestApplyPiiReplacementsToDocx(unittest.TestCase):
    """D-06 + D-07: apply_pii_replacements_to_docx(doc, hit_locations, mode)
    partial 走 mask_for_entity, blackout 写 "[已脱敏]"; 段级 style.name 在 replace 前后不变.
    """

    def _make_hit(self, entity_type, normalized, hit_text):
        from privacyguard.pii.hits import PIIHit
        return PIIHit(
            entity_type=entity_type,
            page_offset=0,
            page_length=len(normalized),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy="(unused)",
            normalized=normalized,
            validator_passed=True,
        )

    def _make_doc_with_paragraph(self, text):
        from docx import Document
        doc = Document()
        doc.add_paragraph(text)
        return doc

    def test_redacted_paragraph_loses_original_secret(self):
        """partial mode 真脱敏: 段内不再含原始敏感字符串。"""
        secret = fake_id_card()
        doc = self._make_doc_with_paragraph(f"测试样本 身份证 {secret} 后面文本")
        text = doc.paragraphs[0].text
        offset = text.find(secret)
        from privacyguard.pii.hits import PIIHit
        hit = PIIHit(
            entity_type="CN_ID_CARD",
            page_offset=offset,
            page_length=len(secret),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy="(unused)",
            normalized=secret,
            validator_passed=True,
        )
        apply_pii_replacements_to_docx(
            doc, {"paragraph_0": [(hit, offset)]}, mode="partial",
        )
        new_text = doc.paragraphs[0].text
        self.assertNotIn(secret, new_text)

    def test_partial_mode_writes_mask_text(self):
        """partial mode 产物含 mask 字符串 ('********')."""
        secret = fake_id_card()
        doc = self._make_doc_with_paragraph(f"身份证 {secret}")
        text = doc.paragraphs[0].text
        offset = text.find(secret)
        from privacyguard.pii.hits import PIIHit
        hit = PIIHit(
            entity_type="CN_ID_CARD",
            page_offset=offset,
            page_length=len(secret),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy="(unused)",
            normalized=secret,
            validator_passed=True,
        )
        apply_pii_replacements_to_docx(
            doc, {"paragraph_0": [(hit, offset)]}, mode="partial",
        )
        new_text = doc.paragraphs[0].text
        from privacyguard.pii.mask import mask_for_entity
        expected_mask = mask_for_entity("CN_ID_CARD", secret)
        self.assertIn(expected_mask, new_text)

    def test_blackout_mode_writes_brackets(self):
        """blackout mode 产物含 '[已脱敏]'. 与 partial mode 区分。"""
        secret = fake_phone()
        doc = self._make_doc_with_paragraph(f"电话 {secret}")
        text = doc.paragraphs[0].text
        offset = text.find(secret)
        from privacyguard.pii.hits import PIIHit
        hit = PIIHit(
            entity_type="CN_PHONE",
            page_offset=offset,
            page_length=len(secret),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy="(unused)",
            normalized=secret,
            validator_passed=True,
        )
        apply_pii_replacements_to_docx(
            doc, {"paragraph_0": [(hit, offset)]}, mode="blackout",
        )
        new_text = doc.paragraphs[0].text
        self.assertIn("[已脱敏]", new_text)
        self.assertNotIn(secret, new_text)

    def test_paragraph_style_preserved_after_replace(self):
        """D-07: 段级 paragraph.style.name 在 replace 前后不变。"""
        secret = fake_id_card()
        from docx import Document
        doc = Document()
        p = doc.add_paragraph(f"测试 {secret} 文本")
        original_style = p.style.name
        text = p.text
        offset = text.find(secret)
        from privacyguard.pii.hits import PIIHit
        hit = PIIHit(
            entity_type="CN_ID_CARD",
            page_offset=offset,
            page_length=len(secret),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy="(unused)",
            normalized=secret,
            validator_passed=True,
        )
        apply_pii_replacements_to_docx(
            doc, {"paragraph_0": [(hit, offset)]}, mode="partial",
        )
        self.assertEqual(doc.paragraphs[0].style.name, original_style)

    def test_replace_across_runs_replaces_full_substring(self):
        """跨 run 命中 (run0='110101' + run1='199001011234') replace 后段内不再含完整子串.

        注: 测试用字面字符串 split 到多个 run, 验证 D-07: 段合并 run + 区间替换.
        D-07 允许段级样式保留 + run 级格式丢失; 这里仅断言子串消失.
        """
        secret = fake_id_card()
        from docx import Document
        doc = Document()
        # 先构造含 secret 的段落, 然后拆 run
        p = doc.add_paragraph(f"前缀 {secret} 后缀")
        # 拆 run: 用文档元素 API 把段落拆成多个 run
        # 简单方法: 找出 secret 起止位置, 重新创建 run
        text = p.text
        offset = text.find(secret)
        # 拆为: prefix_run + secret_run + suffix_run
        prefix_text = text[:offset]
        suffix_text = text[offset + len(secret):]
        # 清空段落并重建 run
        # python-docx 没有 run 拆分 API; 用 _element.clear_content() + 重新 add_run
        p._element.clear_content()  # noqa: SLF001 (内部用法)
        p.add_run(prefix_text)
        p.add_run(secret)
        p.add_run(suffix_text)
        text = p.text
        offset = text.find(secret)
        from privacyguard.pii.hits import PIIHit
        hit = PIIHit(
            entity_type="CN_ID_CARD",
            page_offset=offset,
            page_length=len(secret),
            page_rect=(0.0, 0.0, 0.0, 0.0),
            confidence_tier="HIGH",
            source="text",
            mask_strategy="(unused)",
            normalized=secret,
            validator_passed=True,
        )
        apply_pii_replacements_to_docx(
            doc, {"paragraph_0": [(hit, offset)]}, mode="partial",
        )
        new_text = doc.paragraphs[0].text
        self.assertNotIn(secret, new_text)

    def test_no_hits_for_key_skips_paragraph(self):
        """hit_locations 中无对应 key 时, paragraph 原样保留。"""
        from docx import Document
        doc = Document()
        original_text = "原文 110101199001011234 文本"
        doc.add_paragraph(original_text)
        apply_pii_replacements_to_docx(doc, {"paragraph_NONE": []}, mode="partial")
        self.assertEqual(doc.paragraphs[0].text, original_text)

    def test_round_trip_through_disk_does_not_leak_secret(self):
        """完整 reverse-extraction: 写入磁盘 → 重新读取 → 断言敏感字符串不存在。"""
        secret = fake_id_card()
        from docx import Document
        doc = Document()
        doc.add_paragraph(f"前缀 {secret} 后缀")
        with tempfile.TemporaryDirectory() as tmp:
            in_path = os.path.join(tmp, "in.docx")
            doc.save(in_path)
            doc2 = Document(in_path)
            text = doc2.paragraphs[0].text
            offset = text.find(secret)
            from privacyguard.pii.hits import PIIHit
            hit = PIIHit(
                entity_type="CN_ID_CARD",
                page_offset=offset,
                page_length=len(secret),
                page_rect=(0.0, 0.0, 0.0, 0.0),
                confidence_tier="HIGH",
                source="text",
                mask_strategy="(unused)",
                normalized=secret,
                validator_passed=True,
            )
            apply_pii_replacements_to_docx(
                doc2, {"paragraph_0": [(hit, offset)]}, mode="partial",
            )
            out_path = os.path.join(tmp, "out.docx")
            doc2.save(out_path)
            doc3 = Document(out_path)
            out_text = doc3.paragraphs[0].text
            self.assertNotIn(secret, out_text)


# ----------------------------------------------------------------------
# Test 4: Importability — 懒加载纪律 (OPS-03 / D-13)
# ----------------------------------------------------------------------
class TestWordAdapterImportability(unittest.TestCase):
    """OPS-03 懒加载纪律:

    - 通过 `privacyguard.pii` 顶层访问三函数触发 word_adapter 模块加载 (D-13)
    - 通过 `privacyguard` 顶层导入 (无 PII 触发) 不应触发 word_adapter 加载
    - word_adapter.py 源码不含 `from docx import` / `import docx` 子串 (D-11 + T-03-02)
    """

    def _snapshot_privacyguard_modules(self):
        return {
            name: module
            for name, module in list(sys.modules.items())
            if name == "privacyguard" or name.startswith("privacyguard.")
        }

    def _restore_privacyguard_modules(self, cached):
        for name in list(sys.modules):
            if name == "privacyguard" or name.startswith("privacyguard."):
                sys.modules.pop(name, None)
        sys.modules.update(cached)

    def test_collect_pii_word_hits_triggers_word_adapter_loaded(self):
        """D-13: privacyguard.pii.collect_pii_word_hits 触发 word_adapter 模块加载。"""
        cached = self._snapshot_privacyguard_modules()
        for name in list(cached):
            sys.modules.pop(name, None)
        try:
            module = importlib.import_module("privacyguard")
            self.assertNotIn(
                "privacyguard.pii.word_adapter", sys.modules,
                "import privacyguard 不应触发 privacyguard.pii.word_adapter",
            )
            # 触发已有导出 (验证 import privacyguard 不拉起 PII 子模块)
            _ = module.validate_safe_path
            pii_pkg = importlib.import_module("privacyguard.pii")
            _ = pii_pkg.collect_pii_word_hits
            self.assertIn(
                "privacyguard.pii.word_adapter", sys.modules,
                "通过 privacyguard.pii.collect_pii_word_hits 应触发 word_adapter 加载 (D-13)",
            )
        finally:
            self._restore_privacyguard_modules(cached)

    def test_locate_pii_hits_in_paragraph_triggers_word_adapter_loaded(self):
        """D-13: privacyguard.pii.locate_pii_hits_in_paragraph 触发 word_adapter 懒加载。"""
        cached = self._snapshot_privacyguard_modules()
        for name in list(cached):
            sys.modules.pop(name, None)
        try:
            module = importlib.import_module("privacyguard")
            _ = module.validate_safe_path
            pii_pkg = importlib.import_module("privacyguard.pii")
            self.assertNotIn("privacyguard.pii.word_adapter", sys.modules)
            _ = pii_pkg.locate_pii_hits_in_paragraph
            self.assertIn(
                "privacyguard.pii.word_adapter", sys.modules,
                "通过 privacyguard.pii.locate_pii_hits_in_paragraph 应触发 word_adapter 加载 (D-13)",
            )
        finally:
            self._restore_privacyguard_modules(cached)

    def test_apply_pii_replacements_to_docx_triggers_word_adapter_loaded(self):
        """D-13: privacyguard.pii.apply_pii_replacements_to_docx 触发 word_adapter 懒加载。"""
        cached = self._snapshot_privacyguard_modules()
        for name in list(cached):
            sys.modules.pop(name, None)
        try:
            module = importlib.import_module("privacyguard")
            _ = module.validate_safe_path
            pii_pkg = importlib.import_module("privacyguard.pii")
            self.assertNotIn("privacyguard.pii.word_adapter", sys.modules)
            _ = pii_pkg.apply_pii_replacements_to_docx
            self.assertIn(
                "privacyguard.pii.word_adapter", sys.modules,
                "通过 privacyguard.pii.apply_pii_replacements_to_docx 应触发 word_adapter 加载 (D-13)",
            )
        finally:
            self._restore_privacyguard_modules(cached)

    def test_word_adapter_source_does_not_import_docx(self):
        """D-11 + T-03-02: word_adapter.py 源码不含 'from docx' / 'import docx'.

        仅扫描非 docstring / 注释行, 防止 docstring 中提及 'python-docx' 误触发。
        """
        from pathlib import Path
        word_adapter_path = (
            Path(__file__).resolve().parents[2]
            / "privacyguard" / "pii" / "word_adapter.py"
        )
        if not word_adapter_path.exists():
            self.skipTest(
                "privacyguard/pii/word_adapter.py 不存在; Task 2 GREEN 后此检查生效"
            )
        # 解析 AST, 排除 docstring 行
        import ast
        source = word_adapter_path.read_text(encoding="utf-8")
        tree = ast.parse(source)
        # 取每个顶级语句的代码行 (排除 Module docstring / Expression)
        import_lines = []
        for node in tree.body:
            if isinstance(node, ast.Import):
                for alias in node.names:
                    if alias.name == "docx" or alias.name.startswith("docx."):
                        import_lines.append(node.lineno)
            elif isinstance(node, ast.ImportFrom):
                if node.module and (node.module == "docx" or node.module.startswith("docx.")):
                    import_lines.append(node.lineno)
        self.assertEqual(
            import_lines, [],
            f"word_adapter.py 顶部包含 docx import (D-11): 行号 {import_lines}",
        )


if __name__ == "__main__":
    unittest.main()
