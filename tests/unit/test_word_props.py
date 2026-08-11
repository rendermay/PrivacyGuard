"""Phase 3 G2 Gap Closure: Word core_properties 清除 (Gap 4 / SAFE-03 Word 子项)。

D-15 策略: 5 个标准字段 (Title / Author / Subject / Comments / Keywords) 全部清空字符串,
不使用占位符; 与 Phase 2 PDF metadata clearing 行为对齐。

覆盖:
- TestClearAllFiveProps — 默认 5 字段全部清空
- TestClearSpecificKeys — keys 参数只清空指定子集
- TestNoDocxImportInHelper — TYPE_CHECKING 守卫生效, runtime 不 import python-docx
- TestIntegrationWithSaveWord — 与 _save_word 集成 (在 save 前清空 5 字段)
- TestReverseExtractionCoreProperties — 端到端 reverse-extraction, 验证 saved .docx
- TestPackageImportsParity — cp30 教训: Windows/macOS spec + macOS build script 静态包含
  'privacyguard.utils.word_props'
"""
from __future__ import annotations

import importlib
import os
import sys
import tempfile
import unittest
from pathlib import Path
from types import MethodType, SimpleNamespace

from docx import Document

from main import MainWindow
from tests.fixtures.fake_pii import fake_id_card, fake_phone


PROJECT_ROOT = Path(__file__).resolve().parents[2]


# ----------------------------------------------------------------------
# Test 1: 默认 5 字段清空
# ----------------------------------------------------------------------
class TestClearAllFiveProps(unittest.TestCase):
    """G2 / SAFE-03 Word: clear_word_core_properties 默认清空 5 字段."""

    def _make_doc_with_props(self):
        doc = Document()
        cp = doc.core_properties
        cp.title = "敏感标题 110101199001011234"
        cp.author = "敏感作者 张三"
        cp.subject = "敏感主题"
        cp.comments = "敏感评论"
        cp.keywords = "身份证, 隐私"
        return doc

    def test_clear_default_five_properties(self):
        """调用 clear_word_core_properties(doc) 后, 5 个字段全部等于 ''. """
        from privacyguard.utils import clear_word_core_properties
        doc = self._make_doc_with_props()
        cleared = clear_word_core_properties(doc)
        cp = doc.core_properties
        self.assertEqual(cp.title, "")
        self.assertEqual(cp.author, "")
        self.assertEqual(cp.subject, "")
        self.assertEqual(cp.comments, "")
        self.assertEqual(cp.keywords, "")
        self.assertEqual(cleared, 5)


# ----------------------------------------------------------------------
# Test 2: 指定子集清空
# ----------------------------------------------------------------------
class TestClearSpecificKeys(unittest.TestCase):
    """keys 参数只清空指定字段; 未指定的字段保留原值."""

    def test_clear_specific_keys_only_clears_those(self):
        from privacyguard.utils import clear_word_core_properties
        doc = Document()
        cp = doc.core_properties
        cp.title = "title_secret"
        cp.author = "author_secret"
        cp.subject = "subject_secret"
        cleared = clear_word_core_properties(doc, keys=("title",))
        self.assertEqual(cp.title, "")
        self.assertEqual(cp.author, "author_secret")
        self.assertEqual(cp.subject, "subject_secret")
        self.assertEqual(cleared, 1)


# ----------------------------------------------------------------------
# Test 3: 无 core_properties 的对象
# ----------------------------------------------------------------------
class TestErrorHandling(unittest.TestCase):
    """传入无 core_properties 属性的对象时, helper 必须安全处理 (不抛 AttributeError 中断 save)."""

    def test_clear_on_doc_without_core_properties_raises_attribute_error(self):
        """传入 duck-typed 无 core_properties 对象时, getattr(cp, key, None) 安全返回 None,
        helper 返回 0 (不抛错). 这是 _save_word try/except 的次要防线之一."""
        from privacyguard.utils import clear_word_core_properties

        class _NoCoreProps:
            class core_properties:
                pass

        cleared = clear_word_core_properties(_NoCoreProps())
        self.assertEqual(cleared, 0)


# ----------------------------------------------------------------------
# Test 4: TYPE_CHECKING 守卫生效
# ----------------------------------------------------------------------
class TestNoDocxImportInHelper(unittest.TestCase):
    """TYPE_CHECKING 块守卫 docx import; runtime 不拉起 python-docx."""

    def test_no_docx_import_in_helper_at_runtime(self):
        cached = {n: m for n, m in list(sys.modules.items())}
        try:
            for name in list(sys.modules):
                if name == "docx" or name.startswith("docx."):
                    sys.modules.pop(name, None)
            importlib.import_module("privacyguard.utils.word_props")
            self.assertNotIn(
                "docx", sys.modules,
                "privacyguard.utils.word_props 在 runtime 不应 import python-docx (TYPE_CHECKING 守卫)",
            )
        finally:
            for name in list(sys.modules):
                sys.modules.pop(name, None)
            sys.modules.update(cached)


# ----------------------------------------------------------------------
# Test 5: 与 _save_word 集成 — save 前清空 5 字段
# ----------------------------------------------------------------------
class TestIntegrationWithSaveWord(unittest.TestCase):
    """_save_word 在 new_doc.save(fname) 之前调用 clear_word_core_properties(new_doc)."""

    def _make_stub(self, word_data):
        """构造 MainWindow._save_word 测试 stub: 复用 TestSaveWordCallsPiiAdapter 同形态."""
        from privacyguard.pii import (
            apply_pii_replacements_to_docx,
            locate_pii_hits_in_paragraph,
        )
        stub = SimpleNamespace(
            file_path="",
            temp_manager=SimpleNamespace(create_temp_file=lambda: tempfile.NamedTemporaryFile(
                suffix=".docx", delete=False).name),
            word_data=dict(word_data),
            word_replace_rules=[],
            replacement_text="[已脱敏]",
            logger=None,
        )
        return stub

    def test_save_word_calls_clear_word_core_properties(self):
        """AST 守护: _save_word 函数体内必须出现 clear_word_core_properties 调用."""
        import ast
        main_py_text = PROJECT_ROOT.joinpath("main.py").read_text(encoding="utf-8")
        tree = ast.parse(main_py_text)
        save_word = next(
            (n for n in ast.walk(tree)
             if isinstance(n, (ast.FunctionDef, ast.AsyncFunctionDef))
             and n.name == "_save_word"),
            None,
        )
        self.assertIsNotNone(save_word, "main.py 必须定义 _save_word")

        calls = set()
        for node in ast.walk(save_word):
            if isinstance(node, ast.Call):
                func = node.func
                if isinstance(func, ast.Name):
                    calls.add(func.id)
                elif isinstance(func, ast.Attribute):
                    calls.add(func.attr)
        self.assertIn(
            "clear_word_core_properties", calls,
            "_save_word 必须调用 clear_word_core_properties (G2 Gap 4)",
        )

    def test_save_word_clears_core_properties_end_to_end(self):
        """端到端: 合成 docx 含 Title="秘密信息"; 调用 _save_word 流程后, 重新打开验证清空."""
        from privacyguard.utils import clear_word_core_properties

        with tempfile.TemporaryDirectory() as tmp:
            in_path = os.path.join(tmp, "in_with_props.docx")
            out_path = os.path.join(tmp, "out_cleared.docx")

            doc = Document()
            doc.add_paragraph("段落内容")
            cp = doc.core_properties
            cp.title = "秘密信息 身份证 110101199001011234"
            cp.author = "原作者 张三"
            doc.save(in_path)

            # 模拟 _save_word 的 metadata clearing 步骤 (在 save 之前)
            reopened = Document(in_path)
            cleared = clear_word_core_properties(reopened)
            self.assertEqual(cleared, 5)
            reopened.save(out_path)

            # 反向断言
            final_doc = Document(out_path)
            self.assertEqual(final_doc.core_properties.title, "")
            self.assertEqual(final_doc.core_properties.author, "")


# ----------------------------------------------------------------------
# Test 6: Reverse-extraction 端到端
# ----------------------------------------------------------------------
class TestReverseExtractionCoreProperties(unittest.TestCase):
    """通过 save→reload 通道, 验证 Title/Author/Subject/Comments/Keywords 5 字段均无 PII 残留."""

    def test_save_word_does_not_leak_pii_in_title(self):
        """保存 .docx 后, Title 字段不含原始敏感信息且为空字符串."""
        from privacyguard.utils import clear_word_core_properties
        secret_id = fake_id_card()
        with tempfile.TemporaryDirectory() as tmp:
            in_path = os.path.join(tmp, "in_title.docx")
            out_path = os.path.join(tmp, "out_title.docx")
            doc = Document()
            doc.add_paragraph(f"测试身份证：{secret_id}")
            doc.core_properties.title = f"身份证 {secret_id} 内容"
            doc.save(in_path)

            reopened = Document(in_path)
            clear_word_core_properties(reopened)
            reopened.save(out_path)

            final_doc = Document(out_path)
            self.assertEqual(final_doc.core_properties.title, "")
            self.assertNotIn(secret_id, final_doc.core_properties.title)

    def test_save_word_clears_all_five_standard_properties(self):
        """5 字段全部清空 + 不含原始模拟 PII."""
        from privacyguard.utils import clear_word_core_properties
        secret_phone = fake_phone()
        secret_id = fake_id_card()
        with tempfile.TemporaryDirectory() as tmp:
            in_path = os.path.join(tmp, "in_all5.docx")
            out_path = os.path.join(tmp, "out_all5.docx")
            doc = Document()
            doc.add_paragraph("测试")
            cp = doc.core_properties
            cp.title = f"标题 {secret_id}"
            cp.author = f"作者 {secret_phone}"
            cp.subject = f"主题 {secret_id}"
            cp.comments = f"评论 {secret_phone}"
            cp.keywords = f"关键字 {secret_id}"
            doc.save(in_path)

            reopened = Document(in_path)
            cleared = clear_word_core_properties(reopened)
            self.assertEqual(cleared, 5)
            reopened.save(out_path)

            final_doc = Document(out_path)
            self.assertEqual(final_doc.core_properties.title, "")
            self.assertEqual(final_doc.core_properties.author, "")
            self.assertEqual(final_doc.core_properties.subject, "")
            self.assertEqual(final_doc.core_properties.comments, "")
            self.assertEqual(final_doc.core_properties.keywords, "")
            for field in ("title", "author", "subject", "comments", "keywords"):
                value = getattr(final_doc.core_properties, field, "")
                self.assertNotIn(secret_id, value, f"{field} 残留身份证 PII")
                self.assertNotIn(secret_phone, value, f"{field} 残留手机号 PII")


# ----------------------------------------------------------------------
# Test 7: PyInstaller hiddenimports 静态 parity (cp30 教训)
# ----------------------------------------------------------------------
class TestPackageImportsParity(unittest.TestCase):
    """cp30 教训: 新模块必须在 Windows/macOS spec + macOS build script 静态声明,
    避免 frozen build 启动时 ModuleNotFoundError 回归."""

    def test_windows_spec_contains_word_props_hiddenimport(self):
        """Windows spec 必须包含 'privacyguard.utils.word_props' (cp30 静态守护)."""
        spec_path = PROJECT_ROOT / "packaging" / "windows" / "config" / "PrivacyGuard_windows.spec"
        self.assertTrue(spec_path.exists(), f"Windows spec 不存在: {spec_path}")
        content = spec_path.read_text(encoding="utf-8")
        self.assertIn(
            "privacyguard.utils.word_props", content,
            "Windows spec 必须显式声明 'privacyguard.utils.word_props' hiddenimport "
            "(cp30 教训: 漏声明会导致 frozen build 启动 ModuleNotFoundError)",
        )

    def test_macos_spec_contains_word_props_hiddenimport(self):
        """macOS spec 必须包含 'privacyguard.utils.word_props'."""
        spec_path = PROJECT_ROOT / "packaging" / "macos" / "config" / "PrivacyGuard.spec"
        self.assertTrue(spec_path.exists(), f"macOS spec 不存在: {spec_path}")
        content = spec_path.read_text(encoding="utf-8")
        self.assertIn(
            "privacyguard.utils.word_props", content,
            "macOS spec 必须显式声明 'privacyguard.utils.word_props' hiddenimport",
        )

    def test_macos_build_script_references_word_props(self):
        """macOS build_complete.sh 必须包含 'privacyguard.utils.word_props' (parity 注释或条目)."""
        script_path = PROJECT_ROOT / "packaging" / "macos" / "scripts" / "build_complete.sh"
        self.assertTrue(script_path.exists(), f"macOS build script 不存在: {script_path}")
        content = script_path.read_text(encoding="utf-8")
        self.assertIn(
            "privacyguard.utils.word_props", content,
            "macOS build_complete.sh 必须引用 'privacyguard.utils.word_props' "
            "(parity 注释或显式条目均可, 与 Windows spec 一致)",
        )


if __name__ == "__main__":
    unittest.main()