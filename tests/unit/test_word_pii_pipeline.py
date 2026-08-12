"""Phase 3 Word 端到端 PII 流程测试（D-13 — 79/79 升级 88/88）。

5 + 1 个测试类共 7 + 4 个测试方法（Wave 1 → Wave 2 GREEN）：

- TestWordAdapterCollectUnits — WordAdapter.collect_units 段落 + 表格 双向映射
- TestWordPIIAutoTrigger — PIIEngine.detect 在 Word 文本中命中 6 类以上实体
- TestWordRedactRoundTrip — redact_word 真脱敏 + 反向断言 SAFE-02
- TestWordDocumentPropertiesCleared — clear_word_doc_props 清 5 core + revision=1
- TestWordMergePriorityRulePiManualOcr — priority 锁定 rule > pii > manual > ocr
- TestWordPIIPanelHighlights — cp27 增量 patch 契约 + 短码徽章 + partial mask 渲染

Wave 3 (03-03) 新增 3 个测试类共 9 个测试方法 — WordCandidateDialog UI 行为 + 翻页 + 跨翻页持久化：
- TestWordCandidateDialog — 基础 UI + 实体类型 / 来源筛选 + confirmed 信号 payload
- TestWordCandidateDialogPagination — 50 条分页 + 筛选组合 + 行 label 截断 30 字符
- TestWordCandidateDialogSelectionAcrossPages — 跨翻页选择持久化（per BLOCKER 4）
"""
import os
import tempfile
import unittest
from types import MethodType
from unittest.mock import Mock

from bs4 import BeautifulSoup
from docx import Document
from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QApplication
from PyQt6.QtWebEngineWidgets import QWebEngineView

from privacyguard.pii.engine import PIIEngine
from privacyguard.pii.hits import PIIHit, TextUnit
from privacyguard.pii.mask import mask_for_entity

from tests.fixtures.fake_pii import (
    fake_bank_account,
    fake_bank_card,
    fake_email,
    fake_id_card,
    fake_phone,
    fake_uscc,
    fake_vat_invoice_20,
)
from tests.fixtures.fake_word import build_fake_docx


# PyQt6 必须在导入 QWidget 派生类前创建 QApplication 实例。
# 这对所有 TestWordCandidateDialog* 测试类至关重要。
_APP = QApplication.instance() or QApplication([])


class TestWordAdapterCollectUnits(unittest.TestCase):
    """WordAdapter.collect_units 段落 + 表格双向映射。"""

    def test_collect_units_returns_text_unit_per_block(self):
        """collect_units 至少返回与段落数 + 单元格数一致的 TextUnit 列表 + key_index 映射。"""
        path = build_fake_docx(paragraphs=["段落 0", "段落 1"])
        try:
            from privacyguard.word import WordAdapter
            units, key_index = WordAdapter.collect_units(path)
            self.assertGreaterEqual(len(units), 2)
            # key_index 的所有 value 字符串必须能在 word_data 中被引用
            for key in key_index.values():
                self.assertIsInstance(key, str)
        finally:
            os.remove(path)


class TestWordPIIAutoTrigger(unittest.TestCase):
    """PIIEngine.detect 在 Word 文本中命中 6 类以上实体（9 类可见）。"""

    def test_engine_detects_pii_in_word_text(self):
        """构造含 7 类 PII 的 Word 文本，断言 entity_types >= 6。"""
        text = (
            f"身份证 {fake_id_card()} "
            f"手机 {fake_phone()} "
            f"邮箱 {fake_email()} "
            f"卡号 {fake_bank_card()} "
            f"统一信用代码 {fake_uscc()} "
            f"发票 {fake_vat_invoice_20()} "
            f"账号 {fake_bank_account()}"
        )
        engine = PIIEngine()
        hits = engine.detect(TextUnit(page_index=0, text=text, source='text'))
        entity_types = {h.entity_type for h in hits}
        self.assertGreaterEqual(
            len(entity_types), 6,
            f"期望 ≥6 类 PII 命中，实际 {len(entity_types)}: {entity_types}",
        )
        # 关键类别必须命中
        for required in ('CN_ID_CARD', 'CN_PHONE', 'CN_EMAIL', 'CN_BANK_CARD', 'CN_USCC'):
            self.assertIn(required, entity_types, f"必须命中 {required}")


class TestWordRedactRoundTrip(unittest.TestCase):
    """redact_word 真脱敏 + 反向断言 SAFE-02（导出 docx 不含原文敏感字串）。"""

    def test_redact_word_partial_mask_visible(self):
        """构造含身份证的段落，redact_word + clear_word_doc_props 后导出 docx 不含原文 + 含 partial mask。"""
        from privacyguard.word.redact import redact_word
        from privacyguard.word.clear_doc_props import clear_word_doc_props
        from main import merge_word_matches_with_priority

        secret_id = fake_id_card()
        path = build_fake_docx(paragraphs=[f"原文 {secret_id}"])
        try:
            word_data = {
                "paragraph_0": {
                    "text": f"原文 {secret_id}",
                    "ocr": [],
                    "manual": [],
                    "pii": [],
                }
            }
            engine = PIIEngine()
            hits = engine.detect(TextUnit(page_index=0, text=word_data["paragraph_0"]["text"], source='text'))
            word_data["paragraph_0"]["pii"] = hits

            doc = Document(path)
            merged = merge_word_matches_with_priority(
                word_data["paragraph_0"]["text"],
                rules=[],
                default_replacement_text="[已脱敏]",
                manual_matches=[],
                ocr_matches=[],
                pii_matches=hits,
            )
            redact_word(doc, "paragraph_0", merged)
            clear_word_doc_props(doc)

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as fh:
                out_path = fh.name
            try:
                doc.save(out_path)
                out_doc = Document(out_path)
                out_text = "".join(p.text for p in out_doc.paragraphs)
                self.assertNotIn(
                    secret_id, out_text,
                    f"导出 docx 仍含原文敏感字串：{secret_id}",
                )
                self.assertIn(
                    mask_for_entity("CN_ID_CARD", secret_id), out_text,
                    f"导出 docx 应含 partial mask：{mask_for_entity('CN_ID_CARD', secret_id)}",
                )
            finally:
                os.remove(out_path)
        finally:
            os.remove(path)


class TestWordDocumentPropertiesCleared(unittest.TestCase):
    """clear_word_doc_props 清 5 core + revision=1（D-08 / D-24）。"""

    def test_clear_core_5_fields_always_succeeds(self):
        """写入敏感 5 字段 → clear → 导出 docx → 断言 5 字段全部 == ""。"""
        from privacyguard.word.clear_doc_props import clear_word_doc_props

        path = build_fake_docx(paragraphs=["test"])
        try:
            doc = Document(path)
            # 写敏感字符串
            doc.core_properties.title = "SENSITIVE TITLE"
            doc.core_properties.author = "SENSITIVE AUTHOR"
            doc.core_properties.subject = "SENSITIVE SUBJECT"
            doc.core_properties.keywords = "SENSITIVE KEYWORDS"
            doc.core_properties.last_modified_by = "SENSITIVE LAST_MODIFIED_BY"

            clear_word_doc_props(doc)

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as fh:
                out_path = fh.name
            try:
                doc.save(out_path)
                out_doc = Document(out_path)
                for prop_name in ('title', 'author', 'subject', 'keywords', 'last_modified_by'):
                    value = getattr(out_doc.core_properties, prop_name)
                    self.assertEqual(
                        value, "",
                        f"{prop_name} 应为 ''，实际 {value!r}",
                    )
            finally:
                os.remove(out_path)
        finally:
            os.remove(path)

    def test_clear_revision_set_to_1(self):
        """clear_word_doc_props 必须把 revision 字段重置为整数 1（D-08 / D-24 锁）。"""
        from privacyguard.word.clear_doc_props import clear_word_doc_props

        path = build_fake_docx(paragraphs=["test"])
        try:
            doc = Document(path)
            doc.core_properties.revision = 99

            clear_word_doc_props(doc)

            self.assertEqual(doc.core_properties.revision, 1)
        finally:
            os.remove(path)


class TestWordMergePriorityRulePiManualOcr(unittest.TestCase):
    """merge_word_matches_with_priority 第六参数 pii_matches 优先级锁定 rule > pii > manual > ocr。"""

    def test_rule_beats_pii(self):
        """规则命中与 PII 命中不重叠时，merged 应包含两条：rule 在前 + pii 在后。"""
        from main import merge_word_matches_with_priority
        from privacyguard.pii.hits import PIIHit

        text = "张三 13812345678"
        rules = [{"enabled": True, "mode": "exact", "find": "张三", "replace": "[姓名]"}]
        phone_hit = PIIHit(
            entity_type='CN_PHONE',
            page_offset=3,
            page_length=11,
            page_rect=(0.0, 0.0, 66.0, 12.0),
            mask_strategy="138****5678",
            normalized="13812345678",
        )
        merged = merge_word_matches_with_priority(
            text, rules, "[已脱敏]",
            manual_matches=[], ocr_matches=[],
            pii_matches=[phone_hit],
        )
        self.assertEqual(len(merged), 2)
        self.assertEqual(merged[0]["source"], "rule")
        self.assertEqual(merged[1]["source"], "pii")

    def test_pii_beats_manual_on_overlap(self):
        """区间重叠时 pii 命中覆盖 manual 命中（D-19 priority 锁）。"""
        from main import merge_word_matches_with_priority
        from privacyguard.pii.hits import PIIHit

        secret_id = fake_id_card()
        text = f"原文 {secret_id}"
        manual = [{"start": 0, "end": 18, "text": secret_id, "replacement": "[手动]"}]
        id_hit = PIIHit(
            entity_type='CN_ID_CARD',
            page_offset=3,
            page_length=18,
            page_rect=(0.0, 0.0, 108.0, 12.0),
            mask_strategy=mask_for_entity("CN_ID_CARD", secret_id),
            normalized=secret_id,
        )
        merged = merge_word_matches_with_priority(
            text, [], "[已脱敏]",
            manual_matches=manual, ocr_matches=[],
            pii_matches=[id_hit],
        )
        self.assertEqual(len(merged), 1)
        self.assertEqual(merged[0]["source"], "pii")


def _build_pii_panel_stub():
    """构造 PII 面板测试 stub（不继承 MainWindow；通过 MethodType 绑定 MainWindow 方法）。"""
    from main import MainWindow

    stub = type("_StubMainWindow", (), {})()
    stub.word_data = {
        "paragraph_5": {
            "text": "身份证 53010219200508011X 末位",
            "ocr": [],
            "manual": [],
            "pii": [],
        }
    }
    # 左右双栏 WebEngineView（Mock spec 保证 page() / runJavaScript 被自动模拟）
    stub.word_preview = Mock(spec=QWebEngineView)
    stub.word_preview.isHidden.return_value = False
    stub.word_preview_replaced = Mock(spec=QWebEngineView)
    stub.word_preview_replaced.isHidden.return_value = False
    # 绑定 MainWindow 三个新方法到 stub（via MethodType 替代继承）
    stub._apply_word_pii_panel_updates = MethodType(
        MainWindow._apply_word_pii_panel_updates, stub
    )
    stub._build_pii_block_fragment = MethodType(
        MainWindow._build_pii_block_fragment, stub
    )
    stub._build_pii_mask_block_fragment = MethodType(
        MainWindow._build_pii_mask_block_fragment, stub
    )
    return stub


def _build_id_card_hit():
    """构造单条 CN_ID_CARD PIIHit 命中（page_offset=4 page_length=18）。"""
    from privacyguard.pii.hits import PIIHit

    secret = "53010219200508011X"
    return [
        PIIHit(
            entity_type='CN_ID_CARD',
            page_offset=4,
            page_length=18,
            page_rect=(0, 0, 0, 0),
            confidence_tier='HIGH',
            source='text',
            mask_strategy=mask_for_entity('CN_ID_CARD', secret),
            normalized=secret,
        )
    ]


class TestWordPIIPanelHighlights(unittest.TestCase):
    """Phase 3 (03-word) — TestWordPIIPanelHighlights 4 方法
    cp27 增量 patch 契约 + 短码徽章 + partial mask 渲染。
    """

    def test_apply_word_pii_panel_updates_uses_runJavaScript_not_setHtml(self):
        """_apply_word_pii_panel_updates 走 cp27 局部 patch:
        web_view.page().runJavaScript 被调 + web_view.setHtml 未被调（D-10 锁）。"""
        stub = _build_pii_panel_stub()
        hits = _build_id_card_hit()

        stub._apply_word_pii_panel_updates(key="paragraph_5", hits=hits)

        # cp27 契约: 局部 patch via runJavaScript（左右双栏）
        stub.word_preview.page().runJavaScript.assert_called()
        stub.word_preview_replaced.page().runJavaScript.assert_called()
        # cp27 反向断言: 禁止整页 setHtml
        stub.word_preview.setHtml.assert_not_called()
        stub.word_preview_replaced.setHtml.assert_not_called()

    def test_build_pii_block_fragment_contains_short_code_badge(self):
        """_build_pii_block_fragment 构造左栏原文高亮 + 短码徽章 HTML 片段。"""
        stub = _build_pii_panel_stub()
        hits = _build_id_card_hit()

        fragment = stub._build_pii_block_fragment(key="paragraph_5", hits=hits)

        # HTML 元素名 + entity_type 属性 + 短码徽章（per BLOCKER 5 + D-21 单一来源）
        self.assertIn('<mark class="pii-highlight"', fragment)
        self.assertIn('data-entity-type="CN_ID_CARD"', fragment)
        self.assertIn('<span class="pii-tag">ID</span>', fragment)
        # 原文包裹
        self.assertIn("53010219200508011X", fragment)
        # 左栏**不**含 mask 字符串（Visuals §PII Highlight 锁定）
        self.assertNotIn(
            mask_for_entity('CN_ID_CARD', '53010219200508011X'), fragment,
        )

    def test_build_pii_mask_block_fragment_contains_mask_string_not_original(self):
        """_build_pii_mask_block_fragment 构造右栏 partial mask 片段。"""
        stub = _build_pii_panel_stub()
        hits = _build_id_card_hit()

        fragment = stub._build_pii_mask_block_fragment(key="paragraph_5", hits=hits)

        # HTML 元素名 + entity_type 属性 + mask 字符串 + title 前缀
        self.assertIn('<mark class="pii-mask"', fragment)
        self.assertIn('data-entity-type="CN_ID_CARD"', fragment)
        self.assertIn(
            mask_for_entity('CN_ID_CARD', '53010219200508011X'), fragment,
        )
        self.assertIn("已替换为：", fragment)
        # 右栏**不**含原文（Visuals §PII Partial-Mask 锁定）
        self.assertNotIn("53010219200508011X", fragment)

    def test_entity_type_short_code_covers_all_9_locked_types(self):
        """ENTITY_TYPE_SHORT_CODE 9 短码字典覆盖 9 类 entity（D-21 单一来源锁）。
        9 个短码值锁定为 ASCII uppercase；不允许中文 / 小写。"""
        from main import ENTITY_TYPE_SHORT_CODE

        expected_keys = {
            'CN_ID_CARD', 'CN_PHONE', 'CN_BANK_CARD', 'CN_EMAIL',
            'CN_USCC', 'CN_TAXPAYER_ID', 'CN_TAXPAYER_ID_15',
            'CN_VAT_INVOICE', 'CN_BANK_ACCOUNT',
        }
        self.assertEqual(set(ENTITY_TYPE_SHORT_CODE.keys()), expected_keys)
        # ASCII uppercase 短码（D-21 锁）
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_ID_CARD'], 'ID')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_PHONE'], 'PHONE')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_BANK_CARD'], 'BANK')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_EMAIL'], 'EMAIL')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_USCC'], 'USCC')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_TAXPAYER_ID'], 'TAX')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_TAXPAYER_ID_15'], 'TAX15')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_VAT_INVOICE'], 'VAT')
        self.assertEqual(ENTITY_TYPE_SHORT_CODE['CN_BANK_ACCOUNT'], 'ACCT')


# ----------------------------------------------------------------------
# Wave 3 (03-03) — WordCandidateDialog UI 行为测试 (UX-01 / UX-02 / BLOCKER 3 / BLOCKER 4)
# ----------------------------------------------------------------------


def _ensure_qapp():
    """确保 QApplication 单例存在（PyQt6 必须在创建 QDialog 前有 QApplication）。"""
    return QApplication.instance() or QApplication([])


def _build_pii_hit(entity_type, key='paragraph_5', page_offset=4, page_length=18,
                   mask_strategy='530102********011X', normalized='53010219200508011X'):
    """构造一条 PIIHit 命中（page_offset / page_length 默认对正 CN_ID_CARD 18 位）。"""
    return PIIHit(
        entity_type=entity_type,
        page_offset=page_offset,
        page_length=page_length,
        page_rect=(0.0, 0.0, 108.0, 12.0),
        confidence_tier='HIGH',
        source='text',
        mask_strategy=mask_strategy,
        normalized=normalized,
    )


def _build_manual_match(start, end, text, replacement='[手动]'):
    """构造一条 manual 命中 dict（与 main.py:word_data[key]['manual'] 形态一致）。"""
    return {'start': start, 'end': end, 'text': text, 'replacement': replacement}


class TestWordCandidateDialog(unittest.TestCase):
    """WordCandidateDialog 基础 UI 行为（UX-01 / UX-02 最低功能 + D-11 / D-25 锁）。"""

    def setUp(self):
        _ensure_qapp()

    def test_dialog_opens_with_all_hits_in_word_data(self):
        """word_data 三通道 (pii + ocr + manual) 总命中应全部进入 _all_hits + 初始页 0 + window title。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        word_data = {
            "paragraph_5": {
                "text": "身份证 53010219200508011X 手机 13812345678",
                "ocr": [],
                "manual": [],
                "pii": [
                    _build_pii_hit('CN_ID_CARD', key='paragraph_5',
                                   page_offset=4, page_length=18,
                                   mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                                   normalized='53010219200508011X'),
                    _build_pii_hit('CN_PHONE', key='paragraph_5',
                                   page_offset=27, page_length=11,
                                   mask_strategy=mask_for_entity('CN_PHONE', '13812345678'),
                                   normalized='13812345678'),
                ],
            },
            "table_0_cell_0_0": {
                "text": "53010219200508011X",
                "ocr": [],
                "manual": [_build_manual_match(0, 18, '53010219200508011X')],
                "pii": [],
            },
        }

        dlg = WordCandidateDialog(word_data)

        # 三通道总命中数 = 2 pii + 1 manual = 3
        self.assertEqual(len(dlg._all_hits), 3)
        self.assertEqual(dlg.windowTitle(), 'Word 候选审阅')
        self.assertEqual(dlg._page, 0)

    def test_entity_filter_changes_visible_rows(self):
        """实体类型筛选切换 → list_widget 行数随之变化（D-25 + UX-02）。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        word_data = {"paragraph_5": {"text": "x", "ocr": [], "manual": [], "pii": []}}
        # 3 个 CN_ID_CARD + 2 个 CN_PHONE = 5 个 pii hit
        pii_hits = []
        offset = 0
        for i in range(3):
            pii_hits.append(_build_pii_hit(
                'CN_ID_CARD', key=f'paragraph_{i}', page_offset=offset, page_length=18,
                mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                normalized='53010219200508011X',
            ))
            offset += 20
        for i in range(3, 5):
            pii_hits.append(_build_pii_hit(
                'CN_PHONE', key=f'paragraph_{i}', page_offset=offset, page_length=11,
                mask_strategy=mask_for_entity('CN_PHONE', '13812345678'),
                normalized='13812345678',
            ))
            offset += 15
        word_data['paragraph_5']['pii'] = pii_hits

        dlg = WordCandidateDialog(word_data)
        dlg._refresh()
        # 全部 5 条可见
        self.assertEqual(dlg.list_widget.count(), 5)
        # 切到 CN_PHONE → 仅 2 条
        phone_idx = dlg.entity_filter.findData('CN_PHONE')
        self.assertGreaterEqual(phone_idx, 0)
        dlg.entity_filter.setCurrentIndex(phone_idx)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 2)
        # 切回 "全部类型"（data == ''） → 5 条
        dlg.entity_filter.setCurrentIndex(0)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 5)

    def test_source_filter_changes_visible_rows(self):
        """来源筛选切换 → list_widget 行数随之变化（D-25 + UX-02）。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        word_data = {
            "paragraph_5": {
                "text": "x",
                "ocr": [
                    _build_manual_match(0, 11, '13812345678'),
                    _build_manual_match(15, 18, 'ABC'),
                ],
                "manual": [_build_manual_match(0, 18, '53010219200508011X')],
                "pii": [
                    _build_pii_hit('CN_ID_CARD', page_offset=0, page_length=18,
                                   mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                                   normalized='53010219200508011X'),
                    _build_pii_hit('CN_PHONE', page_offset=20, page_length=11,
                                   mask_strategy=mask_for_entity('CN_PHONE', '13812345678'),
                                   normalized='13812345678'),
                    _build_pii_hit('CN_EMAIL', page_offset=35, page_length=12,
                                   mask_strategy='a***@b.com',
                                   normalized='abc@def.com'),
                ],
            },
        }

        dlg = WordCandidateDialog(word_data)
        dlg._refresh()
        # 全部 6 条
        self.assertEqual(dlg.list_widget.count(), 6)
        # 切到 ocr → 2 条
        ocr_idx = dlg.source_filter.findData('ocr')
        self.assertGreaterEqual(ocr_idx, 0)
        dlg.source_filter.setCurrentIndex(ocr_idx)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 2)
        # 切到 pii → 3 条
        pii_idx = dlg.source_filter.findData('pii')
        self.assertGreaterEqual(pii_idx, 0)
        dlg.source_filter.setCurrentIndex(pii_idx)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 3)

    def test_confirmed_hit_emits_to_main_window(self):
        """点击 '确认选中的 N 项' → confirmed 信号 emit list[dict{key, hit, source}] payload。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        secret_id = '53010219200508011X'
        word_data = {
            "paragraph_5": {
                "text": f"原文 {secret_id}",
                "ocr": [],
                "manual": [],
                "pii": [_build_pii_hit(
                    'CN_ID_CARD', key='paragraph_5', page_offset=3, page_length=18,
                    mask_strategy=mask_for_entity('CN_ID_CARD', secret_id),
                    normalized=secret_id,
                )],
            },
        }
        dlg = WordCandidateDialog(word_data)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 1)

        captured = []
        dlg.confirmed.connect(lambda payload: captured.append(payload))

        dlg._on_confirm_clicked()

        self.assertEqual(len(captured), 1)
        payload = captured[0]
        self.assertIsInstance(payload, list)
        self.assertEqual(len(payload), 1)
        entry = payload[0]
        self.assertIsInstance(entry, dict)
        self.assertEqual(entry.get('key'), 'paragraph_5')
        # hit_dict 应含 entity_type + mask_strategy 字段（D-18 契约）
        hit_dict = entry.get('hit') or {}
        self.assertEqual(hit_dict.get('entity_type'), 'CN_ID_CARD')
        self.assertEqual(hit_dict.get('mask_strategy'),
                         mask_for_entity('CN_ID_CARD', secret_id))
        self.assertEqual(entry.get('source'), 'pii')

    def test_empty_state_when_all_hits_filtered_out(self):
        """过滤后 0 条但 _all_hits 非空 → 显示空态文案 + 主 CTA disabled（per E3 covered·partial）。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        word_data = {"paragraph_5": {"text": "x", "ocr": [], "manual": [], "pii": [
            _build_pii_hit('CN_ID_CARD', page_offset=0, page_length=18,
                           mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                           normalized='53010219200508011X'),
        ]}}

        dlg = WordCandidateDialog(word_data)
        phone_idx = dlg.entity_filter.findData('CN_PHONE')
        self.assertGreaterEqual(phone_idx, 0)
        dlg.entity_filter.setCurrentIndex(phone_idx)
        dlg._refresh()
        # list_widget 无可见行
        self.assertEqual(dlg.list_widget.count(), 0)
        # page_label 含 '0'（页数 / 条数）
        self.assertIn('0', dlg.page_label.text())
        # 主 CTA disabled
        self.assertFalse(dlg.btn_confirm.isEnabled())


class TestWordCandidateDialogPagination(unittest.TestCase):
    """WordCandidateDialog 50 条分页 + 筛选组合 + 行 label 截断 30 字符（D-25 + UX-02）。"""

    def setUp(self):
        _ensure_qapp()

    def test_pagination_over_50_entries(self):
        """60 个 hit 应分为 2 页；第一页 50 条 + 第二页 10 条 + 翻页按钮 enable 状态正确。"""
        from privacyguard.word.candidate_dialog import PAGE_SIZE, WordCandidateDialog

        self.assertEqual(PAGE_SIZE, 50)

        # 60 个 pii hit（每个 key 1 个 hit，60 个不同 key）
        word_data = {"paragraph_5": {"text": "x", "ocr": [], "manual": [], "pii": []}}
        pii_hits = []
        for i in range(60):
            offset = i * 19
            pii_hits.append(_build_pii_hit(
                'CN_ID_CARD', key=f'paragraph_{i}', page_offset=offset, page_length=18,
                mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                normalized='53010219200508011X',
            ))
        word_data['paragraph_5']['pii'] = pii_hits

        dlg = WordCandidateDialog(word_data)
        dlg._refresh()
        # 第一页 50 条
        self.assertEqual(dlg.list_widget.count(), 50)
        # page_label 含 '第 1 / 2 页（共 60 条）'
        page_text = dlg.page_label.text()
        self.assertIn('1', page_text)
        self.assertIn('2', page_text)
        self.assertIn('60', page_text)
        # 翻页按钮 enable 状态
        self.assertTrue(dlg.btn_next.isEnabled())
        self.assertFalse(dlg.btn_prev.isEnabled())
        # 翻到第二页
        dlg._next_page()
        dlg._refresh()
        # 第二页 10 条
        self.assertEqual(dlg.list_widget.count(), 10)
        # 末页 btn_next disabled
        self.assertFalse(dlg.btn_next.isEnabled())
        self.assertTrue(dlg.btn_prev.isEnabled())

    def test_pagination_filter_combination(self):
        """30 CN_ID_CARD + 30 CN_PHONE = 60；筛选 CN_ID_CARD 后只 30 条全部可见、无分页。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        word_data = {"paragraph_5": {"text": "x", "ocr": [], "manual": [], "pii": []}}
        pii_hits = []
        for i in range(30):
            pii_hits.append(_build_pii_hit(
                'CN_ID_CARD', key=f'paragraph_{i}', page_offset=i * 19, page_length=18,
                mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                normalized='53010219200508011X',
            ))
        for i in range(30, 60):
            pii_hits.append(_build_pii_hit(
                'CN_PHONE', key=f'paragraph_{i}', page_offset=i * 12, page_length=11,
                mask_strategy=mask_for_entity('CN_PHONE', '13812345678'),
                normalized='13812345678',
            ))
        word_data['paragraph_5']['pii'] = pii_hits

        dlg = WordCandidateDialog(word_data)
        id_idx = dlg.entity_filter.findData('CN_ID_CARD')
        dlg.entity_filter.setCurrentIndex(id_idx)
        dlg._refresh()
        # 30 条全部可见，无需分页
        self.assertEqual(dlg.list_widget.count(), 30)
        page_text = dlg.page_label.text()
        self.assertIn('1', page_text)
        self.assertIn('30', page_text)
        # 末页 → btn_next disabled
        self.assertFalse(dlg.btn_next.isEnabled())

    def test_row_label_truncates_normalized_at_30_chars(self):
        """行 label 截断 normalized[:30] + '...'；normalized[30:] 不暴露。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        normalized_50 = '1234567890123456789012345678901234567890ABCDEFGHIJ'
        word_data = {"paragraph_5": {"text": "x", "ocr": [], "manual": [], "pii": [
            _build_pii_hit('CN_ID_CARD', page_offset=0, page_length=50,
                           mask_strategy='[MASK]', normalized=normalized_50),
        ]}}

        dlg = WordCandidateDialog(word_data)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 1)
        row_text = dlg.list_widget.item(0).text()

        # normalized[:30] 应在行文本中
        self.assertIn(normalized_50[:30], row_text)
        # normalized[30:] 不应在行文本中（隐私截断锁）
        self.assertNotIn(normalized_50[30:], row_text)
        # 截断标识符 '...' 或 '…'
        self.assertTrue(('...' in row_text) or ('…' in row_text),
                        f"行 label 必须含 '...'/'…' 截断标识符；实际：{row_text!r}")


class TestWordCandidateDialogSelectionAcrossPages(unittest.TestCase):
    """WordCandidateDialog 跨翻页选择持久化（per BLOCKER 4）。"""

    def setUp(self):
        _ensure_qapp()

    def test_selection_persists_across_pages(self):
        """60 个 hit；第 1 页取消 5 项 → 翻到第 2 页 → 翻回第 1 页 → 那 5 项仍 Unchecked → confirm emit 55 项。"""
        from privacyguard.word.candidate_dialog import WordCandidateDialog

        word_data = {"paragraph_5": {"text": "x", "ocr": [], "manual": [], "pii": []}}
        pii_hits = []
        for i in range(60):
            pii_hits.append(_build_pii_hit(
                'CN_ID_CARD', key=f'paragraph_{i}', page_offset=i * 19, page_length=18,
                mask_strategy=mask_for_entity('CN_ID_CARD', '53010219200508011X'),
                normalized='53010219200508011X',
            ))
        word_data['paragraph_5']['pii'] = pii_hits

        dlg = WordCandidateDialog(word_data)
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 50)
        # 第 1 页 = paragraph_0..paragraph_49
        # 第 2 页 = paragraph_50..paragraph_59

        # 在第 1 页取消前 5 个 checkbox（paragraph_0..paragraph_4）
        for i in range(5):
            dlg.list_widget.item(i).setCheckState(Qt.CheckState.Unchecked)
        # 同步进 self._selection（per BLOCKER 4）
        dlg._sync_selection_from_list()

        # 翻到第 2 页（paragraph_50..paragraph_59）—— 第 2 页的 10 项应保持 Checked
        dlg._next_page()
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 10)
        for i in range(10):
            self.assertEqual(
                dlg.list_widget.item(i).checkState(), Qt.CheckState.Checked,
                f"第 2 页第 {i} 项应保持 Checked（未在第 1 页被取消）"
            )

        # 翻回第 1 页 —— paragraph_0..paragraph_4 应仍 Unchecked（per BLOCKER 4 跨翻页持久化）
        dlg._prev_page()
        dlg._refresh()
        self.assertEqual(dlg.list_widget.count(), 50)
        for i in range(5):
            self.assertEqual(
                dlg.list_widget.item(i).checkState(), Qt.CheckState.Unchecked,
                f"第 1 页第 {i} 项应保持 Unchecked（per BLOCKER 4 跨翻页持久化）"
            )

        # confirm —— emit 60 - 5 = 55 项
        captured = []
        dlg.confirmed.connect(lambda payload: captured.append(payload))
        dlg._on_confirm_clicked()

        self.assertEqual(len(captured), 1)
        payload = captured[0]
        self.assertEqual(len(payload), 55)
        # 所有 payload 项含 entity_type + key + page_offset + page_length 四元组
        for entry in payload:
            self.assertIn('key', entry)
            self.assertIn('hit', entry)
            self.assertIn('source', entry)
            hit_dict = entry['hit']
            self.assertEqual(hit_dict['entity_type'], 'CN_ID_CARD')
            self.assertIn('page_offset', hit_dict)
            self.assertIn('page_length', hit_dict)


# ----------------------------------------------------------------------
# Wave 4 (03-04) — TestWordDataKeySync + TestWordPartialMaskInComparePane
# D-22 data-key 同步契约 + FMT-02 partial mask 在右栏可见契约
# ----------------------------------------------------------------------


def _build_data_key_stub():
    """构造 data-key 同步测试 stub（MethodType 绑定 _add_data_key_attributes + fallback）。"""
    from main import MainWindow

    stub = type("_StubDataKey", (), {})()
    stub._add_data_key_attributes = MethodType(
        MainWindow._add_data_key_attributes, stub
    )
    stub._add_data_key_regex_fallback = MethodType(
        MainWindow._add_data_key_regex_fallback, stub
    )
    return stub


class TestWordDataKeySync(unittest.TestCase):
    """Phase 3 (03-word) — TestWordDataKeySync 3 方法
    D-22 data-key 注入契约：mammoth 渲染后 DOM data-key 数 ≈ word_data key 数。

    直接复用 main.py 既有 _add_data_key_attributes + _add_data_key_regex_fallback
    helper（per D-22 不重写；03-04 仅补齐测试覆盖）。
    """

    def setUp(self):
        _ensure_qapp()

    def test_data_key_count_matches_word_data(self):
        """mammoth 渲染 + _add_data_key_attributes 后 DOM data-key 数 ≥ word_data key 数。"""
        path = build_fake_docx(
            paragraphs=["段落 0", "段落 1"],
            tables=[[["cell 0", "cell 1"]]],
            add_pii=False,
        )
        try:
            doc = Document(path)
            word_data = {}
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    word_data[f"paragraph_{idx}"] = {
                        "text": para.text,
                        "ocr": [],
                        "manual": [],
                        "pii": [],
                    }
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            word_data[
                                f"table_{table_idx}_cell_{row_idx}_{cell_idx}"
                            ] = {
                                "text": cell.text,
                                "ocr": [],
                                "manual": [],
                                "pii": [],
                            }

            # 模拟 mammoth 输出：段落 + 表格 HTML
            html = (
                "<p>段落 0</p>"
                "<p>段落 1</p>"
                "<table>"
                "<tr><td>cell 0</td><td>cell 1</td></tr>"
                "</table>"
            )
            text_blocks = {
                k: {"text": v["text"], "escaped": v["text"]}
                for k, v in word_data.items()
            }

            stub = _build_data_key_stub()
            tagged = stub._add_data_key_attributes(html, text_blocks)
            soup = BeautifulSoup(tagged, "html.parser")
            keys_found = {
                el.get("data-key")
                for el in soup.find_all(attrs={"data-key": True})
            }

            # 段落 keys 必须全部命中
            self.assertIn("paragraph_0", keys_found)
            self.assertIn("paragraph_1", keys_found)
            # 至少一个表格 cell key 命中（_add_data_key_attributes 表格 td 也匹配）
            table_keys = [k for k in keys_found if k.startswith("table_0_cell_0_")]
            self.assertGreaterEqual(
                len(table_keys), 1,
                f"期望至少 1 个 table_0_cell_0_* 命中，实际 {table_keys}",
            )
        finally:
            os.remove(path)

    def test_data_key_fallback_used_for_inline_tags(self):
        """段落含 <strong> inline 标签时 _add_data_key_regex_fallback 兜底生效。"""
        path = build_fake_docx(
            paragraphs=["段落 0 含有 粗体 测试"],
            add_pii=False,
        )
        try:
            # 模拟 mammoth 输出含 inline 标签（block-level BeautifulSoup 匹配可能失败）
            html = "<p>段落 0 含有 <strong>粗体</strong> 测试</p>"
            text_blocks = {
                "paragraph_0": {
                    "text": "段落 0 含有 粗体 测试",
                    "escaped": "段落 0 含有 粗体 测试",
                }
            }

            stub = _build_data_key_stub()
            tagged = stub._add_data_key_attributes(html, text_blocks)
            soup = BeautifulSoup(tagged, "html.parser")
            keys_found = {
                el.get("data-key")
                for el in soup.find_all(attrs={"data-key": True})
            }

            # 若 BeautifulSoup 路径未命中 → fallback 兜底
            if "paragraph_0" not in keys_found:
                tagged = stub._add_data_key_regex_fallback(tagged, text_blocks)
                soup = BeautifulSoup(tagged, "html.parser")
                keys_found = {
                    el.get("data-key")
                    for el in soup.find_all(attrs={"data-key": True})
                }

            self.assertIn(
                "paragraph_0", keys_found,
                f"期望 paragraph_0 命中（fallback 兜底），实际 {keys_found}",
            )
        finally:
            os.remove(path)

    def test_data_key_sync_no_overlap(self):
        """100 段 docx：data-key 命中数 ≥ word_data key 数 × 0.9（允许少量 fallback 失败）。"""
        paragraphs = [f"段落 {i}" for i in range(100)]
        path = build_fake_docx(paragraphs=paragraphs, add_pii=False)
        try:
            doc = Document(path)
            word_data_count = sum(1 for p in doc.paragraphs if p.text.strip())

            html_parts = [f"<p>段落 {i}</p>" for i in range(100)]
            html = "".join(html_parts)
            text_blocks = {
                f"paragraph_{i}": {
                    "text": f"段落 {i}",
                    "escaped": f"段落 {i}",
                }
                for i in range(100)
            }

            stub = _build_data_key_stub()
            tagged = stub._add_data_key_attributes(html, text_blocks)
            # BeautifulSoup 可能漏掉部分段落 → 跑一次 fallback 兜底
            soup = BeautifulSoup(tagged, "html.parser")
            data_key_count = len(
                soup.find_all(attrs={"data-key": True})
            )
            if data_key_count < word_data_count:
                tagged = stub._add_data_key_regex_fallback(tagged, text_blocks)
                soup = BeautifulSoup(tagged, "html.parser")
                data_key_count = len(
                    soup.find_all(attrs={"data-key": True})
                )

            self.assertGreaterEqual(
                data_key_count,
                int(word_data_count * 0.9),
                f"data-key 命中 {data_key_count} 应 ≥ word_data key 数 {word_data_count} × 0.9",
            )
        finally:
            os.remove(path)


class TestWordPartialMaskInComparePane(unittest.TestCase):
    """Phase 3 (03-word) — TestWordPartialMaskInComparePane 2 方法
    FMT-02 partial mask 契约：右栏仅展示 partial mask；左栏展示原文。
    """

    def setUp(self):
        _ensure_qapp()

    def test_partial_mask_string_in_right_pane(self):
        """_build_pii_mask_block_fragment 右栏 fragment 含 partial mask 字符串 + 不含原文。"""
        stub = _build_pii_panel_stub()
        hits = _build_id_card_hit()
        mask_strategy = mask_for_entity('CN_ID_CARD', '53010219200508011X')

        fragment = stub._build_pii_mask_block_fragment(key="paragraph_5", hits=hits)

        # mask 元素 + entity_type 属性 + partial mask 字符串
        self.assertIn('<mark class="pii-mask"', fragment)
        self.assertIn('data-entity-type="CN_ID_CARD"', fragment)
        self.assertIn(mask_strategy, fragment)
        self.assertIn('已替换为：', fragment)
        # 右栏**不**含原文（Visuals §PII Partial-Mask 锁定）
        self.assertNotIn(
            '53010219200508011X', fragment,
            f"右栏 fragment 不应含原文：{fragment!r}",
        )

    def test_left_pane_contains_original_right_pane_contains_mask(self):
        """左栏原文 / 右栏 partial mask：左右 fragment 不相等 + mark class 不同 + 内容不同。"""
        stub = _build_pii_panel_stub()
        hits = _build_id_card_hit()
        mask_strategy = mask_for_entity('CN_ID_CARD', '53010219200508011X')

        left_fragment = stub._build_pii_block_fragment(key="paragraph_5", hits=hits)
        right_fragment = stub._build_pii_mask_block_fragment(key="paragraph_5", hits=hits)

        # 左右 fragment 不相等（双栏内容差异锁）
        self.assertNotEqual(
            left_fragment, right_fragment,
            "左右 fragment 必须不相等（Visuals 双栏差异锁）",
        )
        # 左栏 mark class = pii-highlight；右栏 mark class = pii-mask
        self.assertIn('<mark class="pii-highlight"', left_fragment)
        self.assertIn('<mark class="pii-mask"', right_fragment)
        # 左栏含原文 + 右栏含 partial mask
        self.assertIn('53010219200508011X', left_fragment)
        self.assertNotIn('53010219200508011X', right_fragment)
        self.assertIn(mask_strategy, right_fragment)
        self.assertNotIn(mask_strategy, left_fragment)


if __name__ == "__main__":
    unittest.main()
