"""Phase 3 (03-word) — _ModularWordWorker PII 接入端到端测试 (D-12 + D-15 #2).

锁:
- _ModularWordWorker.__init__ 缓存 self._pii_engine = PIIEngine() (Pitfall 8)
- _ModularWordWorker.run() 段落路径写入 word_data[key]['pii'] = collect_pii_word_hits(...)
- _ModularWordWorker.run() 表格 cell 路径同样写入 'pii' 键 (table_X_cell_Y_Z)
- 'pii' 键与既有 'ocr' 键并存, ocr 键不被覆盖 (D-11 + D-12)
- 2 个 WordWorker 实例的 _pii_engine 是不同实例 (Pitfall 8 实例独立性)
- 取消场景下 word_data[key]['pii'] 仍是 list (类型不变, 不抛 KeyError)

不变量 (D-17):
- tests.unit.test_word_replace_rules 与 test_batch_word_replace 不被破坏
- tests.unit.test_word_pii_adapter (Plan 1) 全部 PASS
"""
import unittest

from docx import Document

from privacyguard.workers.word_worker import WordWorker
from tests.fixtures.fake_pii import fake_id_card, fake_phone


def _build_docx_with_paragraphs(paragraphs):
    """构造一个只含段落的 docx (无表格)。"""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


def _build_docx_with_table_and_paragraphs(paragraphs, table_rows=1, table_cols=1):
    """构造一个含段落 + 表格的 docx。"""
    doc = _build_docx_with_paragraphs(paragraphs)
    table = doc.add_table(rows=table_rows, cols=table_cols)
    return doc, table


def _build_word_data_from_doc(doc):
    """按 worker 期望的 key 格式构造 word_data dict (与 main.py:_open_word_docx 对齐)。"""
    word_data = {}
    for idx, para in enumerate(doc.paragraphs):
        word_data[f'paragraph_{idx}'] = {
            'type': 'paragraph',
            'index': idx,
            'text': para.text,
            'ocr': [],
            'manual': [],
            'pii': [],
        }
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                key = f'table_{table_idx}_cell_{row_idx}_{cell_idx}'
                word_data[key] = {
                    'type': 'table_cell',
                    'table': table_idx,
                    'row': row_idx,
                    'cell': cell_idx,
                    'text': cell.text,
                    'ocr': [],
                    'manual': [],
                    'pii': [],
                }
    return word_data


class TestModularWordWorkerPii(unittest.TestCase):
    """D-12: WordWorker 接入 PII 引擎 + word_data 字段扩展验证."""

    # ------------------------------------------------------------------
    # Test 1: 段落路径 PII 写入
    # ------------------------------------------------------------------
    def test_worker_writes_pii_key_after_run(self):
        """段落含 1 个身份证号, run() 后 word_data['paragraph_0']['pii'] 命中非空."""
        secret_id = fake_id_card()
        text = f"测试段落 {secret_id} 内容"
        doc = _build_docx_with_paragraphs([text])
        word_data = _build_word_data_from_doc(doc)

        worker = WordWorker(doc, word_data, [], '', '[已脱敏]')
        worker.run()

        pii_hits = word_data['paragraph_0']['pii']
        self.assertIsInstance(pii_hits, list)
        self.assertGreater(len(pii_hits), 0, "段落含身份证号, PII 命中应非空")
        entity_types = {h.entity_type for h in pii_hits}
        self.assertIn("CN_ID_CARD", entity_types)

    # ------------------------------------------------------------------
    # Test 2: 表格 cell 路径 PII 写入
    # ------------------------------------------------------------------
    def test_worker_writes_pii_for_table_cells(self):
        """表格 cell 含 1 个手机号, run() 后 word_data['table_0_cell_0_0']['pii'] 命中非空."""
        secret_phone = fake_phone()
        cell_text = f"联系方式 {secret_phone}"
        doc, table = _build_docx_with_table_and_paragraphs(["普通段落"], table_rows=1, table_cols=1)
        table.cell(0, 0).text = cell_text
        word_data = _build_word_data_from_doc(doc)

        worker = WordWorker(doc, word_data, [], '', '[已脱敏]')
        worker.run()

        pii_hits = word_data['table_0_cell_0_0']['pii']
        self.assertIsInstance(pii_hits, list)
        self.assertGreater(len(pii_hits), 0, "表格 cell 含手机号, PII 命中应非空")
        entity_types = {h.entity_type for h in pii_hits}
        self.assertIn("CN_PHONE", entity_types)

    # ------------------------------------------------------------------
    # Test 3: 与 ocr 键并存
    # ------------------------------------------------------------------
    def test_worker_preserves_ocr_key_when_writing_pii(self):
        """run() 后 ocr 键与 pii 键并存 — ocr 由 _find_matches 刷新, pii 由 PIIEngine 写入.

        锁: PII 扫描不应破坏既有 ocr 键结构 (类型与字段). 两者在同一 key 下并存.
        """
        text = f"测试段落 {fake_id_card()} 内容"
        doc = _build_docx_with_paragraphs([text])
        word_data = _build_word_data_from_doc(doc)

        worker = WordWorker(doc, word_data, [], '', '[已脱敏]')
        worker.run()

        # ocr 键保留 (类型 list; 由 _find_matches 刷新)
        self.assertIn('ocr', word_data['paragraph_0'])
        self.assertIsInstance(word_data['paragraph_0']['ocr'], list)
        # pii 键并存 (类型 list, 非空)
        self.assertIn('pii', word_data['paragraph_0'])
        self.assertIsInstance(word_data['paragraph_0']['pii'], list)
        self.assertGreater(len(word_data['paragraph_0']['pii']), 0)
        # 两个键独立 — 互不影响类型与存在性
        self.assertNotEqual(id(word_data['paragraph_0']['ocr']),
                            id(word_data['paragraph_0']['pii']))

    # ------------------------------------------------------------------
    # Test 4: 取消场景
    # ------------------------------------------------------------------
    def test_worker_cancellation_still_emits_partial_results(self):
        """请求取消后, run() 不抛 KeyError; 已扫描 key 的 pii 键是 list."""
        paragraphs = [f"段落 {i} 内容" for i in range(5)]
        doc = _build_docx_with_paragraphs(paragraphs)
        word_data = _build_word_data_from_doc(doc)

        worker = WordWorker(doc, word_data, [], '', '[已脱敏]')
        # QThread.run() 同步执行, requestInterruption() 在 run() 之前调用
        worker.requestInterruption()
        worker.run()

        # 已扫描的 paragraph_0 仍应有 'pii' 键 (类型 list)
        self.assertIn('paragraph_0', word_data)
        self.assertIsInstance(word_data['paragraph_0']['pii'], list)
        # 取消后部分段落可能未处理, 但已处理段落的字段类型一致
        for key, value in word_data.items():
            if key.startswith('paragraph_'):
                self.assertIn('pii', value, f"取消后 key={key} 缺 pii 键")
                self.assertIsInstance(value['pii'], list)

    # ------------------------------------------------------------------
    # Test 5: PIIEngine 实例独立性
    # ------------------------------------------------------------------
    def test_worker_pii_engine_loaded_once_per_instance(self):
        """2 个 WordWorker 实例的 _pii_engine id 不同 (Pitfall 8 实例独立)."""
        doc = _build_docx_with_paragraphs([f"段落 {fake_id_card()}"])
        word_data = _build_word_data_from_doc(doc)
        worker_a = WordWorker(doc, word_data, [], '', '[已脱敏]')

        doc2 = _build_docx_with_paragraphs([f"段落 {fake_id_card()}"])
        word_data2 = _build_word_data_from_doc(doc2)
        worker_b = WordWorker(doc2, word_data2, [], '', '[已脱敏]')

        # 实例独立
        self.assertIsNot(worker_a._pii_engine, worker_b._pii_engine)
        self.assertNotEqual(id(worker_a._pii_engine), id(worker_b._pii_engine))

        # 调用 run() 验证 pii 命中非空 (间接证明 _pii_engine 在 run 中被复用)
        worker_a.run()
        self.assertGreater(len(word_data['paragraph_0']['pii']), 0)


if __name__ == "__main__":
    unittest.main()
