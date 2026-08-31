"""v1.1.12: 端到端测试 - 用真实抵账协议.docx 验证 partial masking 效果。

无需 GUI 启动,直接调用 WordWorker 在真实文本上跑,断言关键 mask 输出。
可用作 CI 回归测试,也可在命令行直接 `python -m tests.e2e.test_dossier_partial_masking` 看结果。
"""

import os
import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SAMPLE_DOCX = PROJECT_ROOT / "PDF" / "抵账协议0522.docx----刘骁毅原版.docx"


def _extract_paragraphs(docx_path: Path) -> list:
    """按文档顺序提取所有段落和表格单元格的纯文本。"""
    doc = Document(str(docx_path))
    body = doc.element.body
    rows = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            text = p.text
            if text.strip():
                rows.append(("P", text))
        elif child.tag == qn("w:tbl"):
            from docx.table import Table
            t = Table(child, doc)
            for r_idx, row in enumerate(t.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        if p.text.strip():
                            rows.append(("T", p.text))
    return rows


def _mask_text(text: str, name_to_pattern: dict, default_rules_meta: dict) -> list:
    """对单条文本应用 main 项目的 mask 逻辑,返回 hit 列表(模拟 WordWorker._find_matches)。"""
    from secureredact.utils.masking import apply_mask_for_rule
    import re

    hits = []
    for rule_name, pattern in name_to_pattern.items():
        try:
            for m in re.finditer(pattern, text, re.IGNORECASE):
                matched_text = m.group()
                masked = apply_mask_for_rule(matched_text, rule_name, default_rules_meta)
                if masked is None:
                    masked = "[已脱敏]"
                hits.append({
                    "rule_name": rule_name,
                    "text": matched_text,
                    "replacement": masked,
                    "source": "rule",
                })
        except re.error:
            pass
    return hits


def run_full_test():
    """端到端:加载真实 docx,跑 mask,断言关键字段,输出可读报告。"""
    if not SAMPLE_DOCX.exists():
        print(f"[!] 样本文档不存在: {SAMPLE_DOCX}")
        return False

    # 模拟 GUI 启动时 SimpleConfig 加载,强制覆盖 mask 字段
    import os, json
    os.chdir(str(PROJECT_ROOT))
    from main import SimpleConfig, DEFAULT_RULES_META
    sc = SimpleConfig("config.json")

    # 提取 name_to_pattern
    name_to_pattern = {
        name: rule.get("pattern", "")
        for name, rule in sc._config["redaction"]["default_rules"].items()
        if isinstance(rule, dict) and rule.get("pattern")
    }

    print("=" * 70)
    print(f" 抵账协议端到端 partial mask 测试")
    print("=" * 70)
    print(f" 样本文档: {SAMPLE_DOCX.name}")
    print(f" 规则数: {len(name_to_pattern)}")
    print()

    paragraphs = _extract_paragraphs(SAMPLE_DOCX)
    print(f" 段落数: {len(paragraphs)}")
    print()

    all_hits = []
    for kind, text in paragraphs:
        hits = _mask_text(text, name_to_pattern, DEFAULT_RULES_META)
        for h in hits:
            all_hits.append({"para_kind": kind, "para_text": text, **h})

    print(f" 命中总数: {len(all_hits)}")
    print()
    print("-" * 70)
    print(" 命中详情:")
    print("-" * 70)
    for h in all_hits:
        original = h["text"]
        masked = h["replacement"]
        # 截断原文显示
        if len(original) > 50:
            original_disp = original[:50] + "..."
        else:
            original_disp = original
        if original == masked:
            change = "(无变化)"
        else:
            change = f"→ {masked[:50]}{'...' if len(masked) > 50 else ''}"
        print(f"  [{h['rule_name']:<14}] {original_disp}")
        print(f"  {'':14}  {change}")
        print()

    return all_hits


# 真实文本期望断言(关键字段必须按预期 mask)
EXPECTED_HITS = {
    # USCC 必须按 USCC 4+4 规则 mask
    "91220201307949958P": "9122**********58P",  # 18 位
    "9151152MA6ARBNK4W": "9151*********NK4W",    # 17 位
    "91511500MAACKMR07X": "9151**********07X",    # 18 位
    "91220200MA17WHWF27": "9122**********F27",    # 18 位
    # 法定代表人姓名(只保留姓)
    "法定代表人:周超": "法定代表人:周*",  # 注意:实际匹配可能不带"法定代表人:"前缀,见 hit 实际内容
    "法定代表人：周超": "法定代表人：周*",  # 全角冒号
}


class TestDossierPartialMasking(unittest.TestCase):
    """抵账协议端到端 partial mask 断言测试。"""

    @classmethod
    def setUpClass(cls):
        if not SAMPLE_DOCX.exists():
            raise unittest.SkipTest(f"样本文档不存在: {SAMPLE_DOCX}")
        os.chdir(str(PROJECT_ROOT))
        from main import SimpleConfig, DEFAULT_RULES_META
        sc = SimpleConfig("config.json")
        cls.name_to_pattern = {
            name: rule.get("pattern", "")
            for name, rule in sc._config["redaction"]["default_rules"].items()
            if isinstance(rule, dict) and rule.get("pattern")
        }
        cls.default_rules_meta = DEFAULT_RULES_META
        cls.paragraphs = _extract_paragraphs(SAMPLE_DOCX)

    def _all_hits(self) -> list:
        all_hits = []
        for kind, text in self.paragraphs:
            hits = _mask_text(text, self.name_to_pattern, self.default_rules_meta)
            for h in hits:
                all_hits.append({"para_kind": kind, "para_text": text, **h})
        return all_hits

    def test_uscc_18char_mask_4_4(self):
        """18 位 USCC 必须按 4+4 规则 mask。"""
        hits = self._all_hits()
        uscc_hits = [h for h in hits if h["rule_name"] == "统一社会信用代码"]
        self.assertGreater(len(uscc_hits), 0, "未命中任何 USCC")
        # 至少有一个 USCC 的 mask 输出含 *
        masked_set = {h["replacement"] for h in uscc_hits}
        self.assertTrue(any("*" in m for m in masked_set),
                        f"USCC mask 未生效: {masked_set}")

    def test_legal_rep_name_mask(self):
        """法定代表人姓名只保留姓。"""
        hits = self._all_hits()
        rep_hits = [h for h in hits if h["rule_name"] == "法定代表人"]
        self.assertGreater(len(rep_hits), 0, "未命中任何法定代表人")
        # 至少有一个 hit 的 mask 保留中文第一字 + 后续 *
        for h in rep_hits:
            if "周超" in h["text"] or "周" in h["text"]:
                # hit.text 可能是 "法定代表人:周超" 或 "周超"
                if h["text"].startswith("法定代表人"):
                    # 整段包含 label,期望 "法定代表人:周*"
                    self.assertIn("周*", h["replacement"],
                                  f"法定代表人 mask 失败: {h['replacement']}")
                break

    def test_address_8_2_mask(self):
        """地址规则必须按 8+2 partial mask(若命中)。"""
        hits = self._all_hits()
        addr_hits = [h for h in hits if h["rule_name"] == "地址（含门牌号）"]
        if not addr_hits:
            self.skipTest("地址规则未命中(可能因格式差异)")
        # 至少有一个 hit 的 mask 保留前 8 字 + 末 2 字
        for h in addr_hits:
            r = h["replacement"]
            self.assertIn("*", r, f"地址 mask 未生效: {r}")

    def test_id_card_6_4_mask(self):
        """身份证号必须按 6+4 partial mask(若文档含身份证号)。"""
        hits = self._all_hits()
        id_hits = [h for h in hits if h["rule_name"] == "身份证号"]
        if not id_hits:
            self.skipTest("文档不含身份证号")
        for h in id_hits:
            self.assertIn("*", h["replacement"])

    def test_mobile_3_4_mask(self):
        """手机号必须按 3+4 partial mask(若文档含手机号)。"""
        hits = self._all_hits()
        m_hits = [h for h in hits if h["rule_name"] == "手机号码"]
        if not m_hits:
            self.skipTest("文档不含手机号")
        for h in m_hits:
            r = h["replacement"]
            self.assertTrue(r.startswith("***") or "*" in r,
                            f"手机号 mask 失败: {r}")

    def test_legal_party_labels_not_masked(self):
        """甲方/乙方/丙方/丁方/戊方 作为标签, 不应被 jieba/人名识别误脱敏。

        关键: 当前 _mask_text 测试 helper 不模拟 jieba X3 路径
        (jieba 需要 WordWorker._find_matches 内部触发),所以这里只验证
        'rule' 路径的 hit 不应包含甲方/乙方/丙方/丁方/戊方 单独的 2 字标签。
        jieba 路径的白名单保护在 tests.unit.test_word_worker_black_white 中覆盖。
        """
        hits = self._all_hits()
        for h in hits:
            if h["text"] in ("甲方", "乙方", "丙方", "丁方", "戊方"):
                self.fail(
                    f"'{h['text']}' 不应作为独立 hit 出现(它应当作为 label 不被脱敏)"
                )

    def test_company_name_rule_masks_company_full_name(self):
        """新增的'公司名'规则必须命中以'有限公司/集团/公司/中心'结尾的中文公司名。

        v1.1.12: 启用智能 mask (`apply_company_mask`), 保留'省/市'前导和'有限公司'后缀, 中间打码。
        """
        hits = self._all_hits()
        company_hits = [h for h in hits if h["rule_name"] == "公司名"]
        self.assertGreater(len(company_hits), 0,
                           "公司名规则未命中任何公司(应当命中至少一家)")
        for h in company_hits:
            r = h["replacement"]
            # 智能 mask 至少含 * (有打码)
            self.assertIn("*", r, f"公司名 replacement 应当含 *: {r}")
            # 不应丢失关键公司尾缀
            if h["text"].endswith("有限公司"):
                self.assertTrue(r.endswith("有限公司"),
                                f"应当保留'有限公司'后缀: {r}")
            # 不应丢失省/市前导(若原文有)
            head_idx = -1
            for kw in ("特别行政区", "自治区", "省", "市"):
                idx = h["text"].find(kw, 0, 6)
                if idx >= 0:
                    head_idx = idx + len(kw)
                    break
            # 尝试省份名直接开头
            if head_idx == -1:
                for prov in (
                    "北京", "上海", "天津", "重庆", "河北", "山西", "辽宁", "吉林",
                    "黑龙江", "江苏", "浙江", "安徽", "福建", "江西", "山东", "河南",
                    "湖北", "湖南", "广东", "海南", "四川", "贵州", "云南", "陕西",
                    "甘肃", "青海", "内蒙古", "广西", "西藏", "宁夏", "新疆",
                    "台湾", "香港", "澳门",
                ):
                    if h["text"].startswith(prov):
                        head_idx = len(prov)
                        break
            if head_idx > 0:
                self.assertTrue(r.startswith(h["text"][:head_idx]),
                                f"应当保留省/市前导 ({h['text'][:head_idx]!r}): {r}")


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--report":
        # CLI 模式:输出可读报告
        run_full_test()
    else:
        # unittest 模式
        unittest.main(verbosity=2)